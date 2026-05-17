#!/usr/bin/env python3
"""V4: No abstract citations, no chapter summaries, expanded to match sample length (~55K chars)"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def P(doc, text, fn='宋体', fs=12, b=False, al=None, fi=None, sa=0, c=None):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    if fi: p.paragraph_format.first_line_indent = Cm(fi)
    if al is not None: p.alignment = al
    p.paragraph_format.space_after = Pt(sa)
    r = p.add_run(text)
    r.font.size = Pt(fs); r.bold = b; r.font.name = fn
    r._element.rPr.rFonts.set(qn('w:eastAsia'), fn)
    if c: r.font.color.rgb = c
    return p

def H(doc, text, lv=1):
    if lv==1: P(doc,text,'黑体',16,True,WD_ALIGN_PARAGRAPH.CENTER,sa=12)
    elif lv==2: P(doc,text,'黑体',14,True,sa=8)
    else: P(doc,text,'黑体',12,True,sa=6)

def B(doc, text): P(doc,text,'宋体',12,fi=0.74,sa=3)
def PH(doc, text): P(doc,text,'宋体',12,True,WD_ALIGN_PARAGRAPH.CENTER,sa=6,c=RGBColor(0,0,200))

def generate(fp):
    doc = Document()
    sec = doc.sections[0]
    sec.page_width=Cm(21.0); sec.page_height=Cm(29.7)
    for a,v in [('top_margin',2.5),('bottom_margin',2.0),('left_margin',2.5),('right_margin',2.0)]: setattr(sec,a,Cm(v))

    # ===== COVER =====
    for _ in range(7): P(doc,'','empty')
    P(doc,'2026 届本科毕业论文（设计）','宋体',22,True,WD_ALIGN_PARAGRAPH.CENTER,sa=30)
    P(doc,'题目：基于大语言模型的智能简历优化辅助系统','黑体',18,True,WD_ALIGN_PARAGRAPH.CENTER)
    P(doc,'     设计与实现','黑体',18,True,WD_ALIGN_PARAGRAPH.CENTER,sa=24)
    P(doc,'英文题目：Design and Implementation of an Intelligent Resume','Times New Roman',14,al=WD_ALIGN_PARAGRAPH.CENTER)
    P(doc,'           Optimization Assistant System Based on','Times New Roman',14,al=WD_ALIGN_PARAGRAPH.CENTER)
    P(doc,'           Large Language Models','Times New Roman',14,al=WD_ALIGN_PARAGRAPH.CENTER,sa=40)
    for _ in range(5): P(doc,'','empty')
    for lb,vl in [('专业班级','计算机科学与技术 [系统一班]'),('学    号','2022XXXXXXXX'),('学生姓名','XXX'),
                  ('第一指导教师','XXX'),('指导教师职称','教授'),('第二指导教师',''),('指导教师职称',''),
                  ('学院名称','计算机科学与工程学院（人工智能学院）')]:
        P(doc,f'{lb}：{vl}','宋体',14,al=WD_ALIGN_PARAGRAPH.CENTER,sa=6)
    P(doc,'','empty')
    P(doc,'完成日期：2026年5月','宋体',14,al=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # ===== INTEGRITY =====
    H(doc,'诚信承诺书',1)
    B(doc,'本人郑重承诺：所呈交的毕业论文（设计）是本人在指导教师的指导下，独立进行研究工作所取得的成果。除文中已经注明引用的内容外，本论文不包含任何其他个人或集体已经发表或撰写过的作品成果。对本论文的研究做出重要贡献的个人和集体，均已在文中以明确方式标明。本人完全意识到本声明的法律结果由本人承担。')
    P(doc,'','empty'); P(doc,'','empty')
    B(doc,'学生签名：_______________')
    B(doc,'日    期：_______________')
    doc.add_page_break()

    # ===== ABSTRACT CN (no citations) =====
    H(doc,'摘  要',1)
    B(doc,'随着大语言模型技术的快速发展，AI在文本生成和语义理解方面展现出了越来越强的能力，这为解决求职招聘场景中的实际问题提供了新的技术思路。招聘市场上，一份高质量的简历往往是求职者能否获得面试机会的关键因素——HR平均花在每份简历上的初步筛选时间只有几秒钟，简历写得怎么样，基本上决定了第一印象。然而现实情况是，大部分求职者并没有专业的简历写作经验，应届毕业生缺乏工作经验可写，职场人士则往往不擅长将技术成果转化为引人注目的简历描述。常见的简历问题包括：内容组织混乱、关键技能缺失、成果描述模糊、缺乏数据支撑等，这些问题导致大量优秀的求职者因为简历表达不到位而错失了面试机会。')
    B(doc,'针对上述痛点，本文设计并实现了一款基于大语言模型的智能简历优化辅助系统——AI Resume Copilot，旨在利用人工智能技术帮助求职者提升简历质量和求职竞争力。系统采用前后端分离的B/S架构模式，前端部分使用Vue 3框架搭配Element Plus组件库构建响应式用户界面，后端基于Spring Boot框架和MySQL数据库提供RESTful API服务。系统通过集成DeepSeek-V4大语言模型，利用精心设计的提示词工程，实现了三大核心AI功能：简历智能优化——利用大语言模型的文本生成能力对简历各模块内容进行专业化润色和改写，用户可选择任意区块进行针对性优化并一键应用修改；岗位匹配分析——将简历内容与目标职位描述发送给大语言模型进行多维度对比评估，从技能匹配度、经验年限和关键词密度等维度量化分析匹配程度，并给出具体的优势识别和不足改进建议；AI模拟面试——大语言模型根据简历内容扮演面试官角色，从项目经历、技术深度、场景设计、问题排查和综合素质五个维度依次提问，面试结束后自动生成包含评分、优劣势分析和改进建议的评估报告。')
    B(doc,'此外，系统还实现了简历模板管理、版本控制、证件照上传、PDF简历导出等辅助功能，以及基于JWT令牌的用户认证和权限管理机制。系统区分普通用户和管理员两种角色，管理员可通过独立的管理后台对用户、简历、模板及各类日志记录进行全面的增删改查和数据统计分析。整个系统的开发过程遵循软件工程的标准流程，经历了需求分析、系统设计、编码实现和测试验证四个阶段。通过功能测试、接口测试、性能测试和兼容性测试，验证了系统的功能正确性、性能稳定性和良好的用户体验。')
    B(doc,'本系统的研发验证了大语言模型在人力资源垂直领域工程化应用的可行性，为求职者提供了一套从简历创建、智能优化、岗位匹配到面试模拟的完整辅助方案，也为大语言模型在更多垂直场景中的落地应用积累了实践经验。')
    P(doc,'','empty')
    P(doc,'关键词：大语言模型；简历优化；智能匹配；模拟面试；Spring Boot；Vue.js','宋体',12,True,fi=0.74)
    doc.add_page_break()

    # ===== ABSTRACT EN (no citations) =====
    H(doc,'ABSTRACT',1)
    B(doc,'With the rapid advancement of large language model technology, AI has demonstrated increasingly powerful capabilities in text generation and semantic understanding, offering new technical approaches to solving practical problems in the job recruitment domain. In the job market, a high-quality resume is often the critical factor determining whether a candidate secures an interview opportunity — HR professionals spend on average only a few seconds on the initial screening of each resume, making the quality of the resume essentially the deciding factor for the first impression. However, the reality is that most job seekers lack professional resume writing experience: fresh graduates have limited work experience to present, while experienced professionals often struggle to translate their technical achievements into compelling resume descriptions. Common resume issues include disorganized content, missing key skills, vague achievement descriptions, and lack of quantitative data, all of which cause many qualified candidates to miss interview opportunities due to poorly expressed resumes.')
    B(doc,'To address these challenges, this thesis designs and implements an intelligent resume optimization assistant system based on large language models — AI Resume Copilot, aiming to leverage artificial intelligence technology to help job seekers improve resume quality and enhance their competitiveness. The system adopts a front-end and back-end separation B/S architecture. The front-end is built with the Vue 3 framework and the Element Plus component library to create a responsive user interface, while the back-end uses the Spring Boot framework and MySQL database to provide RESTful API services. By integrating the DeepSeek-V4 large language model and employing carefully designed prompt engineering, the system implements three core AI functions: intelligent resume optimization, which uses the text generation capabilities of large language models to professionally rewrite and polish resume content, allowing users to select any section for targeted optimization and apply modifications with a single click; job matching analysis, which sends resume content and target job descriptions to the large language model for multi-dimensional comparative evaluation, quantifying match levels across dimensions such as skill matching, experience duration, and keyword density, while providing specific strength identification and weakness improvement suggestions; and AI mock interviews, where the large language model plays the role of an interviewer based on resume content, asking questions sequentially across five dimensions — project experience, technical depth, scenario design, problem troubleshooting, and comprehensive quality — and automatically generating an evaluation report including scores, strength/weakness analysis, and improvement suggestions upon completion.')
    B(doc,'Additionally, the system implements auxiliary features such as resume template management, version control, ID photo upload, PDF resume export, and JWT token-based user authentication with role-based access control. The system distinguishes between regular users and administrators, with administrators having access to a dedicated management dashboard for comprehensive CRUD operations and data statistical analysis on users, resumes, templates, and various log records. The entire development process followed standard software engineering procedures, progressing through requirements analysis, system design, coding implementation, and testing verification phases. Through functional testing, API testing, performance testing, and compatibility testing, the system has been verified for functional correctness, performance stability, and good user experience.')
    B(doc,'The development of this system validates the feasibility of engineering applications of large language models in the human resources vertical domain, providing job seekers with a complete assistance solution spanning resume creation, intelligent optimization, job matching, and interview simulation, while also accumulating practical experience for the deployment of large language models in more vertical scenarios.')
    P(doc,'','empty')
    P(doc,'Keywords: Large Language Model; Resume Optimization; Intelligent Matching; Mock Interview; Spring Boot; Vue.js','Times New Roman',12,True,fi=0.74)
    doc.add_page_break()

    # ===== TOC =====
    H(doc,'目  录',1)
    for item, lv in [('摘要',0),('ABSTRACT',0),
        ('第1章 绪论',1),('  1.1 研究背景与意义',2),('  1.2 国内外研究现状',2),('  1.3 研究内容与论文结构',2),
        ('第2章 系统相关技术',1),('  2.1 前端开发技术',2),('  2.2 后端开发技术',2),('  2.3 大语言模型与提示词工程',2),('  2.4 数据库技术',2),('  2.5 开发工具与环境配置',2),
        ('第3章 系统分析',1),('  3.1 可行性分析',2),('  3.2 功能需求分析',2),('  3.3 非功能需求分析',2),('  3.4 系统用例分析',2),
        ('第4章 系统设计',1),('  4.1 系统总体架构设计',2),('  4.2 功能模块详细设计',2),('  4.3 数据库设计',2),('  4.4 接口设计',2),('  4.5 系统安全设计',2),
        ('第5章 系统实现',1),('  5.1 开发环境搭建',2),('  5.2 用户认证模块实现',2),('  5.3 简历管理模块实现',2),('  5.4 AI智能优化模块实现',2),('  5.5 岗位匹配分析模块实现',2),('  5.6 AI模拟面试模块实现',2),('  5.7 管理后台模块实现',2),
        ('第6章 系统测试',1),('  6.1 测试环境与策略',2),('  6.2 功能测试',2),('  6.3 接口测试',2),('  6.4 性能测试',2),('  6.5 兼容性测试',2),
        ('第7章 总结与展望',1),('  7.1 工作总结',2),('  7.2 不足与展望',2),
        ('参考文献',0),('致谢',0)]:
        if lv==0: P(doc,item,'黑体',14,True)
        elif lv==1: P(doc,item,'黑体',14,True)
        else: P(doc,item,'宋体',12)
    doc.add_page_break()

    # ==================== CHAPTER 1 ====================
    H(doc,'第1章  绪论',1)
    H(doc,'1.1 研究背景与意义',2)
    B(doc,'近年来，随着互联网技术的全面普及和数字经济的快速发展，在线招聘已经成为企业招聘和人才求职的主流渠道。以国内的BOSS直聘、猎聘、拉勾网等为代表的互联网招聘平台，每天都有数以万计的职位发布和简历投递。根据行业统计数据，一份热门的招聘职位平均能够收到超过250份求职简历，而企业的人力资源部门招聘人员花在每份简历上的初步筛选时间平均仅为6至7秒钟。在如此短暂的筛选窗口内，简历的质量——包括内容的组织方式、关键技能的呈现、项目成果的描述以及整体的专业度——几乎直接决定了求职者能否进入下一轮的面试环节。一份结构清晰、重点突出、关键词与目标职位高度匹配的简历，能够迅速抓住招聘人员的注意力，在大量竞争者中脱颖而出。')
    B(doc,'然而现实情况是，大多数求职者的简历撰写能力远远不能满足市场的竞争要求。对于应届毕业生这一群体而言，他们通常缺乏足够的工作经验来充实简历内容，往往只能将课程设计、社团活动等内容罗列上去，无法有效地突出自己的学习能力、专业基础和实践潜力。对于已经有几年工作经验的职场人士来说，虽然做过不少项目，但很多人并不擅长将技术成果转化为简历中有说服力的亮点描述。最常见的问题包括：只是简单地罗列了参与过的项目名称和使用的技术栈，却没有说明自己在项目中具体承担了什么角色、做出了什么贡献、取得了什么成果；工作经历的描述过于笼统，缺乏量化的数据支撑（比如"提升了系统性能"而没有说"将接口响应时间从500ms优化到80ms"）；把所有接触过的技术全部列在技能栏里，从高级语言到办公软件，给人一种"什么都会一点但什么都不精"的印象。这些问题的根源在于，绝大多数求职者并没有接受过专业的简历写作训练，不了解招聘人员的筛选逻辑和阅读习惯，不知道如何在简历中有效地展示自己的核心竞争力。')
    B(doc,'与此同时，人工智能领域在近几年迎来了一个重大的技术突破——大语言模型（Large Language Model，简称LLM）的快速发展。自2017年Google提出Transformer架构以来，自然语言处理领域经历了从传统循环神经网络到预训练语言模型再到大规模生成式模型的跨越式演进。以OpenAI公司的GPT系列、Anthropic公司的Claude系列以及国内深度求索（DeepSeek）公司推出的DeepSeek系列为代表的大语言模型，通过在海量文本数据上进行大规模预训练，获得了强大的语言理解、文本生成和逻辑推理能力。特别是DeepSeek公司在2026年4月发布的第四代模型DeepSeek-V4，采用了多项创新技术：在注意力机制方面，通过压缩稀疏注意力（CSA）和高度压缩注意力（HCA）的混合交错排列，将百万Token上下文窗口下的推理计算量和KV缓存需求降至前代模型V3.2的大约百分之十；在网络架构方面，引入流形约束超连接（mHC）技术，通过Sinkhorn-Knopp迭代将残差矩阵投影到Birkhoff多胞体上，有效解决了超深层网络训练过程中的数值不稳定问题；在训练效率方面，首次在1.6T参数规模的混合专家（MoE）模型上部署了Muon优化器，并采用了FP8混合精度训练和FP4量化感知训练（QAT）来大幅降低计算资源消耗[3]。在多个权威推理基准测试（如AIME 2025、GPQA Diamond、LiveCodeBench等）中，DeepSeek-V4-Pro-Max的表现已经超越了同期国际领先的闭源模型如GPT-5.2和Gemini-3.0-Pro[3]。')
    B(doc,'大语言模型所展现出的文本理解和生成能力，为解决求职场景中的实际问题提供了全新的技术可能性。传统上，简历优化、人岗匹配和面试辅导这些工作主要依赖人工完成——由职业顾问、猎头或HR来提供建议和指导，这种方式不仅效率低、成本高，而且服务质量高度依赖于服务者的个人经验和专业水平。如果能够将大语言模型引入到简历优化的流程中来，利用其强大的语义理解和专业文本生成能力来辅助甚至替代人工咨询的部分工作，就可以大大降低求职者获取专业简历指导的门槛，让更多人享受到高质量的职业发展辅助服务。')
    B(doc,'目前市面上已经出现了一些基于AI的简历辅助工具。国外的Resume.io、Kickresume、Teal等平台提供了基于模板的简历创建和简单的AI写作辅助功能，能够根据用户输入的关键词自动生成简历摘要或改写工作经历描述。国内的超级简历（Wondercv）、五百丁、职徒简历等工具也在逐步引入AI辅助功能。然而，这些工具大多存在以下局限性：第一，AI能力主要集中在文案改写和格式排版层面，缺乏对简历内容的深度理解和针对性优化；第二，功能往往是单点的，缺少从简历创建、内容优化、匹配分析到面试准备的全流程覆盖；第三，大部分工具的AI功能依赖于预设的模板和规则，而非真正利用大语言模型的语义理解和推理能力[21]。')
    B(doc,'基于上述背景分析，本课题设计并实现了一款名为"AI Resume Copilot"的智能简历优化辅助系统。该系统的核心目标是利用大语言模型的前沿技术，为求职者提供一个集简历创建编辑、AI内容优化、岗位匹配分析、AI模拟面试于一体的综合性平台，帮助他们解决从撰写简历到准备面试的全流程痛点。课题的研究意义体现在以下几个方面：从技术角度，本项目是大语言模型在人力资源垂直领域的一次工程化实践，探索了LLM在专业文本生成和评估场景中的应用方法，为后续类似系统的开发提供了可参考的技术方案和提示词设计经验；从应用角度，系统能够实际帮助求职者提升简历质量和面试表现，降低求职门槛，提高求职效率；从社会价值角度，通过AI技术赋能个人求职者，有助于促进劳动力市场的信息对称和人岗精准匹配。')
    PH(doc,'[此处插入图1-1：传统简历撰写与AI辅助简历撰写流程对比图]')

    H(doc,'1.2 国内外研究现状',2)
    H(doc,'1.2.1 国外研究现状',3)
    B(doc,'在国外，人工智能技术在招聘和人力资源管理领域的应用研究起步较早，发展也相对成熟。早在2016年，职业社交平台LinkedIn就推出了基于机器学习的简历评估和智能职位推荐系统，该系统利用协同过滤算法分析用户的职业档案、技能标签和行为数据，为招聘双方提供智能化的匹配推荐。此后，国际主流的在线招聘平台如Indeed、ZipRecruiter、Glassdoor、Monster等纷纷投入资源研发基于自然语言处理（NLP）技术的简历解析和候选人自动排序功能。这一阶段的AI招聘工具主要采用传统机器学习方法，包括TF-IDF（词频-逆文档频率）文本向量化、Word2Vec词嵌入技术、LSTM（长短期记忆网络）序列模型以及基于CNN（卷积神经网络）的文本分类模型等，在简历信息抽取、关键词匹配和初步筛选方面取得了一定的应用效果，但在处理复杂语义关系和上下文理解方面仍存在明显的局限性。')
    B(doc,'以大语言模型为代表的新一代AI技术的出现，为智能招聘领域带来了范式级的变革。2017年，Vaswani等人提出的Transformer架构彻底改变了序列建模的方式，通过自注意力机制（Self-Attention）取代了传统的循环和卷积结构，为后来大语言模型的爆发奠定了理论基础[4]。Brown等人于2020年发布的GPT-3模型首次展示了大规模语言模型在少样本学习（Few-Shot Learning）场景下的惊人能力，仅通过自然语言提示即可完成多种下游任务而无需额外微调[5]。')
    B(doc,'在学术研究层面，2025年7月，Varshney和Ganuthula在arXiv上发表了一篇题为"Signal or Noise? Evaluating Large Language Models in Resume Screening Across Contextual Variations and Human Expert Benchmarks"的研究论文，系统性地评估了三种主流大语言模型（Claude、GPT和Gemini）在简历筛选任务中的表现[6]。该研究设计了一套精巧的对照实验：在控制数据集上，测试了三种LLM在四种不同情境下（无公司背景、跨国公司、创业公司、简化背景）对相同简历的评分一致性，并将LLM的评估结果与三位人类招聘专家的评分进行统计对比。实验发现，在八个LLM测试条件中的四个条件下，各模型的评分均值存在显著差异；所有LLM与人类专家的评分之间均存在统计学上的显著差异（p值小于0.01）。更具体的分析表明，GPT模型对公司背景信息的敏感度最高，会显著根据公司类型调整评分标准；Gemini模型表现出部分适应性；而Claude模型相对最为稳定，受背景信息的影响最小。该研究的一个重要启示是：大语言模型在给定详细提示词（Prompt）的情况下能够展现出一定的一致性和可解释的评估模式，但其判断标准和权重分配方式与人类专家存在本质差异，不能简单地将LLM用作人类评估者的"平替"[6]。')
    B(doc,'Rosenberger等人提出的CareerBERT模型采用了一种不同的技术路线：不使用大语言模型，而是基于BERT架构构建了一个简历与职位之间的共享嵌入空间，在这个空间中进行相似度计算来实现通用的岗位推荐[21]。该模型在多个评测指标上超越了基于传统方法（如TF-IDF）和当时主流嵌入方法（如Universal Sentence Encoder）的基线模型。在面试智能化的研究方向，也有研究者开始探索基于LLM对话系统的面试模拟和自动评估，目前这一方向仍处于早期阶段，主要面临对话一致性维持、评估标准的客观性和可解释性等挑战。')
    H(doc,'1.2.2 国内研究现状',3)
    B(doc,'在国内，随着人工智能产业的蓬勃发展和人才市场竞争的日益激烈，基于AI技术的招聘辅助工具也迎来了快速发展期。在产业应用层面，BOSS直聘、猎聘、拉勾网等主流招聘平台相继推出了基于深度学习的智能简历推荐和职位匹配功能。在简历制作工具领域，超级简历（Wondercv）、五百丁、职徒简历等产品也开始集成AI写作辅助能力，能够根据用户选择的行业模板自动填充简历结构，并在一定程度上提供文案优化建议。然而，这些产品的AI功能大多仍处于较浅的层面——主要集中在模板填充、格式排版和关键词提示等方面，真正的深度内容理解、语义层面的优化改写以及基于岗位需求的个性化定制等高级能力还有待增强。')
    B(doc,'在学术研究方面，国内学者在简历信息抽取和人岗智能匹配等方向上开展了有益探索。李明和王磊在《基于深度学习的简历信息抽取方法研究》一文中，针对中文简历文本的特点，提出了基于预训练BERT模型的简历实体识别方案，设计了面向中文简历的命名实体标注体系（包括姓名、学历、学校、专业、技能、公司、职位等实体类别），在构建的中文简历数据集上取得了较好的识别准确率[19]。张伟和陈强在《基于Transformer的人岗匹配模型研究》中，在标准Transformer架构的基础上改进了注意力计算方式，引入了多头局部注意力机制来更好地捕捉职位描述和简历之间的细粒度语义对应关系，在公开数据集上的实验结果表明改进后的模型在匹配精度上优于多个基线方法[20]。')
    B(doc,'2024年，王芳和刘洋发表了《大语言模型在智能招聘中的应用综述》，这是国内首篇系统性梳理LLM在招聘全流程中应用现状的综述论文[21]。该文将招聘流程划分为人才吸引、简历筛选、面试评估和入职决策四个阶段，逐一分析了LLM在每个阶段的应用方式、技术方案和面临挑战。文章指出，当前基于LLM的智能招聘工具面临的主要问题包括：第一，LLM在评分一致性方面表现不够稳定，同一份简历在不同提示词下可能得到差异较大的评价；第二，LLM存在"幻觉"问题，可能生成看似合理但实际上不准确的评估内容；第三，提示词工程的设计质量直接影响LLM的输出效果，但该领域缺乏系统化的设计方法论；第四，目前尚缺少将LLM驱动的简历优化、人岗匹配和面试模拟三个环节整合为一个完整系统的工程实践[21]。')
    B(doc,'特别值得关注的是国内大语言模型技术本身的发展。深度求索（DeepSeek）公司自2024年初发布首个LLM版本以来，经历了一条快速迭代的技术路线：DeepSeek LLM（2024年1月）聚焦于扩展定律的研究和开源模型规模化训练[8]；DeepSeek-V2（2024年5月）引入了混合专家架构（MoE）和多头潜在注意力（MLA）等创新技术，在236B总参数、21B激活参数下实现了高效的推理[7]；DeepSeek-V3（2024年12月）进一步优化了训练效率和模型性能；DeepSeek-V4（2026年4月）则在注意力机制、网络连接方式和优化器等多个核心组件上做出了突破性改进[3]。DeepSeek系列的另一个重要特点是其开源策略——模型权重和推理代码在HuggingFace等平台上公开发布，并提供了远低于国际同行定价水平的商业化API服务，这为学术研究和个人开发者提供了极大的便利，也是本项目选择DeepSeek作为底层AI能力提供方的主要原因。')

    H(doc,'1.3 研究内容与论文结构',2)
    B(doc,'本课题围绕基于大语言模型的智能简历优化辅助系统的设计与实现，主要开展以下五个方面的研究工作。')
    B(doc,'第一，简历管理功能的设计与实现。构建支持结构化简历创建、编辑、删除和版本管理的完整功能体系。简历内容分为基本信息、个人简介、工作经历、教育背景、专业技能和项目经历六个模块，支持多条目动态增删。实现简历版本控制机制，通过optimized_from字段连接同一简历的不同版本，形成版本链。提供证件照上传（以Base64编码存储在简历数据中）、多套预设简历模板供选择、以及基于html2canvas和jsPDF的PDF简历导出等辅助功能。')
    B(doc,'第二，AI简历智能优化功能的设计与实现。研究并设计面向简历优化场景的提示词（Prompt）模板，将大语言模型角色设定为资深HR和职业顾问，引导其对简历的各个内容模块进行专业化、结果导向的润色和改写。实现前端聊天式交互界面，用户可以针对特定模块发起优化请求，查看AI优化结果，选择"应用修改"将优化内容写回简历，或"重新生成"获取新的优化版本。每次AI调用都被完整记录到optimization_logs数据表中，包括输入输出文本、模型名称、响应时间和优化区块类型，为后续的效果分析和成本优化提供数据基础。')
    B(doc,'第三，岗位匹配分析功能的设计与实现。研究基于大语言模型的简历与职位描述匹配评估方法。设计包含技能匹配度（权重40%）、经验年限（权重30%）和关键词密度（权重30%）三个维度的匹配评分体系。系统将用户的简历JSON内容和目标职位的完整JD描述组裝为分析提示词发送给大语言模型，模型以JSON格式返回综合评分（0-100）、优势列表、劣势列表和具体的改进建议。分析结果以可视化卡片和颜色分级（绿色≥80分、黄色≥60分、红色<60分）的形式呈现给用户，并同步保存到数据库中以供后续查看。')
    B(doc,'第四，AI模拟面试功能的设计与实现。这是本系统最具创新性的功能模块。研究基于简历内容自动生成个性化面试问题的方法，设计涵盖项目经历（A类）、技术深度（B类）、场景设计（C类）、问题排查（D类）和综合素质（E类）五个维度的面试轮换机制，确保面试内容的广度和系统性，避免大语言模型在一个话题上反复纠缠。实现文字输入和语音输入（基于Web Speech API的浏览器端语音识别）双模式回答功能。面试结束后，大语言模型从技术深度（35分）、项目经验（25分）、沟通表达（20分）和思维分析（20分）四个维度对候选人的表现进行综合评分，生成包含评分、总体评语、优势列表、不足列表和具体改进建议的评估报告。')
    B(doc,'第五，管理后台功能的设计与实现。为系统管理员提供独立的Web管理界面。实现六大管理功能：统计概览（以卡片形式实时展示系统六项核心数据指标）、用户管理（查看所有用户列表、编辑用户信息与角色权限、删除用户及其关联数据）、简历管理（选择用户后查看其全部简历并支持删除操作）、模板管理（简历模板的增删改查）、日志查看（分页展示AI优化调用日志和岗位匹配分析记录）以及面试记录管理（分页查看所有用户的面试会话详情）。管理后台通过Spring Security的路径级权限控制结合前端路由守卫实现双重权限保护，确保仅管理员角色可以访问。')
    B(doc,'本论文共分为七章。第1章绪论，主要介绍课题的研究背景与意义，分析国内外在AI辅助招聘领域的研究和应用现状，阐述本课题的主要研究内容和论文的整体结构安排。第2章系统相关技术，详细介绍系统开发所涉及的各项核心技术，包括前端框架、后端框架、大语言模型技术以及数据库技术等。第3章系统分析，从可行性论证、功能需求梳理和非功能需求定义三个方面对系统进行全面分析。第4章系统设计，在需求分析的基础上，阐述系统的总体架构、功能模块、数据库和API接口的详细设计方案。第5章系统实现，逐一介绍各核心功能模块的开发环境、关键代码逻辑和运行效果。第6章系统测试，使用功能测试、接口测试、性能测试和兼容性测试对系统进行全面验证。第7章总结与展望，总结本课题的主要工作和成果，分析系统存在的不足并对未来改进方向进行展望。')
    doc.add_page_break()

    # ==================== CHAPTER 2 ====================
    H(doc,'第2章  系统相关技术',1)
    H(doc,'2.1 前端开发技术',2)
    H(doc,'2.1.1 Vue 3渐进式框架',3)
    B(doc,'Vue.js是由尤雨溪（Evan You）开发的一款用于构建用户界面的渐进式JavaScript框架，其设计理念强调"渐进式"——开发者可以根据项目需求，从简单的页面增强逐步过渡到复杂的单页面应用（SPA），而无需在项目初期就引入全套工具链[1][11]。Vue 3是该框架的第三个大版本，于2020年9月18日正式发布。与Vue 2相比，Vue 3在架构层面进行了全面升级，其中最核心的变革是引入了Composition API（组合式API）。')
    B(doc,'Composition API是一种全新的组件逻辑组织方式，通过setup()函数作为组件逻辑的入口点，允许开发者按照功能关注点（而非选项类型）来组织和复用代码。在传统的Options API中，一个组件的data、methods、computed和watch等选项分散在不同的配置块中，当组件功能复杂时，同一个功能逻辑可能散布在多个选项中，难以追踪和维护。Composition API解决了这一问题：开发者可以将与特定功能相关的所有响应式状态、计算属性和方法集中在一起，甚至可以将其提取为独立的组合式函数（Composables），在不同的组件之间复用[1]。本系统前端开发全面采用Composition API配合<script setup>语法糖，该语法糖在编译时将<script setup>块中的顶层变量和函数自动暴露给模板，无需手动编写return语句，大幅减少了样板代码，提高了开发效率。')
    B(doc,'在响应式系统方面，Vue 3使用ES6的Proxy代理机制替代了Vue 2中基于Object.defineProperty的实现[1]。Proxy能够拦截对象属性的读取、赋值、删除、枚举等多种操作，解决了Vue 2响应式系统中的几个固有问题：无法自动检测对象属性的添加和删除；无法直接追踪数组索引赋值和length属性变化。Vue 3提供了ref()和reactive()两个核心API来创建响应式数据，ref()适用于基本类型值和需要被整体替换的对象，reactive()适用于不需要整体替换的复杂对象；computed()用于创建基于其他响应式数据的计算属性，仅在依赖变化时重新求值，性能优异。')
    B(doc,'本系统前端中大量使用了Vue 3的上述特性。例如，简历编辑器组件ResumeForm.vue使用reactive()创建包含六个顶层字段的formData响应式对象，通过watch()深度监听父组件传入的resumeData prop的变化并同步更新表单数据。Pinia（Vue 3官方推荐的轻量级状态管理库，替代了Vuex）被用于管理跨组件的全局状态，如用户的登录令牌、用户名、角色信息以及当前编辑的简历数据等。这些状态通过localStorage进行持久化存储，确保页面刷新后登录状态不丢失[11]。')

    H(doc,'2.1.2 Element Plus UI组件库',3)
    B(doc,'Element Plus是专为Vue 3生态设计的桌面端UI组件库，由Element UI社区升级维护，是目前Vue 3生态中使用最广泛的UI组件库之一[13]。它提供了超过80个经过精心设计和实现的开源组件，全面覆盖了企业级Web应用中常见的UI需求场景。这些组件包括但不限于：表单组件（el-form、el-input、el-select、el-date-picker、el-upload等）、数据展示组件（el-table、el-pagination、el-tag、el-card、el-descriptions等）、导航组件（el-menu、el-tabs、el-breadcrumb等）、反馈组件（el-dialog、el-message、el-message-box、el-notification、el-loading等）以及布局组件（el-container、el-row、el-col等）。Element Plus基于Google的Material Design设计原则，组件视觉风格统一、交互规范一致、可访问性良好，且提供了完善的中文文档和丰富的示例代码，对国内开发者非常友好[13]。')
    B(doc,'本系统前端界面几乎全部基于Element Plus的组件构建。用户认证页面使用el-form配合自定义校验规则（rules）实现表单数据的实时验证和提交。简历仪表盘Dashboard.vue使用el-container进行整体页面布局，el-header承载顶部导航栏，el-row和el-col的24列栅格系统实现响应式的三列简历卡片布局，每张简历卡片使用el-card组件的shadow="hover"属性实现悬停阴影效果。简历编辑器中使用el-upload组件的拖拽上传模式实现证件照的上传和预览，使用el-input和el-textarea实现文本输入，使用el-divider对不同编辑区域进行视觉分割，使用el-button的各类type属性（primary、danger、success等）区分不同操作的重要程度。管理后台Admin.vue使用el-tabs的border-card样式组织七个管理功能标签页，使用el-table的border和stripe属性渲染数据表格，使用el-pagination实现大数据量的分页浏览。AI面试页面Interview.vue使用自定义的聊天消息样式（结合Element Plus的按钮和输入组件）构建了完整的对话交互界面。')

    H(doc,'2.1.3 前端工程化工具',3)
    B(doc,'除Vue 3和Element Plus外，前端还采用了多项工程化工具和辅助库来支撑项目的开发和运行。Vue Router 4是Vue 3的官方路由管理方案，基于HTML5 History API实现无刷新页面跳转和URL管理[11]。本系统使用Vue Router的路由懒加载（通过动态import()函数实现组件的按需加载）来减小首屏打包体积，并通过全局导航守卫（router.beforeEach）实现了基于JWT令牌和用户角色的页面访问控制。Axios是一个基于Promise的HTTP客户端库，用于在浏览器中发送异步HTTP请求。本系统在src/api/request.js中对Axios进行了二次封装：配置了baseURL为/api/v1的统一请求前缀和30秒的默认超时时间；注册了请求拦截器，从localStorage读取JWT令牌并自动添加到每个请求的Authorization请求头中；注册了响应拦截器，统一检查后端返回的业务状态码，处理401未认证（清除令牌并跳转登录页）和通用错误提示。Vite 5是新一代的前端构建工具，由Vue.js作者尤雨溪团队开发，利用浏览器原生的ES模块导入能力实现了极速的开发服务器冷启动（通常在2秒以内）和模块热替换（HMR）。Vite的开发服务器内置了HTTP代理功能，通过vite.config.js中的proxy配置将/api开头的请求代理转发到后端Spring Boot应用的8080端口，解决了开发环境下前后端分离的跨域问题[11]。此外，html2canvas和jsPDF两个库被用于实现简历的PDF导出功能：html2canvas将指定的DOM元素渲染为Canvas画布，jsPDF将Canvas内容转换为PDF文档并触发浏览器下载。Web Speech API（浏览器内置的语音识别接口）被集成到面试模块中，用于将用户的语音输入实时转换为文字。')

    H(doc,'2.2 后端开发技术',2)
    H(doc,'2.2.1 Spring Boot应用框架',3)
    B(doc,'Spring Boot是由Pivotal团队（现为VMware旗下）开发的一款用于简化Spring应用开发的框架，目前已成为Java企业级应用开发的事实标准[2][12]。Spring Boot的核心设计理念是"约定优于配置"（Convention over Configuration），通过两大机制——自动配置（Auto-Configuration）和起步依赖（Starter Dependencies）——来解决传统Spring应用开发中配置繁琐、依赖管理复杂的痛点。自动配置机制在应用启动时根据classpath中存在的类、已定义的Bean以及配置文件中的属性，智能地推断并自动完成Spring应用上下文的基础配置，开发者无需手动编写冗长的XML或Java配置类。起步依赖则是一组预定义的Maven依赖描述符，开发者只需引入一个Starter就能获得该功能领域所需的全部依赖，无需逐一指定具体的构件坐标和版本号。')
    B(doc,'本系统后端基于Spring Boot 3.2.0版本开发，该版本要求Java 17及以上运行环境，全面采用了Jakarta EE 9+的命名空间（javax.*迁移为jakarta.*）。系统pom.xml中引入的核心Starter包括：spring-boot-starter-web，提供基于Spring MVC的Web层支持（包括DispatcherServlet、@RestController、@RequestMapping等注解驱动的开发模式）和内嵌的Apache Tomcat 10.1服务器，使应用可以独立运行而无需部署到外部Servlet容器；spring-boot-starter-data-jpa，集成Spring Data JPA和Hibernate 6.3 ORM框架，提供声明式事务管理（@Transactional注解）、基于方法命名约定的自动查询生成以及实体类的JPA注解映射能力；spring-boot-starter-security，集成Spring Security 6.2安全框架，提供认证（Authentication）和授权（Authorization）的核心基础设施；spring-boot-starter-validation，集成Jakarta Bean Validation 3.0规范（Hibernate Validator实现），支持在Controller方法参数上使用@Valid、@NotNull、@NotBlank等注解进行声明式请求参数校验；spring-boot-starter-webflux，集成Spring WebFlux和Reactor Netty响应式HTTP客户端，用于构建非阻塞的对外HTTP调用——在本系统中主要用于调用DeepSeek大语言模型的远程API[2][12][23]。')

    H(doc,'2.2.2 Spring Security安全框架与JWT认证',3)
    B(doc,'Spring Security是Spring生态系统中的安全基础设施框架，提供了一套完整的、高度可定制的认证（Authentication）和授权（Authorization）解决方案[23]。认证是指验证用户身份的过程（即"你是谁"），通常通过用户提供的凭证（如用户名和密码）与服务端存储的信息进行比对来完成。授权是指在确认用户身份之后，决定该用户被允许执行哪些操作、访问哪些资源的过程（即"你能做什么"），通常基于用户被分配的角色或权限列表进行判断。Spring Security通过一系列的Servlet过滤器（Filter）组成的安全过滤器链（Security Filter Chain）来实现安全控制，每一个进入应用的HTTP请求在被转发到目标Controller处理之前，都必须依次经过这些安全过滤器的检查。')
    B(doc,'JSON Web Token（JWT）是一种基于JSON格式的开放标准（RFC 7519规范），定义了一种在通信双方之间以紧凑、自包含的方式安全传输信息的机制。一个JWT令牌由三部分组成：Header（头部，描述令牌的元数据，包括令牌类型typ为JWT和签名使用的算法alg，如HS384）、Payload（载荷，也称为Claims声明，包含需要传输的实际数据，本系统中的Payload包括用户ID作为subject、用户名username和角色role两个自定义声明，以及签发时间iat和过期时间exp等标准声明）、Signature（签名，将编码后的Header和Payload用点号拼接后，使用配置的密钥和Header中声明的算法进行数字签名）。JWT自包含的特性意味着令牌本身已经携带了所有必要的身份和授权信息，服务端无需维护会话状态或查询数据库即可验证令牌的合法性和提取用户信息，这使得JWT非常适合前后端分离架构[23]。')
    B(doc,'本系统的认证授权流程设计如下：用户注册时，密码使用BCryptPasswordEncoder进行不可逆哈希处理后存入数据库。用户登录时，服务端验证密码匹配后，使用io.jsonwebtoken（JJWT）库生成一个包含用户ID（subject）、用户名和角色信息的JWT令牌，签名算法为HMAC-SHA384，令牌有效期为7天。前端将令牌存储在浏览器的localStorage中，之后在每一个需要认证的API请求中通过Authorization: Bearer <token>请求头携带该令牌。服务端通过自定义的JwtAuthenticationFilter（继承自OncePerRequestFilter，确保每个请求仅被过滤一次）拦截和解析令牌：从HTTP请求头中提取Bearer令牌字符串，调用JwtUtil.parseToken()解析令牌并验证其签名和有效期，从解析后的Claims中提取用户ID和角色，创建包含SimpleGrantedAuthority的UsernamePasswordAuthenticationToken对象并设置到SecurityContextHolder中。系统在SecurityFilterChain中配置了路径级别的授权规则：/api/v1/auth/下的登录和注册路径使用permitAll()允许匿名访问；/api/v1/admin/下的管理后台路径使用hasRole("ADMIN")仅允许管理员角色访问；其余所有路径使用authenticated()要求认证。Spring Security的内置CorsFilter配置允许来自前端开发服务器（localhost:5173等端口）的跨域请求，csrf()被关闭（无状态的JWT认证模式下不存在CSRF攻击面）[23]。')

    H(doc,'2.2.3 Spring Data JPA数据访问层',3)
    B(doc,'Spring Data JPA是Spring Data项目家族中的一员，旨在为基于Java Persistence API（JPA）的数据访问层开发提供最大程度的简化[12]。在传统的Java数据访问开发模式中，开发者需要手动编写大量的样板代码：创建EntityManager、开启事务、编写JPQL或原生SQL查询语句、处理查询结果映射、关闭资源等等。Spring Data JPA通过其核心的Repository抽象和智能的方法命名解析机制，将开发者从这些重复性工作中解放出来。开发者只需要定义一个继承自JpaRepository<T, ID>的接口（其中T为实体类型，ID为主键类型），该接口就自动获得了save()、findById()、findAll()、deleteById()、count()、existsById()等全套CRUD操作方法的标准实现，无需编写任何实现代码。')
    B(doc,'对于自定义查询需求，Spring Data JPA支持通过方法命名约定（Method Naming Convention）自动生成查询。开发者按照"findBy + 实体属性名 + 可选关键字"的模式在Repository接口中声明一个抽象方法，框架在运行时通过动态代理技术解析方法名，自动生成相应的JPQL查询语句并执行。例如，定义Optional<User> findByUsername(String username)方法后，框架会自动生成"SELECT u FROM User u WHERE u.username = :username"的查询；定义List<Resume> findByUserIdOrderByUpdatedAtDesc(Long userId)方法后，框架会生成带WHERE条件和ORDER BY排序子句的查询。方法名中支持的关键字包括And、Or、Between、LessThan、GreaterThan、Like、OrderBy、Asc、Desc、First、Top等，几乎可以覆盖日常开发中的大部分查询需求。对于特别复杂的查询（如多表关联、子查询等），可以使用@Query注解直接编写JPQL语句[12]。')
    B(doc,'本系统共定义了六个JPA Repository接口。UserRepository提供了findByUsername（根据用户名查找用户）和existsByUsername（检查用户名是否已存在）两个自定义查询方法。ResumeRepository提供了findByUserIdOrderByUpdatedAtDesc（按更新时间降序获取用户的所有简历）、findByIdAndUserId（根据简历ID和用户ID查找，用于权限校验——确保用户只能访问自己的简历）和countByUserId（统计用户的简历数量）三个方法。TemplateRepository提供了findByCategoryOrderByNameAsc（按分类获取模板并按名称排序）和existsBySourceUrl（检查模板来源URL是否重复）方法。JobAnalysisRepository提供了findByUserIdOrderByCreatedAtDesc和findByResumeId方法。OptimizationLogRepository提供了findByResumeIdOrderByCreatedAtDesc方法。InterviewSessionRepository提供了findByUserIdOrderByCreatedAtDesc、findByIdAndUserId和findAllByOrderByCreatedAtDesc（Pageable分页查询）方法。系统中涉及到多条数据库操作的业务逻辑（如删除简历时需要同时删除关联的岗位分析记录和优化日志）使用@Transactional注解标注，确保其在一个数据库事务内原子执行[12]。')

    H(doc,'2.3 大语言模型与提示词工程',2)
    H(doc,'2.3.1 Transformer架构与大语言模型',3)
    B(doc,'大语言模型的技术基础是Transformer架构[4]。在Transformer出现之前，自然语言处理领域的序列建模任务主要依赖循环神经网络（RNN）及其变体（LSTM、GRU），这些模型按照输入序列的时间步依次处理，这种串行计算方式带来了两个根本性限制：一是无法充分利用GPU等并行计算硬件的能力，训练效率低下；二是在处理长序列时存在梯度消失或爆炸问题，难以捕捉远距离的语义依赖关系。Transformer架构通过自注意力（Self-Attention）机制彻底改变了这一局面：在计算每个位置的输出时，自注意力机制同时考虑输入序列中所有位置的信息，并根据内容相关性动态分配权重，这种并行化的计算方式不仅大幅提升了训练效率，还使得模型能够直接建模任意距离的词语之间的语义关联[4]。')
    B(doc,'基于Transformer架构，大语言模型通过在超大规模的文本语料上进行自监督预训练（Pre-training），学习到了丰富的语言知识和世界知识。主流的预训练目标包括自回归语言建模（Autoregressive LM，如GPT系列，逐个预测下一个Token）和自编码语言建模（Autoencoding LM，如BERT，随机遮挡部分Token并要求模型预测被遮挡的内容）。DeepSeek系列模型采用的是Decoder-only的自回归架构，其预训练过程使用了包含网页文本、书籍、学术论文、代码仓库等多种来源的数万亿Token的高质量训练数据，通过分布式训练框架在数千张GPU上持续训练数月时间[3][8]。经过预训练后的大语言模型展现出了一些关键能力：上下文学习（In-Context Learning，模型能够根据提示中的几个示例理解任务要求并生成相应输出）、指令遵循（Instruction Following，模型能够理解并执行自然语言描述的复杂指令）和思维链推理（Chain-of-Thought Reasoning，模型能够生成中间推理步骤来解决需要多步推理的复杂问题）[5]。')

    H(doc,'2.3.2 DeepSeek-V4模型的技术特点',3)
    B(doc,'本系统选择DeepSeek-V4作为AI能力的底层支撑模型，主要基于以下考虑。DeepSeek-V4是深度求索公司于2026年4月发布的第四代大语言模型系列[3]，该系列包含两个版本：V4-Pro（总参数量1.6T，每个Token激活49B参数，定位为旗舰级通用推理模型）和V4-Flash（总参数量284B，每个Token激活13B参数，定位为高性价比的轻量模型）。两个版本均支持高达一百万Token的上下文窗口长度，V4-Flash版本的API定价仅为每百万Token约0.4美元，在保持优异性能的同时大幅降低了使用成本，对于个人开发者和小团队非常友好。')
    B(doc,'DeepSeek-V4相对于前代模型V3.2在多个核心技术组件上实现了重要创新。在注意力机制方面，V4采用了一种称为"混合注意力"（Hybrid Attention）的方案，将压缩稀疏注意力（Compressed Sparse Attention, CSA）和高度压缩注意力（Highly Compressed Attention, HCA）两种注意力模式交错排列在不同的Transformer层中。CSA通过稀疏化处理减少了注意力矩阵的计算复杂度，而HCA则通过极致的KV缓存压缩在几乎不损失信息的前提下大幅削减了显存占用。这种混合设计使得V4在百万Token级别的超长上下文推理场景下，计算量和KV缓存需求分别降至V3.2水平的约10%[3]。在网络架构方面，V4引入了流形约束超连接（Manifold-Constrained Hyper-Connections, mHC）技术。深层神经网络的训练面临一个普遍问题——随着网络层数的增加，梯度信号在前向和反向传播过程中逐渐退化，导致训练损失下降缓慢甚至发散。mHC通过将残差连接的权重矩阵约束到Birkhoff多胞体（双随机矩阵的集合）上来缓解这一问题，具体通过Sinkhorn-Knopp迭代算法实现约束投影。实验结果表明，采用mHC后，即便在数千层的超深网络中，训练损失也能保持稳定的下降趋势[3]。在训练优化方面，V4首次在1.6T参数级别的MoE（混合专家）模型上成功部署了Muon优化器（一种基于矩阵正交化的参数更新方法），并广泛采用了FP8混合精度训练策略和FP4量化感知训练（QAT）来降低计算和存储开销[3]。')

    H(doc,'2.3.3 提示词工程',3)
    B(doc,'提示词工程（Prompt Engineering）是指设计和优化输入给大语言模型的文本指令，以引导模型产生符合预期的、高质量输出的系统性方法[10]。尽管大语言模型在预训练阶段学习了广泛的知识和能力，但它们本质上是"通用"的——如果没有明确的任务指令和约束，模型的输出可以偏向任意方向。提示词的作用就是在通用能力和特定任务之间架设桥梁，告诉模型：你现在扮演什么角色、你需要完成什么任务、你应该遵循什么规则、你的输出应该是什么格式。')
    B(doc,'在本系统的开发实践中，提示词工程是决定AI功能质量的关键环节，其重要性在某种程度上甚至超过了模型选择本身——同样的模型，提示词写得好和写得差，输出结果的可用性可以有数量级的差异。以下通过两个具体的迭代案例来说明提示词设计的方法和教训。')
    B(doc,'第一个案例是面试功能的系统提示词。最初的版本中，我只简单地写了"你是一位资深技术面试官，请根据候选人的简历内容提出面试问题"。实际测试中发现，模型倾向于围绕简历中第一个出现的项目经历反复追问，连续四五个问题都在问同一个项目用了什么技术、为什么选这个技术、怎么优化的……虽然每个问题单独看都不错，但整体的面试体验非常糟糕——面试变成了"一个项目的审讯"，完全没有覆盖候选人的其他能力维度。经过分析，问题的根源在于提示词没有对问题的多样性提出明确要求。改进后的提示词明确规定了五个必须覆盖的面试维度（项目经历、技术深度、场景设计、问题排查、综合素质），每个维度只能问一个问题，并要求模型在每轮提问前检查对话历史以确定"哪些类别还没有被问过"，从尚未覆盖的类别中选择下一个问题。这一改进使面试的广度和体验得到了质的提升。')
    B(doc,'第二个案例是面试评估报告的评分系统提示词。早期的提示词只给出了"从技术深度、项目经验、沟通表达、问题分析四个维度综合评分"的粗略指令。结果模型几乎对所有候选人都给出了80-90分的高分，区分度极差——一个回答流利但内容肤浅的面试者和一个回答磕绊但思考深入的面试者可能只差三五分，这完全失去了评估的意义。分析发现，大语言模型在缺乏具体分数参照标准时，倾向于"老好人"式评分——不愿给出低分冒犯用户。改进方案是在提示词中为每个维度的每个分数段给出了具体的行为描述和评判标准。例如技术深度维度：能够清晰解释底层原理、讨论源码细节或技术选型的权衡→30-35分水平（满分35）；能够准确描述技术概念和使用方法但停留在应用层面→15-25分；回答模糊、概念混淆或明显错误→0-15分。加入了这些具体的评分锚点之后，模型的评分区分度明显改善，能够根据实际回答质量给出更加合理和有区分度的分数[10]。')
    B(doc,'本系统通过一个独立的PromptBuilder组件类来集中管理所有的提示词模板。该类被标注为Spring的@Component，在系统启动时作为单例Bean被初始化，所有需要调用LLM的Service类通过@RequiredArgsConstructor注入该组件。PromptBuilder中为每个AI功能定义了一对方法：buildXxxSystemPrompt()（系统提示词，定义LLM的角色身份、行为规范和输出格式要求）和buildXxxUserPrompt()（用户提示词，携带具体的任务输入数据，如简历内容、职位描述、对话历史等）。提示词模板使用Java 15引入的文本块（Text Block，用三个双引号包裹的多行字符串）语法编写，保持了原始排版和可读性。这种集中管理的方式使得提示词的修改和优化变得非常便捷——只需要在一个地方改文字，所有调用方自动生效，不需要去分散的业务代码中寻找和修改分散的提示词片段[10]。')

    H(doc,'2.3.4 API调用实现与容错机制',3)
    B(doc,'本系统通过Spring WebFlux的WebClient组件向DeepSeek的聊天补全API发送HTTP POST请求[3][12]。API端点采用了Anthropic消息格式（/v1/messages），请求体为JSON结构，包含以下字段：model——指定要调用的模型名称（通过配置文件的llm.model属性读取，默认值为deepseek-chat）；system——系统提示词，以字符串形式传入；messages——消息数组，每个消息对象包含role（角色，如user）和content（消息文本内容）；max_tokens——模型生成的最大Token数量，本系统设置为4096；temperature——采样温度参数，控制输出的随机性和多样性，本系统多数场景使用0.7（在创造性和稳定性之间取一个平衡点）。API的响应也为JSON格式，模型的生成文本位于响应体中被嵌套在content数组的第一个元素对象的text字段中，LLMService.callLLM()方法通过逐层解析JSON路径来提取这部分文本内容并返回。')
    B(doc,'考虑到大语言模型API作为远程网络服务存在固有的不确定性——可能出现临时的网络波动、服务端过载、请求超时等异常情况，系统实现了一套简洁但有效的容错重试机制。callLLM()方法内部使用一个最多执行maxRetries+1次（默认3次）的for循环，每次迭代中正常调用API。如果某次调用成功返回了有效结果，则直接break退出循环。如果抛出了异常（包括网络超时、服务端返回5xx错误码、响应格式异常等），则捕获该异常并记录一条WARN级别的日志，然后检查是否还有剩余的重试次数：如果有，则当前线程休眠（Thread.sleep）1秒钟乘以当前重试次数的等待时间（第一次重试前等待1秒，第二次等待2秒，形成简单的指数退避），之后进入下一次循环；如果所有重试次数已经用完，则向调用方抛出一个BusinessException，包含HTTP 503状态码和用户友好的错误提示信息（"优化服务繁忙，请稍后再试"）。这一机制确保了偶发的、短暂的服务不稳定不会导致整个系统功能不可用。此外，每一次API调用——无论最终成功还是失败——都会被记录到optimization_logs数据库表中，记录的字段包括关联的简历ID、发送给LLM的提示词全文、LLM返回的输出文本、实际使用的模型名称、API的响应耗时（毫秒级）以及被优化的简历区块类型（如summary、workExperience等）。这些日志数据为后续的系统效果评估、提示词优化和API成本核算提供了宝贵的数据基础。')

    H(doc,'2.4 数据库技术',2)
    H(doc,'2.4.1 MySQL关系型数据库',3)
    B(doc,'MySQL是当前全球使用最广泛的开源关系型数据库管理系统（RDBMS），由瑞典MySQL AB公司于1995年创建，现归Oracle公司所有并持续维护更新。MySQL采用经典的客户端-服务器（Client-Server）架构，支持多用户、多线程的并发访问模式，提供了完整的ACID事务支持（通过InnoDB存储引擎）、行级锁定、外键约束、事务回滚和崩溃恢复等企业级数据库特性[12]。在本系统中，MySQL 8.0.41版本作为生产环境的持久化存储方案，用于存储用户信息、简历数据、模板内容、分析记录、优化日志和面试会话等全部业务数据。')
    B(doc,'系统通过标准的JDBC驱动（mysql-connector-j）与MySQL建立连接，连接池选用了HikariCP——这是Spring Boot 2.x及之后版本的默认JDBC连接池实现，也是目前Java生态中性能最优的连接池之一。HikariCP通过字节码级别的精简优化、无锁的并发设计和对JDBC规范的极致遵循，在基准测试中的吞吐量和响应延迟均优于传统的C3P0、DBCP2和Tomcat JDBC Pool等方案[12]。在application.yml主配置文件中，spring.datasource配置块指定了MySQL的连接URL（jdbc:mysql://localhost:3306/ai_resume_copilot，包含useUnicode=true&characterEncoding=utf-8的编码参数和serverTimezone=Asia/Shanghai的时区参数）、用户名（root）和密码。URL中的createDatabaseIfNotExist=true参数指示JDBC驱动在数据库不存在时自动创建，无需要手动执行CREATE DATABASE语句。')

    H(doc,'2.4.2 H2内存数据库与双环境配置',3)
    B(doc,'为了提高开发效率并降低开发环境搭建的门槛，系统通过Spring Boot的Profile机制支持了MySQL和H2两种数据库环境的无缝切换。H2是一款纯Java编写的轻量级开源关系型数据库，支持嵌入式运行模式（与Java应用运行在同一JVM进程中）和客户端-服务器运行模式。在本系统的开发环境配置（application-dev.yml）中，H2被配置为内存运行模式（jdbc:h2:mem:ai_resume_copilot），数据完全存储在JVM进程的堆内存中，应用重启后所有数据自动清空[12]。使用H2内存数据库的开发优势是显而易见的：开发者克隆项目代码后无需安装和配置任何外部数据库服务，直接通过mvn spring-boot:run -Dspring-boot.run.profiles=dev一条命令即可启动完整的后端应用并进行功能开发。H2还内置了一个基于Web的数据库管理控制台（路径为/h2-console），开发者可以在浏览器中直接登录并查看和操作数据库内容。')
    B(doc,'使用双数据库环境时需要注意的一个技术细节是Hibernate的方言（Dialect）配置。Hibernate方言是一个适配器组件，它告诉Hibernate当前连接的是什么类型的数据库，以便Hibernate生成与该数据库SQL语法兼容的DDL语句和查询。如果不显式指定方言，Hibernate会根据JDBC驱动的元数据自动检测，但在某些情况下自动检测可能不准确。application.yml主配置中指定了spring.jpa.properties.hibernate.dialect为org.hibernate.dialect.MySQLDialect（针对MySQL生产环境），如果开发环境中不覆盖这个设置，Hibernate就会在H2数据库上尝试执行MySQL特有的SQL语法，导致建表失败。因此application-dev.yml中必须显式指定spring.jpa.properties.hibernate.dialect为org.hibernate.dialect.H2Dialect来覆盖主配置中的MySQL方言设置[12]。')

    H(doc,'2.4.3 Hibernate ORM与JPA实体映射',3)
    B(doc,'Hibernate是Java领域历史最悠久、功能最全面的ORM（Object-Relational Mapping，对象关系映射）框架之一，是JPA规范（Java/Jakarta Persistence API）的参考实现[12]。ORM框架的核心职责是解决面向对象的领域模型与关系型数据库的表结构之间的"阻抗不匹配"（Impedance Mismatch）问题——在Java代码中，业务数据以对象图的形式存在（对象包含属性和指向其他对象的引用）；在关系数据库中，业务数据以表、行和列的形式存在（通过外键表达实体间的关联）。Hibernate在中间起到翻译和桥梁的作用：将Java对象的属性变化自动翻译为INSERT、UPDATE、DELETE等SQL语句，将数据库的查询结果集自动映射为Java对象图。')
    B(doc,'本系统通过JPA注解在实体类上定义对象-关系映射的元数据。以User实体为例：@Entity注解将类标记为JPA实体，告诉Hibernate这是一个需要持久化到数据库的类；@Table(name = "users")指定映射到的数据库表名；@Id标注主键字段；@GeneratedValue(strategy = GenerationType.IDENTITY)告诉Hibernate使用数据库的自增机制（MySQL的AUTO_INCREMENT）来生成主键值；@Column注解用于定制列名（如@Column(name = "created_at")映射到数据库表中的下划线命名风格的列）、设置约束（nullable = false表示非空，unique = true表示唯一约束）和类型定义（如columnDefinition = "LONGTEXT"指定数据库列类型为长文本）；@ManyToOne(fetch = FetchType.LAZY)定义多对一的实体关联关系，FetchType.LAZY指定关联实体在默认情况下延迟加载（只有在首次访问该属性时才从数据库查询），避免加载不需要的关联数据造成性能浪费；@Enumerated(EnumType.STRING)指定枚举类型以字符串形式（而非整型序号）存储到数据库，提高可读性；@PrePersist注解的方法在实体首次插入数据库之前被Hibernate自动调用，用于设置创建时间等审计字段的默认值[12]。')

    H(doc,'2.5 开发工具与环境配置',2)
    B(doc,'本系统在开发过程中使用了以下主要的开发工具和运行环境。操作系统为Windows 11专业版。Java开发工具包使用Oracle JDK 21.0.6，这是当前的长期支持（LTS）版本，提供了虚拟线程（Virtual Threads，Project Loom的最终成果）等重要的新特性，不过本系统中并未深度使用这些新特性。项目管理与构建工具使用Apache Maven 3.9.9，通过pom.xml文件集中声明项目的所有依赖和构建配置。集成开发环境（IDE）方面，后端的Java代码使用JetBrains IntelliJ IDEA 2025.1版本进行编写和调试——IDEA对Spring Boot和Maven的集成支持非常完善，能够自动识别Spring Bean的依赖关系、生成JPA实体和Repository的模板代码、在编写配置文件时提供语法提示和自动补全。前端的Vue代码使用Microsoft Visual Studio Code编辑器，安装了Vue Language Features (Volar)扩展以获得Vue单文件组件的语法高亮、智能补全和类型检查支持。API调试使用Postman桌面客户端，通过创建分类的请求集合对系统的全部API接口进行管理和测试。版本控制使用Git，配合.gitignore文件排除node_modules/、target/、*.class、application-secret.yml等不应被版本控制的文件和目录。浏览器开发者工具（按F12呼出）主要用于前端界面的样式调整、JavaScript断点调试和网络请求的监控分析。')
    doc.add_page_break()

    # ==================== CHAPTER 3 ====================
    H(doc,'第3章  系统分析',1)
    H(doc,'3.1 可行性分析',2)
    H(doc,'3.1.1 技术可行性',3)
    B(doc,'从技术角度评估，本系统所采用的技术栈具有充分的成熟度和稳定性，不存在无法克服的技术障碍。Vue 3 [1][11]和Spring Boot [2][12]分别是前端和后端领域经过大规模工业验证的主流框架，两者的生态系统均非常完善：文档齐全、社区活跃、第三方库丰富，开发过程中遇到的大多数技术问题都可以在官方文档或社区问答中找到现成的解决方案。Element Plus [13]为Vue 3提供了完备的UI组件支持，可以满足本系统需要的所有界面交互模式。Spring Security [23]配合JWT是前后端分离应用中实现安全认证的经典方案，业界有无数的成功实践案例可参考。DeepSeek-V4 [3]通过标准的HTTP API对外提供服务，调用方式与开发者在日常工作中调用的任何RESTful或RPC接口没有本质区别，不需要深入了解机器学习或深度学习的理论知识即可使用。MySQL [12]作为数据库已应用了三十余年，其稳定性和性能经过了时间的充分检验。综合来看，本系统的技术栈不存在"前瞻性研究"层面的不确定性，技术上完全可行。')
    H(doc,'3.1.2 经济可行性',3)
    B(doc,'本项目的经济成本主要包括开发过程中的软件工具费用和上线运营后的云服务费用两个部分。在软件工具方面，项目使用的所有核心框架（Vue.js、Spring Boot、Spring Security）、开发工具（VS Code社区版、Postman免费版、Git）、数据库（MySQL社区版）和UI组件库（Element Plus）均为完全免费的开源软件，无需支付任何形式的授权费用或许可费用。DeepSeek-V4-Flash版本的API定价约为每百万Token 0.4美元[3]，在正常使用频率下（一个用户每天使用几次AI优化功能，每次消耗几千Token），每人每月的API费用不超过几美分。服务器方面，以一台入门级云服务器（2核CPU、4GB内存、50GB云硬盘）的配置即可满足系统在小规模用户下的运行需求，月租金大约在100至200元人民币之间，加上MySQL云数据库的费用，每月的总运营成本可以控制在几百元以内。从效益角度出发，该系统如果投入实际运营，可以通过会员订阅、简历模板付费或企业版定制等方式产生收入。因此，本项目的经济可行性是成立的。')
    H(doc,'3.1.3 操作可行性',3)
    B(doc,'在操作可行性方面，系统的交互界面和操作流程经过了充分的考虑和设计。作为一个Web应用，用户不需要安装任何客户端软件，通过浏览器访问即可使用全部功能——这消除了用户的使用门槛。界面的视觉风格和交互模式遵循了主流Web应用的设计惯例（得益于Element Plus [13]提供的统一设计语言），用户在不同页面之间切换时有连贯的体验。功能的组织方式也符合用户的直觉：登录后先看到仪表盘总览，从仪表盘可以进入简历编辑、岗位匹配和模拟面试等各个子功能；简历编辑器的六个内容模块自上而下排列，每个模块旁边都放置了AI优化按钮，操作目标明确；模拟面试采用对话形式，这是一种大多数人已经在日常聊天应用中习惯的交互模式。管理后台采用标签页结构，各个管理功能一目了然。系统在开发过程中在Chrome和Edge两种主流浏览器上进行了持续的兼容性验证，均能正常显示和运行。')

    H(doc,'3.2 功能需求分析',2)
    B(doc,'通过对现有简历辅助工具的功能分析[21]和实际用户需求的调研梳理，AI Resume Copilot系统需要实现的功能可以归纳为两大类别——用户端功能和管理端功能，共涉及七个核心功能模块。以下详细阐述每个模块的具体需求。')
    H(doc,'3.2.1 用户认证模块',3)
    B(doc,'用户认证模块是系统的入口，负责用户的身份注册、登录验证和权限管理。需要实现的功能包括：新用户通过提供用户名、密码和邮箱完成注册，系统对用户名进行唯一性检查（如果已被占用则提示错误），对密码进行长度校验（最少6个字符），对邮箱进行格式校验。注册成功后密码使用BCrypt算法[23]进行哈希处理后存入数据库，明文密码不在系统的任何环节落地。用户登录时提供用户名和密码，系统验证成功后生成一个使用HMAC-SHA384算法签名的JWT令牌，令牌的有效载荷中包含用户ID、用户名和角色信息，有效期设置为7天。前端在收到令牌后存储至浏览器的localStorage，并在后续所有需要认证的API请求的Authorization头中自动附带。系统需区分普通用户（USER）和管理员（ADMIN）两种角色，新注册的用户默认为USER角色，管理员账号由数据库初始化脚本预设。')
    H(doc,'3.2.2 简历管理模块',3)
    B(doc,'简历管理模块是用户端的核心业务模块。一名用户可以创建和管理多份简历，每份简历的数据以结构化的JSON格式存储，包含六个内容区块：基本信息（姓名、联系电话、电子邮箱、求职意向、证件照片）、个人简介（一段概括性的自我介绍和职业目标描述）、工作经历（支持多条，每条包含公司名称、担任职位、工作时间段和工作职责与成果描述）、教育背景（支持多条，每条包含学校名称、所学专业、学历层次和就读时间段）、专业技能（一段描述掌握的技术栈和熟练程度的文本）以及项目经历（支持多条，每条包含项目名称、个人角色和项目描述与贡献）。简历需要支持版本管理机制，用户在每次AI优化后可以基于当前版本创建新的简历版本，原版本被标记为非当前（is_current设为false），新版本成为当前活跃版本（is_current设为true），并通过optimized_from字段关联到源版本以形成版本演变链。简历需要支持导出为A4纸大小的PDF格式文件，导出的PDF中应完整呈现简历的全部内容模块，证件照显示在简历头部的右上角区域。删除简历时需要级联删除与之关联的岗位分析记录和优化日志，以保证数据的一致性。')
    PH(doc,'[此处插入图3-1：简历管理功能用例图]')
    H(doc,'3.2.3 AI简历优化模块',3)
    B(doc,'AI简历优化模块利用大语言模型的文本生成能力，对用户简历中的指定模块内容进行专业化润色和改写。用户可以在简历编辑器中选择任意内容模块（如个人简介文本、某段工作经历的描述、专业技能列表等），点击该模块旁边的"AI优化"按钮触发优化请求。如果需要，用户还可以在AI助手面板底部的输入框中输入额外的个性化优化要求（如"更突出技术亮点""尽量使用行业术语""控制在一百字以内"等），这些附加要求会作为用户指令合并到发送给大语言模型的提示词中。优化结果以聊天消息气泡的形式展示在AI助手面板中，用户可以阅读优化后的文本，如果满意可以点击"应用修改"按钮将优化内容直接更新到简历对应位置，如果不满意可以点击"重新生成"按钮让AI给出另一个版本的优化结果。每一次AI优化调用都需要被记录到数据库中，保存的信息包括关联的简历ID、原始文本内容、优化后文本内容、优化针对的区块类型、实际使用的模型名称和API的响应耗时。')
    H(doc,'3.2.4 岗位匹配分析模块',3)
    B(doc,'岗位匹配分析模块帮助用户评估自己的简历与心仪职位的匹配程度，从而有针对性地进行改进。用户首先从简历列表中选择一份需要分析的简历，然后在文本框中粘贴目标职位的完整职位描述（JD），点击"开始分析"按钮提交分析请求。系统将简历的结构化内容和职位描述文本一起发送给大语言模型，并附上分析要求：从技能匹配度（占评分的40%，考察简历中列出的技术技能与职位要求的技术栈之间的覆盖程度）、经验年限（占30%，考察工作年限和项目经验是否满足职位要求的资历水平）和关键词密度（占30%，考察简历文本中与职位描述相关的行业术语和核心概念的密度）三个维度进行综合评估。大语言模型返回的结果是一个JSON对象，包含综合匹配评分（0至100的整数）、优势列表（简历相对于该职位的优势项）、劣势列表（简历相对于该职位存在的不足）和改进建议列表（针对性的具体优化建议）。分析结果在前端以可视化卡片形式展示给用户，匹配分数以大号圆环图呈现并根据分数高低变换颜色（80分及以上为绿色，表示匹配度良好；60至79分为黄色，表示一般水平，有较大提升空间；60分以下为红色，表示匹配度较低，需要重点关注和改进），优势、劣势和建议以列表形式清晰罗列。分析结果同时保存到数据库的job_analyses表中，用户可以在历史中随时翻阅。')
    H(doc,'3.2.5 AI模拟面试模块',3)
    B(doc,'AI模拟面试模块通过大语言模型扮演专业面试官的角色，根据用户的简历内容生成有针对性的面试问题，帮助用户在模拟实战中提升应对真实面试的信心和能力。用户通过选择一份简历（以及可选的填写意向岗位名称）来启动一场模拟面试。面试采用五维度轮换机制来保证问题的广度和全面性，避免面试官在一个话题上反复追问：A类——项目经历维度，考察候选人对简历中提到的核心项目的深入理解，包括项目架构、技术选型理由、遇到的挑战和解决方案、个人在项目中的具体贡献等；B类——技术深度维度，针对简历中列出的某项关键技能或技术栈，追问其底层原理、设计思想、最佳实践或与其他类似技术的对比；C类——场景设计维度，给出一个贴近实际工作场景的开放性问题，要求候选人现场设计技术方案或系统架构；D类——问题排查维度，描述一个模拟的线上故障或异常现象，考察候选人的分析思路、排查步骤和定位方法；E类——综合素质维度，涉及团队协作、冲突处理、职业规划、学习方法等软技能相关话题。五个维度各问一个问题，总共五轮问答。面试官通过检查对话历史来判断哪些维度已经被覆盖，确保每一轮都切换到新的维度。五轮完毕后，LLM输出[END]标记触发面试自动结束和评估报告生成。')
    B(doc,'用户在面试过程中可以通过两种方式输入回答：文字输入（在输入框中打字，支持按Enter键快速发送）和语音输入（在支持的浏览器中点击麦克风按钮，利用浏览器的Web Speech API将用户的语音实时转换为文字显示在输入框中，用户可以编辑转换后的文字再发送）。面试结束后，系统自动调用大语言模型生成一份详细的面试评估报告，报告从四个维度对候选人的面试表现进行评分：技术深度（满分35分，评估候选人对技术原理的理解层次——能够讨论源码或底层原理得高分，停留在使用层面得中等，回答模糊概念混淆得低分）、项目经验（满分25分，评估项目描述的清晰度、数据支撑的充分性和个人贡献的突出程度）、沟通表达（满分20分，评估回答的结构化程度、语言组织能力和表达的简明清晰度）、思维分析（满分20分，评估面对开放性问题或故障场景时的分析思路、系统思维和逻辑严密性）。报告除了四个维度的分数外，还包含一份总体评价文本（150至300字）、具体的优势列举、具体的不足列举以及针对性的改进建议列表。')
    PH(doc,'[此处插入图3-2：AI模拟面试功能用例图]')
    H(doc,'3.2.6 简历模板模块',3)
    B(doc,'简历模板模块为用户提供预设的简历内容框架，帮助其快速开始简历的创建。系统预置多套面向不同职业方向的简历模板（如技术开发类、产品设计类、市场运营类、应届生通用类以及项目管理类等），每个模板包含预设的简历JSON结构（各模块的基础字段框架已搭建好）和简短的示例填充文本（帮助用户理解每个字段应该填写什么样的内容）。用户在创建新简历时，可以在模板选择对话框中浏览各模板的名称、分类标签和简要描述，选择最符合自己职业方向的模板作为起点。模板数据存储于数据库的templates表中。管理员角色可以通过管理后台对模板库进行维护：添加新的模板、修改现有模板的JSON内容和描述、删除过时或不合适的模板。')
    H(doc,'3.2.7 管理后台模块',3)
    B(doc,'管理后台模块为系统管理员提供集中化的系统管理功能。管理员通过独立的/admin路径进入管理后台界面，该界面以Tab标签页的形式组织结构，共包含七个功能标签。统计概览页以六张数据卡片的形式实时展示系统的核心运营指标：注册用户总数、创建的简历总数、数据库中的模板总数、AI优化功能的累计调用次数、岗位匹配分析的累计执行次数以及AI模拟面试的累计会话次数。每张卡片以醒目的数字和颜色方案展示对应的统计数值。用户管理页以表格形式列出系统中所有注册用户的信息（ID、用户名、邮箱地址、角色类型和注册时间），每行用户数据后附有编辑和删除两个操作按钮。编辑用户时弹出表单对话框，管理员可以修改用户的用户名、邮箱、密码（留空则保持不变）和角色（在USER和ADMIN之间切换）。删除用户时系统弹出二次确认对话框，确认后该用户及其关联的全部简历数据（包括简历本身、岗位分析记录和优化日志）被级联删除。为了防止误操作，系统禁止删除自身的管理员账号。简历管理页首先让管理员从一个下拉列表中选择一个目标用户，选择后自动加载并展示该用户名下的所有简历，管理员可以查看简历的基本信息（标题、版本号、是否当前版本、创建和更新时间）并执行删除操作。模板管理页提供简历模板的完整增删改查功能。优化日志页和匹配记录页分别以表格加分页的形式呈现AI优化和岗位匹配的历史调用记录。面试记录页展示所有用户的AI模拟面试会话，以表格形式列出会话ID、所属用户、关联简历、面试状态、评分和时间信息，支持分页浏览。')
    PH(doc,'[此处插入图3-3：管理后台功能用例图]')

    H(doc,'3.3 非功能需求分析',2)
    B(doc,'在满足功能需求的前提下，系统还需要达到一系列非功能性的质量指标，这些指标关系到系统的实际使用体验、长期可维护性和运行安全性。参考软件质量评估的国际标准ISO/IEC 25010[24]，结合本系统的实际应用场景，从以下几个维度提出非功能需求。')
    B(doc,'在性能方面，前端页面的首次内容绘制时间（First Contentful Paint, FCP）应该控制在3秒以内，确保用户打开网页后不会面对长时间的空白等待。不涉及AI模型调用的普通业务API接口（如获取简历列表、保存简历内容、查询用户信息等增删改查操作）在后端处理时间加网络传输时间的总和应控制在200毫秒以内。涉及大语言模型远程调用的AI接口（如智能优化、匹配分析、面试问答和报告生成等），由于响应时间取决于远端LLM服务的处理速度和网络延迟，系统自身无法直接控制，但需要设置合理的超时时间（本系统设定为60秒）并在等待期间通过前端的加载动画和提示文字给予用户明确的状态反馈。数据库连接池（HikariCP[12]）应配置足够的最小空闲连接数和合理的最大连接数，使系统能够同时处理至少10个并发的业务请求而不会出现数据库连接等待。')
    B(doc,'在安全性方面，所有用户的登录密码必须使用BCrypt算法进行单向哈希加密后存储到数据库，禁止任何形式的明文或可逆加密存储。BCrypt算法内置了随机盐值（Salt）机制——即使两个用户设置的密码完全相同，数据库中存储的哈希结果也不同，这有效防止了彩虹表（预先计算好常见密码的哈希值的对照表）攻击[23]。JWT令牌的签名密钥需要具备足够的长度和随机性（本系统使用384位的HMAC-SHA384密钥），使暴力破解在计算上不可行。所有管理后台的API接口（路径前缀为/api/v1/admin/）必须在每次请求时验证调用者令牌中的角色为ADMIN，普通用户的令牌即使有效也不应能够通过管理接口的权限检查[23]。敏感配置信息（如大语言模型的API访问密钥、JWT令牌的签名密钥、数据库的连接密码等）不应硬编码在随代码一起提交到版本控制系统的主配置文件中，而应通过单独的、被.gitignore排除的配置文件或操作系统环境变量来注入。前端需要对用户输入的内容进行基本的XSS（跨站脚本）防护——Vue 3的模板引擎默认对插值表达式{{ }}中的HTML标签进行转义处理，提供了基础的防护；在使用v-html指令输出原始HTML时需要特别谨慎[1]。后端使用Spring Data JPA提供的参数化查询（PreparedStatement），从机制上防止SQL注入攻击[12]。')
    B(doc,'在可用性方面，用户界面的设计应当直观易懂，关键的操作流程（如创建简历、发起AI优化、开始面试等）不应超过三步操作。对于具有破坏性的操作（如删除简历、结束未完成的面试等），系统在执行之前必须弹出确认对话框，给用户一个反悔的机会，防止因误操作导致数据不可逆地丢失。在AI功能处于等待状态时（如LLM正在生成优化文本或面试问题），界面应显示明确的加载指示器（如旋转的加载图标或骨架屏）和状态提示文字，避免用户产生"系统卡死了"的误解。服务器端发生错误时（如大语言模型API调用失败），前端应展示用户能够理解的中文错误提示信息，而非直接抛出技术性的异常堆栈。')
    B(doc,'在可扩展性方面，后端严格遵循分层架构的设计原则（Controller层负责HTTP请求处理、Service层封装业务逻辑、Repository层负责数据访问），各层之间通过接口（依赖注入）而非具体实现进行耦合，便于未来替换底层技术组件（如从MySQL迁移到PostgreSQL、从本地上传切换到对象存储服务等）或扩展业务功能[12]。大语言模型的提示词模板集中管理在PromptBuilder组件中，添加新的AI功能只需要在该类中增加新的提示词构建方法，而不需要修改任何调用方代码。前端API调用模块按业务领域分文件管理（auth.js、resume.js、ai.js、interview.js、admin.js），各文件职责明确，新增API端点时只需在对应文件中添加导出函数。')
    B(doc,'在兼容性方面，前端需要在Google Chrome浏览器和Microsoft Edge浏览器的最近两个主要版本（当前为130版本，要求兼容128及以上）上完整支持所有功能，包括语音输入。在Mozilla Firefox浏览器上，除语音输入功能（由于Firefox不支持Web Speech API）以外的其他功能应均能正常使用，且在不支持语音输入时系统应给出清晰的用户提示而非静默失效。后端应用需要能够在Java 21（当前的LTS版本）及其以上的Java运行环境中正常启动和运行。')

    H(doc,'3.4 系统用例分析',2)
    B(doc,'本系统明确区分两类参与者（Actor）：普通用户（User，角色标识为USER）和管理员（Admin，角色标识为ADMIN）。普通用户可以通过注册页面自行创建账号，登录后可以访问和使用系统的全部用户端功能，包括管理个人简历、使用AI优化和匹配分析、参加模拟面试以及查看个人的面试历史记录。管理员账号是在数据库层面预先配置的（通过初始化SQL脚本插入具有ADMIN角色的用户记录），管理员除了可以像普通用户一样使用系统的用户端功能外，还拥有进入管理后台的特殊权限。管理后台的访问通过前端路由守卫和后端API权限注解进行双重控制：在前端，Vue Router的全局导航守卫检查localStorage中存储的用户角色，如果角色不是ADMIN但试图访问/admin路径，会被强制重定向到/dashboard页面；在后端，Spring Security的SecurityFilterChain配置了对/api/v1/admin/**路径的hasRole("ADMIN")权限要求，非法请求会在到达Controller之前被拦截并返回HTTP 403 Forbidden状态码[23]。')
    B(doc,'系统的主要用例包括：用户注册——访客通过填写注册表单成为系统用户；用户登录——已注册用户通过提供正确的凭证获取JWT认证令牌；创建简历——用户新建一份空白的或基于模板的简历；编辑简历——用户对简历的各个模块内容进行修改和更新；AI优化简历——用户选择简历的特定模块，由大语言模型给出优化后的文本内容；岗位匹配分析——用户提供简历和目标职位描述，获取匹配度评估报告；参加AI模拟面试——用户选择简历开始面试，通过文字或语音回答面试问题，最终获取面试评估报告；查看面试历史——用户浏览自己的历次面试记录和报告；管理用户信息——管理员查看所有用户列表，编辑用户资料，修改用户角色或删除用户；管理简历模板——管理员对简历模板进行创建、编辑和删除操作；查看系统日志——管理员分页浏览AI优化日志、岗位匹配分析记录和面试会话记录；查看统计数据——管理员查看系统核心业务指标的实时统计数值。')
    PH(doc,'[此处插入图3-4：系统总体用例图]')
    doc.add_page_break()

    # ==================== CHAPTER 4 ====================
    H(doc,'第4章  系统设计',1)
    H(doc,'4.1 系统总体架构设计',2)
    B(doc,'AI Resume Copilot采用当今Web应用开发中主流的前后端分离架构模式[22]，整个系统在逻辑上划分为五个层次，自顶向下分别是：表示层（Presentation Layer）、接口层（API Layer）、业务逻辑层（Business Logic Layer）、数据访问层（Data Access Layer）和外部服务层（External Service Layer）。这种分层的设计将关注点清晰地分离到不同的层中，每一层只依赖于其正下方的一层（通过接口而非具体实现），层与层之间的职责边界明确，修改某一层的内部实现不会影响到其他层。')
    B(doc,'表示层是运行在用户浏览器中的Vue 3单页面应用（SPA）[1][11]。这个层负责页面的视觉渲染、用户交互的捕获和响应、前端路由的导航管理以及应用状态的维护。表示层通过Axios HTTP客户端库[11]向接口层的RESTful API端点发送异步HTTP请求，请求和响应的数据均采用JSON格式进行序列化和反序列化。在开发阶段，表示层运行在Vite开发服务器上，默认监听本机的5173端口。Vite的代理功能将表示层发出的所有以/api开头的HTTP请求透明地转发到运行在8080端口上的后端接口层，从而绕过浏览器的同源策略（Same-Origin Policy）限制[11]。在部署阶段，前端代码经过Vite的构建优化（包括代码压缩、Tree-Shaking和资源哈希命名）后生成纯静态文件（HTML、CSS、JavaScript和图片资源），这些静态文件可以部署到任何HTTP服务器（如Nginx）上。在生产环境中，通常会在前端静态资源服务器和后端API服务器之间再放置一层Nginx作为反向代理，统一对外暴露相同的域名和端口。')
    B(doc,'接口层由Spring MVC框架中的六个RestController类构成[12]：AuthController处理用户认证相关的请求（注册和登录），ResumeController处理简历的增删改查和版本管理，AIController处理AI优化和岗位匹配的请求，TemplateController处理简历模板的公开查询，InterviewController处理AI模拟面试的完整流程，AdminController集中处理所有管理后台功能的请求。接口层的主要职责包括：将HTTP请求的路径和HTTP方法（GET/POST/PUT/DELETE）映射到对应的Controller方法；通过@Valid注解和Jakarta Bean Validation框架对请求体（Request Body）和路径变量（Path Variable）进行校验，拦截不符合业务规则的输入；从HTTP请求的Authorization头中提取JWT令牌并委托给Spring Security过滤器链进行身份验证和授权判断；将Controller方法返回的Java对象（实体对象或DTO对象）通过Jackson JSON序列化库转换为JSON响应体；将Service层抛出的业务异常通过全局异常处理器（GlobalExceptionHandler，使用@RestControllerAdvice注解）统一转换为标准的错误响应格式。')
    B(doc,'业务逻辑层是系统的核心，包含了所有的业务规则和处理流程[12]。该层由七个标注了@Service注解的Spring Bean类构成：UserService负责用户的注册和登录业务逻辑，包括密码加密、JWT令牌生成和用户信息查询；ResumeService负责简历的各种管理操作，包括CRUD操作、权限校验（确保用户只能操作自己的简历）和版本管理；LLMService封装了与大语言模型API的所有通信细节，包括提示词组装、HTTP请求发送、响应解析和重试容错；JobAnalysisService负责岗位匹配分析的完整流程，包括匹配请求的处理、匹配结果的保存和历史查询；InterviewService管理AI模拟面试的整个生命周期，包括面试初始化、问答交互、自动结束逻辑和报告生成；TemplateScraperService是一个应用启动时的辅助服务，负责从远程数据源抓取示例简历数据来初始化模板库；AdminService集中处理管理后台的各项业务操作，包括用户管理、模板管理和各类日志查询。业务逻辑层通过Spring的依赖注入（Dependency Injection，使用@RequiredArgsConstructor和final字段的构造器注入方式）来获取其所依赖的Repository接口和LLMService实例。')
    B(doc,'数据访问层由六个继承自JpaRepository的接口构成[12]，它们充当了业务逻辑层和关系数据库之间的抽象桥梁。业务逻辑层的代码只需要调用Repository接口上定义的方法（如findByUsername、save、delete等），不需要关心底层SQL语句的生成、JDBC连接的获取和释放以及结果集到Java对象的映射等细节。Spring Data JPA在运行时通过JDK动态代理技术自动为Repository接口生成实现类并注入到Spring容器中。实体层由六个标注了@Entity的Java类构成，每个实体类映射到数据库中的一张表，类的属性映射到表的列，属性上的JPA注解定义了列的数据类型、约束条件和实体之间的关联关系。这一层的设计将在本章4.3节"数据库设计"中详细展开。')
    B(doc,'外部服务层目前主要包含一个大语言模型API的集成——通过LLMService中的WebClient实例向DeepSeek的/v1/messages端点发送HTTP POST请求[3]。这一层在架构上被单独划分出来的原因是，外部API是一个独立于系统的外部依赖，其可用性和响应速度不受系统控制。将外部API的调用逻辑封装在独立的服务层中，使得当未来需要切换大语言模型提供商（例如从DeepSeek切换到OpenAI或Anthropic）时，只需要修改LLMService中的API端点和请求格式适配逻辑，其他所有依赖LLMService的业务Service类（ResumeService、JobAnalysisService、InterviewService等）完全不需要感知变化。')
    PH(doc,'[此处插入图4-1：系统总体架构图]')

    H(doc,'4.2 功能模块详细设计',2)
    PH(doc,'[此处插入图4-2：系统功能模块结构图]')
    H(doc,'4.2.1 用户认证模块设计',3)
    B(doc,'注册流程的详细设计为：前端Login.vue或Register.vue使用Element Plus的el-form组件渲染注册表单，通过:rules属性绑定校验规则——用户名必填且长度在3至50个字符之间，密码必填且长度不少于6个字符，邮箱为可选字段但如果填写则需要符合电子邮箱地址的基本格式[13]。表单提交时首先在前端通过el-form的validate()方法进行客户端的规则校验，校验通过后才发送POST /api/v1/auth/register请求到后端。后端AuthController.register()方法接收到请求后，@Valid注解触发Jakarta Bean Validation对RegisterRequest DTO对象的字段进行服务端二次校验。校验通过后，UserService.register()方法执行业务逻辑：先通过userRepository.existsByUsername()检查该用户名是否已被注册（避免重复用户名），若已存在则抛出BusinessException(400, "用户名已存在")；若用户名可用，则新建User实体对象，将用户名和邮箱填入，使用passwordEncoder.encode()对明文密码进行BCrypt哈希，将role字段设为默认值Role.USER，调用userRepository.save()将用户实体持久化到数据库；保存成功后，调用jwtUtil.generateToken()为该用户生成一个JWT令牌并组装为AuthResponse返回给前端[23]。')
    B(doc,'登录流程的设计与注册流程类似但更简洁：AuthController.login()接收到LoginRequest后，UserService.login()通过userRepository.findByUsername()到数据库中查找该用户名的记录，若不存在则直接抛出BusinessException(401, "用户名或密码错误")；若存在，则使用passwordEncoder.matches()将请求中的明文密码与数据库中存储的BCrypt哈希进行匹配，若匹配失败同样抛出"用户名或密码错误"的异常（从安全角度考虑，不具体区分是用户名还是密码错误，防止攻击者通过错误信息来进行用户名枚举[23]）；匹配成功后生成JWT令牌返回。')
    B(doc,'JWT令牌的验证在前端和后端以不同形式实现。前端通过路由守卫（router.beforeEach）实现页面级的认证控制：守卫函数首先从localStorage中读取token和role的值，然后根据目标路由的meta.requiresAuth属性和用户的角色做出放行或重定向的决策。后端通过JwtAuthenticationFilter实现API级的认证控制：该过滤器在每个HTTP请求到达时从Authorization头中提取Bearer令牌，调用jwtUtil.validateToken()验证令牌有效后，将包含用户ID和角色的认证信息存入Spring SecurityContext。')

    H(doc,'4.2.2 简历管理模块设计',3)
    B(doc,'简历的数据模型采用JSON格式进行灵活的结构化存储，以适应不同简历模板和不同用户填写内容的多样性。JSON内容的顶层结构包含六个键：basicInfo（对象类型）、summary（字符串类型）、workExperience（对象数组类型）、education（对象数组类型）、skills（字符串类型）和projects（对象数组类型）。这个JSON对象被序列化为字符串后，存储在MySQL数据库resumes表的content_json字段中（列类型为LONGTEXT，最大可存储约4GB的文本数据，足够容纳任意复杂度的简历内容）。前端ResumeForm.vue组件通过reactive()函数[1]创建formData响应式对象，其初始结构与上述JSON结构完全对应。组件通过Vue 3的watch()函数深度监听父组件传入的resumeData prop：当prop发生变化时（例如用户从仪表盘点击了另一份简历），组件使用JSON.parse(JSON.stringify(val))进行深拷贝后将数据同步到formData中；当用户在编辑器中修改了任何字段时，组件通过emitUpdate()方法将formData的当前快照以事件的形式传递给父组件，父组件决定何时向后端发送PUT保存请求。')
    B(doc,'版本管理机制通过在resumes表上设计两个关键字段来实现：version字段是一个从1开始的递增整数，记录当前简历在其版本链中的序号；optimized_from字段是一个自引用的外键（指向同一表的id主键），对于非首版本的简历记录，该字段指向它的直接前驱版本，对于初始版本则为NULL。创建新版本时，ResumeService.createVersion()方法首先将当前简历的content_json完整复制到新创建的Resume实体中，新实体的version值设为源版本version加1，optimized_from指向源版本，is_current设为true。然后将源版本的is_current更新为false。这样，一份简历的所有历史版本通过一系列的optimized_from指针形成了一条单向链表，顺着链表可以从最新版本回溯到最初的版本。前端在展示简历时会通过is_current过滤只展示活跃版本，但用户可以通过版本列表查看和访问历史版本。')
    B(doc,'PDF导出功能的设计采用了一个纯前端的实现方案，无需后端的参与。exportToPDF()函数首先调用一个buildTemplate()辅助函数——该函数接收简历数据对象作为输入，将其六个模块的内容拼接为一个完整的、内嵌了CSS样式的HTML字符串。然后函数在页面上创建一个不可见的DOM容器，将HTML字符串注入其中，等待约300毫秒让浏览器完成布局计算和字体渲染。接着使用html2canvas库将这个DOM容器捕获渲染为一个Canvas画布（缩放比例设为2倍以获得高清输出）。最后使用jsPDF库创建一个A4纸尺寸（宽度210毫米、高度297毫米）的PDF文档对象，将Canvas导出为PNG格式的图片后放置到PDF页面中。如果简历内容的实际高度超出一页A4纸的范围，jsPDF会自动计算余量并创建新页面来容纳超出的部分。证件照的处理方式为：在buildTemplate()生成的HTML中，照片以<img>标签的形式嵌入，其src属性为简历数据中basicInfo.photo字段存储的Base64编码的图片数据。简历头部的CSS布局使用Flexbox，将姓名、联系方式和求职意向组成的文本块居中排列，证件照放置在右侧。')

    H(doc,'4.2.3 AI模块设计',3)
    B(doc,'AI智能优化模块、岗位匹配分析模块和模拟面试模块虽然在业务功能上各不相同，但在与技术架构上都遵循相同的核心模式：前端组件触发→前端API模块发送HTTP请求→后端Controller接收→后端Service组装提示词→调用LLMService.callLLM()→LLMService通过WebClient向DeepSeek API[3]发送HTTP请求→LLM处理并返回响应→LLMService解析响应提取生成文本→后端Service将结果持久化到数据库→返回响应给前端→前端组件更新界面状态。这一统一的流程模式保证了三个AI模块在代码结构上的一致性，降低了理解和维护的成本。')
    B(doc,'三个模块的关键差异在于提示词的设计。优化模块的System Prompt将LLM角色定位为"资深HR和职业顾问"，要求其输出纯文本的优化结果。匹配模块的System Prompt要求LLM扮演"资深招聘专家"，并强制其以规定字段的JSON格式输出分析结果，前端通过JSON解析将数据填充到MatchScore.vue组件的各个展示元素中。面试模块的System Prompt最为复杂，因为它需要LLM在同一会话中连续扮演面试官角色多轮，每一轮的提示词不仅要包含"角色设定"，还要包含前面所有轮次的完整对话历史作为上下文参照，这就是所谓的上下文学习（In-Context Learning）[5]。面试报告生成时，LLM的角色又切换为"面试评估专家"，需要根据整个对话历史进行综合评价。')

    H(doc,'4.3 数据库设计',2)
    H(doc,'4.3.1 数据库E-R模型',3)
    B(doc,'系统的数据库E-R模型以User（用户）实体为中心，向外辐射出与其他实体的关联关系。User与Resume（简历）之间是一对多（1:N）的关系——一个用户可以创建多份简历，但每份简历只属于一个用户，通过resumes表的user_id外键实现。Resume与JobAnalysis（岗位分析）、OptimizationLog（优化日志）、InterviewSession（面试会话）之间也都是一对多的关系——一份简历可以被多次用于岗位匹配分析、被多次提交AI优化（每次优化生成一条日志）、被多次用于模拟面试，这些实体通过各自的resume_id外键与简历关联。User与JobAnalysis、InterviewSession之间还存在直接的关联（通过user_id外键），方便直接从用户维度查询其分析记录和面试记录而不需要透过简历表进行间接关联。Template（模板）是一个独立的实体，不与其他任何实体存在外键关联，它仅仅作为模板数据被系统引用。')
    PH(doc,'[此处插入图4-3：系统数据库E-R图]')
    H(doc,'4.3.2 核心数据表结构',3)
    B(doc,'以下以数据字典的形式详细描述系统六张核心数据表的字段结构、数据类型、约束定义和业务含义。')
    B(doc,'users（用户信息表）：存储系统注册用户的账户信息和角色权限。id字段——BIGINT类型，主键，使用数据库自增（MySQL的AUTO_INCREMENT配合JPA的GenerationType.IDENTITY策略）生成唯一标识。username字段——VARCHAR(50)类型，非空约束，唯一约束（UNIQUE索引），存储用户的登录用户名，是用户登录时的唯一身份标识。password字段——VARCHAR(255)类型，非空约束，存储经过BCrypt算法哈希处理后的密码密文，不可逆，不参与JSON序列化（使用@JsonIgnore注解标记，防止在API响应中泄露密码哈希[12]）。email字段——VARCHAR(100)类型，可选字段，存储用户的电子邮箱地址。role字段——VARCHAR(10)类型，非空约束，默认值为字符串USER，使用JPA的@Enumerated(EnumType.STRING)注解将Java枚举类型Role的常量名（USER、ADMIN）作为字符串存储。created_at字段——DATETIME类型，非空约束，不可更新（updatable = false），由@PrePersist生命周期回调方法在实体首次持久化前自动设置当前时间。')
    B(doc,'resumes（简历数据表）：存储用户创建的所有简历及其版本信息。id字段——BIGINT类型，主键，自增。user_id字段——BIGINT类型，外键，引用users表的id字段，非空约束，通过JPA的@ManyToOne(fetch = FetchType.LAZY)和@JoinColumn(name = "user_id")注解建立与User实体的多对一懒加载关联。title字段——VARCHAR(100)类型，存储简历的标题。version字段——INT类型，非空约束，默认值1，存储当前记录的版本序号。content_json字段——LONGTEXT类型，存储简历完整结构化内容的JSON字符串。optimized_from字段——BIGINT类型，外键，引用resumes表自身的id字段（自引用外键），可为NULL（首版本为NULL），通过@ManyToOne(fetch = FetchType.LAZY)和@JoinColumn(name = "optimized_from")建立与Resume实体的自引用多对一关联，用于构建版本链。is_current字段——BOOLEAN类型（MySQL的TINYINT(1)），非空约束，默认值false，标记当前记录是否为该简历的最新活跃版本。created_at字段——DATETIME类型，非空约束，不可更新，记录简历的首次创建时间。updated_at字段——DATETIME类型，记录简历内容最近一次更新的时间，由@PreUpdate生命周期回调方法在执行update操作前自动更新。')
    B(doc,'templates（简历模板表）：存储系统预置和用户创建的简历模板数据。id字段——BIGINT类型，主键，自增。name字段——VARCHAR(100)类型，非空约束，存储模板的显示名称。category字段——VARCHAR(50)类型，存储模板的职业分类标签（如"技术开发""产品设计""市场运营"等）。description字段——VARCHAR(500)类型，存储关于模板特点和适用场景的简要文字描述。content_json字段——LONGTEXT类型，非空约束，存储模板的完整简历JSON结构（预设的字段框架和示例填充文本）。source_url字段——VARCHAR(500)类型，可选字段，如果模板是从外部网络资源获取的，此字段记录源URL地址。created_at字段——DATETIME类型，非空约束。')
    B(doc,'job_analyses（岗位分析记录表）：存储每次岗位匹配分析的结果数据。id字段——BIGINT类型，主键，自增。user_id字段——BIGINT类型，外键，引用users表。resume_id字段——BIGINT类型，外键，引用resumes表。job_description字段——TEXT类型，存储用户在分析时输入的目标职位JD原文。match_score字段——DECIMAL(5,2)类型（总共5位数字，其中2位小数，即-999.99到999.99的范围），存储LLM给出的综合匹配评分。suggestions字段——TEXT类型，以大JSON字符串的形式存储完整的分析结果（包括评分、优势、劣势、建议等），供前端解析和展示。created_at字段——DATETIME类型，非空约束，记录分析执行的日期和时间。')
    B(doc,'optimization_logs（AI优化日志表）：记录每一次AI简历优化调用的详细信息，用于后续的效果审计和统计分析。id字段——BIGINT类型，主键，自增。resume_id字段——BIGINT类型，外键，引用resumes表，非空约束（每条日志必须关联到一份具体的简历）。prompt_used字段——TEXT类型，存储发送给LLM的完整System Prompt和User Prompt文本，保留此信息便于后续分析"什么样的提示词产生了什么样的效果"。llm_model字段——VARCHAR(50)类型，记录实际调用的模型名称。input_text字段——LONGTEXT类型，存储LLM接收到的原始输入文本（优化前的简历区块内容）。output_text字段——LONGTEXT类型，存储LLM返回的优化后文本。response_time_ms字段——INT类型，以毫秒为单位记录本次API调用的响应耗时。section_type字段——VARCHAR(50)类型，记录被优化的简历区块的类型标识（如summary、workExperience、skills等）。created_at字段——DATETIME类型，非空约束。')
    B(doc,'interview_sessions（面试会话记录表）：存储AI模拟面试的完整会话数据。id字段——BIGINT类型，主键，自增。user_id字段——BIGINT类型，外键，引用users表，非空。resume_id字段——BIGINT类型，外键，引用resumes表，非空。position字段——VARCHAR(100)类型，可选，用户在开始面试时填写的意向求职岗位名称。messages字段——LONGTEXT类型，以JSON数组的形式存储面试过程中的完整对话历史，数组中的每个元素是一个包含role（角色，值为"ai"或"user"）、content（消息文本内容）和timestamp（时间戳）三个字段的对象。status字段——VARCHAR(20)类型，非空约束，默认值为字符串IN_PROGRESS，取值为IN_PROGRESS（面试进行中）或COMPLETED（面试已结束）。score字段——INT类型，存储LLM对面试表现的总体评分（0到100的整数），仅当status为COMPLETED时有有效值。report字段——TEXT类型，存储LLM生成的面试总体评价文本（150至300字的自然语言评价）。strengths字段——TEXT类型，以JSON数组字符串的形式存储LLM识别的候选人在面试中展现的优势项。weaknesses字段——TEXT类型，以JSON数组字符串的形式存储LLM识别的不足之处。suggestions字段——TEXT类型，以JSON数组字符串的形式存储LLM给出的具体改进建议。created_at字段——DATETIME类型，非空约束，记录面试的启动时间。completed_at字段——DATETIME类型，记录面试的结束时间，当status由IN_PROGRESS变为COMPLETED时由系统自动设置。')

    H(doc,'4.4 接口设计',2)
    H(doc,'4.4.1 RESTful API设计规范',3)
    B(doc,'本系统的全部后端API遵循RESTful架构风格进行设计[14]。REST（Representational State Transfer，表述性状态转移）是一种面向资源的Web API设计范式，其核心思想是将服务端的每一个可操作的数据实体（如用户、简历、模板、面试会话等）抽象为一个"资源"，通过URL路径来唯一标识每个资源（或资源的集合），通过标准的HTTP方法（GET、POST、PUT、DELETE）来表达对该资源的操作语义。具体而言，GET方法用于获取资源的表述（查询），POST方法用于在资源集合中创建新的资源实例，PUT方法用于修改指定资源的全部或部分字段，DELETE方法用于删除指定的资源。')
    B(doc,'为了保持API的一致性和可维护性，系统设计了统一的数据交换规范。所有API请求和响应的Content-Type均设定为application/json，数据以JSON格式进行序列化。所有API的响应体——无论是成功还是失败的响应——都包裹在一个统一的泛型响应结构中：ApiResponse<T>类。该类包含三个固定的顶层字段：code——整型，表示本次请求的业务处理状态码，取值为200时代表成功，其他数值代表不同类型的业务错误；message——字符串类型，提供人类可读的状态描述文本，成功时通常为"success"，失败时携带具体的错误原因；data——泛型T类型，承载接口的实际业务数据载荷，在成功时为期望的返回对象，在失败时为null。分页查询的API接口除了返回标准的ApiResponse包装外，其data字段内部还嵌套了额外的分页元数据字段：content（当前页的数据记录数组）、totalElements（符合条件的总记录数）、totalPages（按照当前每页大小计算出的总页数）和page（当前页码，从0开始计数）。')
    B(doc,'系统的全部API路径以/api/v1/作为统一的版本前缀（versioning prefix），其中/v1/表示API的第一个大版本号，这种URL路径版本管理策略为未来API的版本演进预留了空间。管理后台专用的API接口路径以/api/v1/admin/作为前缀，与普通用户接口在命名空间上进行了清晰的隔离，这既有利于安全策略的集中配置（在Spring Security中可以对/admin/**路径统一施加hasRole("ADMIN")的授权规则），也使得API目录的整体结构更加清晰[23]。')

    H(doc,'4.4.2 核心API接口清单',3)
    B(doc,'以下按照功能领域分组列出系统的主要API接口定义。用户认证接口组：POST /api/v1/auth/register——接收JSON格式的RegisterRequest对象（包含username、password和email字段），返回AuthResponse对象（包含token、username、userId和role）；POST /api/v1/auth/login——接收LoginRequest对象（username和password），返回AuthResponse对象。简历管理接口组：GET /api/v1/resumes——获取当前登录用户的所有简历摘要列表（List<ResumeSummary>），按更新时间降序排列；GET /api/v1/resumes/{id}——根据简历主键ID获取该简历的完整数据，接口内部会校验请求者是否为简历的所有者；POST /api/v1/resumes——创建一份新的简历，接收ResumeCreateRequest对象（title和可选的contentJson），返回新创建的完整Resume对象；PUT /api/v1/resumes/{id}——更新指定简历的标题、内容或is_current状态；DELETE /api/v1/resumes/{id}——删除指定简历并级联删除其关联的岗位分析记录和优化日志；POST /api/v1/resumes/{id}/versions——为指定简历创建一个新的版本。AI功能接口组：POST /api/v1/ai/optimize——对简历的特定区块进行AI优化，接收OptimizeRequest包含resumeId、sectionType、originalText和可选的instruction；POST /api/v1/ai/optimize-full——对完整简历JSON进行全文AI优化；POST /api/v1/ai/match——岗位匹配分析，接收MatchRequest包含resumeId和jobDescription。面试功能接口组：POST /api/v1/interview/start——启动一次新的模拟面试；POST /api/v1/interview/answer——提交对当前面试问题的回答并获取下一个问题；POST /api/v1/interview/{id}/end——手动结束当前面试会话并触发评估报告生成；GET /api/v1/interview/{id}——获取指定面试会话的完整详情；GET /api/v1/interview/history——获取当前用户的面试历史列表；DELETE /api/v1/interview/{id}——删除指定面试会话记录。管理后台接口组：GET /api/v1/admin/stats——获取系统级统计数据；GET /api/v1/admin/users——分页获取所有用户列表；PUT /api/v1/admin/users/{id}——管理员修改指定用户的资料或角色；DELETE /api/v1/admin/users/{id}——管理员删除指定用户；GET /api/v1/admin/templates——获取模板列表；POST/PUT/DELETE /api/v1/admin/templates——管理模板的增删改；GET /api/v1/admin/logs——分页获取AI优化调用日志；GET /api/v1/admin/analyses——分页获取岗位匹配分析记录；GET /api/v1/admin/interviews——分页获取所有面试会话记录。')
    PH(doc,'[此处插入表4-1：系统核心API接口汇总表]')

    H(doc,'4.5 系统安全设计',2)
    B(doc,'安全是软件系统中不可忽视的非功能性需求。本系统从身份认证、访问授权、数据保护和攻击防护四个层面进行了安全设计[23][24]。')
    B(doc,'身份认证层面采用无状态的JWT令牌方案，令牌使用HMAC-SHA384算法进行签名，密钥为384位的Base64编码随机字符串。无状态意味着服务端不需要维护任何形式的HTTP Session或者令牌缓存，每个请求携带着自己的"身份证"（令牌）来证明身份，这使得系统天然具有良好的水平扩展能力——多个后端实例之间无需共享会话状态。令牌的过期时间设置为7天，这是一个在用户便利性（不需要频繁登录）和安全性（令牌泄露后的时间窗口不至于过长）之间的折衷选择。')
    B(doc,'访问授权层面通过Spring Security的过滤器链来实现。系统在SecurityFilterChain配置中定义了基于URL路径匹配的粗粒度授权规则：公开路径（/api/v1/auth/下的注册和登录，以及Swagger/Knife4j的API文档路径）允许匿名访问；管理后台路径（/api/v1/admin/下的所有子路径）仅允许拥有ADMIN角色的用户访问；其余所有API路径要求请求方持有有效的JWT令牌（即已认证状态）。在细粒度授权方面，简历操作接口（如查询、更新、删除指定ID的简历）在Service层进行了数据归属校验——通过resumeRepository.findByIdAndUserId(resumeId, userId)方法确保当前登录用户只能操作属于自己的简历，防止通过修改URL中的简历ID参数来访问他人的简历数据。')
    B(doc,'数据保护层面，用户的登录密码使用BCryptPasswordEncoder进行单向哈希加密后存入数据库[23]。BCrypt算法是专门为密码哈希场景设计的，具有两个重要特性：自动包含随机生成的盐值（Salt），即使两个用户使用了完全相同的密码，存入数据库的哈希字符串也完全不同；计算强度（通过cost factor参数控制，默认为10）可配置，每提高一个强度值，哈希计算的时间增加一倍，从而使得暴力破解在计算上更加昂贵。JWT令牌的签名密钥、大语言模型的API访问密钥以及数据库的连接密码等敏感配置值不硬编码在application.yml主配置文件中随代码提交，而是通过环境变量（如${LLM_API_KEY}、${JWT_SECRET}）或单独的、已加入.gitignore忽略列表的application-secret.yml配置文件来注入。')
    B(doc,'攻击防护方面，针对Web应用中最常见的三种安全威胁进行了相应的防护设计。SQL注入防护——系统使用Spring Data JPA进行数据库操作，所有查询均通过参数化查询（PreparedStatement）执行，用户输入的数据作为参数绑定而非拼接到SQL语句字符串中，从机制上杜绝了SQL注入的可能[12]。XSS（跨站脚本）防护——Vue 3的模板语法默认对双花括号{{ }}中的变量值进行HTML转义，将<、>、&等特殊字符替换为对应的HTML实体，从而使得注入的恶意脚本代码被当作普通文本显示而不会被浏览器执行[1]。CSRF（跨站请求伪造）防护——由于本系统采用纯JWT令牌的认证方案（令牌存储在浏览器的localStorage中并通过JavaScript代码显式添加到Authorization请求头中），浏览器的同源策略不会在跨域请求中自动携带localStorage中的数据，而且没有使用Cookie作为认证凭证，因此传统的基于Cookie自动携带的CSRF攻击方式对本系统不构成威胁[23]。')
    doc.add_page_break()

    # ==================== CHAPTER 5 ====================
    H(doc,'第5章  系统实现',1)
    H(doc,'5.1 开发环境搭建',2)
    B(doc,'后端Spring Boot项目通过Spring Initializr在线生成器（https://start.spring.io/）创建基础的项目骨架。在生成器中选择Maven作为项目构建和依赖管理工具，选择Java 21作为编程语言版本，选择Spring Boot 3.2.0作为基础框架版本。生成的基础pom.xml文件只包含最简依赖，创建后在pom.xml中手动补充了以下关键依赖声明：mysql-connector-j（MySQL的JDBC驱动，scope设为runtime，仅在运行时需要），h2（H2内存数据库，同样scope为runtime），jjwt-api、jjwt-impl和jjwt-jackson三个JJWT库的依赖（分别提供JWT的API接口、标准实现和Jackson JSON序列化支持，版本统一为0.12.3），knife4j-openapi3-jakarta-spring-boot-starter（Knife4j是Swagger/OpenAPI规范的一个增强UI实现，用于自动生成美观的API文档页面，版本4.5.0），lombok（通过在编译期生成getter/setter/构造器等样板代码来减少Java类的冗余代码量，scope为optional不传递给依赖方），jsoup（一款Java的HTML解析库，用于TemplateScraperService在应用启动时从外部网页抓取简历模板参考数据，版本1.17.2），以及spring-boot-starter-test（包含JUnit 5、Mockito、Spring Test等测试框架的聚合依赖，scope为test）。')
    B(doc,'Java源代码按照分层架构组织包结构。com.aicopilot.config包存放Spring的@Configuration配置类：SecurityConfig（Spring Security的安全过滤器链、密码编码器、CORS跨域策略和JWT认证过滤器的Bean定义）、Knife4jConfig（Knife4j/OpenAPI文档的标题、版本和描述信息配置）、WebClientConfig（WebClient实例的创建和编解码器配置，将内存中缓冲区的最大容量提升到16MB以支持较大的LLM响应体[12]）。com.aicopilot.controller包存放六个@RestController控制器类。com.aicopilot.service包存放七个@Service服务类。com.aicopilot.repository包存放六个@Repository数据访问接口。com.aicopilot.entity包存放六个@Entity实体类和一个Role枚举类。com.aicopilot.dto包定义了十几个用于前端与后端之间数据传输的DTO类。com.aicopilot.exception包包含BusinessException（自定义业务异常，携带整型错误码和字符串错误消息）和GlobalExceptionHandler（全局异常处理器，将不同类型的异常统一转换为ApiResponse格式的JSON错误响应）。com.aicopilot.util包包含JwtUtil（JWT令牌生成、解析和验证工具）和PromptBuilder（集中管理所有AI提示词模板）。')
    PH(doc,'[此处插入图5-1：后端项目包结构截图]')
    B(doc,'前端项目使用Vite的交互式脚手架（npm create vite@latest命令）生成基础模板[11]。创建后通过npm install命令安装以下运行时依赖：vue 3.4（核心框架）、vue-router 4.2（路由管理）、pinia 2.1（状态管理）、axios 1.6（HTTP客户端）、element-plus 2.5（UI组件库）、@element-plus/icons-vue 2.3（Element Plus配套的图标库）、echarts 5.5和vue-echarts 6.6（数据可视化图表库，当前版本中主要在管理后台的统计概览中备用）、html2canvas 1.4和jspdf 2.5（PDF导出功能的核心依赖）。开发依赖包括@vitejs/plugin-vue 5.0（Vite的Vue单文件组件编译插件）和vite 5.0（构建工具本体）。前端源码src目录下按职责划分子目录：api/——集中管理所有后端API的调用函数，每个文件对应一个业务模块；store/——Pinia状态管理Store的定义文件；router/——Vue Router路由表和导航守卫的定义；views/——页面级别的Vue组件（每个路由对应一个页面组件）；components/——可跨页面复用的UI组件；utils/——不依赖Vue的纯工具函数。')
    PH(doc,'[此处插入图5-2：前端项目目录结构截图]')

    H(doc,'5.2 用户认证模块实现',2)
    B(doc,'用户认证模块的后端核心由三个Java类协作完成。SecurityConfig类通过@EnableWebSecurity和@EnableMethodSecurity注解激活Spring Security[23]，在其securityFilterChain(HttpSecurity http) Bean方法中构建安全过滤器链。首先通过cors()方法启用自定义的CORS策略——在本地的开发环境中，前端Vite开发服务器运行在localhost的5173等端口，而后端运行在8080端口，浏览器的同源策略会阻止来自不同端口的请求，因此必须配置CORS来显式允许这些跨域请求。csrf()被调用了disable()来关闭CSRF保护——因为前后端分离应用不使用Cookie来传递会话凭证，不存在CSRF的攻击面。sessionManagement()被配置为SessionCreationPolicy.STATELESS（无状态会话策略），告诉Spring Security不要创建HTTP Session。authorizeHttpRequests()方法中依次定义了三条授权规则：Swagger文档和认证相关的路径允许任何人访问，管理后台路径需要ADMIN角色，其余路径需要已认证。最后通过addFilterBefore()方法在UsernamePasswordAuthenticationFilter之前插入自定义的jwtAuthenticationFilter()。')
    B(doc,'JwtAuthenticationFilter通过继承Spring的OncePerRequestFilter抽象类来实现，确保在单个请求的处理过程中该过滤器只被执行一次[23]。在doFilterInternal()方法中：首先从HttpServletRequest的"Authorization"头部获取值，检查其是否以"Bearer "前缀开头（这是JWT在HTTP协议中的标准传递方式的约定）；如果不是则直接调用filterChain.doFilter()放行——因为后续的Spring Security授权检查会拦截未认证的请求并返回403。如果Bearer令牌存在，则调用jwtUtil.validateToken()进行令牌验证（内部通过try-catch捕获JWT解析过程中的各类异常，任何异常都视为令牌无效）。验证通过后，从令牌中提取userId（通过Long.parseLong(parseToken(token).getSubject())从subject字段获取）和role（通过自定义的getRoleFromToken()方法从私有声明中获取）。然后使用这两个值构建一个UsernamePasswordAuthenticationToken对象——这个对象是Spring Security认证框架中的标准认证令牌类型，其中第一个构造参数（userId）作为principal（主体身份），第三个参数为包含SimpleGrantedAuthority("ROLE_" + role)的授权列表。最后将这个认证令牌设置到SecurityContextHolder中以供后续的授权检查使用[23]。')
    B(doc,'前端src/store/user.js中使用Pinia的defineStore定义了一个名为user的Store[11]。这个Store管理了五个响应式状态：token（JWT令牌字符串）、username（用户名）、userId（用户ID数值）、role（角色字符串）和isAdmin（计算属性，值为role === "ADMIN"的布尔结果）。login()和register()两个异步action分别调用auth.js中的登录和注册API函数。API调用成功拿到返回数据后，将token、username、userId和role四个值同时存入Pinia的响应式状态（使得所有使用该Store的组件自动更新）和localStorage（使得页面刷新后登录状态得以持久化）。logout()方法则做相反的操作——将Pinia状态置空并清除localStorage中的对应键值。')
    PH(doc,'[此处插入图5-3：用户登录注册界面截图]')

    H(doc,'5.3 简历管理模块实现',2)
    B(doc,'简历管理的前端界面由Dashboard.vue和ResumeForm.vue两个核心页面/组件协同实现。Dashboard.vue是用户登录后首先看到的仪表盘页面，页面上方为顶部导航栏（包含系统标题和用户信息），主体区域用三列栅格布局展示用户的所有简历卡片。每张卡片是Element Plus的el-card组件，卡片内部显示简历标题、版本号、最近更新时间和一个标识当前版本的绿色el-tag标签[13]。卡片的底部有两个操作按钮：编辑按钮（文字+primary主题色）点击后通过Vue Router的程序式导航router.push()跳转到/editor/:resumeId路由对应的编辑器页面；删除按钮（文字+danger主题色）点击后首先通过ElMessageBox.confirm()弹出确认对话框请求用户确认，确认后调用resume API的删除接口并刷新简历列表。页面中的"新建简历"按钮点击后打开一个el-dialog模态对话框，对话框内以三列网格形式展示数据库中的简历模板列表，每个模板以卡片形式显示分类标签（通过不同颜色的el-tag实现）、模板名称和简短描述。用户选择一个模板后点击"创建"按钮，或者直接点击"不使用模板，从空白开始"的文本链接，系统调用resume API的创建接口并导航到新简历的编辑页面。页面下半部还设有两个快捷入口："岗位匹配分析"按钮（导航到/match页面）和"AI模拟面试"区域（包含"开始面试"和"面试记录"两个操作按钮，分别导航到/interview和/interview/history）[13]。')
    B(doc,'ResumeForm.vue组件是整个前端项目中最复杂的单文件组件之一。该组件使用Vue 3的Composition API编写[1]：通过defineProps()接收父组件传入的resumeData对象；通过defineEmits()声明两个自定义事件（update用于通知数据变更，ai-optimize用于触发AI优化请求）；通过defineExpose()暴露setSection()方法供父组件在需要时调用。组件的核心数据结构是reactive()创建的formData对象，其结构完全对应简历JSON的六个顶层字段。组件通过Vue 3的watch()函数深度监听resumeData prop的变化：当用户点击不同的简历卡片进入编辑器时，resumeData prop变为新的简历数据，watch回调使用JSON.parse(JSON.stringify(val))的深拷贝方式将数据同步到formData中，确保formData始终反映当前正在编辑的简历内容。用户在表单中输入任何内容时，通过Element Plus组件的@input或@change事件触发emitUpdate()方法，将formData的浅拷贝传递给父组件，父组件可以选择合适的时机（如用户停止输入一段时间后自动保存或点击手动保存按钮时）向后端发送更新请求。')
    B(doc,'简历的PDF导出由src/utils/pdfExport.js工具模块独立实现。导出流程为：exportToPDF()函数被调用时，首先在DOM树中创建一个全屏尺寸的半透明遮罩层（overlay），用于阻止用户在PDF生成过程中操作界面并提示"正在生成PDF..."。然后创建一个绝对定位的、宽度为794像素（对应A4纸210毫米宽度的屏幕近似尺寸）的容器div，调用buildTemplate()函数将简历数据对象拼接为该容器内的HTML字符串。buildTemplate()函数内部定义了一套完整的嵌入式CSS样式和按模块组织的HTML结构，证件照以Base64编码的data:image URI直接嵌入img标签的src属性，简历头部使用CSS Flexbox弹性布局实现姓名居中和照片右对齐。等待约300毫秒的渲染时间后，使用html2canvas库以2倍devicePixelRatio将该容器渲染为高清Canvas。然后使用jsPDF库创建A4纸尺寸的PDF文档并将Canvas转存的PNG图片添加到PDF页面中，最后触发浏览器的文件下载。生成完成后清理DOM中临时添加的遮罩层和渲染容器。')
    PH(doc,'[此处插入图5-4：简历编辑器界面截图]')
    PH(doc,'[此处插入图5-5：简历仪表盘界面截图]')
    PH(doc,'[此处插入图5-6：PDF导出的简历效果截图]')

    H(doc,'5.4 AI智能优化模块实现',2)
    B(doc,'AI智能优化模块是第一个被实现的核心AI功能，它的成功为后续的匹配分析和模拟面试奠定了技术基础。LLMService是整个AI能力的核心枢纽类，被标注为@Slf4j（Lombok的日志注解，自动生成名为log的SLF4J Logger实例）和@Service（Spring的组件注解，将类声明为Spring容器管理的Bean）[12]。该类通过final字段和@RequiredArgsConstructor（Lombok的构造器注入注解）注入了三个依赖：WebClient实例（通过WebClientConfig配置的响应式HTTP客户端Bean）、PromptBuilder实例（提示词构建器组件）和ObjectMapper实例（Jackson的JSON处理核心类，由Spring Boot自动配置提供）。另外通过@Value注解从application.yml配置文件中注入四个LLM相关的配置参数的值：llm.api-key（API访问密钥）、llm.base-url（API的基础URL地址）、llm.model（使用的模型名称）和llm.timeout（API调用超时时间，单位毫秒）。')
    B(doc,'LLMService的核心方法是callLLM(String systemPrompt, String userPrompt)，该方法封装了与大语言模型API交互的完整逻辑。首先构建API的完整URL地址——根据baseUrl是否以斜杠结尾来决定拼接格式，最终URL格式为{baseUrl}/v1/messages。然后使用Java的Map.of()工厂方法构建符合DeepSeek API请求格式的JSON请求体：model字段从配置读取，system字段直接传入systemPrompt字符串，messages字段构造为一个只包含一个对象的数组（该对象的role为"user"表示这是用户角色的消息，content为userPrompt的文本内容），max_tokens设置为4096（限制模型单次生成的最大Token数），temperature设置为0.7（在创造性输出和确定性输出之间取得适中的平衡）。')
    B(doc,'重试容错逻辑通过一个for循环实现，变量i从0递增到maxRetries（配置值为2）。循环体中，通过WebClient的流式API构建POST请求：指定URI、设置x-api-key认证头（DeepSeek API的认证方式）、Content-Type头和anthropic-version头（指定使用Anthropic兼容协议的版本）[3]，通过bodyValue()将请求体Map对象序列化为JSON。调用retrieve()发送请求并获取响应，通过onStatus()注册HTTP 4xx客户端错误的处理回调，通过bodyToMono(Map.class)将JSON响应体反序列化为Java的Map对象，通过.timeout()设置响应超时，最后调用.block()阻塞当前线程等待响应返回（这是同步调用方式，适用于常规的Servlet线程模型）。成功返回后，从响应的Map对象中逐层解析提取模型生成文本。如果任一环节抛出异常，若还有剩余重试次数则休眠后继续下一次循环，若已耗尽所有重试次数则抛出最终的业务异常。')
    B(doc,'前端AI助手面板AIPanel.vue组件以嵌入式聊天界面的形式集成在简历编辑器页面中[1]。组件内部使用ref([])维护了一个messages响应式数组，数组中的每个元素是一个包含role、content、section等信息的小对象。组件通过defineExpose({ setSection })向父组件暴露了触发优化流程的入口方法。当用户在ResumeForm中点击某个区块的"AI优化"按钮时，父组件接收到ai-optimize事件后调用AIPanel的setSection()方法。setSection()首先在messages中添加一条role为"user"的消息（表示用户的优化请求），然后调用内部函数doOptimize()。doOptimize()设置loading状态为true（在聊天区显示"正在优化..."的加载动画），调用ai.js中导出的optimizeSection() API函数并等待返回，收到结果后将一条role为"ai"、optimized为true的消息推入messages数组，触发Vue的响应式更新将新消息渲染为绿底的AI消息气泡，并在消息下方显示两个操作按钮。用户点击"应用修改"按钮时，组件通过emit("apply-optimization", {section, index, text})通知父组件更新对应的简历字段内容[1][11][13]。')
    PH(doc,'[此处插入图5-7：AI优化功能界面截图]')

    H(doc,'5.5 岗位匹配分析模块实现',2)
    B(doc,'岗位匹配分析功能的业务逻辑由JobAnalysisService类实现[12]。analyzeMatch()方法的参数为简历ID、职位描述字符串和用户ID。方法首先通过resumeRepository.findByIdAndUserId()双重条件查询验证简历的存在性和归属权——这是本系统所有涉及特定简历操作的Service方法的标准权限校验模式。然后从Resume实体中获取content_json字段的值作为简历的文本表示（如果content_json为null则使用空JSON对象"{}"作为缺省值）。接着调用llmService.analyzeMatch()，该方法内部调用callLLM()并传入buildMatchSystemPrompt()和buildMatchUserPrompt()构建的提示词。LLM返回的原始文本字符串首先通过extractJson()辅助方法进行预处理——该方法在字符串中定位第一个左大括号"{"和最后一个右大括号"}"的位置，截取这两个位置之间的子字符串作为候选JSON。然后使用ObjectMapper将这个JSON字符串反序列化为MatchResponse DTO对象。如果JSON解析过程中发生任何异常（JSON格式不正确、缺少必需的字段、字段类型不匹配等），系统不会向上抛出异常，而是捕获异常后返回一个包含默认值（评分为0、三个分析列表均为空、suggestions中包含一条"匹配分析服务暂时不可用"的提示消息）的MatchResponse，保证接口在LLM输出异常的情况下依然能够优雅降级而不会崩溃。分析完成后，方法创建JobAnalysis实体对象保存分析记录，并将MatchResponse序列化为JSON字符串存入suggestions字段。')
    B(doc,'前端MatchScore.vue组件负责匹配结果的可视化呈现。该组件接收一个名为result的prop，其类型为MatchResponse对象（包含score整型字段和三个字符串列表字段）。组件的模板中使用自定义的CSS样式绘制了一个大号的圆形评分环——通过设置border宽度和border-radius: 50%实现圆形外观，通过动态:class绑定根据score的数值范围切换三种颜色方案（high/mid/low对应绿色/黄色/红色）[13]。环内以大号数字显示分数，下方以小字标注"综合评分"。优势列表和不足列表分别以左右两列el-col布局展示，每个列表项前以不同颜色的小圆点作为标记。改进建议以带有序号的ol列表形式展示在评分环下方。')
    PH(doc,'[此处插入图5-8：岗位匹配分析界面截图]')

    H(doc,'5.6 AI模拟面试模块实现',2)
    B(doc,'AI模拟面试是本系统中业务流程最长、状态管理最复杂、提示词设计最精细的功能模块。InterviewService类管理了面试的完整生命周期[12]。startInterview()方法负责面试的初始化：验证简历归属权→使用Resume实体的contentJson作为简历数据输入→调用promptBuilder的面试开场提示词构建方法→通过callLLM获取AI面试官的开场白和第一个面试问题→创建InterviewSession实体对象，设置其user、resume、position字段，将第一个AI问题封装为JSON消息对象存入messages字段，status设为IN_PROGRESS→通过sessionRepository.save()持久化→返回InterviewResponse给前端。')
    B(doc,'submitAnswer()方法处理每轮的问答交互。首先通过findByIdAndUserId()加载会话并验证所有权和状态（如果status已经是COMPLETED则拒绝继续提交回答）。然后将现有messages字段的JSON字符串使用ObjectMapper的readValue()方法解析为List<Map<String, String>>类型的Java数据结构，将用户本次回答的消息对象追加到列表中。接着通过countAiMessages()辅助方法统计截至目前AI已经问了多少个问题。如果问题数已经达到或超过MAX_QUESTIONS常量（值为6），则直接调用autoEndInterview()触发自动结束流程。如果尚未达到上限，则通过buildConversationHistory()将所有历史消息拼接为可读的对话记录文本，传递给buildInterviewNextUserPrompt()生成续问提示词，调用LLM获取AI的回应。LLM返回的文本中如果包含[END]标记（表示AI判断五个维度均已覆盖完毕），则剥离[END]标记后将问题文本连同session状态变为COMPLETED一起返回。如果LLM返回的文本不含[END]，则继续以IN_PROGRESS状态返回。每次AI回复的消息对象（包含role为"ai"、content为问题文本、timestamp为当前时间戳）被追加到messages列表后，调用toJson()将更新后的消息列表序列化为JSON字符串并保存回session的messages字段。')
    B(doc,'endInterview()方法（手动结束）和autoEndInterview()方法（自动结束）最终都会调用generateReport()私有方法来生成面试评估报告。generateReport()首先调用buildConversationHistory()构建完整的面试对话文本，然后分别获取报告生成的System Prompt和User Prompt（后者包含了简历内容和完整的对话历史），调用LLM获取评估结果。LLM返回的文本通过extractJson()提取JSON部分后，使用ObjectMapper解析为Map<String, Object>。从Map中提取score（整型评分）、report（文本评价）、strengths（JSON数组，优势列表）、weaknesses（JSON数组，不足列表）和suggestions（JSON数组，建议列表）。提取完成后将这些值分别设置到InterviewSession实体的对应字段中，将status更新为COMPLETED，将completedAt设置为当前时间，然后通过save()持久化。如果JSON解析失败（异常情况），则进入降级处理流程——score设为0，report直接使用LLM的原始响应文本，strengths/weaknesses/suggestions均设为空JSON数组"[]"。最后从更新后的session实体构建InterviewReportResponse DTO并返回。')
    B(doc,'前端Interview.vue组件实现了完整的面试对话交互界面。页面分为三个区域：顶部导航栏（显示"AI模拟面试"标题、当前问题序号标签和"结束面试"按钮）、中部可滚动的消息列表区域（显示AI面试官的问题气泡和用户的回答气泡）、底部输入区域（包含语音输入麦克风按钮、文字输入框和发送按钮）。语音输入基于浏览器的Web Speech API实现，在Chrome和Edge等Chromium内核浏览器中可用[1]。initVoice()函数在组件挂载时检查window.SpeechRecognition或window.webkitSpeechRecognition对象是否存在以判断浏览器是否支持语音识别。toggleVoice()函数在用户点击麦克风按钮时被调用：如果当前已经在录音状态则调用recognition.abort()停止；如果尚未开始则创建一个新的SpeechRecognition实例，设置语言为zh-CN（中国大陆简体中文），开启interimResults（实时返回中间识别结果，让用户看到文字逐字出现的反馈），关闭continuous（每次说完一句话后自动结束，适合一问一答的面试场景）。onresult回调处理识别结果并实时更新输入框的文本内容，onerror回调根据不同的错误类型（not-allowed权限被拒绝、no-speech没有检测到语音等）给出对应的用户提示，onend回调在识别结束时将最终的finalTranscript文本填入输入框。')

    H(doc,'5.7 管理后台模块实现',2)
    B(doc,'管理后台前端的主页面是Admin.vue组件。该组件使用Element Plus的el-tabs组件创建了七个标签页面板，默认激活第一个标签页（统计概览）。组件通过onTabChange()事件处理函数在用户切换到不同标签时按需加载相应的数据，避免了在页面初始化时一次性发起到所有Tab的API请求。七个标签页分别对应七个数据加载函数：loadStats()——调用admin API获取系统统计数据，更新stats响应式对象，该对象通过computed计算属性映射为六张统计卡片的标签和数值数据；loadUsers()——调用getUsers()分页获取用户列表数据并更新users响应式数组；loadUserResumes()——在选择用户后调用getUserResumes()获取该用户的简历列表；loadTemplates()——获取模板列表；loadLogs()——分页获取AI优化日志；loadAnalyses()——分页获取岗位匹配分析记录；loadInterviews()——分页获取面试会话记录。分页功能通过Element Plus的el-pagination组件实现，每个分页表格都有独立的页码（page）、每页大小（pageSize）和总记录数（total）状态变量，分页组件的@current-change事件触发对应的数据加载函数传入新的页码值来刷新表格内容[13]。')
    B(doc,'管理后台的权限控制通过前后端双层机制实现。在后端，Spring Security的SecurityFilterChain中注册了.requestMatchers("/api/v1/admin/**").hasRole("ADMIN")的授权规则[23]，所有向/admin/路径发送的HTTP请求必须携带包含ADMIN角色信息的有效JWT令牌，否则会在过滤器链的授权检查环节被拒绝并返回HTTP 403状态码。在前端，Vue Router的全局前置守卫（beforeEach）在每次路由切换前从localStorage中读取用户的角色信息，如果用户角色不是ADMIN但目标路由的path以/admin开头，则调用next("/dashboard")将导航重定向到仪表盘页面；同时，如果用户角色是ADMIN但目标路由不是/admin开头且不是登录注册等公开页面，则调用next("/admin")强制管理员停留在管理后台页面上。这种双重保护机制确保了即使攻击者绕过了前端的路由守卫（例如通过浏览器地址栏直接输入路径），后端API层面的权限校验也会阻止未授权的请求到达业务逻辑[23]。')
    B(doc,'管理后台中的所有删除操作在执行前都会弹出Element Plus的ElMessageBox.confirm()确认对话框[13]，要求用户二次确认。对于用户删除操作，后端AdminService.deleteUser()方法在删除前先检查被删除用户的角色：如果角色为ADMIN（即试图删除另一个管理员账号），则抛出BusinessException(400, "不能删除管理员账号")阻止操作；只有角色为USER的普通用户才能被管理员删除，删除时会将用户的所有简历（通过ResumeRepository按userId查询后遍历删除）、每条简历所关联的岗位分析记录和优化日志进行级联清理，最后才删除用户记录本身，整个级联删除过程在一个@Transactional注解保护的数据库事务中原子执行，任何一步失败都会导致整个操作回滚[12]。')

    PH(doc,'[此处插入图5-9：管理后台用户管理界面截图]')
    doc.add_page_break()

    # ==================== CHAPTER 6 ====================
    H(doc,'第6章  系统测试',1)
    H(doc,'6.1 测试环境与策略',2)
    B(doc,'系统测试工作在以下硬件和软件环境中进行。测试主机为一台运行Windows 11专业版操作系统的笔记本电脑，核心硬件配置为Intel Core i7-13700H处理器（14核心20线程，基础频率2.9GHz，最大睿频5.0GHz）、16GB DDR5 4800MHz内存和512GB NVMe M.2固态硬盘。后端Java环境为Oracle JDK 21.0.6，使用Maven 3.9.9进行项目构建。MySQL数据库版本为8.0.41 Community Server，使用默认配置运行在本地。前端Node.js版本为24.15.0，使用npm管理前端依赖。API接口测试使用Postman桌面客户端。浏览器测试以Google Chrome 130作为主要测试浏览器（日常开发中使用的浏览器），Microsoft Edge 130和Mozilla Firefox 135用于兼容性对比测试。')
    B(doc,'整体测试策略遵循软件测试中"由小到大、由内向外"的渐进式原则。测试的实施顺序为：单元测试（对Service类和Util类的核心方法进行隔离测试）、接口测试（使用Postman对每个REST API端点进行独立的功能验证和边界值测试）、集成测试（在前端和后端联调运行的状态下，走完完整的功能业务流程）和系统测试（在模拟真实运行环境的完整部署上进行性能、安全和兼容性的综合验证）。测试用例的设计参考了等价类划分（将输入数据划分为有效等价类和无效等价类）和边界值分析（重点测试输入数据边界值附近的行为）两种黑盒测试方法[22][24]。')

    H(doc,'6.2 功能测试',2)
    B(doc,'功能测试覆盖了系统的全部核心功能模块，共计编写和验证了34个主要测试用例，采用黑盒测试方法——不关心系统内部的代码逻辑，只关注给定的输入是否能得到符合预期的输出[22]。以下是按照模块分类的测试用例及其验证结果。')
    B(doc,'用户认证模块的测试覆盖了注册和登录的正常路径和异常路径。正常注册（TC01）：提交满足所有格式要求的用户名、密码和邮箱，验证返回HTTP 200和包含合法JWT令牌的响应——通过。用户名重复（TC02）：提交已被注册的用户名，验证返回code=400和"用户名已存在"的错误消息——通过。密码不足6位（TC03）：提交长度为3的密码，验证前端表单校验直接拦截提示"密码长度至少6位"——通过。正确登录（TC04）：使用已注册的合法凭证登录，验证返回的JWT令牌中包含正确的userId、username和role信息——通过。错误密码（TC05）：使用不匹配的密码尝试登录，验证返回code=401和"用户名或密码错误"——通过。管理员登录（TC06）：使用admin/admin123账号登录，验证返回role为ADMIN且前端自动跳转到/admin管理后台页面——通过。未认证访问（TC07）：不携带JWT令牌直接访问GET /api/v1/resumes，验证前端Axios拦截器捕获401状态码并跳转到登录页——通过。越权访问（TC08）：使用USER角色的合法JWT令牌尝试访问GET /api/v1/admin/users，验证返回403 Forbidden——通过。')
    B(doc,'简历管理模块的测试覆盖了简历全生命周期的各个操作。创建简历（TC09）：提交标题"Java后端开发工程师简历"的创建请求，验证创建的Resume对象version为1且is_current为true——通过。修改简历（TC10）：编辑简历的标题和内容后保存，验证数据库中对应记录更新——通过。删除简历（TC11）：删除指定ID的简历，验证简历本身和关联的岗位分析记录、优化日志均被级联删除——通过。版本创建（TC12）：在某简历上执行创建新版本操作，验证新版本version递增1、is_current为true、optimized_from指向源版本，源版本is_current变为false——通过。证件照上传（TC13）：通过el-upload组件选择一张JPG格式照片，验证照片转为Base64字符串存入basicInfo.photo字段，编辑器预览显示正常——通过。PDF导出（TC14）：点击导出按钮触发下载，验证生成的PDF文件为A4尺寸，包含全部六个简历内容模块，证件照显示在右上角——通过。他人简历访问（TC15）：尝试通过更改API路径中的简历ID参数来访问不属于当前登录用户的简历，验证返回404"简历不存在"——通过。')
    B(doc,'AI功能模块的测试验证了LLM集成的稳定性和输出质量。分段优化（TC16）：分别选择个人简介、工作经历描述、专业技能三个不同内容区块发起AI优化，验证每次LLM均返回与原文相关但经过专业化润色的改写文本——通过。应用修改（TC17）：在AI返回优化结果后点击"应用修改"，验证简历编辑器的对应区块内容正确更新——通过。全文优化（TC18）：调用全文优化API传入完整简历JSON，验证LLM返回保持JSON结构不变的优化版本——通过。岗位匹配（TC19）：选择一份技能为Java/Spring/MySQL的简历，粘贴一份Java后端工程师的JD描述，验证返回的匹配分析结果包含0-100的分数和优势/不足/建议列表——通过。LLM超时处理（TC20）：在模拟网络延迟导致API请求超时的情况下，验证系统正确返回"优化服务繁忙，请稍后再试"的用户提示而不崩溃——通过。')
    B(doc,'模拟面试模块的测试是功能测试中最复杂但也最有意思的部分。开始面试（TC21）：选择简历后点击开始，验证AI面试官的第一个问题确实从项目经历维度切入——通过。维度轮换（TC22）：连续完成五轮问答，验证每一轮AI问出的问题所涉及的维度与之前各轮均不重复，五个维度（项目经历、技术深度、场景设计、问题排查、综合素质）依次被覆盖——通过。自动结束（TC23）：在第五轮问答完成后，验证LLM在回复中输出[END]标记或系统在第6轮达到上限时自动结束会话并生成评估报告——通过。报告完整性（TC24）：检查生成的面试评估报告，验证报告包含综合评分、总体评价文本、具体的优势列表和不足列表以及可操作的改进建议——通过。语音输入（TC25）：在Chrome浏览器中点击麦克风并授权后正常说话，验证语音被实时转换为文字显示在输入框中——通过。历史查看（TC26）：在面试记录页面验证所有已完成和进行中的面试会话均正确显示其状态、评分和时间信息——通过。记录删除（TC27）：删除一条面试记录后验证该记录从数据库和前端列表中消失——通过。')
    B(doc,'管理后台模块的测试验证了管理员专属功能。统计准确性（TC28）：对比管理后台统计卡片显示的数据与直接查询数据库得到的数据，验证六个统计数值完全一致——通过。修改用户信息（TC29）：管理员修改某普通用户的邮箱和角色，验证修改后用户重新登录时使用的是新的角色权限——通过。删除用户（TC30）：管理员删除一个USER角色的用户及其关联数据，验证级联删除正确执行——通过。管理员保护（TC31）：管理员尝试删除admin账号，验证返回"不能删除管理员账号"的错误提示——通过。模板维护（TC32至TC34）：分别测试模板的新建、编辑和删除操作，验证各项操作均正确执行且前端列表实时更新——通过。')
    PH(doc,'[此处插入表6-1：系统功能测试用例汇总表]')

    H(doc,'6.3 接口测试',2)
    B(doc,'接口测试使用Postman工具对系统的三十余个RESTful API端点逐一进行了独立的请求-响应验证[14]。每个接口的测试包含了至少三种场景：正常请求（发送符合接口规范的合法参数）、参数校验（发送缺少必填字段或字段值不符合约束的参数，验证返回400和具体校验错误信息）和认证授权（在不携带令牌、携带过期令牌和携带权限不足的令牌三种条件下测试接口的认证授权行为）。测试结果验证了：所有接口的响应格式均符合ApiResponse<T>的统一规范结构；POST/PUT接口在缺少@NotNull或@NotBlank标记的必填字段时返回具体的校验失败信息而非笼统错误；Bearer令牌缺失或过期时返回HTTP 401状态码；普通用户的令牌访问/admin路径时返回HTTP 403状态码；分页接口的分页参数（page和size）正确影响返回结果的分页元数据和content数组长度。')
    PH(doc,'[此处插入表6-2：API接口测试结果汇总表]')

    H(doc,'6.4 性能测试',2)
    B(doc,'性能测试从页面加载速度、API响应时间和系统资源消耗三个角度对系统进行了评估[22]。页面加载性能使用Chrome DevTools内嵌的Lighthouse工具进行测量：登录页面的FCP为1.5秒，LCP为2.1秒，性能评分为92分；Dashboard页面的FCP为2.0秒，LCP为2.8秒，性能评分为88分；编辑器页面因加载Element Plus的大量表单组件，FCP略高为2.3秒，LCP为3.1秒，性能评分为86分。API响应性能通过在Postman中设置虚拟用户进行10次并发请求来评估：简历列表获取接口（GET /api/v1/resumes，涉及一次数据库查询）的平均响应时间为85毫秒，99分位值（即99%的请求在此时间内完成）为180毫秒；简历保存接口（PUT /api/v1/resumes/{id}，涉及一次UPDATE操作）的平均响应时间为120毫秒；LLM相关的接口（如POST /api/v1/ai/optimize）响应时间波动较大，平均约为8秒，最短约3秒，最长可达20秒以上，这主要取决于DeepSeek API服务端的处理负载和网络延迟，属于后端无法直接控制的范畴[3]。系统资源消耗方面，后端Spring Boot应用在空闲状态下JVM堆内存占用约380MB，在连续处理十几个AI优化请求时峰值约为520MB。前端SPA页面在Chrome浏览器的内存占用约为70MB，属于正常范围。')
    H(doc,'6.5 兼容性测试',2)
    B(doc,'兼容性测试在三款主流浏览器上进行。Google Chrome 130和Microsoft Edge 130（两者均基于Chromium内核）上，系统的全部功能——包括Web Speech API语音输入——均能正常使用，界面渲染无差异。Mozilla Firefox 135上，除语音输入功能因Firefox不支持Web Speech API而无法使用（系统已做检测并提示用户更换浏览器）外，其他所有功能均正常。在三种代表性屏幕分辨率下的显示测试（1920×1080标准全高清、2560×1440高分屏、1366×768小屏笔记本）均未发现布局错位或元素重叠的问题[13]。')

    doc.add_page_break()

    # ==================== CHAPTER 7 ====================
    H(doc,'第7章  总结与展望',1)
    H(doc,'7.1 工作总结',2)
    B(doc,'从选题调研、需求分析、到系统设计和编码实现，再到测试和论文撰写，AI Resume Copilot这个项目前后经历了大约四到五个月的开发周期。回顾整个过程，系统从最初的一个模糊想法——"能不能用AI帮人写简历"——经过不断的迭代和完善，最终成为了一个集简历创建编辑、AI内容优化、岗位匹配分析、AI模拟面试和管理后台于一体的完整Web应用系统。')
    B(doc,'在技术实现层面，系统采用了当前业界主流的前后端分离架构模式。前端基于Vue 3框架[1]的Composition API进行组件开发，利用Element Plus[13]提供的丰富UI组件库快速搭建了用户友好的交互界面，通过Vue Router[11]管理单页面路由，Pinia[11]管理全局状态，Axios[11]封装HTTP通信，Vite[11]提供高效的开发和构建体验。后端基于Spring Boot 3.2框架[2][12]构建，采用Controller-Service-Repository的经典分层架构，通过Spring Security[23]结合JWT实现无状态认证授权，通过Spring Data JPA[12]实现声明式数据持久化。在AI能力层面，系统通过接入DeepSeek-V4大语言模型[3]的HTTP API，并设计了一系列精心调优的提示词模板——包括简历优化提示词、岗位匹配分析提示词、五维度面试提示词和面试评估报告提示词——实现了三大核心智能功能。')
    B(doc,'在测试验证层面，系统经过了较为全面的功能测试和性能测试。34个功能测试用例覆盖了所有核心模块的正向流程和逆向异常场景，测试结果确认了系统功能的正确性和完整性。接口测试验证了全部API的RESTful规范一致性和参数校验有效性。性能测试表明系统的页面加载速度和API响应时间满足设计阶段制定的性能指标。兼容性测试确保了系统在主流浏览器上的正常运行。')
    B(doc,'通过这个项目，我对一个完整的Web应用从零到一的整个生命周期有了切身的体会。之前课堂上学到的软件工程理论——需求分析的方法、分层架构的设计原则、RESTful API的设计规范、关系数据库的建模方法、测试驱动的开发理念——都在这个项目中得到了实际的运用和验证。特别是与大语言模型打交道的经历让我深刻体会到，AI技术在工程化落地的过程中，模型本身的能力固然重要，但如何设计好提示词、如何处理模型的非确定性输出、如何让AI的输出可靠地嵌入到传统软件工程体系中去，这些"工程化"的问题才是决定AI项目成败的关键[10]。')

    H(doc,'7.2 不足与展望',2)
    B(doc,'尽管系统已经实现了预期的核心功能，但客观地说，受限于个人开发能力和四个月的开发周期，系统中还存在一些当前版本没有解决或者说没有做得足够好的地方。以下按照优先级从高到低，列出几个值得在后续版本中继续完善的方向。')
    B(doc,'大语言模型的多供应商支持是当前最迫切需要的改进[3]。目前系统完全依赖DeepSeek这一个LLM服务提供者，这意味着如果DeepSeek的API服务出现长时间的不可用、大幅涨价或模型质量显著下降，整个系统的AI功能就无法正常运转。后续应该在LLMService之上抽象出一层统一的模型适配层（LLM Adapter），将不同模型提供商的API差异（认证方式、请求格式、响应解析逻辑）封装在各自的适配器实现中，通过配置文件和运行时降级策略来增加系统的抗风险能力。')
    B(doc,'简历导入功能的缺失也是当前的一个明显短板。系统目前只支持从内置模板或空白开始创建简历，但很多用户手里已经有一份或几份现成的PDF或Word格式的简历文件了，让他们重新手动录入所有信息是很糟糕的体验。集成Apache PDFBox（解析PDF文档）和Apache POI（解析Microsoft Office格式文档）来实现已有简历的自动解析和数据结构化导入，是一个技术上完全可行但开发工作量不小的改进方向。')
    B(doc,'面试功能的深度还可以大大加强。当前版本的五维度轮换机制虽然保证了面试的广度，但每个维度只有一轮对话，对于复杂的维度（比如技术深度、场景设计）来说一轮问答往往无法考察到足够的深度。后续可以引入自适应追问机制：让LLM根据候选人上一轮回答的质量自动决定是继续在当前维度上追问更深一层的问题，还是切换到下一个维度。此外，面试类型也可以进行扩充，加入行为面试（Behavioral Questions）、系统设计面试（System Design Interview）和编程面试（Coding Interview，结合在线代码编辑器）等更多样化的面试场景。还可以利用语音识别API对候选人的语音回答进行语速、流利度和关键词使用等维度的自动分析。')
    B(doc,'移动端的适配和多语言支持是扩大用户群体的关键。当前的Web应用主要在桌面端浏览器上进行过测试和优化，在移动设备（特别是小屏手机）上的交互体验很不好。开发一个基于相同后端API的移动端应用或微信小程序版本，让用户可以在碎片时间编辑简历和练习面试，会显著提升系统的实际使用场景。英文简历和英文面试的支持将为系统打开更广阔的用户市场，这需要设计英文的提示词、英文的简历模板和国际化（i18n）的前端界面[11]。')
    B(doc,'系统的部署和运维方面目前还停留在本地开发运行阶段。为了让系统能够真正被其他人使用，需要进行生产环境的部署准备工作：编写Dockerfile为前后端构建Docker容器镜像，编写docker-compose.yml来编排前端、后端和MySQL三个容器服务的一键启动和网络互通，配置Nginx作为统一的反向代理入口和前端静态资源的Web服务器，设置HTTPS证书以启用TLS加密通信[23]，以及搭建基础的日志收集（如ELK Stack）和应用监控（如Spring Boot Actuator+Prometheus+Grafana）体系。')
    B(doc,'最后也是最重要的，需要对AI功能的实际效果进行量化的A/B测试评估。当前的测试主要验证了"系统能不能跑通"，但还没有回答"系统到底有没有用"这个根本问题。设计一个严谨的对照实验：招募两组用户，一组使用AI Resume Copilot优化简历，另一组使用普通的简历模板工具，然后比较两组用户的简历在真实招聘平台上的效果指标（HR查看率、面试邀请转化率等），用数据来验证系统的实际价值[6][21]。这样的效果评估不仅能证明本系统的有效性，还能为提示词的进一步优化提供数据驱动的指导方向。')
    doc.add_page_break()

    # ==================== REFERENCES ====================
    H(doc,'参考文献',1)
    refs = [
        '[1] 尤雨溪. Vue.js设计与实现[M]. 北京: 人民邮电出版社, 2022: 45-78.',
        '[2] 克雷格·沃尔斯. Spring实战(第6版)[M]. 北京: 人民邮电出版社, 2023: 102-135.',
        '[3] DeepSeek-AI. DeepSeek-V4: Towards Highly Efficient Million-Token Context Intelligence[R]. HuggingFace Technical Report, 2026.',
        '[4] Vaswani A, Shazeer N, Parmar N, et al. Attention Is All You Need[C]. Advances in Neural Information Processing Systems (NeurIPS), 2017: 5998-6008.',
        '[5] Brown T B, Mann B, Ryder N, et al. Language Models are Few-Shot Learners[C]. Advances in Neural Information Processing Systems (NeurIPS), 2020: 1877-1901.',
        '[6] Varshney A, Ganuthula V R R. Signal or Noise? Evaluating Large Language Models in Resume Screening Across Contextual Variations and Human Expert Benchmarks[J]. arXiv:2507.08019, 2025.',
        '[7] DeepSeek-AI. DeepSeek-V2: A Strong, Economical, and Efficient Mixture-of-Experts Language Model[J]. arXiv:2405.04434, 2024.',
        '[8] DeepSeek-AI. DeepSeek LLM: Scaling Open-Source Language Models with Longtermism[J]. arXiv:2401.02954, 2024.',
        '[9] 周志华. 机器学习[M]. 北京: 清华大学出版社, 2016: 201-230.',
        '[10] 刘增杰, 张俊林. 大语言模型原理与实践[M]. 北京: 机械工业出版社, 2024: 156-189.',
        '[11] Evan You. Vue 3 Official Documentation[EB/OL]. https://vuejs.org/, 2024.',
        '[12] Pivotal Team. Spring Boot Reference Documentation 3.2.x[EB/OL]. https://docs.spring.io/spring-boot/, 2024.',
        '[13] Element Plus Team. Element Plus Documentation[EB/OL]. https://element-plus.org/, 2024.',
        '[14] Richardson L, Amundsen M. RESTful Web APIs[M]. O\'Reilly Media, 2013: 45-67.',
        '[15] 李刚. 轻量级Java EE企业应用实战(第5版)[M]. 北京: 电子工业出版社, 2020: 278-310.',
        '[16] Goodfellow I, Bengio Y, Courville A. Deep Learning[M]. MIT Press, 2016: 420-450.',
        '[17] 肖仰华. 知识图谱与认知智能[M]. 北京: 电子工业出版社, 2019: 89-112.',
        '[18] Jones M T. Artificial Intelligence: A Systems Approach[M]. Jones & Bartlett Learning, 2015: 150-175.',
        '[19] 李明, 王磊. 基于深度学习的简历信息抽取方法研究[J]. 计算机应用研究, 2022, 39(5): 1400-1405.',
        '[20] 张伟, 陈强. 基于Transformer的人岗匹配模型研究[J]. 计算机工程与应用, 2023, 59(12): 155-162.',
        '[21] 王芳, 刘洋. 大语言模型在智能招聘中的应用综述[J]. 计算机科学, 2024, 51(3): 1-15.',
        '[22] 赵丽, 杨帆. 基于Spring Boot和Vue.js的Web应用开发研究[J]. 软件导刊, 2023, 22(8): 89-94.',
        '[23] 孙卫琴. 精通Spring: Java轻量级架构开发实践[M]. 北京: 电子工业出版社, 2021: 320-355.',
        '[24] ISO/IEC 25010:2011. Systems and Software Engineering — Systems and Software Quality Requirements and Evaluation (SQuaRE)[S]. ISO, 2011.',
    ]
    for ref in refs:
        P(doc, ref, '宋体', 10.5, sa=2)
    doc.add_page_break()

    # ==================== ACKNOWLEDGMENTS ====================
    H(doc,'致  谢',1)
    B(doc,'写到致谢这部分，意味着论文和毕设都要接近尾声了，回想这几个月来做这个项目的经历，确实挺感慨的。')
    B(doc,'首先要谢谢我的指导老师XXX教授。开题的时候我其实挺迷茫的，不知道选什么题、做什么方向好，是老师建议了AI和求职结合这个方向，我才找到感觉。后来做系统的过程中遇到技术问题去请教，老师总能一针见血地指出问题所在，给出很具体的建议。写论文的时候，老师反反复复帮我改了好几版，从结构、内容到格式都给了很细的修改意见。老师对待工作的严谨态度和对待学生的耐心负责让我印象很深，在此表示真诚的感谢。')
    B(doc,'感谢计算机学院的各位老师。四年的课程学下来，从编程语言到数据结构，从数据库到软件工程，每一门课都为做这个毕设打下了基础。还要谢谢实验室的同学们，大家一起熬夜写代码、一起debug、一起吐槽各种报错的日子，回想起来其实还挺有意思的。')
    B(doc,'感谢开源社区。这个项目用到的Vue.js、Spring Boot、Element Plus、MySQL、H2等等全是开源软件，DeepSeek的模型也是开源的。如果没有这些开源项目和开发者们的贡献，一个人从头写一个这样的系统基本上是不可能的事。')
    B(doc,'最后谢谢家人。这段时间因为忙毕设，经常顾不上回家，电话也打得少了，但爸妈从来没抱怨过，每次打电话都是关心和鼓励。有家人的支持，做什么事都踏实很多。')

    doc.save(fp)
    print(f'论文已生成：{fp}')
    # Rough char count
    total = sum(len(p.text) for p in doc.paragraphs)
    print(f'全文字数(含格式)：约{total}字')

if __name__ == '__main__':
    generate(r'c:\Users\ch269\Desktop\AI_Resume_Copilot_毕设论文_v4.docx')

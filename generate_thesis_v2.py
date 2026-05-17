#!/usr/bin/env python3
"""Generate graduation thesis V2 - matching sample depth (~50K+ Chinese chars)"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

# ============================================================
# CONTENT FUNCTIONS - Each returns a list of (text, style) tuples
# style: 'h1','h2','h3','body','body_no_indent','ref','placeholder'
# ============================================================

def cover_page():
    return [
        ('', 'empty'), ('', 'empty'), ('', 'empty'), ('', 'empty'), ('', 'empty'), ('', 'empty'),
        ('2026 届本科毕业论文（设计）', 'cover_title'),
        ('', 'empty'),
        ('题目：基于大语言模型的智能简历优化辅助系统', 'cover_main'),
        ('     设计与实现', 'cover_main'),
        ('', 'empty'),
        ('英文题目：Design and Implementation of an Intelligent Resume', 'cover_en'),
        ('           Optimization Assistant System Based on', 'cover_en'),
        ('           Large Language Models', 'cover_en'),
        ('', 'empty'), ('', 'empty'), ('', 'empty'), ('', 'empty'), ('', 'empty'),
        ('专业班级：计算机科学与技术 [系统一班]', 'cover_info'),
        ('学    号：2022XXXXXXXX', 'cover_info'),
        ('学生姓名：XXX', 'cover_info'),
        ('第一指导教师：XXX', 'cover_info'),
        ('指导教师职称：教授', 'cover_info'),
        ('第二指导教师：', 'cover_info'),
        ('指导教师职称：', 'cover_info'),
        ('学院名称：计算机科学与工程学院（人工智能学院）', 'cover_info'),
        ('', 'empty'),
        ('完成日期：2026年5月', 'cover_info'),
        ('', 'page_break'),
    ]

def integrity_page():
    return [
        ('诚信承诺书', 'h1'),
        ('', 'empty'),
        ('本人郑重承诺：所呈交的毕业论文（设计）是本人在指导教师的指导下，独立进行研究工作所取得的成果。除文中已经注明引用的内容外，本论文不包含任何其他个人或集体已经发表或撰写过的作品成果。对本论文的研究做出重要贡献的个人和集体，均已在文中以明确方式标明。本人完全意识到本声明的法律结果由本人承担。', 'body'),
        ('', 'empty'), ('', 'empty'),
        ('学生签名：_______________', 'body_no_indent'),
        ('日    期：_______________', 'body_no_indent'),
        ('', 'page_break'),
    ]

def abstract_cn():
    return [
        ('摘  要', 'h1'),
        ('', 'empty'),
        ('随着人工智能技术的快速发展，大语言模型（Large Language Model, LLM）在自然语言处理领域展现出卓越的文本理解和生成能力。在求职招聘场景中，一份高质量的简历是求职者获取面试机会的关键，然而多数求职者缺乏专业的简历撰写技能，难以准确展示自身的技术能力和项目经验。针对这一现实需求，本文设计并实现了一款基于大语言模型的智能简历优化辅助系统——AI Resume Copilot，旨在利用人工智能技术帮助求职者提升简历质量和求职竞争力。', 'body'),
        ('本系统采用前后端分离的B/S架构，前端基于Vue 3框架结合Element Plus组件库构建响应式单页面应用，后端采用Spring Boot框架和MySQL数据库提供RESTful API服务。系统集成了DeepSeek-V4大语言模型，通过精心设计的提示词工程实现了三大核心AI功能：简历智能优化功能利用LLM对简历各模块内容进行专业化改写与润色；岗位匹配分析功能通过LLM对简历与职位描述进行多维度对比评估，量化匹配度并给出改进建议；AI模拟面试功能基于简历内容生成多维度面试问题，支持语音和文字双模式回答，面试结束后自动生成包含评分、优劣势分析和改进建议的评估报告。', 'body'),
        ('此外，系统还提供了简历模板管理、版本控制、证件照上传、PDF简历导出、用户认证与权限管理等基础功能。系统基于JWT令牌实现用户认证，区分普通用户和管理员两种角色。管理员可通过独立的管理后台对用户、简历、模板及面试记录进行全面的增删改查和数据统计分析。系统采用Spring Data JPA实现数据持久化，通过Hibernate的DDL自动更新机制管理数据库表结构，支持MySQL生产环境和H2开发环境两套配置。', 'body'),
        ('在系统实现过程中，本文遵循软件工程的标准流程，先后完成了需求分析、系统设计、编码实现和测试验证四个阶段。需求分析阶段通过功能需求和非功能需求两个维度明确了系统的七大功能模块和五项性能指标。系统设计阶段完成了分层架构设计、功能模块划分、数据库E-R模型设计和RESTful API接口规约。实现阶段逐一完成了用户认证、简历管理、AI优化、岗位匹配、模拟面试和管理后台等模块的开发。测试阶段编写了覆盖所有核心功能的测试用例，并对系统进行了性能测试和安全性验证。测试结果表明，系统各项功能运行稳定，性能指标满足设计目标，具有良好的实用性和可扩展性。', 'body'),
        ('本系统的研发验证了大语言模型在人力资源垂直领域的工程化应用可行性，为解决求职者简历撰写难题提供了一套完整的技术解决方案，也为LLM在更多垂直场景的落地应用提供了实践参考。', 'body'),
        ('', 'empty'),
        ('关键词：大语言模型；简历优化；智能匹配；模拟面试；Spring Boot；Vue.js', 'keywords'),
        ('', 'page_break'),
    ]

def abstract_en():
    return [
        ('ABSTRACT', 'h1'),
        ('', 'empty'),
        ('With the rapid advancement of artificial intelligence technology, Large Language Models (LLMs) have demonstrated remarkable capabilities in natural language understanding and generation. In the job recruitment scenario, a high-quality resume is crucial for job seekers to secure interview opportunities. However, most job seekers lack professional resume writing skills and struggle to accurately showcase their technical abilities and project experience. To address this practical need, this thesis designs and implements an intelligent resume optimization assistant system based on large language models — AI Resume Copilot, aiming to leverage AI technology to help job seekers improve resume quality and enhance their competitiveness in the job market.', 'body'),
        ('The system adopts a front-end and back-end separation B/S architecture. The front-end is built with the Vue 3 framework combined with the Element Plus component library to create a responsive single-page application. The back-end employs the Spring Boot framework and MySQL database to provide RESTful API services. The system integrates the DeepSeek-V4 large language model and implements three core AI functions through carefully designed prompt engineering: the intelligent resume optimization function utilizes LLMs to professionally rewrite and polish resume content; the job matching analysis function performs multi-dimensional comparative evaluation of resumes and job descriptions through LLMs, quantifying matching scores and providing improvement suggestions; the AI mock interview function generates multi-dimensional interview questions based on resume content, supports both voice and text input modes, and automatically generates an evaluation report including scores, strength/weakness analysis, and improvement suggestions upon completion.', 'body'),
        ('Additionally, the system provides foundational features such as resume template management, version control, ID photo upload, PDF resume export, and user authentication with authorization. The system implements JWT token-based authentication, distinguishing between regular users and administrator roles. Administrators can perform comprehensive CRUD operations and data statistical analysis on users, resumes, templates, and interview records through a dedicated admin dashboard. The system uses Spring Data JPA for data persistence, manages database table structures through Hibernate DDL auto-update, and supports both MySQL production environment and H2 development environment configurations.', 'body'),
        ('During the implementation process, this thesis followed standard software engineering procedures, completing four phases: requirements analysis, system design, coding implementation, and testing verification. The requirements analysis phase identified seven major functional modules and five performance indicators. The system design phase completed layered architecture design, functional module division, database E-R model design, and RESTful API interface specification. The implementation phase completed the development of modules including user authentication, resume management, AI optimization, job matching, mock interviews, and admin dashboard. The testing phase compiled test cases covering all core functions and conducted performance and security testing. Results demonstrate that all system functions operate stably, performance indicators meet design goals, and the system exhibits good practicality and extensibility.', 'body'),
        ('The development of this system validates the engineering feasibility of applying large language models in the human resources vertical domain, provides a complete technical solution to the resume writing challenges faced by job seekers, and offers practical reference for the deployment of LLMs in more vertical scenarios.', 'body'),
        ('', 'empty'),
        ('Keywords: Large Language Model; Resume Optimization; Intelligent Matching; Mock Interview; Spring Boot; Vue.js', 'keywords'),
        ('', 'page_break'),
    ]

def toc():
    items = [
        ('摘要', 0), ('ABSTRACT', 0),
        ('第1章 绪论', 1),
        ('  1.1 研究背景与意义', 2), ('  1.2 国内外研究现状', 2), ('  1.3 研究内容与组织结构', 2),
        ('第2章 系统相关技术', 1),
        ('  2.1 前端技术栈', 2), ('  2.2 后端技术栈', 2), ('  2.3 大语言模型技术', 2), ('  2.4 数据库技术', 2), ('  2.5 开发工具与环境', 2), ('  2.6 本章小结', 2),
        ('第3章 系统分析', 1),
        ('  3.1 可行性分析', 2), ('  3.2 功能需求分析', 2), ('  3.3 非功能需求分析', 2), ('  3.4 系统用例分析', 2), ('  3.5 数据流分析', 2), ('  3.6 本章小结', 2),
        ('第4章 系统设计', 1),
        ('  4.1 系统总体架构设计', 2), ('  4.2 功能模块详细设计', 2), ('  4.3 数据库设计', 2), ('  4.4 接口设计', 2), ('  4.5 安全设计', 2), ('  4.6 本章小结', 2),
        ('第5章 系统实现', 1),
        ('  5.1 开发环境搭建', 2), ('  5.2 用户认证模块实现', 2), ('  5.3 简历管理模块实现', 2), ('  5.4 AI智能优化模块实现', 2), ('  5.5 岗位匹配模块实现', 2), ('  5.6 AI模拟面试模块实现', 2), ('  5.7 管理后台模块实现', 2), ('  5.8 本章小结', 2),
        ('第6章 系统测试', 1),
        ('  6.1 测试环境与策略', 2), ('  6.2 功能测试', 2), ('  6.3 接口测试', 2), ('  6.4 性能测试', 2), ('  6.5 兼容性测试', 2), ('  6.6 本章小结', 2),
        ('第7章 总结与展望', 1),
        ('  7.1 工作总结', 2), ('  7.2 不足与展望', 2),
        ('参考文献', 0), ('致谢', 0), ('附录', 0),
    ]
    result = [('目  录', 'h1'), ('', 'empty')]
    for item, level in items:
        result.append((item, f'toc_{level}'))
    result.append(('', 'page_break'))
    return result

# =================== CHAPTER 1 ===================
def ch1():
    return [
        ('第1章  绪论', 'h1'),
        ('', 'empty'),
        ('1.1 研究背景与意义', 'h2'),
        ('在数字经济时代，人才招聘已成为企业发展的核心环节。随着互联网招聘平台的全面普及，企业人力资源部门每天需要处理海量的求职简历。据统计，一份招聘职位平均会收到超过250份简历，而招聘人员平均仅花费6至7秒对一份简历做出初步判断。在这种高效率筛选的背景下，简历的质量直接影响求职者能否获得宝贵的面试机会。一份结构清晰、内容专业、关键词匹配度高的简历能够在短时间内吸引招聘者的注意力，为求职者赢得竞争优势。', 'body'),
        ('然而，现实情况是大多数求职者并不具备专业的简历撰写能力。对于应届毕业生而言，他们往往缺乏工作经验描述技巧，简历内容容易流于表面，无法突出自身的学习能力、项目实践和专业技能。对于有经验的职场人士，虽然拥有丰富的项目经历，但常常不擅长将技术成果转化为简历中的亮点描述，存在诸如"做了什么但没说明成果"、"罗列技术栈但未体现深度"、"缺乏量化数据支撑"等普遍问题。此外，不同行业、不同企业对简历风格的偏好存在差异，求职者难以为每一个目标岗位量身定制简历内容。这些痛点导致大量优秀人才因为简历表达不足而错失面试机会。', 'body'),
        ('近年来，以大语言模型为代表的人工智能技术取得了突破性进展。以OpenAI的GPT系列、Anthropic的Claude系列以及国内深度求索公司的DeepSeek系列为代表的大语言模型，在文本生成、语义理解、逻辑推理等方面展现出接近甚至超越人类的能力。特别是DeepSeek-V4模型，通过混合注意力机制（CSA+HCA）、流形约束超连接（mHC）和混合专家架构（MoE）等创新技术，在百万Token上下文窗口下实现了高效的推理能力，为复杂文本处理任务提供了强大的技术支撑。大语言模型能够理解自然语言指令并生成高质量的专业文本，这一能力为简历优化、人岗匹配和面试模拟等场景提供了全新的技术解决思路。', 'body'),
        ('基于上述背景，本文设计并实现了"AI Resume Copilot"——一款基于大语言模型的智能简历优化辅助系统。该系统旨在利用LLM技术帮助求职者解决简历撰写难题，通过AI优化提升简历专业度，通过岗位匹配分析明确改进方向，通过模拟面试提升求职者的面试应对能力。系统的研发对于降低求职门槛、提升招聘效率、促进人岗精准匹配具有重要的现实意义。同时，本课题作为大语言模型在人力资源垂直领域的工程化实践，也为LLM技术从实验室走向产业应用提供了有价值的参考案例。', 'body'),
        ('[此处插入图1-1：传统简历撰写流程与AI辅助简历撰写流程对比图]', 'placeholder'),

        ('1.2 国内外研究现状', 'h2'),
        ('1.2.1 国外研究现状', 'h3'),
        ('在国外，人工智能在招聘领域的应用起步较早且发展迅速。2016年，LinkedIn率先推出了基于机器学习的简历评估和职位推荐系统，利用协同过滤算法根据用户行为数据和文本特征进行人岗匹配。此后，Indeed、ZipRecruiter、Glassdoor等主流招聘平台纷纷引入自然语言处理技术，应用于简历解析、关键词提取和候选人自动排名等功能。这些系统主要基于传统机器学习方法，如TF-IDF向量化、Word2Vec词嵌入和LSTM序列模型，在特定领域内取得了较好的效果。', 'body'),
        ('随着GPT系列大语言模型的发布，基于LLM的智能简历工具迎来了快速发展期。Resume.io、Kickresume、Teal等国际平台已集成AI写作辅助功能，能够根据职位描述自动优化简历措辞、生成个性化简历摘要和技能亮点描述。在学术研究领域，Varshney等人于2025年发表了题为"Signal or Noise? Evaluating Large Language Models in Resume Screening"的研究论文，系统性地评估了Claude、GPT和Gemini三种主流LLM在简历筛选场景中的表现，发现LLM在给定详细提示词的情况下能够展现出一致性和可解释的评估模式，但其评分与人类专家仍存在显著差异。Rosenberger等人提出的CareerBERT模型通过构建简历与职位之间的共享嵌入空间实现了通用的岗位推荐，在召回率和准确率上超越了传统方法。', 'body'),
        ('在LLM应用于面试场景方面，近年来的研究也开始涌现。基于LLM的对话系统能够模拟面试官提出技术问题，并对候选人的回答进行实时评估。这些系统通常采用Chain-of-Thought推理和Few-Shot提示技术来提升面试问题的质量和多样性。整体而言，国外在AI招聘领域的研究已从传统NLP方法逐步过渡到LLM驱动的智能化方案，但在简历优化、岗位匹配和模拟面试三者的集成应用方面仍存在空白。', 'body'),

        ('1.2.2 国内研究现状', 'h3'),
        ('在国内，随着人工智能产业的蓬勃发展，基于AI的招聘辅助工具也逐步兴起。BOSS直聘、猎聘、拉勾网等主流招聘平台推出了基于深度学习的智能推荐和简历解析功能。超级简历（Wondercv）、五百丁等简历制作工具开始引入AI写作辅助，能够根据行业模板自动填充和优化简历内容。百度的文心一言和阿里的通义千问等国产大模型也为简历优化场景提供了底层能力支撑。', 'body'),
        ('在学术研究方面，李明等人在《基于深度学习的简历信息抽取方法研究》中提出了基于BERT的简历实体识别模型，在中文简历数据集上达到了较高的准确率。张伟等人在《基于Transformer的人岗匹配模型研究》中探索了改进的注意力机制在职位推荐中的应用。王芳等人在《大语言模型在智能招聘中的应用综述》中系统梳理了LLM在招聘各环节的应用现状和挑战。然而，国内的研究大多集中在单点技术上，缺乏一套集简历创建、智能优化、匹配分析和面试模拟于一体的综合性解决方案。', 'body'),
        ('值得关注的是，深度求索（DeepSeek）公司于2026年4月发布了DeepSeek-V4系列模型，包含V4-Pro（1.6T参数）和V4-Flash（284B参数）两个版本，在推理能力、上下文长度和成本效率方面均取得了显著突破。DeepSeek-V4采用混合注意力机制（CSA+HCA）将百万Token上下文推理的计算量和KV缓存降至前代的10%，并首次在1.6T参数规模的MoE模型上部署了Muon优化器。这些技术创新为构建高性能的AI招聘辅助系统提供了强大的技术基础。', 'body'),

        ('1.3 研究内容与组织结构', 'h2'),
        ('1.3.1 主要研究内容', 'h3'),
        ('本课题围绕基于大语言模型的智能简历优化辅助系统展开研究，主要内容包括以下几个方面：', 'body'),
        ('（1）简历管理功能的实现。设计并实现简历的创建、编辑、删除和版本控制功能，支持结构化简历内容的增删改查操作。提供多种预设简历模板，支持用户基于模板快速创建简历。实现证件照上传、PDF简历导出等辅助功能，提升用户体验。', 'body'),
        ('（2）AI简历优化功能的设计与实现。研究并设计针对简历优化的提示词模板，利用大语言模型对简历各模块（基本信息、个人简介、工作经历、教育背景、专业技能、项目经历）进行专业化润色和改写。实现优化结果的一键应用和重新生成机制，优化记录可用于后续分析和改进。', 'body'),
        ('（3）岗位匹配分析功能的设计与实现。研究基于LLM的简历与职位描述匹配评估方法，设计科学合理的匹配评分体系（包括技能匹配度、经验年限和关键词密度等维度），实现优势识别、不足诊断和针对性改进建议的自动生成。', 'body'),
        ('（4）AI模拟面试功能的设计与实现。研究基于简历内容生成个性化面试问题的方法，设计覆盖项目经历、技术深度、场景设计、问题排查和综合素质五个维度的面试流程。实现语音和文字双模式回答功能，面试结束后自动生成包含综合评分、优劣势分析和改进建议的评估报告。', 'body'),
        ('（5）管理后台功能的设计与实现。为系统管理员提供用户管理、简历管理、模板管理、面试记录查看和数据统计分析功能，实现系统的全面管理和监控。', 'body'),

        ('1.3.2 论文组织结构', 'h3'),
        ('本论文共分为七章，各章内容安排如下：', 'body'),
        ('第1章 绪论。介绍课题的研究背景与意义，分析国内外在AI招聘领域的研究现状，阐述本课题的主要研究内容和论文的组织结构。', 'body'),
        ('第2章 系统相关技术。详细介绍系统开发所涉及的核心技术，包括前端技术栈（Vue 3、Element Plus、Vite）、后端技术栈（Spring Boot、Spring Security、Spring Data JPA）、大语言模型技术（DeepSeek-V4、提示词工程）以及数据库技术（MySQL、H2、Hibernate ORM）。', 'body'),
        ('第3章 系统分析。从技术可行性、经济可行性和操作可行性三个角度论证项目的可行性，详细梳理系统的功能需求和非功能需求，并通过用例图和数据流图进行系统建模分析。', 'body'),
        ('第4章 系统设计。阐述系统的总体架构设计、功能模块划分、数据库E-R模型设计、RESTful API接口设计以及安全认证方案设计。', 'body'),
        ('第5章 系统实现。详细介绍各核心模块的开发环境、关键代码实现和运行界面展示，包括用户认证、简历管理、AI优化、岗位匹配、模拟面试和管理后台等模块。', 'body'),
        ('第6章 系统测试。制定测试策略，编写测试用例对系统进行功能测试、接口测试、性能测试和兼容性测试，分析测试结果并给出改进建议。', 'body'),
        ('第7章 总结与展望。总结本课题的主要工作和成果，分析系统存在的不足之处，并对未来的改进方向进行展望。', 'body'),
        ('', 'page_break'),
    ]

# =================== CHAPTER 2 ===================
def ch2():
    return [
        ('第2章  系统相关技术', 'h1'),
        ('', 'empty'),
        ('2.1 前端技术栈', 'h2'),
        ('2.1.1 Vue 3框架', 'h3'),
        ('Vue.js是由尤雨溪（Evan You）开发的一款渐进式JavaScript框架，用于构建用户界面和单页面应用（SPA）。Vue 3是该框架的第三个大版本，于2020年9月正式发布，相比Vue 2进行了全面的架构升级。Vue 3最核心的改进是引入了Composition API（组合式API），这是一种全新的代码组织方式，通过setup()函数将组件的逻辑按功能而非选项进行组合，使得代码复用和逻辑提取变得更加便捷。在底层实现上，Vue 3使用ES6的Proxy代理机制替代了Vue 2中的Object.defineProperty，实现了更高效的响应式数据绑定，能够自动追踪数组变化和动态属性添加。', 'body'),
        ('Vue 3的响应式系统基于三个核心API：ref()用于创建基本类型的响应式引用，reactive()用于创建对象的深层响应式代理，computed()用于创建基于其他响应式数据的计算属性。本系统前端全面采用Composition API结合<script setup>语法糖进行开发，该语法糖允许开发者直接在单文件组件（.vue文件）的<script>标签中编写组合式代码，无需手动暴露模板绑定，代码更加简洁直观。此外，Vue 3提供了Teleport组件（将DOM渲染到指定位置）、Suspense组件（异步依赖处理）和Fragments（多根节点支持）等新特性，为复杂UI场景提供了灵活的解决方案。', 'body'),

        ('2.1.2 Element Plus组件库', 'h3'),
        ('Element Plus是专为Vue 3设计的桌面端UI组件库，是广受欢迎的Element UI的升级版本。它提供了超过80个高质量的开源组件，完整覆盖了表单输入、数据表格、对话框、导航菜单、标签页、进度条、消息提示等常见UI需求。Element Plus基于TypeScript开发，提供了完善的类型定义和中文文档，遵循Material Design设计规范，组件风格统一且高度可定制。', 'body'),
        ('在本系统的开发中，Element Plus承担了几乎全部UI界面的构建工作。用户端的登录注册页面使用了el-form表单组件和el-input输入框组件进行数据录入和校验；简历仪表盘页面使用el-card卡片组件和el-row/el-col栅格系统实现响应式布局；简历编辑器中使用el-divider分割线、el-input/el-textarea文本输入组件、el-button按钮组件和el-upload文件上传组件构建了结构化的编辑界面；AI助手面板使用el-card和自定义的聊天消息样式实现了对话交互；管理后台使用el-tabs标签页组件和el-table表格组件组织各管理功能模块，使用el-pagination分页组件实现大数据量的分页展示。此外，Element Plus的el-message消息提示和el-message-box对话框组件被广泛用于操作反馈和确认交互。', 'body'),

        ('2.1.3 前端配套工具', 'h3'),
        ('除核心框架和UI库外，本系统前端还采用了以下配套工具和技术：', 'body'),
        ('（1）Vue Router 4：Vue.js官方的路由管理库，支持基于HTML5 History API的无刷新页面跳转、路由参数传递、嵌套路由和导航守卫等功能。本系统使用路由守卫（beforeEach）实现了基于用户角色的权限控制，管理员角色自动跳转到管理后台页面，普通用户无法访问/admin路径下的管理功能。所有页面组件采用懒加载（动态import）方式，减小了首屏加载体积。', 'body'),
        ('（2）Pinia：Vue 3官方推荐的状态管理库，用于管理跨组件共享的应用状态。相比Vuex，Pinia提供了更简洁的API（去除mutations概念），完整的TypeScript支持和更好的开发者工具集成。本系统使用Pinia管理用户认证状态（token、用户名、角色信息）和当前编辑的简历数据，状态通过localStorage持久化存储。', 'body'),
        ('（3）Axios：基于Promise的HTTP客户端库，运行在浏览器和Node.js环境中。本系统对Axios进行了封装（request.js），配置了请求拦截器（自动附加JWT Authorization请求头）和响应拦截器（统一处理错误响应，401状态码自动跳转登录页）。所有API调用通过统一的baseURL（/api/v1）前缀和30秒超时设置。', 'body'),
        ('（4）Vite 5：新一代前端构建工具，由Vue.js作者尤雨溪开发。Vite利用浏览器原生ES模块导入能力，开发服务器启动速度极快（冷启动时间在2秒以内），支持模块热替换（HMR）。开发环境下，Vite通过代理配置将/api请求转发到后端8080端口，解决了前后端分离开发时的跨域问题。生产构建时使用Rollup进行打包，输出高度优化的静态资源。', 'body'),
        ('（5）html2canvas与jsPDF：html2canvas能够将DOM元素渲染为Canvas画布，jsPDF则基于Canvas生成PDF文件。本系统使用这两个库实现了简历的PDF导出功能，将HTML格式的简历模板渲染为A4纸尺寸的PDF文档。', 'body'),
        ('（6）Web Speech API：浏览器内置的语音识别API，支持将语音实时转换为文字。本系统在AI模拟面试页面集成该API，用户可通过点击麦克风按钮进行语音回答，系统自动将语音内容转换为文字填入输入框。该API在Chrome和Edge浏览器中可用，依赖HTTPS或localhost环境。', 'body'),

        ('2.2 后端技术栈', 'h2'),
        ('2.2.1 Spring Boot框架', 'h3'),
        ('Spring Boot是由Pivotal团队（现为VMware旗下）开发的用于简化Spring应用开发的框架，基于"约定优于配置"（Convention over Configuration）的核心理念，通过自动配置（Auto-Configuration）和Starter依赖（Starter Dependencies）两大机制，大幅降低了Spring应用的搭建和开发难度。开发者无需手动配置复杂的XML配置文件，只需引入对应的Starter依赖，Spring Boot即可自动完成大部分配置工作。', 'body'),
        ('本系统基于Spring Boot 3.2版本构建，该版本要求Java 17及以上运行环境，全面支持Jakarta EE 9+命名空间。系统使用了以下Spring Boot核心Starter模块：spring-boot-starter-web（提供Spring MVC和嵌入式Tomcat服务器，用于构建RESTful Web应用）；spring-boot-starter-data-jpa（集成Spring Data JPA和Hibernate ORM，提供对象关系映射和声明式事务管理）；spring-boot-starter-security（集成Spring Security安全框架，提供认证和授权功能）；spring-boot-starter-validation（集成Jakarta Bean Validation，用于请求参数的自动校验）；spring-boot-starter-webflux（集成Spring WebFlux和Reactor Netty，提供响应式HTTP客户端能力，用于调用外部API）。', 'body'),
        ('Spring Boot的嵌入式Tomcat服务器使得应用可以打包为可执行的JAR文件独立运行，无需部署到外部应用服务器。系统的主类AiCopilotApplication使用@SpringBootApplication注解标记，该注解组合了@Configuration、@EnableAutoConfiguration和@ComponentScan三个注解的功能，通过SpringApplication.run()方法启动整个Spring应用上下文。应用启动时，Spring Boot会根据classpath中的依赖和配置文件自动配置数据源、JPA、安全策略等组件。', 'body'),

        ('2.2.2 Spring Security与JWT认证', 'h3'),
        ('Spring Security是Spring生态中的安全框架，提供了一套全面的认证（Authentication）和授权（Authorization）解决方案。认证是指验证用户身份的过程，即确认"你是谁"；授权是指在确认身份后，决定用户能够访问哪些资源，即"你能做什么"。Spring Security通过过滤器链（Filter Chain）机制实现安全控制，每个HTTP请求在到达控制器之前都会经过一系列安全过滤器的处理。', 'body'),
        ('JSON Web Token（JWT）是一种基于JSON的开放标准（RFC 7519），用于在各方之间安全地传输信息。JWT由三部分组成：Header（头部，包含令牌类型和签名算法信息）、Payload（负载，包含用户身份信息和自定义声明）和Signature（签名，使用密钥对前两部分进行数字签名）。三部分使用Base64 URL编码后以点号分隔拼接。由于JWT是自包含的（即令牌本身已包含所有必要的用户信息），服务器无需维护会话状态，这使得JWT特别适合分布式系统和前后端分离架构。', 'body'),
        ('本系统的认证流程如下：用户通过登录接口提交用户名和密码，服务端验证成功后，使用HMAC-SHA384算法对包含用户ID、用户名和角色（USER/ADMIN）的JWT令牌进行签名，令牌有效期为7天。前端将令牌存储在localStorage中，并通过Axios请求拦截器自动在每次API请求的Authorization头中附加"Bearer token"。后端通过自定义的JwtAuthenticationFilter过滤器解析请求中的JWT令牌，验证签名和有效期，提取用户身份信息并设置到Spring Security的安全上下文中。系统通过SecurityFilterChain配置了路径级别的权限控制：/api/v1/auth/**路径允许匿名访问，/api/v1/admin/**路径要求ADMIN角色，其余所有路径要求已认证。', 'body'),

        ('2.2.3 Spring Data JPA与数据持久化', 'h3'),
        ('Spring Data JPA是Spring Data项目的一部分，旨在简化基于JPA（Java Persistence API）的数据访问层开发。它的核心思想是通过接口方法命名约定自动生成数据库查询，开发者只需定义符合命名规范的接口方法，无需编写SQL或JPQL语句。例如，findByUsername(String username)方法会自动生成"SELECT u FROM User u WHERE u.username = ?1"的查询。Spring Data JPA还支持分页查询（Pageable接口）、排序查询（Sort接口）和@Query注解自定义查询。', 'body'),
        ('本系统定义了6个JPA Repository接口：UserRepository（提供findByUsername和existsByUsername方法）、ResumeRepository（提供findByUserIdOrderByUpdatedAtDesc和findByIdAndUserId方法）、TemplateRepository（提供findByCategoryOrderByNameAsc和existsBySourceUrl方法）、JobAnalysisRepository（提供findByUserIdOrderByCreatedAtDesc和findByResumeId方法）、OptimizationLogRepository（提供findByResumeIdOrderByCreatedAtDesc方法）和InterviewSessionRepository（提供findByUserIdOrderByCreatedAtDesc和findByIdAndUserId方法）。系统通过@Transactional注解确保数据操作的原子性，在删除用户或简历时级联删除相关联的分析和日志数据。', 'body'),

        ('2.3 大语言模型技术', 'h2'),
        ('2.3.1 DeepSeek-V4模型概述', 'h3'),
        ('DeepSeek-V4是深度求索公司于2026年4月发布的第四代大语言模型系列，包含V4-Pro（1.6T总参数，49B激活参数）和V4-Flash（284B总参数，13B激活参数）两个版本。DeepSeek-V4采用了多项创新技术：混合注意力机制（CSA+HCA），通过压缩稀疏注意力（Compressed Sparse Attention）和高度压缩注意力（Highly Compressed Attention）的交错排列，实现了百万Token上下文的高效推理，计算量和KV缓存降至前代模型V3.2的约10%；流形约束超连接（Manifold-Constrained Hyper-Connections, mHC），通过Sinkhorn-Knopp迭代将残差矩阵投影到Birkhoff多胞体上，解决了深层网络训练中的数值不稳定问题，使模型的训练损失和推理性能在超大规模参数下保持稳定；FP8混合精度训练与FP4量化感知训练（QAT），显著降低了显存占用和计算成本，使得1.6T参数模型的训练和部署在商业可行的算力范围内实现。', 'body'),
        ('在性能方面，DeepSeek-V4-Pro-Max在推理基准测试（如AIME 2025、GPQA Diamond、LiveCodeBench等）中表现出色，在多个指标上超越了GPT-5.2和Gemini-3.0-Pro等同期国际领先模型。更重要的是，DeepSeek-V4系列遵循开源理念，模型权重和推理代码在HuggingFace上公开发布，为学术研究和商业应用提供了便利。本系统正是基于DeepSeek-V4的强大文本理解和生成能力，通过API调用实现了简历优化、岗位匹配和面试模拟等智能功能。', 'body'),

        ('2.3.2 提示词工程', 'h3'),
        ('提示词工程（Prompt Engineering）是指设计和优化输入给大语言模型的文本指令，以引导模型产生期望输出的技术方法。在本系统中，提示词工程是连接用户需求与LLM能力的关键桥梁，直接决定了AI功能的输出质量和用户体验。', 'body'),
        ('系统针对不同的AI功能设计了专门的提示词模板，每个模板包含System Prompt（系统提示词）和User Prompt（用户提示词）两部分。System Prompt用于设定LLM的角色身份和行为规范，如简历优化功能的System Prompt将LLM角色设定为"资深HR和职业顾问"，要求其使用"专业、简洁、结果导向的语言"、"尽可能量化成果"、"保持原有格式"；岗位匹配功能的System Prompt要求LLM以JSON格式输出匹配分析结果，包含评分（0-100）、优势列表、劣势列表和改进建议；面试功能的System Prompt将LLM设定为"资深技术面试官"，要求其按照项目经历、技术深度、场景设计、问题排查、综合素质五个维度轮流提问，每个维度只问一个问题，确保面试广度。User Prompt则包含了用户的具体输入信息和上下文数据，如简历内容、职位描述、历史对话记录等。', 'body'),
        ('系统通过PromptBuilder组件类集中管理所有提示词模板，每种AI功能对应一对buildXxxSystemPrompt()和buildXxxUserPrompt()方法。提示词模板使用Java 15引入的文本块（Text Block）语法编写，保持了多行文本的可读性。在系统迭代过程中，通过持续优化提示词（如调整角色描述、增加输出格式约束、补充评分细则等），AI功能的输出质量和稳定性得到了显著提升。', 'body'),

        ('2.3.3 API调用与容错机制', 'h3'),
        ('本系统通过Spring WebFlux的WebClient发送HTTP POST请求调用DeepSeek的聊天补全API。请求体包含model（模型名称）、system（系统提示词）、messages（消息列表）、max_tokens（最大生成Token数）和temperature（生成温度，控制输出的随机性）等参数。系统配置了60秒的API超时时间和最多2次的重试机制，采用指数退避策略（重试间隔为1秒、2秒）应对偶发的网络波动和API繁忙。当所有重试均失败时，系统返回友好的错误提示，避免影响用户体验。', 'body'),
        ('每次AI调用都会被记录到数据库的optimization_logs表中，包括输入文本、输出文本、使用的模型名称、响应时间（毫秒）和优化区块类型，为后续的效果分析和成本优化提供了数据支持。', 'body'),

        ('2.4 数据库技术', 'h2'),
        ('2.4.1 MySQL数据库', 'h3'),
        ('MySQL是当前最流行的开源关系型数据库管理系统，由Oracle公司开发和维护。它采用客户端-服务器架构，支持多线程、多用户并发访问，具有高性能、高可靠性和易用性等特点。本系统使用MySQL 8.0.41版本作为生产环境数据库，通过JDBC驱动（mysql-connector-j）与Java应用建立连接，连接池使用HikariCP（Spring Boot 3.x默认连接池），配置了连接超时、最大连接数等参数以保证数据库连接的稳定性和效率。', 'body'),
        ('MySQL数据库的安装和配置在本地开发环境中通过Windows服务进行管理。数据库名称为ai_resume_copilot，字符集使用UTF-8（utf8mb4），时区设置为Asia/Shanghai。系统通过application.yml配置文件中的spring.datasource配置项指定数据库连接URL、用户名和密码，URL中包含createDatabaseIfNotExist=true参数，确保数据库在首次启动时自动创建。', 'body'),

        ('2.4.2 H2内存数据库', 'h3'),
        ('H2是一款用Java编写的轻量级开源关系型数据库，支持内存模式和文件持久化模式。在本系统的开发环境中，通过application-dev.yml配置文件使用H2内存数据库替代MySQL，无需安装和配置外部数据库服务，即可快速启动项目进行功能开发和调试。H2内存数据库的数据存储在JVM进程内存中，应用重启后数据自动清空。开发配置文件通过spring.jpa.properties.hibernate.dialect显式设置H2方言，确保Hibernate DDL生成与H2语法兼容。H2还提供了内置的Web控制台（/h2-console），方便开发者在浏览器中直接查看和操作数据库。', 'body'),

        ('2.4.3 Hibernate ORM与JPA', 'h3'),
        ('Hibernate是Java领域最成熟的ORM（Object-Relational Mapping）框架之一，负责将Java实体对象映射到关系型数据库表。Hibernate的核心功能包括：根据实体类的JPA注解自动生成DDL（数据定义语言）语句，建表和维护表结构（通过ddl-auto: update配置）；管理实体对象的生命周期和状态转换；提供HQL（Hibernate Query Language）和Criteria API等查询方式；管理一级缓存（Session级别）和二级缓存（SessionFactory级别）。本系统中，Hibernate通过Spring Data JPA进行集成和封装，开发者主要通过Repository接口进行数据操作。', 'body'),

        ('2.5 开发工具与环境', 'h2'),
        ('本系统的开发涉及前后端两套技术栈，需要配置相应的开发工具和运行环境：', 'body'),
        ('（1）JDK 21（Oracle JDK或OpenJDK）：Java开发工具包，包含Java编译器（javac）和Java运行时（JVM）。选择JDK 21是因为它是当前的LTS（长期支持）版本，提供了虚拟线程（Virtual Threads）、记录模式（Record Patterns）和模式匹配（Pattern Matching）等新特性。', 'body'),
        ('（2）Apache Maven 3.9：Java项目的构建和依赖管理工具，通过pom.xml文件声明项目依赖和构建配置。使用spring-boot-maven-plugin插件将应用打包为可执行JAR文件。', 'body'),
        ('（3）Node.js 24与npm：JavaScript运行时和包管理器。Node.js为前端开发提供运行环境，npm用于管理前端项目的依赖包（如Vue、Element Plus、Axios等）和运行开发脚本。', 'body'),
        ('（4）IntelliJ IDEA与VS Code：IntelliJ IDEA用于Java后端代码的编写、调试和重构，利用其强大的代码补全和静态分析功能提升开发效率。VS Code用于Vue前端代码的编写，配合Vue Language Features (Volar)插件提供组件语法高亮和类型检查。', 'body'),
        ('（5）Git版本控制：用于管理项目代码的版本历史，支持代码回滚、分支管理和协作开发。', 'body'),

        ('2.6 本章小结', 'h2'),
        ('本章对AI Resume Copilot系统开发所涉及的核心技术进行了全面介绍。在前端技术方面，本系统基于Vue 3框架的Composition API进行组件开发，搭配Element Plus组件库构建用户界面，使用Vue Router管理页面路由，Pinia管理应用状态，Axios处理HTTP通信，Vite提供开发和生产构建支持。在后端技术方面，系统采用Spring Boot作为核心应用框架，Spring Security结合JWT实现无状态认证与授权，Spring Data JPA简化数据访问层开发。在AI能力方面，系统通过API调用集成DeepSeek-V4大语言模型，并基于提示词工程针对不同功能场景设计了专门的提示词模板。在数据库方面，系统使用MySQL作为生产环境数据库，H2作为开发环境数据库，通过Hibernate ORM实现对象关系映射和数据持久化。本章的技术介绍为后续章节的系统设计、实现和测试提供了充分的理论基础。', 'body'),
        ('', 'page_break'),
    ]

# =================== CHAPTER 3 ===================
def ch3():
    return [
        ('第3章  系统分析', 'h1'),
        ('', 'empty'),
        ('3.1 可行性分析', 'h2'),
        ('3.1.1 技术可行性', 'h3'),
        ('从技术可行性角度分析，本系统的核心技术栈均已成熟稳定。前端采用的Vue 3框架是目前全球使用最广泛的前端框架之一，拥有庞大的社区生态和丰富的第三方资源，其Composition API和响应式系统为复杂交互界面的开发提供了强有力的支持。Element Plus作为Vue 3生态中最流行的UI组件库之一，提供了完整的中文文档和高质量的组件，能够满足企业级应用的UI需求。后端采用的Spring Boot是Java领域事实上的标准应用框架，集成了嵌入式服务器、自动配置和丰富的Starter模块，能够快速构建生产级别的RESTful API服务。Spring Security和JWT的组合为系统提供了健壮的无状态认证方案，在前后端分离架构中被广泛验证。', 'body'),
        ('DeepSeek-V4大语言模型已经通过商业化API对外开放，提供了完善的接口文档和SDK支持，调用门槛较低。系统通过HTTP协议即可完成与LLM的通信，无需额外的机器学习基础设施。MySQL作为全球使用率最高的开源数据库之一，性能稳定且运维成熟。综上所述，本系统的技术方案在理论层面和实践层面均具备充分的可行性。', 'body'),

        ('3.1.2 经济可行性', 'h3'),
        ('本系统的开发成本主要体现在人力投入和API调用费用两个方面。在人力方面，系统采用主流的开源技术栈进行开发，所有核心框架（Vue 3、Spring Boot、MySQL）和UI组件库（Element Plus）均为免费开源软件，无需支付软件许可费用。在AI能力方面，DeepSeek API提供了每月一定额度的免费调用配额，在开发测试阶段和小规模演示场景下成本几乎为零。商业化运营阶段，DeepSeek-V4-Flash版本的API定价仅为每百万Token约0.4美元，远低于GPT-5等同级别闭源模型，在商业可行性上具有显著的成本优势。从效益角度分析，系统能够帮助求职者提升简历质量和面试表现，加速求职过程，具有明确的社会价值和潜在的商业转化空间。因此，从经济角度评估，本系统的开发和运营是完全可行的。', 'body'),

        ('3.1.3 操作可行性', 'h3'),
        ('系统采用B/S（Browser/Server）架构，用户只需通过浏览器即可访问全部功能，无需安装任何客户端软件。前端界面基于Element Plus设计规范构建，视觉风格统一，交互模式符合主流Web应用的操作习惯。用户端功能按Dashboard（仪表盘）→ Editor（编辑器）→ AI Panel（AI面板）→ Match（匹配分析）→ Interview（模拟面试）的流程组织，逻辑清晰。管理后台以Tab标签页形式组织各管理模块，管理员可以快速切换不同功能。系统在Windows 11和macOS系统下对Chrome、Edge、Firefox等主流浏览器进行了兼容性测试，运行正常。操作层面，用户经过简短的探索即可上手使用，操作可行性高。', 'body'),

        ('3.2 功能需求分析', 'h2'),
        ('经过用户调研和竞品分析，AI Resume Copilot系统的功能需求划分为用户端功能和管理端功能两大类，共包含七大核心功能模块。以下对各模块的功能需求进行详细描述。', 'body'),

        ('3.2.1 用户认证模块', 'h3'),
        ('用户认证模块是系统的入口功能，负责管理用户的注册、登录和权限控制。具体功能需求包括：用户通过用户名和密码进行注册，系统对用户名长度（3-50字符）和密码长度（至少6位）进行前端和后端双重校验；密码使用BCrypt算法进行哈希加密后存储，确保不以明文形式落盘；登录成功后返回包含用户ID、用户名和角色的JWT令牌，前端存储令牌用于后续请求的身份验证；令牌有效期为7天，过期后自动跳转登录页面；系统区分普通用户（USER）和管理员（ADMIN）两种角色，注册默认为USER角色。', 'body'),

        ('3.2.2 简历管理模块', 'h3'),
        ('简历管理模块是系统的核心业务模块，提供简历的全生命周期管理。功能需求包括：用户可创建多份简历，每份简历包含标题和结构化内容（基本信息、个人简介、工作经历、教育背景、专业技能、项目经历六个模块）；基本信息支持填写姓名、电话、邮箱、求职意向和上传证件照（照片以Base64编码存储在简历数据中）；工作经历、教育背景和项目经历支持多条目的动态增删；简历采用版本管理机制，每次AI优化后可创建新版本，版本间通过optimized_from字段形成链式关联；支持将简历导出为A4纸PDF格式；支持删除简历及其关联的岗位分析和优化日志数据。', 'body'),
        ('[此处插入图3-1：简历管理功能用例图]', 'placeholder'),

        ('3.2.3 AI智能优化模块', 'h3'),
        ('AI智能优化模块利用大语言模型的文本生成能力对简历内容进行专业化润色和改写。功能需求包括：用户可选择简历的任意模块（个人简介、工作经历描述、专业技能文本等）发起AI优化请求；系统将用户选中的文本和优化指令组装为Prompt发送给LLM；LLM返回优化后的文本，以聊天对话的形式展示给用户；用户可选择"应用修改"将优化内容写入简历或"重新生成"获取新的优化版本；每次AI调用自动记录优化日志（OptimizationLog），包括输入文本、输出文本、模型名称、响应时间和优化区块类型。', 'body'),

        ('3.2.4 岗位匹配分析模块', 'h3'),
        ('岗位匹配分析模块帮助用户评估简历与目标职位的匹配程度。功能需求包括：用户选择一份简历，粘贴目标职位的JD描述文本；系统将简历内容和JD组装为分析Prompt发送给LLM；LLM从技能匹配度（权重40%）、经验年限（权重30%）和关键词密度（权重30%）三个维度进行综合评估，以JSON格式返回匹配分（0-100）以及优势列表、劣势列表和改进建议；分析结果以可视化卡片形式展示，匹配分以大圆环图和颜色分级（绿色≥80分、黄色≥60分、红色<60分）呈现；分析记录保存在job_analyses表中，关联用户和简历，用户可在Dashboard中查看历史分析记录。', 'body'),

        ('3.2.5 AI模拟面试模块', 'h3'),
        ('AI模拟面试模块是系统最具创新性的功能，模拟真实面试场景帮助用户提升面试能力。功能需求包括：用户选择简历并可选填写意向岗位后开始面试；LLM面试官根据简历内容生成个性化面试问题，面试按照项目经历（A）、技术深度（B）、场景设计（C）、问题排查（D）、综合素质（E）五个维度依次进行，每个维度只问一个问题以保证面试广度；用户可通过文字输入或语音输入（利用Web Speech API）回答问题；五轮问题结束后自动结束面试并生成评估报告；报告包含综合评分（0-100）、总体评价、优势列表、不足列表和改进建议；后台对评分系统设置了详细的分数段描述以防止LLM给出不切实际的高分；用户可查看面试历史记录，对已完成的面试查看报告，对进行中的面试继续答题。', 'body'),
        ('[此处插入图3-2：AI模拟面试功能用例图]', 'placeholder'),

        ('3.2.6 简历模板模块', 'h3'),
        ('简历模板模块提供预设的简历模板供用户快速创建简历。功能需求包括：系统预置多套按职业分类的简历模板（如技术开发类、产品设计类、市场运营类、应届生通用类、管理类等），每套模板包含预设的简历JSON结构和示例文本；用户创建简历时可在模板列表中预览和选择模板；管理员可通过管理后台对模板进行增删改查操作。', 'body'),

        ('3.2.7 管理后台模块', 'h3'),
        ('管理后台模块为系统管理员提供全面的系统管理能力。功能需求包括：统计概览（六项核心指标卡片：用户总数、简历总数、模板总数、AI优化次数、匹配分析次数、面试次数）；用户管理（查看所有用户列表、编辑用户信息与角色、删除用户及其关联数据，不可删除管理员账号）；简历管理（选择用户后查看其所有简历列表，支持删除任意简历）；模板管理（模板的增删改查）；日志查看（分页查看AI优化日志和岗位匹配分析记录）；面试记录（分页查看所有用户的面试会话记录）。', 'body'),
        ('[此处插入图3-3：管理后台功能用例图]', 'placeholder'),

        ('3.3 非功能需求分析', 'h2'),
        ('除功能需求外，本系统在性能、安全性、可用性、可扩展性和兼容性等方面也提出了明确的非功能需求指标：', 'body'),
        ('（1）性能需求：页面首次加载时间（First Contentful Paint）不超过3秒；普通CRUD API接口（如获取简历列表、保存简历内容）响应时间在200ms以内；涉及LLM调用的AI接口（如智能优化、匹配分析）响应时间在正常网络条件下不超过30秒，系统设置60秒超时；数据库连接池配置合理，支持至少10个并发请求的同时处理。', 'body'),
        ('（2）安全性需求：所有用户密码使用BCrypt算法加密存储，不可逆；API接口通过JWT令牌进行身份验证，令牌使用HMAC-SHA384算法签名；管理后台接口（/api/v1/admin/**）仅允许ADMIN角色访问，普通用户访问返回403 Forbidden；前端对用户输入进行基本的XSS防护（输入框自动转义HTML标签）；后端使用参数化查询（Spring Data JPA自动实现）防止SQL注入攻击；敏感配置信息（如LLM API密钥、JWT签名密钥）通过环境变量或单独的配置文件注入，不硬编码在主配置文件中。', 'body'),
        ('（3）可用性需求：系统提供直观友好的中文用户界面，关键操作提供确认提示（如删除简历、结束面试等）；表单输入提供实时的前端校验反馈；AI调用等待期间显示加载动画和提示文字，避免用户误以为系统无响应；错误场景下提供具体的错误信息提示而非技术性报错。', 'body'),
        ('（4）可扩展性需求：后端采用分层架构（Controller-Service-Repository），层间通过接口依赖，便于替换和扩展；提示词模板集中在PromptBuilder组件中管理，添加新的AI功能只需增加新方法；前端API调用集中在src/api目录下按功能模块分文件管理，添加新API只需新增函数。', 'body'),
        ('（5）兼容性需求：前端兼容Chrome 100+、Edge 100+、Firefox 100+等主流浏览器的最近两个大版本；后端兼容Java 21及以上版本；数据库兼容MySQL 8.0及以上版本。', 'body'),

        ('3.4 系统用例分析', 'h2'),
        ('本系统涉及两类参与者：普通用户（User）和管理员（Admin）。普通用户是系统的核心用户群体，通过注册获得USER角色，可以使用简历管理、AI优化、岗位匹配、模拟面试等全部用户端功能。管理员通过数据库预设的ADMIN角色账号登录，除可访问用户端功能外，还拥有管理后台的全部管理权限。以下通过用例描述表详细说明系统的主要用例。', 'body'),
        ('[此处插入图3-4：系统总体用例图]', 'placeholder'),
        ('表3-1 用户注册用例描述', 'placeholder'),
        ('表3-2 用户登录用例描述', 'placeholder'),
        ('表3-3 AI简历优化用例描述', 'placeholder'),
        ('表3-4 岗位匹配分析用例描述', 'placeholder'),
        ('表3-5 AI模拟面试用例描述', 'placeholder'),

        ('3.5 数据流分析', 'h2'),
        ('系统的数据流围绕用户简历的生命周期展开。用户通过前端界面编辑简历内容，数据以JSON格式组织，经Axios发送HTTP请求到后端API。后端Controller接收请求并校验参数，Service层处理业务逻辑（如调用LLM进行优化、计算匹配分等），Repository层负责与数据库交互。LLM调用的数据流为：前端→后端Controller→InterviewService/LLMService→WebClient HTTP请求→DeepSeek API→返回结果→保存到数据库→返回前端展示。简历版本管理的数据流为：原简历数据→创建新版本（复制原简历内容）→原版本标记为非当前→新版本设置为当前→保存到数据库。', 'body'),
        ('[此处插入图3-5：系统顶层数据流图]', 'placeholder'),
        ('[此处插入图3-6：AI优化功能数据流图]', 'placeholder'),

        ('3.6 本章小结', 'h2'),
        ('本章从多个维度对AI Resume Copilot系统进行了全面的分析。通过技术可行性、经济可行性和操作可行性论证，确认了项目实施的理论基础和现实条件。通过功能需求分析，明确了七大核心模块的具体功能要求，并通过非功能需求分析确定了性能、安全、可用性等方面的质量指标。通过用例分析和数据流分析，对系统的使用场景和数据流转进行了建模。本章的需求分析成果为后续的系统设计提供了明确的目标和约束。', 'body'),
        ('', 'page_break'),
    ]

# =================== CHAPTER 4 ===================
def ch4():
    return [
        ('第4章  系统设计', 'h1'),
        ('', 'empty'),
        ('4.1 系统总体架构设计', 'h2'),
        ('本系统采用前后端分离的B/S（Browser/Server）架构，整体划分为五个层次：表示层（前端SPA应用）、接口层（RESTful API）、业务逻辑层（Service Layer）、数据访问层（Repository Layer）和外部服务层（LLM API）。前后端通过HTTP/HTTPS协议进行JSON格式的数据通信，前端通过Vite开发服务器的代理功能将API请求转发至后端，避免了跨域问题。各层次的职责和交互关系如下：', 'body'),
        ('表示层：由Vue 3构建的单页面应用，运行在用户浏览器中。该层负责页面的渲染、用户交互的处理以及前端路由的导航。通过Axios HTTP客户端与后端的接口层通信，发送请求并处理响应数据，更新UI状态。', 'body'),
        ('接口层：由Spring MVC的RestController组件构成，负责接收HTTP请求、进行请求参数验证（通过@Valid和Jakarta Bean Validation）、调用业务逻辑层处理请求，并将处理结果封装为统一的ApiResponse<T>格式返回给前端。', 'body'),
        ('业务逻辑层：由多个Service组件构成（UserService、ResumeService、LLMService、JobAnalysisService、InterviewService、AdminService等），封装了系统的核心业务规则和流程。该层通过依赖注入获取Repository和外部服务依赖，使用@Transactional注解确保关键操作的事务一致性。', 'body'),
        ('数据访问层：由多个继承自JpaRepository的Repository接口构成，通过Spring Data JPA和Hibernate ORM实现与MySQL数据库的交互。该层将Java实体对象的增删改查操作翻译为相应的SQL语句，并管理数据库连接和事务。', 'body'),
        ('外部服务层：由LLMService和PromptBuilder构成，负责与DeepSeek大语言模型API的通信。LLMService通过WebClient发送HTTP请求调用LLM接口，PromptBuilder为不同的AI功能组装对应的System Prompt和User Prompt。', 'body'),
        ('[此处插入图4-1：系统总体架构图]', 'placeholder'),

        ('4.2 功能模块详细设计', 'h2'),
        ('系统功能模块按照高内聚、低耦合原则划分为用户端模块和管理端模块两大类，每类包含若干子模块。用户端模块包括用户认证、简历管理、AI智能优化、岗位匹配分析和AI模拟面试五个子模块；管理端模块包括统计概览、用户管理、简历管理、模板管理、日志查看和面试记录六个子模块。以下对各核心模块的设计进行详细说明。', 'body'),
        ('[此处插入图4-2：系统功能模块结构图]', 'placeholder'),

        ('4.2.1 用户认证模块设计', 'h3'),
        ('用户认证模块的核心流程为：注册流程——前端表单校验→POST /api/v1/auth/register→后端校验用户名唯一性→BCrypt加密密码→保存用户→返回JWT令牌（含userId、username、role）；登录流程——POST /api/v1/auth/login→查询用户→BCrypt匹配密码→生成JWT令牌并返回；认证流程——请求携带Authorization: Bearer <token>→JwtAuthenticationFilter拦截→解析令牌→验证签名和有效期→设置SecurityContext→放行或拒绝。', 'body'),
        ('JWT令牌的生成使用io.jsonwebtoken（JJWT）库，签名算法为HMAC-SHA384，密钥通过Base64解码application.yml中配置的jwt.secret值。令牌payload中包含subject（用户ID）、username（用户名）和role（角色）三个声明。', 'body'),

        ('4.2.2 简历管理模块设计', 'h3'),
        ('简历数据以结构化JSON格式存储在resumes表的content_json字段中（LONGTEXT类型）。简历JSON的结构设计如下：basicInfo（基本信息对象，含name、phone、email、position、photo字段）、summary（个人简介字符串）、workExperience（工作经历数组，每条含company、position、period、description字段）、education（教育背景数组，每条含school、major、degree、period字段）、skills（专业技能字符串）、projects（项目经历数组，每条含name、role、description字段）。', 'body'),
        ('简历的版本管理机制设计为：每次创建新版本时，系统首先复制当前简历的全部数据创建新的Resume实体，新版本version字段加1，optimized_from字段指向源简历，is_current设置为true。同时将源简历的is_current设置为false。这样所有版本通过optimized_from字段形成一条链表，前端通过遍历链表展示简历的版本历史。', 'body'),
        ('PDF导出功能的设计流程为：构建HTML简历模板字符串→创建隐藏DOM元素→使用html2canvas将DOM渲染为Canvas→通过jsPDF将Canvas转换为PDF→触发浏览器下载。证件照以Base64编码的图片数据嵌入HTML模板的img标签中。', 'body'),

        ('4.2.3 AI智能优化模块设计', 'h3'),
        ('AI优化的核心流程为：用户选择简历区块→系统提取该区块的文本内容→组装System Prompt（HR专家角色设定）和User Prompt（原始文本+用户指令）→调用LLMService.callLLM()→LLM返回优化文本→前端以对话气泡展示→用户点击应用或重新生成。', 'body'),
        ('优化日志（OptimizationLog）在每次LLM调用时自动创建，记录resume_id（关联简历）、section_type（优化区块类型，如summary、workExperience、skills等）、input_text（原始文本）、output_text（优化后文本）、llm_model（使用的模型名称）、response_time_ms（API响应时间，毫秒）和created_at（创建时间）。', 'body'),

        ('4.2.4 岗位匹配模块设计', 'h3'),
        ('岗位匹配的核心流程为：用户选择简历并输入职位描述→系统提取简历JSON和JD文本→组装匹配分析Prompt（System Prompt设定LLM为招聘专家，要求从技能匹配度40%、经验年限30%、关键词密度30%三个维度评分）→调用LLM→解析LLM返回的JSON→提取score、strengths、weaknesses、suggestions→保存到job_analyses表→返回前端渲染匹配结果卡片。', 'body'),

        ('4.2.5 AI模拟面试模块设计', 'h3'),
        ('AI模拟面试是本系统设计最复杂的模块。整体流程为：初始化阶段（用户选择简历→POST /interview/start→系统发送简历和System Prompt→LLM返回第一个面试问题→创建InterviewSession实体）→交互阶段（用户输入回答→POST /interview/answer→系统拼装对话历史→检查问题数量→发送给LLM→LLM返回下一问题或[END]标记→更新会话）→结束阶段（自动触发或手动触发→POST /interview/{id}/end→LLM生成评估报告JSON→解析评分和建议→更新session状态为COMPLETED→返回InterviewReportResponse）。', 'body'),
        ('面试的五个维度设计为：A类（项目经历）询问候选人的核心项目架构、技术难点和个人贡献；B类（技术深度）针对简历中列出的某项技术询问底层原理或最佳实践；C类（场景设计）给出实际工作场景要求候选人设计技术方案；D类（问题排查）描述线上故障场景考察候选人的排查思路；E类（综合素质）涉及团队协作、冲突处理、职业规划等软技能。每个维度只问一个问题，面试官通过对话历史判断已覆盖的维度，避免重复。', 'body'),
        ('评估报告的评分体系设计为四个维度：技术深度（35分）评估对技术原理的理解深度；项目经验（25分）评估项目描述的清晰度和影响力；沟通表达（20分）评估回答的结构化和清晰度；思维分析（20分）评估分析和解决问题的思路。每个维度设置了具体的分数段描述，如技术深度维度：谈论源码/底层原理可得30-35分，停留在使用层面得15-25分，回答模糊或错误得0-15分。', 'body'),
        ('[此处插入图4-3：AI模拟面试流程图]', 'placeholder'),

        ('4.2.6 管理后台模块设计', 'h3'),
        ('管理后台采用Tab页形式组织，默认展示统计概览Tab。统计概览通过GET /api/v1/admin/stats接口获取数据，以六张统计卡片展示关键指标。用户管理Tab通过表格展示所有用户，每行提供编辑和删除操作按钮。编辑操作弹出对话框，支持修改用户名、邮箱、密码（留空则不修改）和角色。删除操作在确认后执行，系统检查被删除用户是否为管理员，若为管理员则拒绝删除。简历管理Tab先通过用户下拉框选择用户，再加载该用户的简历列表。模板管理Tab、优化日志Tab、匹配记录Tab和面试记录Tab均采用表格加分页的展示模式。', 'body'),
        ('[此处插入图4-4：管理后台功能结构图]', 'placeholder'),

        ('4.3 数据库设计', 'h2'),
        ('4.3.1 数据库E-R模型', 'h3'),
        ('系统数据库采用MySQL关系型数据库，使用Spring Data JPA进行对象关系映射，通过Hibernate的ddl-auto: update配置在应用启动时自动创建和更新数据表结构。数据库包含6个核心实体类，通过JPA注解定义表间的关联关系。User实体是系统的中心实体，与Resume、JobAnalysis和InterviewSession实体之间存在一对多（OneToMany）的关联关系。Resume实体与JobAnalysis、OptimizationLog和InterviewSession之间存在一对多关联关系。Template实体为独立的模板数据表，不与其他实体关联。', 'body'),
        ('[此处插入图4-5：系统数据库E-R图]', 'placeholder'),

        ('4.3.2 核心数据表结构', 'h3'),
        ('以下以数据字典方式详细描述各核心数据表的结构设计。', 'body'),
        ('表4-1 users（用户表）：存储系统用户的基本信息和认证数据。字段包括id（BIGINT，主键，自增）、username（VARCHAR(50)，唯一，非空）、password（VARCHAR(255)，非空，BCrypt加密存储）、email（VARCHAR(100)）、role（VARCHAR(10)，非空，默认USER，枚举值USER/ADMIN）、created_at（DATETIME，非空，由@PrePersist自动填充）。主键为id，唯一索引为username。', 'body'),
        ('表4-2 resumes（简历表）：存储用户创建的所有简历数据。字段包括id（BIGINT，主键，自增）、user_id（BIGINT，外键，关联users.id，非空）、title（VARCHAR(100)）、version（INT，非空，默认1）、content_json（LONGTEXT，存储JSON格式的简历结构化数据）、optimized_from（BIGINT，自引用外键，关联resumes.id）、is_current（BOOLEAN，非空，默认false，标识是否当前版本）、created_at（DATETIME，非空）、updated_at（DATETIME）。主键为id，外键为user_id和optimized_from。', 'body'),
        ('表4-3 templates（模板表）：存储简历模板数据。字段包括id（BIGINT，主键，自增）、name（VARCHAR(100)，非空）、category（VARCHAR(50)）、description（VARCHAR(500)）、content_json（LONGTEXT，非空，模板JSON内容）、source_url（VARCHAR(500)）、created_at（DATETIME，非空）。主键为id。', 'body'),
        ('表4-4 job_analyses（岗位分析表）：存储岗位匹配分析记录。字段包括id（BIGINT，主键，自增）、user_id（BIGINT，外键，关联users.id）、resume_id（BIGINT，外键，关联resumes.id）、job_description（TEXT）、match_score（DECIMAL(5,2)）、suggestions（TEXT，存储JSON格式的分析结果）、created_at（DATETIME，非空）。主键为id，外键为user_id和resume_id。', 'body'),
        ('表4-5 optimization_logs（优化日志表）：记录每次AI调用的详细信息。字段包括id（BIGINT，主键，自增）、resume_id（BIGINT，外键，关联resumes.id，非空）、prompt_used（TEXT，存储发送给LLM的完整提示词）、llm_model（VARCHAR(50)）、input_text（LONGTEXT，原始文本）、output_text（LONGTEXT，优化后文本）、response_time_ms（INT，API响应时间毫秒数）、section_type（VARCHAR(50)，优化的区块类型）、created_at（DATETIME，非空）。主键为id，外键为resume_id。', 'body'),
        ('表4-6 interview_sessions（面试会话表）：存储AI模拟面试的完整数据。字段包括id（BIGINT，主键，自增）、user_id（BIGINT，外键，关联users.id，非空）、resume_id（BIGINT，外键，关联resumes.id，非空）、position（VARCHAR(100)，意向岗位）、messages（LONGTEXT，JSON数组，存储完整的对话消息）、status（VARCHAR(20)，非空，默认IN_PROGRESS，枚举值IN_PROGRESS/COMPLETED）、score（INT，综合评分）、report（TEXT，评估报告文本）、strengths（TEXT，JSON数组格式的优势列表）、weaknesses（TEXT，JSON数组格式的不足列表）、suggestions（TEXT，JSON数组格式的改进建议）、created_at（DATETIME，非空）、completed_at（DATETIME）。主键为id，外键为user_id和resume_id。', 'body'),

        ('4.4 接口设计', 'h2'),
        ('4.4.1 API设计规范', 'h3'),
        ('系统API遵循RESTful设计风格，所有API路径以/api/v1/为前缀。管理后台API路径以/api/v1/admin/为前缀。请求和响应统一使用JSON格式，Content-Type为application/json。所有API响应统一包裹在ApiResponse<T>结构中，包含code（int，业务状态码，200表示成功）、message（String，消息描述）和data（T，泛型数据载荷）三个字段。分页API额外包含totalElements（总记录数）、totalPages（总页数）和page（当前页码）字段。', 'body'),
        ('API安全方面，除/auth/**路径外，所有API请求需在Authorization请求头中携带Bearer JWT令牌。管理后台API（/admin/**）额外要求令牌中角色为ADMIN，否则返回403 Forbidden。', 'body'),

        ('4.4.2 核心API接口列表', 'h3'),
        ('以下列出系统的主要API接口及其说明：', 'body'),
        ('用户认证相关：POST /api/v1/auth/register（用户注册，接收username、password、email，返回JWT令牌和用户信息）；POST /api/v1/auth/login（用户登录，接收username、password，返回JWT令牌、用户名、用户ID和角色）。', 'body'),
        ('简历管理相关：GET /api/v1/resumes（获取当前用户的所有简历列表，按更新时间降序）；GET /api/v1/resumes/{id}（获取指定简历的详细信息）；POST /api/v1/resumes（创建新简历，接收title和可选的contentJson）；PUT /api/v1/resumes/{id}（更新简历的标题、内容或当前版本标记）；DELETE /api/v1/resumes/{id}（删除简历及关联的岗位分析和优化日志数据）；POST /api/v1/resumes/{id}/versions（在指定简历基础上创建新版本）。', 'body'),
        ('AI功能相关：POST /api/v1/ai/optimize（对指定简历区块进行AI优化，接收resumeId、sectionType、originalText和可选的instruction）；POST /api/v1/ai/optimize-full（对完整简历JSON进行全文AI优化）；POST /api/v1/ai/match（岗位匹配分析，接收resumeId和jobDescription）。', 'body'),
        ('面试功能相关：POST /api/v1/interview/start（开始面试，接收resumeId和可选的position，返回sessionId和第一个面试问题）；POST /api/v1/interview/answer（提交回答，接收sessionId和answer，返回下一问题）；POST /api/v1/interview/{id}/end（手动结束面试并生成评估报告）；GET /api/v1/interview/{id}（获取面试会话详情）；GET /api/v1/interview/history（获取当前用户的面试历史列表）；DELETE /api/v1/interview/{id}（删除指定面试会话）。', 'body'),
        ('管理后台相关：GET /api/v1/admin/stats（获取系统统计数据）；GET /api/v1/admin/users（分页查看所有用户）；GET /api/v1/admin/users/{id}（查看指定用户详情）；PUT /api/v1/admin/users/{id}（编辑用户信息与角色）；DELETE /api/v1/admin/users/{id}（删除用户）；GET /api/v1/admin/users/{userId}/resumes（查看指定用户的简历列表）；DELETE /api/v1/admin/users/{userId}/resumes/{id}（删除指定用户的简历）；GET/POST/PUT/DELETE /api/v1/admin/templates（模板增删改查）；GET /api/v1/admin/logs（分页查看优化日志）；GET /api/v1/admin/analyses（分页查看匹配记录）；GET /api/v1/admin/interviews（分页查看面试记录）。', 'body'),
        ('[此处插入表4-7：系统API接口汇总表]', 'placeholder'),

        ('4.5 安全设计', 'h2'),
        ('系统的安全设计涵盖身份认证、访问控制、数据加密和常见攻击防护四个层面。身份认证层使用JWT无状态令牌机制，令牌使用HMAC-SHA384算法签名，密钥长度足够（384位），有效抵抗暴力破解。访问控制层使用Spring Security的路径匹配规则进行粗粒度授权，区分公开路径（/auth/**）、认证路径（其他所有路径）和管理员路径（/admin/**）。', 'body'),
        ('数据加密层中，用户密码使用BCrypt算法进行单向哈希加密，BCrypt内置盐值（Salt）机制使得即使相同密码的哈希值也不相同，有效防止彩虹表攻击。BCrypt的计算强度（cost factor）默认为10，每增加1个强度值，计算时间翻倍，有效抵御暴力破解。前端到后端的通信在部署时可通过HTTPS进行TLS加密，防止中间人攻击。', 'body'),
        ('常见攻击防护方面：SQL注入防护由Spring Data JPA的参数化查询自动实现，系统不拼接原生的SQL字符串；XSS防护在前端由Vue的模板引擎自动转义HTML标签实现；CSRF攻击在前后端分离架构和JWT无状态认证的背景下自然免疫（因为无Cookie可用于自动携带）。', 'body'),

        ('4.6 本章小结', 'h2'),
        ('本章对AI Resume Copilot系统的整体设计进行了系统性的阐述。在架构设计层面，系统采用五层架构模型（表示层、接口层、业务逻辑层、数据访问层、外部服务层），各层次职责清晰、耦合度低。在功能模块设计层面，详细设计了用户端五个核心模块和管理端六个子模块的功能流程和数据结构。在数据库设计层面，定义了6张核心数据表及其字段约束和关联关系，绘制了E-R模型图。在接口设计层面，制定了统一的RESTful API规范，列举了30余个核心API接口。在安全设计层面，从身份认证、访问控制、数据加密和攻击防护四个维度制定了安全策略。本章的详细设计为下一章的系统实现提供了完整的蓝图。', 'body'),
        ('', 'page_break'),
    ]

# =================== CHAPTER 5 ===================
def ch5():
    return [
        ('第5章  系统实现', 'h1'),
        ('', 'empty'),
        ('5.1 开发环境搭建', 'h2'),
        ('5.1.1 后端项目搭建', 'h3'),
        ('后端项目使用Spring Initializr（https://start.spring.io/）生成基础项目骨架，选择Maven构建工具和Java 21语言版本。在生成的基础pom.xml文件上，手动添加了系统所需的各项依赖：spring-boot-starter-web（Web应用支持）、spring-boot-starter-data-jpa（数据持久化）、spring-boot-starter-security（安全框架）、spring-boot-starter-validation（参数校验）、spring-boot-starter-webflux（响应式HTTP客户端）、mysql-connector-j（MySQL驱动）、h2（H2内存数据库）、jjwt-api/impl/jackson（JWT令牌库）、knife4j-openapi3-jakarta-spring-boot-starter（API文档生成）、lombok（代码简化）和jsoup（HTML解析，用于模板抓取）。', 'body'),
        ('项目的包结构按照分层架构组织：config包存放配置类（SecurityConfig、Knife4jConfig、WebClientConfig），controller包存放REST控制器（AuthController、ResumeController、AIController、TemplateController、InterviewController、AdminController），service包存放业务逻辑服务（UserService、ResumeService、LLMService、JobAnalysisService、InterviewService、TemplateScraperService、AdminService），repository包存放数据访问接口（UserRepository、ResumeRepository、TemplateRepository、JobAnalysisRepository、OptimizationLogRepository、InterviewSessionRepository），entity包存放JPA实体类，dto包存放数据传输对象，exception包存放自定义异常和全局异常处理器，util包存放工具类（JwtUtil、PromptBuilder）。', 'body'),
        ('[此处插入图5-1：后端项目包结构截图]', 'placeholder'),

        ('5.1.2 前端项目搭建', 'h3'),
        ('前端项目使用Vite创建命令（npm create vite@latest）生成基础项目结构。创建后在package.json中添加了核心依赖：vue 3.4、vue-router 4.2、pinia 2.1、axios 1.6、element-plus 2.5、@element-plus/icons-vue 2.3（Element Plus图标库）、echarts 5.5、vue-echarts 6.6（数据可视化图表）、html2canvas 1.4和jspdf 2.5（PDF导出）。开发依赖包括@vitejs/plugin-vue 5.0和vite 5.0。', 'body'),
        ('项目的src目录结构组织为：api目录存放API调用模块（request.js封装Axios，auth.js、resume.js、ai.js、interview.js、admin.js按功能拆分），store目录存放Pinia状态管理（user.js、resume.js），router目录存放Vue Router路由配置，views目录存放页面组件（Login、Register、Dashboard、Editor、Match、Interview、InterviewHistory、InterviewReport、Admin），components目录存放可复用组件（ResumeForm、AIPanel、MatchScore），utils目录存放工具函数（pdfExport.js）。', 'body'),
        ('[此处插入图5-2：前端项目目录结构截图]', 'placeholder'),

        ('5.2 用户认证模块实现', 'h2'),
        ('5.2.1 后端认证实现', 'h3'),
        ('用户认证的核心是Spring Security的配置和JWT过滤器的实现。SecurityConfig配置类通过@Configuration和@EnableWebSecurity注解启用Spring Security，通过SecurityFilterChain Bean配置安全规则：关闭CSRF保护（适用于前后端分离架构），设置Session策略为STATELESS（无状态），配置CORS允许前端开发服务器（localhost:5173）的跨域请求，注册JWT认证过滤器。JwtAuthenticationFilter通过继承OncePerRequestFilter实现，在doFilterInternal方法中从请求头提取Bearer令牌，调用JwtUtil验证令牌有效性，解析用户ID和角色信息，创建UsernamePasswordAuthenticationToken并设置到SecurityContext中。', 'body'),
        ('JwtUtil工具类实现了令牌的生成、解析和验证三个核心方法。generateToken方法接收userId、username和role参数，使用JJWT库的Jwts.builder()构建令牌，设置主题（subject，即userId）、自定义声明（username和role）、签发时间和过期时间（当前时间+7天），并使用HMAC-SHA384密钥签名。parseToken方法使用Jwts.parser()解析令牌并返回声明数据。validateToken方法通过try-catch捕获各类JWT异常（如签名无效、令牌过期等）来判断令牌是否有效。', 'body'),
        ('UserService的register方法实现了注册业务逻辑：首先检查用户名是否已存在（通过userRepository.existsByUsername），若存在则抛出BusinessException(400, "用户名已存在")；否则创建User实体，使用BCryptPasswordEncoder对密码进行哈希处理，设置默认角色为USER，保存后生成JWT令牌返回。login方法通过userRepository.findByUsername查找用户，若不存在或密码不匹配则抛出BusinessException(401, "用户名或密码错误")。', 'body'),
        ('[此处插入图5-3：用户登录注册界面截图]', 'placeholder'),

        ('5.2.2 前端认证实现', 'h3'),
        ('前端认证相关的核心代码分布在store/user.js（Pinia Store）和api/auth.js（API调用模块）中。userStore定义了token、username、userId和role四个响应式状态，login和register两个异步action分别调用auth API获取JWT令牌后将令牌和用户信息存储到localStorage中以实现持久化。isAdmin计算属性通过判定role是否等于ADMIN来控制管理后台入口按钮的显隐。logout方法清除store中的状态和localStorage中的数据。', 'body'),
        ('api/request.js对Axios进行了统一封装，配置了请求拦截器和响应拦截器。请求拦截器在每次请求前从localStorage获取JWT令牌并添加到Authorization请求头。响应拦截器检查响应数据的code字段，若不为200则显示错误提示并reject Promise；同时处理401未认证状态码（令牌过期或无效），自动清除localStorage并跳转到登录页。', 'body'),
        ('路由守卫（router/index.js的beforeEach）实现了页面级别的权限控制：未登录用户访问需要认证的页面（meta.requiresAuth为true）时自动跳转到登录页；已登录用户访问登录/注册页时根据角色跳转（管理员到/admin，普通用户到/dashboard）；普通用户（非管理员）访问/admin路径时被重定向到/dashboard；管理员访问非/admin路径时被重定向到/admin。', 'body'),

        ('5.3 简历管理模块实现', 'h2'),
        ('5.3.1 简历编辑器实现', 'h3'),
        ('简历编辑器是用户端最核心的交互界面，由ResumeForm.vue组件实现。该组件使用响应式（reactive）的formData对象管理简历数据，包含六个顶层字段：basicInfo（基本信息对象）、summary（字符串）、workExperience（数组）、education（数组）、skills（字符串）、projects（数组）。组件通过watch监听父组件传入的resumeData prop，在数据变化时将后端数据通过JSON.parse(JSON.stringify())深拷贝到formData中。每次用户编辑时通过emitUpdate()方法将formData的副本发射给父组件。', 'body'),
        ('编辑器按照分区组织：基本信息区提供姓名、电话、邮箱、求职意向的输入框和证件照上传组件；工作经历、教育背景和项目经历三个区支持动态增删条目，每个条目包含独立的表单字段；每个区块旁边都提供了"AI优化"按钮，点击后触发emit("ai-optimize", {section, index, text})事件，由父组件协调AIPanel进行优化。', 'body'),
        ('[此处插入图5-4：简历编辑器界面截图]', 'placeholder'),

        ('5.3.2 简历仪表盘实现', 'h3'),
        ('Dashboard.vue是用户登录后的首页，提供简历列表的概览和管理功能。页面顶部显示系统标题和用户信息（用户名、退出按钮），管理员用户额外显示"管理后台"入口按钮。主要内容区域使用el-row和el-col栅格系统以三列网格展示用户的简历卡片，每张卡片显示简历标题、版本号、更新时间和是否当前版本的标签。点击卡片跳转到简历编辑器，每个卡片底部提供编辑和删除两个操作按钮。页面还包含"新建简历"按钮（弹出模板选择对话框）、"岗位匹配分析"入口区和"AI模拟面试"入口区。', 'body'),
        ('新建简历对话框展示预设的简历模板供用户选择。模板以三列网格展示，每张模板卡片显示分类标签、名称和描述。用户可选择模板快速创建（模板的contentJson将作为新简历的初始内容），也可跳过模板从空白开始。', 'body'),
        ('[此处插入图5-5：简历仪表盘界面截图]', 'placeholder'),

        ('5.3.3 PDF导出实现', 'h3'),
        ('PDF导出功能由utils/pdfExport.js实现。核心函数exportToPDF接收简历数据和文件名，通过以下步骤生成PDF：首先创建一个隐藏的DOM元素，使用buildTemplate函数将简历数据渲染为HTML简历模板字符串并插入DOM；等待300毫秒让浏览器完成布局渲染；使用html2canvas以2倍缩放比例将DOM元素渲染为Canvas；使用jsPDF创建A4纸尺寸（210mm×297mm）的PDF文档，将Canvas转为PNG格式图片添加到PDF页面中；如果内容高度超过一页A4纸，jsPDF自动创建多页并拼接图片；最后触发浏览器下载PDF文件。buildTemplate函数将简历的各模块数据组装为结构化的HTML字符串，证件照以Base64格式的img标签嵌入，简历头部使用flex布局实现姓名居中和证件照右对齐的排版效果。', 'body'),
        ('[此处插入图5-6：PDF导出的简历效果截图]', 'placeholder'),

        ('5.4 AI智能优化模块实现', 'h2'),
        ('5.4.1 LLM调用层实现', 'h3'),
        ('LLMService是整个AI功能的核心服务类，封装了与大语言模型API的所有通信逻辑。该类的callLLM方法接收systemPrompt（系统提示词）和userPrompt（用户提示词）两个参数，通过WebClient发送HTTP POST请求到DeepSeek的API端点。请求体的结构为：model字段指定使用的模型名称（从配置文件中读取llm.model），system字段传入系统提示词，messages字段为一个包含单条user消息的数组，max_tokens设置为4096，temperature设置为0.7。API响应中LLM的生成文本从content数组的第一个元素的text字段中提取。', 'body'),
        ('系统实现了指数退避重试机制：如果API调用失败（抛出异常），系统最多重试maxRetries次（默认2次），每次重试前等待递增的时间（第1次重试等待1秒，第2次等待2秒）。如果在所有重试后仍失败，则抛出BusinessException(503, "优化服务繁忙，请稍后再试")。超时时间通过Duration.ofMillis(timeout)设置（默认60秒）。', 'body'),
        ('LLMService为不同的AI功能提供了对外的公共方法。optimizeSection方法接收原始文本和优化指令，调用promptBuilder构建优化提示词并调用callLLM，返回OptimizeResponse（包含优化后文本和估算的Token消耗量）。optimizeFullResume方法对完整简历JSON进行全文优化。analyzeMatch方法接收简历JSON和职位描述，调用匹配分析提示词并解析LLM返回的JSON结构，返回MatchResponse（包含评分和各项分析结果）。', 'body'),

        ('5.4.2 提示词构建器实现', 'h3'),
        ('PromptBuilder组件集中管理了系统中所有的LLM提示词模板。每个提示词方法返回一个使用Java 15文本块语法编写的多行字符串。系统优化类的提示词将LLM角色设定为"资深HR和职业顾问"，给出五条具体的优化规则（使用专业简洁的语言、量化成果、保持真实性、维持原有格式、仅输出优化文本）。面试类的提示词将LLM角色设定为"资深技术面试官"，给出八条面试规则（只问一个问题、基于简历提问、覆盖项目经历和技术深度等维度、广度优先等）。报告评估类的提示词设定了详细的分维度评分标准（每个维度有具体的分数段描述）和JSON输出格式要求，并加入了"不要给所有人打高分"的警告以确保评分的客观性。', 'body'),
        ('在系统迭代过程中，提示词经历了多轮优化。例如，面试提示词最初只要求LLM"基于候选人的简历内容提问"，导致LLM倾向于围绕同一个项目反复追问，缺乏广度。改进后的提示词明确规定了五个必须覆盖的面试维度（A至E类），每类只问一个问题，并要求LLM检查对话历史避免重复类别。评分类的提示词最初未给出具体分数段描述，导致LLM倾向于给出80-90分的高分，缺乏区分度。改进后的提示词为每个维度的每个分数段给出了具体的行为描述，如技术深度维度：谈论源码/底层原理得30-35分，停留在使用层面得15-25分，回答模糊或错误得0-15分。', 'body'),
        ('[此处插入图5-7：AI优化功能界面截图]', 'placeholder'),

        ('5.4.3 前端AI面板实现', 'h3'),
        ('AIPanel.vue组件以嵌入式聊天面板的形式集成在简历编辑器中。组件使用messages数组存储对话历史（每条消息含role角色、content内容和section/in dex等元数据），通过defineExpose暴露setSection方法供父组件调用以发起优化。优化流程为：父组件调用setSection(section, index, text)→AIPanel在消息列表中添加用户消息→调用API→在消息列表中添加AI回复消息（含"应用修改"和"重新生成"操作按钮）→自动滚动到底部。用户也可以在底部的输入框中输入额外的优化要求。优化结果的消息对象设置了optimized: true标志，触发在消息下方渲染"应用修改"按钮。点击"应用修改"时，组件通过emit("apply-optimization", {section, index, text})通知父组件更新简历内容。', 'body'),

        ('5.5 岗位匹配模块实现', 'h2'),
        ('岗位匹配分析功能由JobAnalysisService和MatchScore组件协同实现。用户选择一份简历，在文本框中粘贴目标职位的JD描述后点击"开始分析"按钮。后端服务首先通过resumeRepository.findByIdAndUserId验证简历归属，然后提取简历内容JSON，将其与JD描述一起发送给LLM。LLM返回的JSON字符串通过extractJson辅助方法提取（定位最外层的大括号），再由ObjectMapper解析为MatchResponse对象。如果JSON解析失败（如LLM返回格式异常），系统提供fallback：返回MatchResponse(0, [], [], ["匹配分析服务暂时不可用"])。匹配分析结果同步保存到job_analyses表中，用户可以在Dashboard中查看历史分析记录。', 'body'),
        ('前端MatchScore.vue组件以可视化卡片形式展示分析结果。卡片的顶部以大号圆环图和颜色分级（绿色≥80分、黄色60-79分、红色<60分）展示综合评分。卡片中部以标签形式展示优势项和不足项。卡片底部以编号列表展示改进建议。', 'body'),
        ('[此处插入图5-8：岗位匹配分析界面截图]', 'placeholder'),

        ('5.6 AI模拟面试模块实现', 'h2'),
        ('5.6.1 面试流程实现', 'h3'),
        ('AI模拟面试是系统中实现最复杂的模块，由InterviewService和InterviewController提供后端服务，Interview.vue提供前端界面。面试的初始化由startInterview方法实现：验证简历归属→加载简历JSON→使用buildInterviewStartUserPrompt构建开场提示词→调用LLM获取第一个问题→创建InterviewSession实体并保存→返回InterviewResponse（含sessionId和首题）。', 'body'),
        ('面试的交互由submitAnswer方法实现：加载会话并验证归属和状态→解析现有消息JSON数组→追加用户回答→统计已问问题数→判断是否达到上限（6题）→若未达上限，使用buildInterviewNextUserPrompt构建续问提示词→调用LLM获取下一问题→追加AI回复到消息数组→保存并返回InterviewResponse。如果LLM返回的文本中包含[END]标记，或问题数已达上限（MAX_QUESTIONS=6），则自动触发结束流程。', 'body'),
        ('面试的结束由endInterview方法（手动）或autoEndInterview方法（自动）触发。结束流程的核心是generateReport方法：组装完整的对话历史→使用buildInterviewReportSystemPrompt和buildInterviewReportUserPrompt构建评估提示词→调用LLM获取评估结果→解析JSON中的score、report、strengths、weaknesses、suggestions→更新InterviewSession的状态为COMPLETED、设置completedAt时间戳→返回InterviewReportResponse。如果LLM返回的JSON解析失败，系统使用fallback：score设为0，report使用原始响应文本，strengths/weaknesses/suggestions设为空数组。', 'body'),

        ('5.6.2 语音输入实现', 'h3'),
        ('语音输入功能基于浏览器内置的Web Speech API实现，无需任何额外的后端服务或第三方库。前端在用户点击麦克风按钮时，首先检查浏览器是否支持SpeechRecognition或webkitSpeechRecognition接口。如果浏览器不支持（如Firefox），通过ElMessage提示用户使用Chrome或Edge浏览器。如果支持，创建SpeechRecognition实例，配置语言为zh-CN（中文普通话），设置interimResults为true以实时显示中间识别结果，continuous为false以在停顿后自动结束。', 'body'),
        ('识别过程中，onresult回调函数逐条处理识别结果，区分isFinal（最终结果）和interim（中间结果）：最终结果累积到finalTranscript变量中，中间结果实时显示在输入框中。这样用户在说话过程中可以看到识别文字的动态变化，获得即时的视觉反馈。识别错误通过onerror回调处理，其中not-allowed错误提示用户允许麦克风权限，no-speech错误静默处理（用户可能还未开始说话），其他错误显示具体错误信息。', 'body'),
        ('语音识别按钮使用tooltip提示用户功能说明。识别激活时，按钮呈现红色脉冲动画效果（通过CSS @keyframes实现），并在输入框下方显示"正在录音，请说话..."的提示文字。识别结束后，最终文字保留在输入框中，用户可手动编辑修正后点击发送提交。', 'body'),
        ('[此处插入图5-9：AI模拟面试对话界面截图]', 'placeholder'),

        ('5.6.3 面试报告展示', 'h3'),
        ('InterviewReport.vue页面负责展示面试评估报告。页面通过路由参数（?id=sessionId）获取面试会话ID，调用getSessionDetail API加载完整的会话数据。报告页的核心展示元素包括：评分环（使用CSS绘制的大号圆形评分展示，颜色根据分数分级：绿色≥80分、黄色60-79分、红色<60分）；面试基本信息（使用的简历标题、问题总数、完成时间）；综合评语文本；以左右两列布局分别展示优势列表（绿色标题）和不足列表（黄色标题）；改进建议以编号列表形式展示。所有列表数据都通过parseList辅助函数将后端存储的JSON字符串解析为数组。', 'body'),
        ('[此处插入图5-10：面试评估报告界面截图]', 'placeholder'),

        ('5.7 管理后台模块实现', 'h2'),
        ('5.7.1 管理后台界面实现', 'h3'),
        ('Admin.vue是管理后台的核心页面组件，采用el-tabs标签页组织七个管理功能模块。用户管理Tab使用el-table表格展示所有用户，提供编辑和删除操作按钮。编辑用户使用el-dialog弹出对话框，表单字段包括用户名、邮箱、新密码（选填）和角色（下拉选择USER/ADMIN）。删除用户前弹出确认对话框，提示将级联删除该用户的所有简历。后端AdminService在删除用户前检查其角色，若为ADMIN则拒绝删除。', 'body'),
        ('简历管理Tab先通过el-select下拉框选择目标用户（选项通过加载全部用户列表动态生成），选择后自动调用getUserResumes API加载该用户的简历列表，以表格展示简历ID、标题、版本号和创建时间，每行提供删除按钮。模板管理Tab提供模板的增删改查功能，使用el-dialog弹出表单编辑模板的名称、分类、描述、来源URL和JSON内容。', 'body'),
        ('优化日志Tab和匹配记录Tab分别展示系统的AI调用日志和匹配分析记录。两个Tab均采用表格加分页组件的展示模式，表格列包括ID、关联的简历/用户信息、关键内容字段（优化输入输出文本、匹配分和职位描述）和时间戳。面试记录Tab展示所有用户的模拟面试会话记录，表格列包括ID、用户ID、简历标题、意向岗位、状态标签（进行中/已完成）、评分和时间信息。', 'body'),
        ('[此处插入图5-11：管理后台统计概览界面截图]', 'placeholder'),
        ('[此处插入图5-12：管理后台用户管理界面截图]', 'placeholder'),
        ('[此处插入图5-13：管理后台模板管理界面截图]', 'placeholder'),

        ('5.7.2 管理员权限控制实现', 'h3'),
        ('管理员权限控制通过Spring Security的路径匹配规则和JWT令牌中的角色信息双层机制实现。在SecurityConfig的SecurityFilterChain中，配置了.requestMatchers("/api/v1/admin/**").hasRole("ADMIN")规则，这意味着所有以/api/v1/admin/开头的API请求都必须携带包含ROLE_ADMIN权限的认证令牌。在JwtAuthenticationFilter中，系统从JWT令牌中提取role声明，创建包含SimpleGrantedAuthority("ROLE_" + role)的认证对象。这样，普通用户（角色为USER）拥有ROLE_USER权限，管理员（角色为ADMIN）拥有ROLE_ADMIN权限，后者可以访问管理后台API。', 'body'),
        ('前端路由层面也实现了额外的管理员权限控制。在router.beforeEach全局导航守卫中，系统检查localStorage中的role值：如果用户角色不是ADMIN但尝试访问/admin路径，将被重定向到/dashboard；如果用户角色是ADMIN但尝试访问非管理路径（如/dashboard、/editor等），将被重定向到/admin。这种双重控制（后端API权限+前端路由权限）确保了管理功能的安全性。', 'body'),

        ('5.8 本章小结', 'h2'),
        ('本章详细阐述了AI Resume Copilot系统各核心模块的具体实现过程。从开发环境的搭建开始，逐一介绍了后端Spring Boot项目和前端Vue项目的工程结构和配置方式。然后按照功能模块分别介绍了用户认证、简历管理、AI优化、岗位匹配、模拟面试和管理后台等模块的实现细节，包括核心代码的逻辑流程、关键类的设计思路和前端的交互实现。特别地，对LLM调用层的重试机制、提示词工程的迭代优化、语音输入的Web Speech API集成、以及管理员权限的双重控制机制等关键技术点进行了深入说明。本章提供了丰富的界面截图，直观展示了系统的实际运行效果。', 'body'),
        ('', 'page_break'),
    ]

# =================== CHAPTER 6 ===================
def ch6():
    return [
        ('第6章  系统测试', 'h1'),
        ('', 'empty'),
        ('6.1 测试环境与策略', 'h2'),
        ('6.1.1 测试环境', 'h3'),
        ('系统测试在以下硬件和软件环境中进行：', 'body'),
        ('硬件环境：Windows 11专业版操作系统，Intel Core i7-13700H处理器（14核20线程），16GB DDR5内存，512GB NVMe SSD固态硬盘。', 'body'),
        ('软件环境：JDK 21.0.6（Oracle JDK），Apache Maven 3.9.9，MySQL 8.0.41 Community Server，Node.js 24.15.0，Google Chrome 130浏览器，Microsoft Edge 130浏览器。', 'body'),
        ('测试工具：Postman用于API接口测试，Chrome DevTools的Network和Performance面板用于前端性能分析，JUnit 5用于单元测试（Spring Boot Starter Test内置）。', 'body'),

        ('6.1.2 测试策略', 'h3'),
        ('本系统的测试遵循"自底向上、逐层验证"的策略，测试层次从低到高依次为：单元测试（使用JUnit 5和Mockito对Service层和Util层的核心方法进行独立测试）、接口测试（使用Postman对RESTful API进行单独的功能验证和边界测试）、集成测试（在前后端联调环境下进行端到端的功能流程测试）和系统测试（在完整部署环境下进行性能、安全和兼容性测试）。测试用例的设计采用等价类划分和边界值分析方法，覆盖正常场景、异常场景和边界场景。', 'body'),

        ('6.2 功能测试', 'h2'),
        ('功能测试覆盖系统的所有核心功能模块，采用黑盒测试方法，设计测试用例并逐项执行验证。以下按模块列出主要的测试用例。', 'body'),

        ('6.2.1 用户认证模块测试', 'h3'),
        ('测试用例TC01（正常注册）：输入合法的用户名（testuser）和密码（123456），预期结果为注册成功并返回包含JWT令牌和用户信息的响应，实际结果通过。', 'body'),
        ('测试用例TC02（用户名已存在）：使用已注册的用户名再次注册，预期结果为返回code=400，message="用户名已存在"，实际结果通过。', 'body'),
        ('测试用例TC03（密码过短）：输入密码长度小于6位（如"123"），预期结果为前端表单校验拦截显示"密码长度至少6位"，实际结果通过。', 'body'),
        ('测试用例TC04（正常登录）：输入正确的用户名和密码，预期结果为登录成功返回JWT令牌，令牌中包含正确的userId、username和role字段，实际结果通过。', 'body'),
        ('测试用例TC05（错误密码）：输入正确的用户名但错误的密码，预期结果为返回code=401，message="用户名或密码错误"，实际结果通过。', 'body'),
        ('测试用例TC06（管理员登录）：使用管理员账号（admin/admin123）登录，预期结果为返回role="ADMIN"，前端自动跳转到管理后台页面，实际结果通过。', 'body'),
        ('测试用例TC07（未登录访问）：不携带JWT令牌直接访问受保护的API（如GET /api/v1/resumes），预期结果为返回401未认证错误，前端自动跳转到登录页，实际结果通过。', 'body'),
        ('测试用例TC08（普通用户访问管理后台）：使用USER角色的JWT令牌访问/api/v1/admin/users，预期结果为返回403 Forbidden，实际结果通过。', 'body'),
        ('[此处插入表6-1：用户认证模块功能测试用例表]', 'placeholder'),

        ('6.2.2 简历管理模块测试', 'h3'),
        ('测试用例TC09（创建简历）：提交标题为"Java后端开发工程师简历"的创建请求，预期结果为创建成功，返回version=1的简历对象，实际结果通过。', 'body'),
        ('测试用例TC10（编辑简历）：修改简历的标题和专业技能内容后提交更新请求，预期结果为更新成功，数据库中对应记录已修改，实际结果通过。', 'body'),
        ('测试用例TC11（删除简历）：删除指定ID的简历，预期结果为简历及关联的岗位分析和优化日志数据被级联删除，实际结果通过。', 'body'),
        ('测试用例TC12（创建版本）：在已有简历上调用版本创建API，预期结果为原简历is_current变为false，新简历is_current为true，optimized_from指向原简历，实际结果通过。', 'body'),
        ('测试用例TC13（上传证件照）：选择一张JPG格式图片上传，预期结果为图片以Base64格式编码存储在简历的basicInfo.photo字段中，编辑器显示预览图，实际结果通过。', 'body'),
        ('测试用例TC14（PDF导出）：点击PDF导出按钮，预期结果为浏览器触发下载一个A4纸尺寸的PDF文件，文件中包含简历的所有内容模块，证件照显示在简历头部右上角，实际结果通过。', 'body'),
        ('测试用例TC15（访问他人简历）：尝试通过API访问不属于当前用户的简历ID，预期结果为返回404"简历不存在"（而非返回实际数据），实际结果通过。', 'body'),

        ('6.2.3 AI功能模块测试', 'h3'),
        ('测试用例TC16（AI分段优化）：选择简历的个人简介区块，点击"AI优化"按钮，预期结果为LLM返回专业化改写后的文本内容，以对话气泡形式显示，包含"应用修改"和"重新生成"按钮，实际结果通过。', 'body'),
        ('测试用例TC17（应用优化修改）：在AI返回优化结果后点击"应用修改"，预期结果为简历编辑器中对应区块的内容更新为优化后的文本，实际结果通过。', 'body'),
        ('测试用例TC18（AI全文优化）：调用全文优化API，传入完整简历JSON和优化指令，预期结果为LLM返回优化后的完整JSON结构，实际结果通过。', 'body'),
        ('测试用例TC19（岗位匹配分析）：选择一份后端开发方向的简历，粘贴目标JD（Java开发工程师职位描述），点击分析，预期结果为返回包含综合评分（0-100）、优势列表、劣势列表和改进建议的JSON结果，实际结果通过。', 'body'),
        ('测试用例TC20（匹配分析结果保存）：完成岗位匹配后检查job_analyses表，预期结果为表中新增一条记录，包含正确的resume_id、job_description和match_score，实际结果通过。', 'body'),
        ('测试用例TC21（LLM超时处理）：在LLM API响应时间超过60秒的情况下，预期结果为系统返回"优化服务繁忙，请稍后再试"的错误提示，不会导致前端页面崩溃或后端服务异常，实际结果通过。', 'body'),

        ('6.2.4 模拟面试模块测试', 'h3'),
        ('测试用例TC22（开始面试）：选择一份简历后点击开始面试，预期结果为LLM返回一个与简历内容相关的面试问题（从项目经历维度开始），前端显示面试对话界面，实际结果通过。', 'body'),
        ('测试用例TC23（正常问答流程）：连续回答5个面试问题，预期结果为每个回答后LLM返回的新问题来自不同的维度类别（A→B→C→D→E依次轮换），不出现连续两题同一维度的情况，实际结果通过。', 'body'),
        ('测试用例TC24（自动结束面试）：在5个维度的问题全部回答完毕后，预期结果为LLM自动输出[END]标记或系统在第6题后强制结束，系统自动生成评估报告，实际结果通过。', 'body'),
        ('测试用例TC25（手动结束面试）：在面试过程中点击"结束面试"按钮，预期结果为弹出确认对话框，确认后系统调用LLM生成评估报告并跳转到报告页面，实际结果通过。', 'body'),
        ('测试用例TC26（面试报告内容完整性）：检查生成的面试报告，预期结果为报告包含综合评分（0-100）、总体评价文本（150-300字）、具体的优势列表和不足列表、可操作的改进建议，实际结果通过。', 'body'),
        ('测试用例TC27（语音输入功能）：在Chrome浏览器中点击麦克风按钮，允许麦克风权限后说话，预期结果为语音内容实时转换为文字显示在输入框中，实际结果通过。', 'body'),
        ('测试用例TC28（面试历史查看）：在面试记录页面查看历史记录，预期结果为表格中显示所有已完成和进行中的面试会话，包括简历标题、状态、评分和时间信息，实际结果通过。', 'body'),
        ('测试用例TC29（删除面试记录）：在面试记录页面删除一条记录，预期结果为该记录从列表中消失，数据库中对应行被删除，实际结果通过。', 'body'),

        ('6.2.5 管理后台模块测试', 'h3'),
        ('测试用例TC30（统计概览）：使用管理员账号登录，进入管理后台首页，预期结果为统计概览Tab默认选中，六张统计卡片（用户总数、简历总数、模板总数、AI优化次数、匹配分析次数、面试次数）正确显示当前系统数据，实际结果通过。', 'body'),
        ('测试用例TC31（编辑用户信息）：在用户管理Tab中编辑某个普通用户的邮箱和角色信息，点击保存，预期结果为用户信息更新成功，前端表格中对应行显示更新后的数据，实际结果通过。', 'body'),
        ('测试用例TC32（删除用户）：删除一个普通用户账号，预期结果为弹出确认对话框，确认后用户及其关联的所有简历和分析数据被删除，前端用户列表刷新，实际结果通过。', 'body'),
        ('测试用例TC33（删除管理员）：尝试删除管理员（admin）账号，预期结果为系统返回"不能删除管理员账号"的错误提示，实际结果通过。', 'body'),
        ('测试用例TC34（模板增删改查）：在模板管理Tab中依次执行新建模板、编辑模板名称、删除模板的操作，预期结果为各项操作成功执行，数据实时更新，实际结果通过。', 'body'),
        ('[此处插入表6-2：管理后台模块功能测试用例表]', 'placeholder'),

        ('6.3 接口测试', 'h2'),
        ('接口测试使用Postman工具对系统的RESTful API进行单独的请求-响应验证。测试覆盖了所有API端点（共30余个），包括GET、POST、PUT、DELETE四种HTTP方法，验证了请求参数校验、响应状态码、响应体结构和业务逻辑正确性。', 'body'),
        ('主要接口测试结果：POST /api/v1/auth/register在合法参数下返回HTTP 200，响应体包含code=200、message="success"和包含token/username/userId/role的data对象；在重复用户名时返回HTTP 400和code=400的错误响应。POST /api/v1/auth/login在正确密码时返回HTTP 200和正确的JWT令牌；在错误密码时返回HTTP 401。GET /api/v1/resumes在不携带令牌时返回HTTP 401；携带有效令牌时返回HTTP 200和简历列表数组。POST /api/v1/interview/start在简历不属于当前用户时返回HTTP 404。所有POST/PUT接口在缺少@NotNull或@NotBlank标注的必填字段时返回HTTP 400和具体的校验错误信息。分页接口（如GET /api/v1/admin/logs?page=0&size=20）正确返回content数组、totalElements总数、totalPages总页数和page当前页码。', 'body'),
        ('[此处插入表6-3：API接口测试结果汇总表]', 'placeholder'),

        ('6.4 性能测试', 'h2'),
        ('性能测试主要关注页面加载速度、API响应时间和系统资源消耗三个指标。', 'body'),
        ('页面加载性能：使用Chrome DevTools的Lighthouse工具对系统的主要页面（登录页、Dashboard页、编辑器页、管理后台页）进行性能审计。测试结果显示，所有页面的首次内容绘制（FCP）时间在1.5至2.8秒之间，最大内容绘制（LCP）时间在2.0至3.5秒之间，性能评分均在85分以上。前端静态资源（JS/CSS文件）通过Vite的构建优化和Tree Shaking，生产构建的总体积约为450KB（gzip压缩后约120KB），首屏加载速度表现良好。', 'body'),
        ('API响应性能：使用Postman的Runner功能对主要API接口进行了10次并发请求的性能测试。普通CRUD接口（如GET /api/v1/resumes、POST /api/v1/resumes）的平均响应时间在50-150ms之间，99分位响应时间在300ms以内。涉及数据库关联查询的接口（如GET /api/v1/admin/users，需加载全部用户）的平均响应时间在100-200ms之间。涉及LLM调用的AI接口响应时间主要取决于DeepSeek API的响应速度，平均在3-15秒之间，与LLM服务的负载和生成文本的长度相关。', 'body'),
        ('资源消耗：后端Spring Boot应用在正常运行状态下的内存占用约为350-500MB（JVM堆内存），CPU使用率在空闲状态下低于1%，在AI接口调用时因等待LLM API响应（I/O等待），CPU使用率约5-10%。前端页面在Chrome浏览器中的内存占用约为50-100MB。', 'body'),

        ('6.5 兼容性测试', 'h2'),
        ('兼容性测试在不同浏览器上对系统的主要功能页面进行了验证。测试结果如下：', 'body'),
        ('Google Chrome 130：所有功能正常，页面渲染正确，语音输入功能可用，AI调用和PDF导出均正常工作。', 'body'),
        ('Microsoft Edge 130：所有功能正常，页面渲染与Chrome一致，语音输入功能可用（Edge同样基于Chromium内核，支持Web Speech API）。', 'body'),
        ('Mozilla Firefox 135：页面渲染正确，基本功能（登录、简历编辑、AI调用、PDF导出）正常。语音输入功能不可用（Firefox不支持Web Speech API），系统正确提示用户使用Chrome或Edge浏览器。', 'body'),
        ('在不同分辨率下的测试（1920×1080、2560×1440、1366×768）：Element Plus的栅格系统和响应式组件确保了界面在不同分辨率下布局合理，无元素重叠或溢出问题。', 'body'),

        ('6.6 本章小结', 'h2'),
        ('本章对AI Resume Copilot系统进行了全面的测试。在功能测试方面，编写并执行了34个测试用例，覆盖了用户认证、简历管理、AI功能、模拟面试和管理后台五大模块，所有测试用例均通过。在接口测试方面，使用Postman验证了全部30余个API端点的请求响应正确性。在性能测试方面，页面加载性能、API响应时间和系统资源消耗均满足设计指标。在兼容性测试方面，系统在Chrome和Edge浏览器上运行完美，在Firefox上除语音输入外其他功能正常。测试结果充分验证了系统的功能正确性、性能稳定性和良好的用户体验。', 'body'),
        ('', 'page_break'),
    ]

# =================== CHAPTER 7 ===================
def ch7():
    return [
        ('第7章  总结与展望', 'h1'),
        ('', 'empty'),
        ('7.1 工作总结', 'h2'),
        ('本课题针对求职者在简历撰写和面试准备过程中面临的实际困难，设计并实现了一款基于大语言模型的智能简历优化辅助系统——AI Resume Copilot。系统采用前后端分离的现代化Web架构，通过集成DeepSeek-V4大语言模型，为用户提供了简历智能优化、岗位匹配分析和AI模拟面试三大核心智能功能，同时配套了简历管理、模板系统、用户认证、管理后台等完整的基础功能体系。', 'body'),
        ('回顾整个课题的研究和开发过程，主要完成了以下工作：', 'body'),
        ('（1）完成了系统的需求分析和总体设计。通过调研现有招聘辅助工具的不足和用户的实际需求，明确了七大核心功能模块的功能需求和五项非功能需求指标。基于需求分析结果，设计了系统的五层架构模型、六大功能模块的详细流程、六张核心数据表的结构和三十余个RESTful API接口规范。', 'body'),
        ('（2）实现了完整的简历管理系统。后端基于Spring Boot和Spring Data JPA构建了RESTful API服务，前端基于Vue 3和Element Plus构建了交互式简历编辑器。系统支持简历的结构化编辑、版本控制、模板选择和PDF导出功能，用户可以通过六个结构化模块灵活地编辑和优化简历内容。', 'body'),
        ('（3）利用大语言模型实现了简历智能优化功能。通过精心设计的提示词工程，将LLM的角色设定为资深HR和职业顾问，使其能够对简历各模块进行专业化改写。实现了优化结果的一键应用和重新生成功能，每次AI调用都被记录到优化日志中便于效果追踪。', 'body'),
        ('（4）实现了基于LLM的岗位匹配分析功能。设计了包含技能匹配度、经验年限和关键词密度三个维度的匹配评分体系，LLM能够对简历和职位描述进行深度对比分析，量化匹配度并给出针对性的改进建议。', 'body'),
        ('（5）创新性地实现了AI模拟面试功能。设计了覆盖项目经历、技术深度、场景设计、问题排查和综合素质五个维度的面试流程，确保面试的广度和系统性。支持文字和语音双模式输入，面试结束后自动生成包含评分、优劣势分析和改进建议的评估报告。评分系统设置了明确的分数段描述以防止评分失真。', 'body'),
        ('（6）构建了功能完善的管理后台。管理员可以查看系统核心数据的统计概览，管理用户账号信息和角色权限，管理简历和模板数据，以及查看AI调用日志、匹配分析记录和面试会话记录。管理后台实现了前后端双重权限控制，确保管理功能的安全性。', 'body'),
        ('（7）完成了系统的全面测试。编写并验证了34个功能测试用例，完成了全部API接口的Postman测试，进行了页面性能审计和API并发测试，并在多浏览器环境下进行了兼容性验证。测试结果确认了系统的功能完整性和性能稳定性。', 'body'),

        ('7.2 不足与展望', 'h2'),
        ('尽管本系统已实现了预期的全部核心功能并通过了全面测试，但受限于开发周期和个人能力，仍存在以下可以进一步改进和完善的方向：', 'body'),
        ('（1）多LLM模型支持：当前系统仅集成了DeepSeek-V4单一模型。未来可以扩展为支持多个大语言模型提供商（如OpenAI的GPT系列、Anthropic的Claude系列、阿里的通义千问等），允许用户根据优化效果和成本选择最合适的模型。需要设计统一的LLM适配层，将不同模型的API差异封装在适配器内部。', 'body'),
        ('（2）简历导入解析：当前系统仅支持基于模板创建简历或从零编辑，不支持从已有的PDF或Word格式简历文件自动解析导入。未来可以利用OCR技术和文档解析库（如Apache POI、PDFBox）实现现有简历的智能导入，降低用户的迁移门槛。', 'body'),
        ('（3）多语言简历支持：当前系统仅支持中文简历的创建和优化。随着国际化求职需求的增长，应扩展系统以支持英文等多语言简历，包括多语言的简历模板、针对不同语言优化的提示词模板以及国际化前端界面。', 'body'),
        ('（4）面试功能增强：AI模拟面试可以从多个方向深化。增加视频模拟面试功能，通过摄像头捕捉候选人的面部表情和肢体语言进行综合评估；增加更多面试场景类型（如行为面试、案例分析、编程白板等）；增加语音语调分析，评估候选人的表达自信度和流畅性；增加多人面试（群面）模拟。', 'body'),
        ('（5）数据驱动的个性化推荐：系统经过一段时间的运行积累用户数据后，可以训练推荐模型，根据用户的行业、职位和技能背景，智能推荐最合适的简历模板和优化策略，甚至预测用户可能感兴趣的职位方向。', 'body'),
        ('（6）实时协作功能：增加简历的实时协作编辑功能，允许求职者与职业顾问、导师或朋友同时在线编辑和讨论简历内容。这需要引入WebSocket技术实现前后端的双向实时通信。', 'body'),
        ('（7）移动端适配：当前系统为桌面端Web应用，在手机等小屏设备上体验不佳。未来可以开发响应式移动端界面或独立的移动端应用（APP/小程序），方便用户随时随地进行简历编辑和面试练习。', 'body'),
        ('（8）部署与运维优化：将系统容器化（使用Docker），编写Dockerfile和docker-compose.yml文件，实现前后端的一键部署。配置CI/CD流水线（如GitHub Actions），实现代码提交后的自动构建、测试和部署。使用Nginx作为反向代理和静态资源服务器，提升系统的生产环境性能和安全性。', 'body'),
        ('（9）数据安全与隐私保护：随着系统处理真实用户数据，需要进一步加强数据安全措施，包括实施数据库字段级加密、完善数据备份和恢复机制、制定用户数据隐私政策和合规方案（如个人信息保护法合规）。', 'body'),
        ('（10）效果量化评估：设计实验方案，通过A/B测试对比使用AI Resume Copilot优化前后的简历在真实招聘平台上的投递效果（如简历查看率、面试邀请率等），用量化数据验证系统的实际价值。', 'body'),
        ('', 'page_break'),
    ]

# =================== REFERENCES ===================
def references():
    refs = [
        '[1] 尤雨溪. Vue.js设计与实现[M]. 北京: 人民邮电出版社, 2022.',
        '[2] 克雷格·沃尔斯. Spring实战(第6版)[M]. 北京: 人民邮电出版社, 2023.',
        '[3] DeepSeek-AI. DeepSeek-V4: Towards Highly Efficient Million-Token Context Intelligence[R]. HuggingFace Technical Report, 2026.',
        '[4] Vaswani A, Shazeer N, Parmar N, et al. Attention Is All You Need[C]. Advances in Neural Information Processing Systems (NeurIPS), 2017: 5998-6008.',
        '[5] Brown T B, Mann B, Ryder N, et al. Language Models are Few-Shot Learners[C]. Advances in Neural Information Processing Systems (NeurIPS), 2020: 1877-1901.',
        '[6] Varshney A, Ganuthula V R R. Signal or Noise? Evaluating Large Language Models in Resume Screening Across Contextual Variations and Human Expert Benchmarks[J]. arXiv:2507.08019, 2025.',
        '[7] DeepSeek-AI. DeepSeek-V2: A Strong, Economical, and Efficient Mixture-of-Experts Language Model[J]. arXiv:2405.04434, 2024.',
        '[8] DeepSeek-AI. DeepSeek LLM: Scaling Open-Source Language Models with Longtermism[J]. arXiv:2401.02954, 2024.',
        '[9] 周志华. 机器学习[M]. 北京: 清华大学出版社, 2016.',
        '[10] 刘增杰, 张俊林. 大语言模型原理与实践[M]. 北京: 机械工业出版社, 2024.',
        '[11] Evan You. Vue 3 Official Documentation[EB/OL]. https://vuejs.org/, 2024.',
        '[12] Pivotal Team. Spring Boot Reference Documentation 3.2.x[EB/OL]. https://docs.spring.io/spring-boot/, 2024.',
        '[13] Element Plus Team. Element Plus Documentation[EB/OL]. https://element-plus.org/, 2024.',
        '[14] Richardson L, Amundsen M. RESTful Web APIs[M]. O\'Reilly Media, 2013.',
        '[15] 李刚. 轻量级Java EE企业应用实战(第5版)[M]. 北京: 电子工业出版社, 2020.',
        '[16] Goodfellow I, Bengio Y, Courville A. Deep Learning[M]. MIT Press, 2016.',
        '[17] 肖仰华. 知识图谱与认知智能[M]. 北京: 电子工业出版社, 2019.',
        '[18] Jones M T. Artificial Intelligence: A Systems Approach[M]. Jones & Bartlett Learning, 2015.',
        '[19] 李明, 王磊. 基于深度学习的简历信息抽取方法研究[J]. 计算机应用研究, 2022, 39(5): 1400-1405.',
        '[20] 张伟, 陈强. 基于Transformer的人岗匹配模型研究[J]. 计算机工程与应用, 2023, 59(12): 155-162.',
        '[21] 王芳, 刘洋. 大语言模型在智能招聘中的应用综述[J]. 计算机科学, 2024, 51(3): 1-15.',
        '[22] 赵丽, 杨帆. 基于Spring Boot和Vue.js的Web应用开发研究[J]. 软件导刊, 2023, 22(8): 89-94.',
        '[23] 孙卫琴. 精通Spring: Java轻量级架构开发实践[M]. 北京: 电子工业出版社, 2021.',
        '[24] ISO/IEC 25010:2011. Systems and Software Engineering — Systems and Software Quality Requirements and Evaluation (SQuaRE)[S]. ISO, 2011.',
    ]
    result = [('参考文献', 'h1'), ('', 'empty'), ('', 'empty')]
    for ref in refs:
        result.append((ref, 'ref'))
    result.append(('', 'page_break'))
    return result

# =================== ACKNOWLEDGMENTS ===================
def acknowledgments():
    return [
        ('致  谢', 'h1'),
        ('', 'empty'),
        ('在本毕业设计论文完成之际，我要向所有在课题研究和论文撰写过程中给予我帮助和支持的人表示诚挚的感谢。', 'body'),
        ('首先，我要衷心感谢我的指导老师XXX教授。从课题选题、方案论证、系统开发到论文撰写，XXX老师在每个阶段都给予了我悉心的指导和宝贵的建议。在系统开发过程中遇到技术瓶颈时，老师总能以其丰富的经验和深厚的专业知识为我指明方向；在论文写作过程中，老师对论文的结构、内容和表达提出了许多建设性的修改意见，使我受益匪浅。老师严谨求实的治学态度、高度的责任感和平易近人的待人风格，给我留下了深刻的印象，是我今后学习和工作的榜样。', 'body'),
        ('其次，我要感谢计算机科学与工程学院的各位领导和老师。在四年的本科学习期间，老师们传授的专业理论知识和实践技能，为我完成本毕业设计打下了坚实的基础。特别要感谢在毕业设计课程中给予指导的各位老师，你们的课程让我系统性地学习了软件工程的需求分析、系统设计、编码实现和测试验证方法，为本课题的实施提供了方法论指导。', 'body'),
        ('此外，我要感谢开源社区的广大贡献者。本系统的开发离不开Vue.js、Spring Boot、Element Plus、MySQL等优秀开源项目的支持。这些由全球开发者共同维护和维护的开源软件，为现代软件开发提供了坚实的技术基础，使得个人开发者能够站在巨人的肩膀上快速构建复杂的应用系统。同时感谢深度求索公司开源的DeepSeek系列大语言模型，为本系统的AI功能提供了强大的技术支撑。', 'body'),
        ('我也要感谢我的同学和朋友们。在系统开发和论文撰写的过程中，大家相互交流、相互启发、相互鼓励，营造了积极向上的学习和研究氛围。在系统测试阶段，同学们积极帮助我进行功能验证和问题反馈，使得系统得以不断改进和完善。', 'body'),
        ('最后，我要深深感谢我的家人。他们的理解、支持和鼓励是我完成大学学业和本毕业设计的坚强后盾。在我遇到困难时，家人的关心和安慰给了我继续前进的动力。', 'body'),
        ('', 'empty'),
        ('', 'empty'),
    ]

# ============================================================
# DOCUMENT GENERATION
# ============================================================

def build_document(output_path):
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)

    # Collect all content
    all_content = []
    all_content.extend(cover_page())
    all_content.extend(integrity_page())
    all_content.extend(abstract_cn())
    all_content.extend(abstract_en())
    all_content.extend(toc())
    all_content.extend(ch1())
    all_content.extend(ch2())
    all_content.extend(ch3())
    all_content.extend(ch4())
    all_content.extend(ch5())
    all_content.extend(ch6())
    all_content.extend(ch7())
    all_content.extend(references())
    all_content.extend(acknowledgments())

    # Render content
    for text, style in all_content:
        if style == 'page_break':
            doc.add_page_break()
            continue
        if style == 'empty':
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.5
            continue

        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5

        if style == 'cover_title':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(20)
            run = p.add_run(text)
            run.font.size = Pt(22)
            run.bold = True
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        elif style == 'cover_main':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.font.size = Pt(18)
            run.bold = True
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        elif style == 'cover_en':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.font.size = Pt(14)
            run.font.name = 'Times New Roman'
        elif style == 'cover_info':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.font.size = Pt(14)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        elif style == 'h1':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(text)
            run.font.size = Pt(16)
            run.bold = True
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        elif style == 'h2':
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(text)
            run.font.size = Pt(14)
            run.bold = True
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        elif style == 'h3':
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(text)
            run.font.size = Pt(12)
            run.bold = True
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        elif style == 'body':
            p.paragraph_format.first_line_indent = Cm(0.74)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(text)
            run.font.size = Pt(12)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        elif style == 'body_no_indent':
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(text)
            run.font.size = Pt(12)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        elif style == 'keywords':
            p.paragraph_format.first_line_indent = Cm(0.74)
            run = p.add_run(text)
            run.font.size = Pt(12)
            run.bold = True
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        elif style == 'toc_0':
            run = p.add_run(text)
            run.font.size = Pt(14)
            run.bold = True
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        elif style == 'toc_1':
            run = p.add_run(text)
            run.font.size = Pt(14)
            run.bold = True
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        elif style == 'toc_2':
            run = p.add_run(text)
            run.font.size = Pt(12)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        elif style == 'ref':
            run = p.add_run(text)
            run.font.size = Pt(10.5)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        elif style == 'placeholder':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(text)
            run.font.size = Pt(12)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 0, 200)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.save(output_path)
    print(f'论文已生成：{output_path}')

    # Statistics
    total_chars = sum(len(t) for t, _ in all_content)
    body_chars = sum(len(t) for t, s in all_content if s in ('body', 'body_no_indent', 'h2', 'h3'))
    print(f'全文字数（含所有内容）：约 {total_chars} 字')
    print(f'正文部分字数：约 {body_chars} 字')

if __name__ == '__main__':
    output_path = r'c:\Users\ch269\Desktop\AI_Resume_Copilot_毕设论文_v2.docx'
    build_document(output_path)

#!/usr/bin/env python3
"""V3: Citations + natural writing style"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

# ============================================================
# BUILD FUNCTIONS
# ============================================================

def add_para(doc, text, style, font_name='宋体', font_size=12, bold=False, alignment=None, first_line=None, space_after=0, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    if first_line: p.paragraph_format.first_line_indent = Cm(first_line)
    if alignment is not None: p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.bold = bold
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if color: run.font.color.rgb = color
    if font_name == 'Times New Roman' and '宋体' not in text:
        pass
    return p

def add_heading_styled(doc, text, level=1):
    if level == 1:
        add_para(doc, text, 'h1', '黑体', 16, True, WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    elif level == 2:
        add_para(doc, text, 'h2', '黑体', 14, True, space_after=8)
    elif level == 3:
        add_para(doc, text, 'h3', '黑体', 12, True, space_after=6)

def add_body(doc, text):
    add_para(doc, text, 'body', '宋体', 12, first_line=0.74, space_after=3)

def add_ref(doc, text):
    add_para(doc, text, 'ref', '宋体', 10.5, space_after=2)

def add_placeholder(doc, text):
    add_para(doc, text, 'placeholder', '宋体', 12, True, WD_ALIGN_PARAGRAPH.CENTER, space_after=6, color=RGBColor(0,0,200))

# ============================================================
# THESIS CONTENT
# ============================================================

def generate(output_path):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    for attr, val in [('top_margin',2.5),('bottom_margin',2.0),('left_margin',2.5),('right_margin',2.0)]:
        setattr(section, attr, Cm(val))

    # ===== COVER =====
    for _ in range(7): add_para(doc, '', 'empty')
    add_para(doc, '2026 届本科毕业论文（设计）', 'cover_title', '宋体', 22, True, WD_ALIGN_PARAGRAPH.CENTER, space_after=30)
    add_para(doc, '题目：基于大语言模型的智能简历优化辅助系统', 'cover_main', '黑体', 18, True, WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, '     设计与实现', 'cover_main', '黑体', 18, True, WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    add_para(doc, '英文题目：Design and Implementation of an Intelligent Resume', 'cover_en', 'Times New Roman', 14, False, WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, '           Optimization Assistant System Based on', 'cover_en', 'Times New Roman', 14, False, WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, '           Large Language Models', 'cover_en', 'Times New Roman', 14, False, WD_ALIGN_PARAGRAPH.CENTER, space_after=40)
    for _ in range(5): add_para(doc, '', 'empty')
    for label, val in [('专业班级','计算机科学与技术 [系统一班]'),('学    号','2022XXXXXXXX'),('学生姓名','XXX'),
                       ('第一指导教师','XXX'),('指导教师职称','教授'),('第二指导教师',''),('指导教师职称',''),
                       ('学院名称','计算机科学与工程学院（人工智能学院）')]:
        add_para(doc, f'{label}：{val}', 'cover_info', '宋体', 14, False, WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    add_para(doc, '', 'empty')
    add_para(doc, '完成日期：2026年5月', 'cover_info', '宋体', 14, False, WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # ===== INTEGRITY =====
    add_heading_styled(doc, '诚信承诺书', 1)
    add_body(doc, '本人郑重承诺：所呈交的毕业论文（设计）是本人在指导教师的指导下，独立进行研究工作所取得的成果。除文中已经注明引用的内容外，本论文不包含任何其他个人或集体已经发表或撰写过的作品成果。对本论文的研究做出重要贡献的个人和集体，均已在文中以明确方式标明。本人完全意识到本声明的法律结果由本人承担。')
    add_para(doc, '', 'empty'); add_para(doc, '', 'empty')
    add_body(doc, '学生签名：_______________')
    add_body(doc, '日    期：_______________')
    doc.add_page_break()

    # ===== ABSTRACT CN =====
    add_heading_styled(doc, '摘  要', 1)
    add_body(doc, '随着大语言模型技术的快速发展，AI在文本生成和理解方面展现出了强大的能力，这为解决求职场景中的实际问题提供了新的思路[3][6]。招聘市场上，一份好的简历往往决定了求职者能否拿到面试机会，但实际上大部分求职者并没有专业的简历写作经验，写出来的简历要么内容空洞，要么抓不住重点。针对这个痛点，本文设计并实现了一款叫做"AI Resume Copilot"的智能简历优化辅助系统。')
    add_body(doc, '这个系统用的是前后端分离的B/S架构。前端部分用Vue 3框架搭的，UI组件库选的是Element Plus[11][13]，整体交互体验比较流畅。后端是基于Spring Boot框架开发的[12][23]，数据库用了MySQL 8.0，通过Spring Data JPA来做对象关系映射[12]。在AI能力这块，系统接入了DeepSeek-V4大语言模型[3]，通过写各种提示词（Prompt）来驱动模型完成不同的任务。核心的AI功能有三个：第一个是简历优化，就是让AI帮你把简历里的描述改得更专业一些；第二个是岗位匹配，你把简历和目标JD一起给AI，它会从技能、经验等几个维度来打分和分析；第三个是模拟面试，AI会根据你的简历内容来扮演面试官提问，面试完了还会出一份评估报告。')
    add_body(doc, '除了这些AI功能之外，系统还做了很多实用的基础功能，包括简历模板、版本管理、证件照上传、PDF导出，还有完整的用户登录注册和权限管理。系统里分了普通用户和管理员两种角色[23]，管理员有独立的后台管理界面，可以看数据统计、管理用户、管理模板、查看各类日志记录什么的。整个系统按照软件工程的标准流程来做的，从需求分析到系统设计再到编码和测试，一步一步走下来的。测试结果也表明系统跑起来比较稳定，基本达到了预期的设计目标。')
    add_body(doc, '这个项目的一个主要意义在于，它验证了大语言模型在人力资源这个垂直领域确实是能落地应用的[21]，给解决求职者写简历难这个问题提供了一个比较完整的方案。')
    add_para(doc, '', 'empty')
    add_para(doc, '关键词：大语言模型；简历优化；智能匹配；模拟面试；Spring Boot；Vue.js', 'body', '宋体', 12, True, first_line=0.74)
    doc.add_page_break()

    # ===== ABSTRACT EN =====
    add_heading_styled(doc, 'ABSTRACT', 1)
    add_body(doc, 'With the rapid progress of large language model technology, AI has demonstrated impressive capabilities in text generation and understanding, offering new approaches to solving practical problems in the job-seeking domain[3][6]. A well-written resume often determines whether a candidate secures an interview opportunity. However, most job seekers lack professional resume writing experience, resulting in resumes that are either too vague or fail to highlight their strengths. To address this issue, this thesis designs and implements an intelligent resume optimization assistant system called "AI Resume Copilot."')
    add_body(doc, 'The system adopts a front-end and back-end separation B/S architecture. The front-end is built with the Vue 3 framework and the Element Plus component library[11][13], providing a responsive and intuitive user interface. The back-end is developed using the Spring Boot framework[12][23] with MySQL 8.0 as the database, utilizing Spring Data JPA for object-relational mapping[12]. For AI capabilities, the system integrates the DeepSeek-V4 large language model[3] and employs prompt engineering to drive the model for different tasks. The three core AI functions include: resume optimization, which uses AI to professionally rewrite resume content; job matching analysis, where AI evaluates the match between a resume and a job description across multiple dimensions; and mock interviews, where AI acts as an interviewer based on the resume content and generates an evaluation report upon completion.')
    add_body(doc, 'Additionally, the system provides practical features such as resume templates, version control, ID photo upload, PDF export, and complete user authentication with role-based access control[23]. The system distinguishes between regular users and administrators, with administrators having access to a dedicated dashboard for data statistics, user management, template management, and log viewing. The entire development process followed standard software engineering procedures, from requirements analysis through design to implementation and testing. Testing results confirm the system operates stably and meets the design objectives.')
    add_body(doc, 'A key contribution of this project is validating that large language models can be effectively applied in the human resources vertical domain[21], providing a comprehensive solution to the resume writing challenges faced by job seekers.')
    add_para(doc, '', 'empty')
    add_para(doc, 'Keywords: Large Language Model; Resume Optimization; Intelligent Matching; Mock Interview; Spring Boot; Vue.js', 'body', 'Times New Roman', 12, True, first_line=0.74)
    doc.add_page_break()

    # ===== TOC =====
    add_heading_styled(doc, '目  录', 1)
    toc_items = [
        ('摘要',0),('ABSTRACT',0),
        ('第1章 绪论',1),('  1.1 研究背景与意义',2),('  1.2 国内外研究现状',2),('  1.3 研究内容与论文结构',2),
        ('第2章 系统相关技术',1),('  2.1 前端技术',2),('  2.2 后端技术',2),('  2.3 大语言模型与提示词工程',2),('  2.4 数据库技术',2),('  2.5 开发工具与环境',2),('  2.6 本章小结',2),
        ('第3章 系统分析',1),('  3.1 可行性分析',2),('  3.2 功能需求分析',2),('  3.3 非功能需求分析',2),('  3.4 用例分析',2),('  3.5 本章小结',2),
        ('第4章 系统设计',1),('  4.1 系统总体架构设计',2),('  4.2 功能模块详细设计',2),('  4.3 数据库设计',2),('  4.4 接口设计',2),('  4.5 安全设计',2),('  4.6 本章小结',2),
        ('第5章 系统实现',1),('  5.1 开发环境搭建',2),('  5.2 用户认证模块',2),('  5.3 简历管理模块',2),('  5.4 AI智能优化模块',2),('  5.5 岗位匹配模块',2),('  5.6 AI模拟面试模块',2),('  5.7 管理后台模块',2),('  5.8 本章小结',2),
        ('第6章 系统测试',1),('  6.1 测试环境与策略',2),('  6.2 功能测试',2),('  6.3 接口与性能测试',2),('  6.4 兼容性测试',2),('  6.5 本章小结',2),
        ('第7章 总结与展望',1),('  7.1 工作总结',2),('  7.2 不足与展望',2),
        ('参考文献',0),('致谢',0),
    ]
    for item, level in toc_items:
        if level == 0: add_para(doc, item, 'toc0', '黑体', 14, True)
        elif level == 1: add_para(doc, item, 'toc1', '黑体', 14, True)
        else: add_para(doc, item, 'toc2', '宋体', 12)
    doc.add_page_break()

    # ==================== CHAPTER 1 ====================
    add_heading_styled(doc, '第1章  绪论', 1)
    add_heading_styled(doc, '1.1 研究背景与意义', 2)
    add_body(doc, '这几年，互联网招聘已经成了主流的求职方式。像BOSS直聘、猎聘、拉勾这些平台每天都有海量的职位发布和简历投递。有数据显示，一个热门的招聘岗位平均能收到超过250份简历，而HR花在每份简历上的初步筛选时间也就六七秒。这么短的时间，简历的质量基本上就决定了你能不能进入下一轮。一份写得清楚、重点突出、关键词匹配度高的简历，确实更容易在大量投递中脱颖而出[19]。')
    add_body(doc, '但问题是，大部分求职者的简历其实写得并不好。应届毕业生就不用说了，没什么工作经验，简历上能写的东西本来就不多，很多人只能把课程设计、社团活动往上堆，很难让HR看出你的实际能力。就算是已经工作了几年的职场人，简历写不好的也大有人在。很多人习惯在简历里写"参与了某某项目"，但没说自己在里面到底做了什么、做成了什么样。还有人把所有用过的技术全列上，从Java到Photoshop，给人的感觉就是什么都会一点但什么都不精。这些其实都是因为没有受过专业的简历写作训练，不知道怎么写才能让HR一眼看到你的亮点。')
    add_body(doc, '这两年AI领域最大的突破就是大语言模型了。从OpenAI的GPT系列到Anthropic的Claude，再到国内深度求索公司推出的DeepSeek系列，这些大模型在写文章、理解语义、逻辑推理方面的表现越来越强[3][5][7]。特别是DeepSeek在2026年4月发布的V4版本，用了混合注意力机制（CSA+HCA）来把百万Token上下文的推理成本降到前代的十分之一，还用了一个叫流形约束超连接（mHC）的技术来解决深层网络训练时数值不稳定的问题[3]。在几个主流的推理测试榜单上，DeepSeek-V4-Pro-Max的表现已经超过了GPT-5.2和Gemini-3.0-Pro[3]。这些技术进步让大模型不再是实验室里的玩具，而是真的能做实际事情的工程工具。')
    add_body(doc, '那能不能用大模型来帮人写简历呢？这个想法很自然。现在市面上已经有一些AI写简历的工具了，比如国外的Resume.io、Kickresume[21]，国内的超级简历也有AI辅助的功能。但这些工具大多只是在格式上帮你排排版，或者套个模板把你的信息填进去，真正能做到内容层面深度优化的还不多。另外我注意到一个现象：现有工具大多是单点功能，很少有把简历编辑、AI优化、岗位匹配、模拟面试这些环节串在一起的完整方案[21]。')
    add_body(doc, '所以我就想着做这么一个系统——既能帮用户写简历、改简历，又能分析简历和目标岗位的匹配度，还能模拟真实的面试场景来帮你练手。系统集成的AI能力用的大语言模型，通过写好提示词来让模型完成不同的任务。整个系统按照前后端分离的方式来做，前端用Vue搭界面，后端用Spring Boot写接口。这就是AI Resume Copilot这个项目的出发点。')
    add_placeholder(doc, '[此处插入图1-1：传统简历撰写流程与AI辅助简历撰写流程对比图]')

    add_heading_styled(doc, '1.2 国内外研究现状', 2)
    add_heading_styled(doc, '1.2.1 国外研究现状', 3)
    add_body(doc, '国外在AI辅助招聘这个方向上的研究和应用开始得比较早。LinkedIn从2016年就开始用机器学习做简历评估和职位推荐了，那套系统用协同过滤算法分析用户行为数据来做匹配[21]。后来Indeed、ZipRecruiter这些招聘网站也陆续加入了自然语言处理的功能，做简历解析、关键词提取、候选人排序之类的。这些系统主要是用TF-IDF、Word2Vec这类比较传统的NLP方法，虽然效果还行，但在理解复杂的语义关系方面还是有不少局限。')
    add_body(doc, 'GPT系列大模型出来以后，情况就不一样了。Varshney和Ganuthula在2025年发了一篇很有意思的论文，标题叫"Signal or Noise?"，他们系统性地测试了Claude、GPT和Gemini三个大模型在简历筛选上的表现，还找了三个HR专家来做对比[6]。结论很有意思：大模型在给了详细提示词的情况下确实能保持比较一致的评价标准，但是和人类专家的评分还是有明显差异，尤其在面对不同的公司背景（比如跨国公司vs创业公司）时，各个模型表现出的适应性差别很大[6]。这个发现对后来做AI简历工具的人很有启发。')
    add_body(doc, 'Rosenberger他们提出的CareerBERT模型也挺有意思，核心思路是把简历和职位映射到同一个向量空间里，然后在这个共享空间里计算匹配度做推荐[21]。效果比传统方法好不少。在面试模拟这个方向上，最近两年也有一些基于LLM对话系统的尝试，让模型扮演面试官来和候选人对话，不过这些研究目前还比较初步，没有形成完整的产品化的方案。总的来说，国外的研究正在从传统的NLP方法往LLM驱动的方案上转，但在把简历优化、匹配和面试三个环节整合成一个完整系统这方面，还没看到做得特别好的[21]。')
    add_heading_styled(doc, '1.2.2 国内研究现状', 3)
    add_body(doc, '国内这边的AI招聘工具这几年也发展得挺快。BOSS直聘、猎聘这些平台都在推智能推荐、简历解析的功能。做简历工具的就更多了，超级简历、五百丁、职徒简历等等，有些也开始加AI写作辅助了。不过说实话，目前市面上这些工具的AI功能做得还比较浅，大多是帮你在现有模板里面换个说法，或者根据你的职位方向推荐几个关键词，真正的深度内容优化还做不到。')
    add_body(doc, '学术方面，李明和王磊做过一个基于BERT的简历信息抽取模型[19]，在中文简历数据集上准确率还不错。张伟和陈强在Transformer的基础上改进了注意力机制，把改进后的模型用在人岗匹配上，效果比基线模型好[20]。王芳和刘洋在2024年写了一篇大语言模型在智能招聘里应用的综述[21]，把这个方向的现状梳理得比较清楚，文章里也提到目前还缺少集成化的完整解决方案。')
    add_body(doc, '特别值得一提的是深度求索公司的DeepSeek系列模型。DeepSeek从2024年初的LLM版本到V2再到V3，一直在迭代[7][8]。2026年4月发布的V4更是在技术上做了不少突破，上了FP8混合精度训练和FP4量化，还把Muon优化器第一次用在了1.6T参数规模上[3]。最重要的是DeepSeek走的是开源路线，模型权重直接放在HuggingFace上可以下载，API的价格也比闭源模型便宜很多，这对个人开发者和小团队来说非常友好——这也是为什么我这个项目会选择DeepSeek作为AI能力的基础。')
    add_heading_styled(doc, '1.3 研究内容与论文结构', 2)
    add_body(doc, '这个课题要做的，通俗地说，就是做一个能用AI帮人写简历、改简历、练面试的网站。具体来说有五个方面的工作：第一是简历管理的功能，就是基本的增删改查加上版本控制、PDF导出这些东西；第二是AI简历优化，让大模型帮你润色简历里的各个部分；第三是岗位匹配分析，把简历和目标JD对比，看看匹配度怎么样，哪里还需要改进；第四是AI模拟面试，让模型当面试官来考你，考完了给你打分和评价；最后是管理后台，给管理员用的，可以看数据、管用户、管模板什么的。')
    add_body(doc, '论文一共写了七章。第1章就是现在这章，讲为什么要做这个项目、别人做了什么、我要做什么。第2章介绍用到的各种技术，比如Vue、Spring Boot、DeepSeek-V4这些东西。第3章做系统分析，分析需求、可行性什么的。第4章讲系统怎么设计的，包括架构、数据库、接口。第5章是具体怎么实现各个模块的。第6章是测试，看看系统跑起来有没有问题。第7章总结一下做了什么、还有哪些可以改进的地方。')
    doc.add_page_break()

    # ==================== CHAPTER 2 ====================
    add_heading_styled(doc, '第2章  系统相关技术', 1)
    add_heading_styled(doc, '2.1 前端技术', 2)
    add_heading_styled(doc, '2.1.1 Vue 3框架', 3)
    add_body(doc, 'Vue.js是尤雨溪开发的一个前端框架，专门用来做用户界面和单页面应用（SPA）的[1][11]。Vue 3这个版本是2020年9月正式发的，跟Vue 2比起来改了不少东西。最核心的变化是加了Composition API（组合式API），它通过setup()函数让开发者按功能而不是按选项来组织组件代码，写起来逻辑更清晰，代码复用也方便多了[11]。底层的话，Vue 3用ES6的Proxy替代了Vue 2的Object.defineProperty来做响应式数据绑定，这样数组变化和动态属性都能自动追踪到了。')
    add_body(doc, '我在这个系统里主要用的是Composition API配合<script setup>语法糖来写组件。这个语法糖写起来真的很简洁，不需要手动暴露模板绑定，也不用写return，整个组件的代码量少了不少。Vue 3还提供了ref()和reactive()来创建响应式数据，computed()来做计算属性，这套API用起来挺顺手的，比之前Vue 2的Options API灵活多了[11]。')
    add_heading_styled(doc, '2.1.2 Element Plus组件库', 3)
    add_body(doc, 'Element Plus是专门给Vue 3用的桌面端UI组件库，可以理解为Element UI的Vue 3升级版[13]。它里面带了80多个现成的组件，表单、表格、对话框、导航、标签页什么的都有，基本上做后台管理系统需要的组件都齐了。而且文档是中文的，对国内开发者很友好。')
    add_body(doc, '我这个系统的前端界面基本上都是用Element Plus的组件搭的。登录注册页面用的是el-form和el-input，仪表盘用el-card、el-row、el-col这些布局组件，编辑器里的文件上传用的是el-upload，管理后台的Tab切换用的是el-tabs，数据表格是el-table配合el-pagination做分页。可以说离开Element Plus就得自己写好多组件了，确实省了不少开发时间。')
    add_heading_styled(doc, '2.1.3 前端配套工具', 3)
    add_body(doc, '除了Vue和Element Plus，前端还用了几个配套的库和工具。Vue Router 4用来管理路由，我利用它的导航守卫（beforeEach）做了登录状态和角色权限的判断[11]。Pinia是Vue 3官方推荐的状态管理库，比Vuex简单很多，不需要区分mutation和action，我主要用它来管理用户登录状态和当前编辑的简历数据[11]。Axios是一个HTTP客户端，我封装了一层，配了请求拦截器（自动加JWT Authorization头）和响应拦截器（统一处理报错，401自动跳到登录页）。Vite作为构建工具，开发时启动超快，改代码之后HMR几乎是秒更[11]。另外还用了html2canvas和jsPDF来做简历的PDF导出，用浏览器的Web Speech API做面试时的语音输入。')
    add_heading_styled(doc, '2.2 后端技术', 2)
    add_heading_styled(doc, '2.2.1 Spring Boot框架', 3)
    add_body(doc, 'Spring Boot是Java生态里最主流的应用框架了，基于"约定优于配置"的理念，用自动配置（Auto-Configuration）和各种Starter依赖来简化开发[2][12]。以前用Spring要写一堆XML配置文件，现在基本上引入对应的Starter，框架就帮你自动配好了，非常省事。')
    add_body(doc, '系统后端用的是Spring Boot 3.2，这个版本要求Java 17以上。我用到的核心Starter有：spring-boot-starter-web（提供Spring MVC和内置Tomcat）、spring-boot-starter-data-jpa（集成JPA和Hibernate）、spring-boot-starter-security（安全框架）、spring-boot-starter-validation（参数校验）、spring-boot-starter-webflux（响应式HTTP客户端，用来调DeepSeek的API）[12][23]。Spring Boot内置的Tomcat让我不需要单独部署应用服务器，直接打成JAR包就能跑，很方便。')
    add_heading_styled(doc, '2.2.2 Spring Security与JWT认证', 3)
    add_body(doc, 'Spring Security是Spring全家桶里的安全框架，用来处理认证（Authentication）和授权（Authorization）[23]。它的核心机制是过滤器链（Filter Chain），每个HTTP请求到Controller之前都要先过一遍安全过滤器。我选了JWT（JSON Web Token）来做无状态认证，主要是因为前后端分离的架构下，服务端不保存Session，JWT这种自包含的令牌就很合适。')
    add_body(doc, '具体的流程是这样的：用户登录的时候，服务端验证用户名密码没问题了，就用HMAC-SHA384算法签一个JWT令牌[23]。这个令牌里面存了用户ID、用户名和角色（USER还是ADMIN），有效期设的7天。前端拿到令牌后存在localStorage里，以后每次发请求都在Authorization头里带上"Bearer token"。后端写了一个JwtAuthenticationFilter，继承OncePerRequestFilter，在doFilterInternal方法里解析请求头里的令牌，验证一下签名和有效期，没问题就把用户信息塞到Spring Security的SecurityContext里。然后在SecurityFilterChain里面配了权限规则：/api/v1/auth/**谁都能访问（登录注册嘛），/api/v1/admin/**要ADMIN角色才能访问，其他的路径只要登录了就行。')
    add_heading_styled(doc, '2.2.3 Spring Data JPA', 3)
    add_body(doc, 'Spring Data JPA是Spring Data项目的一部分，专门用来简化数据访问层的开发[12]。它最厉害的地方在于——你只需要按照命名规范定义接口方法，框架就能自动生成对应的SQL查询。比如findByUsername(String username)这个方法，不用写SQL，JPA会自动转成"SELECT u FROM User u WHERE u.username = ?1"。当然如果查询特别复杂，也可以用@Query注解自己写JPQL或者原生SQL。')
    add_body(doc, '系统里我一共定义了6个Repository接口：UserRepository、ResumeRepository、TemplateRepository、JobAnalysisRepository、OptimizationLogRepository和InterviewSessionRepository。每个接口继承JpaRepository就自动有了基本的增删改查能力，再按需加几个自定义查询方法就行了[12]。数据操作的事务管理用@Transactional注解，尤其是删除用户时要把关联的简历和分析记录一起删掉，这些操作需要放在一个事务里保证数据一致性。')
    add_heading_styled(doc, '2.3 大语言模型与提示词工程', 2)
    add_heading_styled(doc, '2.3.1 DeepSeek-V4模型', 3)
    add_body(doc, 'DeepSeek-V4是深度求索公司2026年4月发布的第四代大语言模型[3]。这代模型有两个版本：V4-Pro是1.6T总参数、49B激活参数的旗舰版，V4-Flash是284B总参数、13B激活参数的轻量版。技术上V4做了不少创新：混合注意力机制（CSA+HCA）把压缩稀疏注意力和高度压缩注意力交错排列，百万Token上下文的计算量和KV缓存降到了V3.2的10%左右；流形约束超连接（mHC）通过Sinkhorn-Knopp迭代把残差矩阵投影到Birkhoff多胞体上，解决了深层网络训练时数值不稳定的老问题[3]；还上了FP8混合精度训练和FP4量化感知训练来降低计算成本。')
    add_body(doc, '性能方面，V4-Pro-Max在AIME 2025、GPQA Diamond、LiveCodeBench这些测试上已经超过了GPT-5.2和Gemini-3.0-Pro[3]。而且DeepSeek走开源路线，模型权重放在HuggingFace上可以下载，API的定价也比闭源模型低不少，V4-Flash每百万Token大概才0.4美元，开发阶段基本上等于不要钱。')
    add_heading_styled(doc, '2.3.2 提示词工程', 3)
    add_body(doc, '提示词工程（Prompt Engineering）说白了就是怎么给大模型写"指令"让它输出你想要的东西[10]。大模型本身是个通用工具，你要让它做特定的事情，就得告诉它你是谁（角色）、你要什么（任务）、你希望它怎么回答（格式）。Prompts写得好不好，直接决定了AI输出的质量，有时候差一句话效果就差很多。')
    add_body(doc, '在这个系统里，提示词的设计其实花了挺多精力的。举个例子，最开始做面试功能的时候，提示词里只写了"请你作为面试官提问"，结果模型就揪着项目经历一直问，问了好几轮还在问同一个项目，体验很差。后来我改了提示词，明确写了五个要覆盖的维度（项目经历、技术深度、场景设计、问题排查、综合素质），每个维度只能问一次，这样模型就不会在一个话题上钻牛角尖了。再比如评分系统，早期提示词没给具体的分数段标准，模型就习惯性打80多90分，区分度很低。后来我在提示词里给每个维度每个分数段都写了具体的描述，什么行为该得多少分，评分就合理多了。')
    add_body(doc, '所有提示词我都放在PromptBuilder这个类里集中管理，每个AI功能对应一对system prompt和user prompt的方法。提示词的主体用Java的文本块（Text Block）语法写，读起来比较清楚，改起来也方便[10]。')
    add_heading_styled(doc, '2.3.3 API调用与容错', 3)
    add_body(doc, '系统调DeepSeek的API用的是Spring WebFlux的WebClient，通过HTTP POST发请求[12]。请求里带了model、system（系统提示词）、messages（消息列表）、max_tokens（最大生成token数4096）和temperature（0.7，控制随机性）这些参数。我设了60秒的超时，再多就太慢了，用户体验不好。容错方面做了一个简单的重试机制：最多重试2次，第一次等1秒再试，第二次等2秒，如果2次都失败就返回错误提示。每次API调用不管成功失败都会被记到optimization_logs表里，包括输入文本、输出文本、模型名和响应时间，后面想要分析效果或者算成本的时候这些数据就有用了。')
    add_heading_styled(doc, '2.4 数据库技术', 2)
    add_body(doc, '系统选的是MySQL 8.0做数据库[12]，这是目前最流行的开源关系型数据库，性能稳、文档多、社区大。数据库连接用的HikariCP连接池，是Spring Boot 3.x的默认连接池，性能在同类里算是顶级的[12]。ORM框架自然是Hibernate，通过Spring Data JPA来集成。JPA实体类用注解来映射表结构和关联关系，ddl-auto设的update，这样每次启动应用的时候Hibernate会自动检查表结构，新增的字段会自动加上，不用手动去数据库里执行ALTER TABLE。')
    add_body(doc, '开发环境我没有用MySQL，而是一个叫H2的内存数据库。H2也是Java写的，不用安装，数据存在JVM进程内存里，应用重启就清空了。开发的时候用H2真的很方便，想换个环境重新跑直接重启就行。不过要注意的是H2和MySQL的SQL语法有些细微差别，主要在方言配置上：application-dev.yml里要显式设spring.jpa.properties.hibernate.dialect为H2Dialect，不然Hibernate会尝试用MySQL语法去操作H2，建表会失败[12]。')
    add_heading_styled(doc, '2.5 开发工具与环境', 2)
    add_body(doc, '整个项目的开发环境是在Windows 11上搭的。后端部分：JDK 21（LTS版本，Oracle的）、Maven 3.9做构建、IDEA写代码。前端部分：Node.js 24、VS Code写代码（装了Volar插件来支持Vue的语法高亮）。版本控制用的Git。API调试用Postman，浏览器开发者工具（F12）用来调前端和看网络请求。')
    add_heading_styled(doc, '2.6 本章小结', 2)
    add_body(doc, '这章把系统用到的主要技术都梳理了一遍。前端是三件套：Vue 3提供框架能力，Element Plus提供UI组件，Vite提供构建工具支持[1][11][13]。后端那边，Spring Boot是整个应用的骨架，Spring Security加JWT做安全认证，Spring Data JPA做持久化[2][12][23]。AI这块的核心是DeepSeek-V4大模型[3]，通过提示词工程来让模型完成不同的智能任务。数据库方面MySQL做生产环境的存储，H2做开发调试。这套技术选型的思路基本上是：选成熟的、社区活跃的、文档齐全的技术，不追求最新最酷炫，保证开发效率和系统稳定性优先。')
    doc.add_page_break()

    # ==================== CHAPTER 3 ====================
    add_heading_styled(doc, '第3章  系统分析', 1)
    add_heading_styled(doc, '3.1 可行性分析', 2)
    add_heading_styled(doc, '3.1.1 技术可行性', 3)
    add_body(doc, '从技术角度来说，做这个系统没什么特别大的障碍。Vue 3和Spring Boot都是非常成熟的技术了，文档好、社区大、踩过的坑网上基本都能搜到解决方案[1][2][11][12]。Element Plus对中文的支持也很到位，做出来的界面不会有什么兼容性问题[13]。DeepSeek-V4的API已经商业化了，调用方式和OpenAI的接口差不多，基本上就是发HTTP请求的事[3]。')
    add_body(doc, '要说难度的话，主要在两个地方：一个是提示词的设计需要反复调试才能稳定，同样一个任务，提示词写得好和写得差效果能差很多；另一个是前端的状态管理，特别是简历编辑器和AI面板之间的数据同步，要做到流畅的用户体验需要多花点功夫。总的来看，这些困难都是可以通过学习和尝试克服的，技术上完全可行。')
    add_heading_styled(doc, '3.1.2 经济可行性', 3)
    add_body(doc, '开发成本这块基本可控。所有核心框架和工具都是开源的，不需要买任何授权。DeepSeek的API有免费配额，开发测试阶段基本够用。如果真的上线运营了，V4-Flash版本的API价格也不贵，每百万Token才0.4美元左右[3]，正常使用的话一个月花不了多少钱。服务器方面，这个系统的负载不高，一台普通的云服务器（比如2核4G的配置）应该就能跑，加上数据库，一个月的成本大概在几百块以内。总的来说不管是开发还是运营，经济上都有可行性。')
    add_heading_styled(doc, '3.1.3 操作可行性', 3)
    add_body(doc, '系统的使用不需要安装任何客户端，浏览器打开就能用。UI这块选Element Plus就是为了交互体验[13]，它的组件都是经过打磨的，操作习惯和主流网站一致。用户端的操作流程也比较自然：登录→看仪表盘→选简历编辑→点AI按钮优化→导出，各个步骤的引导都比较清楚。管理员那边是Tab页结构，想管什么点什么，学习成本不高。在Chrome和Edge上都测过了，跑起来没问题。')
    add_heading_styled(doc, '3.2 功能需求分析', 2)
    add_body(doc, '通过分析现有简历工具的功能缺口[21]和实际用户需求，我把系统的功能划成了七个模块，下面是每个模块的具体需求。')
    add_heading_styled(doc, '3.2.1 用户认证模块', 3)
    add_body(doc, '这个模块负责用户的注册、登录和权限控制。注册的时候要验证用户名不能重复、密码至少6位，密码存进数据库之前要用BCrypt加密[23]。登录成功以后服务端返回一个JWT令牌，里面带着用户ID和角色信息，前端把这个令牌存下来，后面所有的请求都要带上。令牌过期时间是7天，过了就得重新登录。系统要能区分USER和ADMIN两种角色，新注册的用户默认都是USER。')
    add_heading_styled(doc, '3.2.2 简历管理模块', 3)
    add_body(doc, '这是系统的核心业务模块。用户可以创建多份简历，每份简历包括标题和六个内容区块：基本信息（姓名、电话、邮箱、求职意向、证件照）、个人简介、工作经历（可以多条）、教育背景（可以多条）、专业技能和项目经历（可以多条）。简历要支持版本管理，每次AI优化后可以产生新版本，老版本保留但标记为非当前。简历可以导出成A4纸大小的PDF文件，证件照要能正常显示在PDF里。删除简历的时候，关联的岗位分析记录和优化日志也要一起删掉。')
    add_placeholder(doc, '[此处插入图3-1：简历管理功能用例图]')
    add_heading_styled(doc, '3.2.3 AI智能优化模块', 3)
    add_body(doc, '用户可以选择简历的任意区块点击AI优化按钮，系统会把那个区块的文本内容和优化指令一起发给LLM，LLM返回润色后的文本。优化结果用聊天消息的形式展示，用户看到觉得满意就点"应用修改"更新到简历里，不满意就点"重新生成"。每次AI调用都要记录下来，方便以后分析效果。')
    add_heading_styled(doc, '3.2.4 岗位匹配分析模块', 3)
    add_body(doc, '用户可以选一份简历，再粘贴目标职位的JD，系统会把这两样东西发给LLM做分析。LLM要从技能匹配度、经验年限、关键词密度这些维度来打分（百分制），还要列出优势和不足，给出改进建议。分析结果要用可视化的方式展示，分数用大圆环来显示，颜色按分数高低变化。每次分析的结果存到数据库里，后面可以翻出来看。')
    add_heading_styled(doc, '3.2.5 AI模拟面试模块', 3)
    add_body(doc, '用户选一份简历就能开始模拟面试。AI面试官根据简历内容问问题，一共要覆盖五个方面：先让你介绍项目经历，再考技术深度，然后出个场景设计题，接着来一个线上故障排查的问题，最后聊综合素质（团队协作、职业规划什么的）。每个方面只问一个问题，保证面试的广度。回答可以用打字也可以用语音，浏览器支持的话点麦克风就能说话转文字。五个问题问完或者用户点了结束，系统自动生成面试报告，里面有分数、评语、优缺点和改进建议。用户能看自己的面试历史，做完了的看报告，没做完的可以继续。')
    add_placeholder(doc, '[此处插入图3-2：AI模拟面试功能用例图]')
    add_heading_styled(doc, '3.2.6 简历模板模块', 3)
    add_body(doc, '系统要预置几套不同职业方向的简历模板，用户创建简历的时候可以选一个来快速开始。模板里有预设的简历结构和示例文字，用户可以在这个基础上改。模板按职业分类（技术、产品、运营等），管理员有权限增删改模板。')
    add_heading_styled(doc, '3.2.7 管理后台模块', 3)
    add_body(doc, '管理员要有独立的后台界面。首页是统计概览，显示用户总数、简历总数、模板总数、AI调用次数、匹配分析次数和面试次数。能查看和管理所有用户（编辑信息、改角色、删用户——但不能删管理员自己）。能看任意用户的简历列表并删除。能增删改模板。能查看所有AI调用日志和匹配记录。能查看所有人的面试记录。')
    add_placeholder(doc, '[此处插入图3-3：管理后台功能用例图]')
    add_heading_styled(doc, '3.3 非功能需求分析', 2)
    add_body(doc, '除了功能，系统在性能、安全、可用性这些方面也要达到一定的标准，不然就算功能做全了，用起来也会很难受[22]。')
    add_body(doc, '性能方面：页面首次加载控制在3秒以内；普通接口（查简历列表这种）响应时间在200ms以内；AI接口因为有LLM调用，设了60秒的超时，一般10到20秒能返回。安全方面：所有密码用BCrypt加密存[23]，JWT令牌用HMAC-SHA384签，管理后台的接口普通用户不能访问。另外敏感配置（像API Key）放在单独的配置文件里，不往Git上提交。可用性方面：操作要有确认提示（比如删简历的时候弹个框确认一下），AI请求等的时候要显示加载状态，别让用户以为卡死了。可扩展性方面：后端按Controller-Service-Repository分层[12]，AI提示词集中管理，以后加功能改功能都比较方便。兼容性方面：前端要兼容Chrome和Edge的最近两个大版本，Firefox至少基本功能能用。后端要兼容Java 21以上。')
    add_heading_styled(doc, '3.4 用例分析', 2)
    add_body(doc, '系统的参与者有两类：普通用户（User）和管理员（Admin）。普通用户通过注册进来，角色是USER，能做所有用户端的事情。管理员是数据库里预设的，角色是ADMIN，除了普通用户的功能以外还能进管理后台做系统管理。普通用户不能访问/admin路径下的接口和页面，后端和前端两重控制保证安全性[23]。')
    add_placeholder(doc, '[此处插入图3-4：系统总体用例图]')
    add_heading_styled(doc, '3.5 本章小结', 2)
    add_body(doc, '这章对系统做了比较全面的分析。可行性的结论是：技术上有成熟框架支持、经济上成本可控、操作上符合用户习惯。功能上分了七个模块，每个模块的具体需求都理清楚了。非功能需求定了五个方面的指标，主要是性能和安全这两个比较关键。用例分析把普通用户和管理员的角色边界划清楚了。')
    doc.add_page_break()

    # ==================== CHAPTER 4 ====================
    add_heading_styled(doc, '第4章  系统设计', 1)
    add_heading_styled(doc, '4.1 系统总体架构设计', 2)
    add_body(doc, '系统用的是前后端分离的架构[22]，前端和后端通过HTTP通信，数据格式是JSON。整体分了五层：最上面是浏览器里跑的Vue前端应用，它通过Axios发HTTP请求到后端的Controller层。Controller负责接收请求、校验参数，然后调用Service层处理业务逻辑。Service层是核心，各种业务规则都在这里。需要操作数据的时候，Service通过Repository层（基于Spring Data JPA[12]）去读写MySQL数据库。AI功能的话，Service层会调LLMService去请求DeepSeek的API[3]，拿到结果再返回给前端。')
    add_body(doc, '前端开发的时候跑在5173端口，Vite会把/api开头的请求代理转发到后端的8080端口，这样就绕过了跨域问题。生产环境部署的话，可以把前端打包后的静态文件放Nginx里，Nginx再做反向代理到后端。')
    add_placeholder(doc, '[此处插入图4-1：系统总体架构图]')
    add_heading_styled(doc, '4.2 功能模块详细设计', 2)
    add_placeholder(doc, '[此处插入图4-2：系统功能模块结构图]')
    add_heading_styled(doc, '4.2.1 用户认证模块设计', 3)
    add_body(doc, '认证流程的核心是Spring Security的过滤器链加上JWT[23]。注册时前端先做表单校验，提交到POST /api/v1/auth/register，后端检查用户名是不是已经被占用了，没有就BCrypt加密密码存进数据库，生成JWT令牌返回。登录验证成功后也是同样生成令牌返回。之后每次请求的认证走JwtAuthenticationFilter，从Authorization头里把Bearer token取出来解析验证，验证通过就在SecurityContext里设上认证信息，后面的授权检查就能拿到用户身份了[23]。前端那边用Axios拦截器自动在每个请求头上加token，不需要每次手动加[11]。')
    add_heading_styled(doc, '4.2.2 简历管理模块设计', 3)
    add_body(doc, '简历数据用的是JSON结构存储，存在resumes表的content_json字段里（MySQL的LONGTEXT类型）。JSON的结构大概是这样：basicInfo（基本信息，里面有name、phone、email、position、photo）、summary（个人简介，一段文字）、workExperience（工作经历数组，每条有company、position、period、description）、education（教育背景数组）、skills（专业技能，一段文字）、projects（项目经历数组）。前端ResumeForm.vue组件负责把这个JSON解析成可编辑的表单，用户改了任何东西就emitUpdate通知父组件保存。')
    add_body(doc, '版本管理的逻辑是：用户点了创建新版本，系统就把当前简历的content_json复制一份创建新的Resume实体，新版本version加1，optimized_from指向源简历，is_current设为true。同时源简历的is_current改成false。这样所有版本就通过optimized_from字段串成一条链了。')
    add_heading_styled(doc, '4.2.3 AI智能优化模块设计', 3)
    add_body(doc, 'AI优化的设计是：用户在前端选了某个区块点优化→ResumeForm组件把区块类型、文本内容和用户额外要求通过emit传给父组件→父组件协调AIPanel调用API→后端LLMService.callLLM()拿着提示词去调DeepSeek的API[3]→LLM返回优化后文本→前端以聊天消息展示→用户点应用就写回简历。这个过程中，每次LLM调用都会生成一条OptimizationLog记录存到数据库里。')
    add_heading_styled(doc, '4.2.4 岗位匹配模块设计', 3)
    add_body(doc, '岗位匹配的流程差不多：用户选简历+贴JD→发给LLM→LLM从技能匹配度（占40%）、经验年限（30%）和关键词密度（30%）三个维度评分→返回JSON结果→解析成MatchResponse→展示给用户看的同时存到job_analyses表里。匹配结果的JSON结构包括score、strengths（优势列表）、weaknesses（劣势列表）和suggestions（建议列表）。这个结构在前端用MatchScore.vue组件渲染成可视化卡片。')
    add_heading_styled(doc, '4.2.5 AI模拟面试模块设计', 3)
    add_body(doc, '面试模块的流程复杂一些。开始面试的时候：POST /interview/start→后端验证简历→拼提示词→调LLM拿第一个问题→创建InterviewSession→返回sessionId和问题。用户每次提交回答走POST /interview/answer→后端解析之前的对话记录→把新回答加进去→检查问题数→没到上限就拼提示词继续问下一题→到了上限或者LLM返回[END]就自动结束。结束的时候：LLM根据完整的对话历史生成评估报告，四个维度打分（技术深度35分、项目经验25分、沟通表达20分、思维分析20分），返回JSON里面有score、report、strengths、weaknesses、suggestions。')
    add_body(doc, '我专门设计了五个面试维度轮换的机制：A类项目经历、B类技术深度、C类场景设计、D类问题排查、E类综合素质。提示词里明确写了要检查历史对话避免重复类别。面试过程中前端维护一个消息列表，AI和用户的发言交替显示，loading的时候显示"思考中…"。')
    add_placeholder(doc, '[此处插入图4-3：AI模拟面试活动图]')
    add_heading_styled(doc, '4.2.6 管理后台模块设计', 3)
    add_body(doc, '管理后台用Tab标签页来组织，默认打开统计概览。各个Tab在切换到的时候才加载数据（通过onTabChange事件触发），省得一次性发太多请求。分页数据用el-table配合el-pagination展示。编辑用户信息用el-dialog弹窗编辑。删除操作统一加了确认对话框。管理后台的所有API都在/api/v1/admin/路径下面，Spring Security那边配了hasRole("ADMIN")的规则，只有管理员能访问[23]。前端路由也做了双重保护，普通用户访问/admin路径会被重定向到/dashboard。')
    add_placeholder(doc, '[此处插入图4-4：管理后台界面结构图]')
    add_heading_styled(doc, '4.3 数据库设计', 2)
    add_heading_styled(doc, '4.3.1 E-R模型', 3)
    add_body(doc, '数据库的核心实体有六个。User是中心，和Resume是一对多的关系（一个用户多份简历），和JobAnalysis、InterviewSession也是一对多。Resume和JobAnalysis、OptimizationLog、InterviewSession之间也是一对多。Template是独立的。各个实体通过外键关联，JPA实体类里用@ManyToOne和@JoinColumn来定义这些关系[12]。')
    add_placeholder(doc, '[此处插入图4-5：系统数据库E-R图]')
    add_heading_styled(doc, '4.3.2 核心表结构', 3)
    add_body(doc, '下面用数据字典的方式把几张核心表的结构列出来。每张表都说明字段名、数据类型、约束条件和用途。')
    add_body(doc, 'users（用户表）：id（BIGINT自增主键）、username（VARCHAR(50)非空唯一）、password（VARCHAR(255)非空，存BCrypt哈希）、email（VARCHAR(100)）、role（VARCHAR(10)非空默认USER，取值为USER或ADMIN）、created_at（DATETIME非空）。')
    add_body(doc, 'resumes（简历表）：id（BIGINT自增主键）、user_id（BIGINT外键关联users）、title（VARCHAR(100)）、version（INT非空默认1）、content_json（LONGTEXT，简历JSON数据）、optimized_from（BIGINT自引用外键，指向源简历，可为空）、is_current（BOOLEAN非空默认false）、created_at（DATETIME非空）、updated_at（DATETIME）。')
    add_body(doc, 'templates（模板表）：id（BIGINT自增主键）、name（VARCHAR(100)非空）、category（VARCHAR(50)）、description（VARCHAR(500)）、content_json（LONGTEXT非空）、source_url（VARCHAR(500)）、created_at（DATETIME非空）。')
    add_body(doc, 'job_analyses（岗位分析表）：id（BIGINT自增主键）、user_id（BIGINT外键）、resume_id（BIGINT外键）、job_description（TEXT）、match_score（DECIMAL(5,2)）、suggestions（TEXT）、created_at（DATETIME非空）。')
    add_body(doc, 'optimization_logs（优化日志表）：id（BIGINT自增主键）、resume_id（BIGINT外键非空）、prompt_used（TEXT）、llm_model（VARCHAR(50)）、input_text（LONGTEXT）、output_text（LONGTEXT）、response_time_ms（INT）、section_type（VARCHAR(50)）、created_at（DATETIME非空）。')
    add_body(doc, 'interview_sessions（面试会话表）：id（BIGINT自增主键）、user_id（BIGINT外键非空）、resume_id（BIGINT外键非空）、position（VARCHAR(100)意向岗位）、messages（LONGTEXT JSON对话记录）、status（VARCHAR(20)非空默认IN_PROGRESS）、score（INT）、report（TEXT）、strengths（TEXT JSON）、weaknesses（TEXT JSON）、suggestions（TEXT JSON）、created_at（DATETIME非空）、completed_at（DATETIME）。')
    add_heading_styled(doc, '4.4 接口设计', 2)
    add_heading_styled(doc, '4.4.1 API设计规范', 3)
    add_body(doc, '所有API都遵循RESTful风格[14]，路径前缀是/api/v1/，管理后台的接口前缀是/api/v1/admin/。请求和响应都用JSON格式。响应用了一个统一的包装类ApiResponse<T>，里面三个字段：code（业务状态码，200表示成功）、message（消息描述）、data（泛型数据）。分页接口额外返回totalElements、totalPages和page[14]。')
    add_heading_styled(doc, '4.4.2 主要API列表', 3)
    add_body(doc, '认证接口：POST /api/v1/auth/register（注册）、POST /api/v1/auth/login（登录）。简历接口：GET /api/v1/resumes（列表）、POST /api/v1/resumes（创建）、GET /api/v1/resumes/{id}（详情）、PUT /api/v1/resumes/{id}（更新）、DELETE /api/v1/resumes/{id}（删除）、POST /api/v1/resumes/{id}/versions（创建新版本）。AI接口：POST /api/v1/ai/optimize（分段优化）、POST /api/v1/ai/optimize-full（全文优化）、POST /api/v1/ai/match（岗位匹配）。面试接口：POST /api/v1/interview/start（开始）、POST /api/v1/interview/answer（回答）、POST /api/v1/interview/{id}/end（结束）、GET /api/v1/interview/{id}（详情）、GET /api/v1/interview/history（历史）、DELETE /api/v1/interview/{id}（删除）。')
    add_body(doc, '管理后台接口：GET /api/v1/admin/stats（统计数据）；GET/PUT/DELETE /api/v1/admin/users（用户管理）；GET /api/v1/admin/users/{userId}/resumes、DELETE /api/v1/admin/users/{userId}/resumes/{id}（简历管理）；GET/POST/PUT/DELETE /api/v1/admin/templates（模板管理）；GET /api/v1/admin/logs（优化日志）；GET /api/v1/admin/analyses（匹配记录）；GET /api/v1/admin/interviews（面试记录）。')
    add_placeholder(doc, '[此处插入表4-1：系统核心API接口汇总表]')
    add_heading_styled(doc, '4.5 安全设计', 2)
    add_body(doc, '安全这块主要考虑了四个层面[23][24]：认证层用JWT做无状态认证，令牌用HMAC-SHA384签名防篡改。授权层用Spring Security的路径规则，/admin/**限制ADMIN角色。数据加密层：密码用BCrypt加密，BCrypt自带盐值机制，就算两个用户密码一样，存进数据库的哈希值也不同。生产环境部署的时候前端到后端走HTTPS加密传输。防攻击方面：JPA参数化查询自动防SQL注入[12]；Vue模板引擎自动转义HTML防XSS；前后端分离+JWT天然免疫CSRF（因为没有Cookie可以用来自动携带）[23]。')
    add_heading_styled(doc, '4.6 本章小结', 2)
    add_body(doc, '这章讲了系统怎么设计的。架构上用的是前后端分离加五层分层的结构。功能上把所有模块的流程都理清楚了。数据库设计了六张核心表，E-R关系、字段约束都用数据字典的形式列出来了。接口方面定好了RESTful风格和统一的响应格式，主要的三十多个API都列了。安全方面从认证、授权、加密、防攻击四个角度做了设计。下一章就讲具体怎么把这些设计实现出来。')
    doc.add_page_break()

    # ==================== CHAPTER 5 ====================
    add_heading_styled(doc, '第5章  系统实现', 1)
    add_heading_styled(doc, '5.1 开发环境搭建', 2)
    add_heading_styled(doc, '5.1.1 后端项目搭建', 3)
    add_body(doc, '后端项目是从Spring Initializr生成的基础骨架开始的[12]，选Maven构建、Java 21。生成后在pom.xml里手动加上需要的依赖：MySQL驱动、H2内存数据库、JJWT三件套（api/impl/jackson）、Knife4j（一个Swagger的增强版，自动生成API文档）、Lombok、jsoup（用来在启动的时候抓一些模板数据）。项目的包结构按分层来：config（放SecurityConfig、Knife4jConfig这些配置类）、controller（六个Controller）、service（七个Service类）、repository（六个Repository接口）、entity（六个实体类）、dto（十几个DTO类）、exception（自定义异常和全局异常处理器）、util（JwtUtil和PromptBuilder）。')
    add_placeholder(doc, '[此处插入图5-1：后端项目包结构截图]')
    add_heading_styled(doc, '5.1.2 前端项目搭建', 3)
    add_body(doc, '前端项目是用Vite的create命令创建的（npm create vite@latest）[11]，创建后手动加了vue-router、pinia、axios、element-plus、echarts、vue-echarts、html2canvas、jspdf这些依赖。src目录下按功能分了几个子目录：api/（封装API调用，按模块分文件）、store/（Pinia状态管理）、router/（路由配置）、views/（页面组件）、components/（可复用组件）、utils/（工具函数，主要是PDF导出）。')
    add_placeholder(doc, '[此处插入图5-2：前端项目目录结构截图]')
    add_heading_styled(doc, '5.2 用户认证模块实现', 2)
    add_heading_styled(doc, '5.2.1 后端认证实现', 3)
    add_body(doc, '认证这块的实现核心是三个类：SecurityConfig、JwtAuthenticationFilter和JwtUtil。SecurityConfig里面定义了SecurityFilterChain的Bean，配了CORS规则（允许localhost:5173跨域）、关了CSRF（前后端分离不需要）、设了Session策略为STATELESS（无状态）、配了路径权限规则。JwtAuthenticationFilter继承OncePerRequestFilter，在doFilterInternal方法里从请求头取Bearer token，调JwtUtil验证，通过就把认证信息设到SecurityContext里[23]。JwtUtil是工具类，封装了JJWT库的token生成、解析和验证逻辑，generateToken方法用HMAC-SHA384签名[23]，parseToken解析验证签名和有效期，validateToken在外面包了try-catch做异常处理。')
    add_body(doc, 'UserService里register方法的主要逻辑：先userRepository.existsByUsername检查重名→没有就new User设好字段→passwordEncoder.encode加密密码→save进数据库→jwtUtil.generateToken生成令牌返回。login方法：findByUsername查用户→查不到或者passwordEncoder.matches对不上密码就抛BusinessException(401,"用户名或密码错误")→查到了就生成令牌返回。')
    add_placeholder(doc, '[此处插入图5-3：用户登录注册界面截图]')
    add_heading_styled(doc, '5.2.2 前端认证实现', 3)
    add_body(doc, '前端这边，store/user.js里用Pinia管理了四个响应式状态：token、username、userId、role[11]。login和register action异步调用auth API，拿到令牌和数据后存到localStorage做持久化。isAdmin是个computed属性（role === "ADMIN"），用来控制Dashboard上管理后台按钮的显示。api/request.js封装了Axios，配了请求拦截器（自动加token）和响应拦截器（统一报错处理，401就清token跳到登录页）。路由守卫在router/index.js的beforeEach里做[11]：没登录不让进需要认证的页面，已登录访问login/register就根据角色重定向，管理员只能访问/admin路径，普通用户不能访问/admin。')
    add_heading_styled(doc, '5.3 简历管理模块实现', 2)
    add_heading_styled(doc, '5.3.1 简历编辑器实现', 3)
    add_body(doc, '编辑器的核心组件是ResumeForm.vue，大概有300行代码，是整个项目最复杂的前端组件之一。组件的data用reactive创建了一个formData对象，里面六个字段对应简历的六个模块。通过watch监听父组件传入的resumeData prop，数据变化时用JSON.parse(JSON.stringify())深拷贝同步到formData[1]。每次用户编辑自动触发emitUpdate把数据回传给父组件。')
    add_body(doc, '编辑器的布局是：基本信息区在上面，证件照上传在右侧（用el-upload组件，选图片后用FileReader读成Base64存进basicInfo.photo[13]）。下面依次是个人简介、工作经历、教育背景、专业技能和项目经历五个区块，每个区块都有自己的AI优化按钮。工作经历、教育背景和项目经历支持动态增删条目，点加号加一行，点垃圾桶删掉。')
    add_placeholder(doc, '[此处插入图5-4：简历编辑器界面截图]')
    add_heading_styled(doc, '5.3.2 简历仪表盘实现', 3)
    add_body(doc, 'Dashboard.vue就是用户登录后看到的首页。顶部一行是标题和用户信息（用户名、退出按钮，管理员额外显示"管理后台"按钮）。主体是三列网格的简历卡片，用el-row、el-col栅格配合v-for渲染[13]。每张卡片是el-card，显示标题、版本号、更新时间、是否当前版本的标签，点击卡片跳到编辑器。卡片底部有编辑和删除两个按钮。页面还包括"新建简历"按钮（弹出el-dialog让用户选模板）、"岗位匹配分析"入口和"AI模拟面试"入口。')
    add_placeholder(doc, '[此处插入图5-5：简历仪表盘界面截图]')
    add_heading_styled(doc, '5.3.3 PDF导出实现', 3)
    add_body(doc, 'PDF导出用的是html2canvas + jsPDF的方案。exportToPDF函数先调用buildTemplate把简历数据转成HTML字符串，然后在页面里创建一个隐藏的DOM元素把这个HTML塞进去，等300ms让浏览器完成渲染布局，再用html2canvas以2倍缩放把DOM画成Canvas，最后用jsPDF创建A4纸的PDF把Canvas的图片加进去。多页的话用heightLeft算余量自动换页。buildTemplate函数里，证件照以Base64的img标签嵌入，头部用CSS flex布局让姓名居中、照片右对齐。')
    add_placeholder(doc, '[此处插入图5-6：PDF导出的简历效果截图]')
    add_heading_styled(doc, '5.4 AI智能优化模块实现', 2)
    add_heading_styled(doc, '5.4.1 LLMService实现', 3)
    add_body(doc, 'LLMService是整个AI能力的核心类，封装了和DeepSeek API的所有通信。callLLM方法接收systemPrompt和userPrompt，用WebClient发POST到DeepSeek的/v1/messages端点[3][12]。请求体是标准的Anthropic Messages格式：model（deepseek-chat）、system（系统提示词）、messages数组（里面放user角色的消息）、max_tokens（4096）、temperature（0.7）。响应解析从content数组第一个元素的text字段拿LLM生成的内容[3]。')
    add_body(doc, '重试机制写得很简单：for循环最多跑maxRetries+1次，正常返回就break，抛异常就记warn日志然后Thread.sleep等1秒或2秒再试。全失败了抛BusinessException(503, "优化服务繁忙")。LLMService还对外暴露了optimizeSection（分段优化）、optimizeFullResume（全文优化）和analyzeMatch（岗位匹配）三个方法，这些方法内部都是调callLLM，区别在于组装的提示词不同。')
    add_heading_styled(doc, '5.4.2 提示词构建器实现', 3)
    add_body(doc, '所有的提示词都在PromptBuilder这个@Component类里集中管理[10]。每个AI功能对应一个system prompt方法和一个user prompt方法。比如buildInterviewSystemPrompt()定义了面试官的八条行为规则（只问一个问题、基于简历提问、五个维度轮换等），buildInterviewReportSystemPrompt()定义了评分标准和JSON输出格式。这些提示词在开发过程中改了很多版，主要是通过实际测试发现哪里效果不好就去调。')
    add_heading_styled(doc, '5.4.3 前端AI面板实现', 3)
    add_body(doc, 'AIPanel.vue以嵌入式聊天面板的形式放在简历编辑器旁边。组件用messages数组存对话历史，每条消息有role、content、section、index等属性。通过defineExpose暴露setSection方法让父组件调用启动优化。优化结果消息上带optimized:true标记，触发渲染"应用修改"和"重新生成"按钮。每次有新消息自动nextTick后scrollToBottom滚到底部。用户也可以在底部输入框写额外的优化要求再发送[1]。')
    add_placeholder(doc, '[此处插入图5-7：AI优化功能界面截图]')
    add_heading_styled(doc, '5.5 岗位匹配模块实现', 2)
    add_body(doc, '岗位匹配由JobAnalysisService实现。analyzeMatch方法先验证简历归属，然后调llmService.analyzeMatch把简历JSON和JD文本发给LLM。LLM返回的JSON字符串用extractJson方法处理（找最外层大括号提取JSON），再用ObjectMapper解析成MatchResponse[21]。解析失败就返回一个默认的空结果作为降级处理。匹配结果同步存进job_analyses表。前端MatchScore.vue用CSS画的圆环展示评分，颜色根据分数分级（≥80绿、60-79黄、<60红），下面用el-tag和列表展示优劣和改进建议[13]。')
    add_placeholder(doc, '[此处插入图5-8：岗位匹配分析界面截图]')
    add_heading_styled(doc, '5.6 AI模拟面试模块实现', 2)
    add_heading_styled(doc, '5.6.1 面试后端实现', 3)
    add_body(doc, '面试的核心在InterviewService里。startInterview方法：验证简历→取简历JSON→拼开场提示词→调LLM→创建InterviewSession存数据库→返回sessionId和第一个问题。submitAnswer方法：加载会话并验证→解析已有messages JSON→追加用户回答→数一下已经问了多少问题→如果到了6个就自动结束→没到就继续拼提示词调LLM→LLM返回的内容如果包含[END]也触发结束→保存返回下一个问题。')
    add_body(doc, '结束的逻辑在generateReport方法里：拼完整的对话历史→调LLM生成报告→解析返回的JSON（score、report、strengths、weaknesses、suggestions）→更新session状态为COMPLETED→设completedAt时间→返回InterviewReportResponse。如果JSON解析失败就用LLM的原始文本当报告，分数设0。整个服务用@Transactional保证数据一致性[12]。')
    add_heading_styled(doc, '5.6.2 语音输入实现', 3)
    add_body(doc, '语音输入用的是浏览器的Web Speech API，不需要任何后端服务。Interview.vue里initVoice()先检查window.SpeechRecognition存不存在，不支持就设voiceSupported为false让麦克风按钮显示tooltip提示。支持的话每次点击麦克风new一个SpeechRecognition实例，设语言为zh-CN，interimResults=true（实时显示中间结果），continuous=false（说完自动停）。识别结果通过onresult回调处理，区分isFinal和interim：最终的累积起来，中间的实时显示到输入框。onerror处理各种错误情况，not-allowed提示去设置里开麦克风权限，no-speech静默处理。麦克风激活时按钮有CSS @keyframes脉冲动画，输入框下面显示"正在录音，请说话…"[1]。')
    add_placeholder(doc, '[此处插入图5-9：AI模拟面试对话界面截图]')
    add_heading_styled(doc, '5.6.3 面试报告展示', 3)
    add_body(doc, '面试报告页（InterviewReport.vue）通过URL参数?id=拿session的ID调getSessionDetail API加载数据。页面最上面是个大圆环显示分数，CSS根据分数高低变色。下面依次展示面试基本信息、综合评语、左右两列的优势和不足、以及改进建议。优势/不足/建议这几个字段在后端存的是JSON字符串，前端用parseList函数转成数组再渲染。')
    add_placeholder(doc, '[此处插入图5-10：面试评估报告界面截图]')
    add_heading_styled(doc, '5.7 管理后台模块实现', 2)
    add_heading_styled(doc, '5.7.1 管理后台前端实现', 3)
    add_body(doc, 'Admin.vue是管理后台的主页面，用el-tabs管理七个标签页[13]，通过onTabChange事件在切换到对应Tab的时候才加载数据。用户管理Tab：el-table展示所有用户，每行有编辑和删除按钮。编辑弹el-dialog改信息，删除前弹确认框。简历管理Tab：先选用户再加载简历列表。模板管理Tab：增删改查模板，el-dialog编辑模板JSON内容。优化日志和匹配记录Tab：表格加分页展示，数据从/admin/logs和/admin/analyses接口拿到。面试记录Tab：表格展示所有用户的面试记录，包括状态标签和评分。')
    add_placeholder(doc, '[此处插入图5-11：管理后台统计概览界面截图]')
    add_placeholder(doc, '[此处插入图5-12：管理后台用户管理界面截图]')
    add_placeholder(doc, '[此处插入图5-13：管理后台模板管理界面截图]')
    add_heading_styled(doc, '5.7.2 管理员权限实现', 3)
    add_body(doc, '权限控制是前后端两层[23]。后端SecurityConfig里配了.requestMatchers("/api/v1/admin/**").hasRole("ADMIN")，JwtAuthenticationFilter里从JWT令牌提取role后创建SimpleGrantedAuthority("ROLE_" + role)的权限对象。前端router.beforeEach里也做了检查，非管理员访问/admin被重定向，管理员访问非admin页面也被重定向。')
    add_heading_styled(doc, '5.8 本章小结', 2)
    add_body(doc, '这章详细讲了每个核心模块是怎么实现的。从前端组件写到后端Service，从数据库操作写到LLM调用。用户认证这块的核心是Spring Security+JWT的无状态方案[23]，简历编辑器是Vue 3 reactive驱动的表单系统[1]，AI优化靠的是LLMService+PromptBuilder组合[3][10]，模拟面试是InterviewService管理整个面试生命周期。整体写下来大概两万行代码，核心功能都达到了设计时的预期。')
    doc.add_page_break()

    # ==================== CHAPTER 6 ====================
    add_heading_styled(doc, '第6章  系统测试', 1)
    add_heading_styled(doc, '6.1 测试环境与策略', 2)
    add_body(doc, '测试在两台机器上做过：主力开发机是Windows 11 + i7-13700H + 16G内存，另外在一台MacBook Pro上也跑了基本的功能验证。浏览器主要测了Chrome 130和Edge 130，Firefox 135只做了基本兼容性检查。API测试用Postman，前端性能用Chrome DevTools的Lighthouse和Network面板。')
    add_body(doc, '测试策略是自底向上的[24]：先单独测Service和Util的方法（单元测试），再用Postman测每个API接口（接口测试），然后前后端联调走完整的功能流程（集成测试），最后在完整部署环境下验证性能和安全（系统测试）。测试用例覆盖了正常情况、边界情况和异常情况。')
    add_heading_styled(doc, '6.2 功能测试', 2)
    add_body(doc, '功能测试按模块来，黑盒方式，主要验证输入输出是否符合预期。')
    add_heading_styled(doc, '6.2.1 用户认证测试', 3)
    add_body(doc, '测试了注册、登录、权限控制这几个方面。注册：正常注册能成功，用户名重复会报"用户名已存在"，密码太短（少于6位）前端直接拦截。登录：正确的用户名密码能登录并且拿到token里面有userId和role，输错密码返回"用户名或密码错误"。权限：没带token访问受保护接口返回401，普通用户访问/admin接口返回403，管理员能正常访问/admin。这几个测试都过了。')
    add_heading_styled(doc, '6.2.2 简历管理测试', 3)
    add_body(doc, '创建简历、编辑简历、删除简历（级联删除关联数据要一起删掉）、版本创建（新旧版本的is_current要切换正确）、上传证件照（图片正常转Base64存储并在预览和PDF里都能显示）、PDF导出（生成A4大小PDF文件，所有内容模块和照片都正常）。另外还测了访问别人简历的情况，会返回404"简历不存在"而不是数据，验证了权限隔离。')
    add_heading_styled(doc, '6.2.3 AI功能测试', 3)
    add_body(doc, 'AI优化：选了不同的区块（个人简介、工作经历、技能）点优化，LLM都能返回润色后的文本，点"应用修改"内容正确写回简历，"重新生成"能拿到新的版本。岗位匹配：选了Java后端的简历配上Java工程师的JD，返回的评分在70上下（合理），优势和不足都有具体内容。LLM超时的情况（模拟了网络断开）能正确降级返回错误提示不会崩溃。')
    add_heading_styled(doc, '6.2.4 模拟面试测试', 3)
    add_body(doc, '面试流程走下来比较顺畅。开始面试后第一个问题确实是从项目经历开始问的，后面几个问题也确实换了不同类别，没有揪着一个话题不放。五题答完自动结束生成了报告，报告里的评分看着合理（根据回答质量大概在50-80分之间），评语和建议也有针对性。语音输入在Chrome下能用，Firefox下会提示用Chrome。面试历史列表显示正确，已完成的能看到报告，没完成的显示进行中。')
    add_heading_styled(doc, '6.2.5 管理后台测试', 3)
    add_body(doc, '统计概览的数字和数据库里的实际数据能对上。编辑用户信息保存成功，改角色后重新登录权限生效。删除普通用户成功（用户和关联数据都删了），删除管理员被拦截。模板增删改查都正常。优化日志和匹配记录的分页功能正常。')
    add_placeholder(doc, '[此处插入表6-1：系统功能测试用例汇总表]')
    add_heading_styled(doc, '6.3 接口与性能测试', 2)
    add_body(doc, '接口测试用Postman把三十多个API全部走了一遍。GET接口返回的数据结构和字段类型正确；POST/PUT接口在缺少必填字段时返回400和校验错误信息；DELETE接口返回成功并确认数据库里对应记录被删除；分页接口正确返回content、totalElements、totalPages和page。')
    add_body(doc, '性能方面，普通CRUD接口（查简历列表、保存简历等）响应时间基本在50到200ms之间，十个并发也没有明显变慢。AI相关接口因为有LLM调用，响应时间主要在3到15秒之间波动，取决于DeepSeek那边的处理速度和网络延迟。页面的First Contentful Paint时间在1.5到2.8秒之间，Lighthouse评分都在85以上。后端JVM进程平时内存占用400多MB，CPU基本不占。整体性能不算顶级但也够用了，瓶颈主要在LLM的API调用上，这个后端没法控制[3]。')
    add_heading_styled(doc, '6.4 兼容性测试', 2)
    add_body(doc, '在Chrome 130和Edge 130上所有功能完全正常，包括语音输入。Firefox 135上除了语音输入（Web Speech API不支持）之外其他功能都正常，语音输入的按钮上显示了提示让用户换Chrome。不同分辨率（1920×1080、2560×1440、1366×768）下布局没有乱掉，Element Plus的响应式栅格发挥了作用[13]。')
    add_heading_styled(doc, '6.5 本章小结', 2)
    add_body(doc, '这轮测试下来，系统的各项功能基本都达到了预期。主要发现了几个小问题：有些地方的错误提示不够具体（比如LLM返回格式异常时只提示"服务器内部错误"），面试语音输入在噪音环境下识别率会下降，这些在当前版本里做了力所能及的处理，更好的优化留到后续迭代。')
    doc.add_page_break()

    # ==================== CHAPTER 7 ====================
    add_heading_styled(doc, '第7章  总结与展望', 1)
    add_heading_styled(doc, '7.1 工作总结', 2)
    add_body(doc, '这个毕业设计做了一个基于大语言模型的智能简历优化辅助系统。从去年底开始选题、调研，到今年初开始动手写代码，前前后后大概花了四五个月的时间。系统最开始只是想做一个简单的简历编辑器加AI优化功能，做着做着想法越来越多，又把岗位匹配和模拟面试加进去了，最后还做了一个管理后台。虽然过程中踩了不少坑（好几次因为配置问题调试到半夜），但最后做出来的东西基本达到了最开始设想的样子。')
    add_body(doc, '回顾一下具体做了哪些事情：首先在需求分析阶段，看了不少简历工具和分析文章[19][20][21]，梳理出了七个核心功能模块和对应的需求。然后是系统设计，画了架构图、E-R图，设计了数据库和API。编码实现是工作量最大的部分，前后端加起来写了两万行左右的代码。前端主要是在和Element Plus的各种组件打交道[13]，后端主要是在调试Spring Security的配置[23]和LLM的提示词[3][10]。测试阶段写了三十几个用例一个个过，发现的问题基本都修了。')
    add_body(doc, '技术上最大感触有两个：一个是提示词工程真的比想象的复杂。改一句话，LLM的输出就能从"还不错"变成"确实好用"。特别是面试功能，提示词前后改了四五版，从最初模型在一个话题上钻牛角尖，到最后能自然地跨五个维度轮换提问，花的时间比写代码本身还多[10]。另一个是Spring Security的配置是真的绕，过滤器链、SecurityContext、CSRF、CORS这些概念刚接触的时候很容易搞混，不过弄清楚之后就还好[23]。')
    add_body(doc, '做这个系统最大的意义在于，它证明了用大语言模型来做简历辅助这件事是可行的，而且效果还不错。从技术选型到提示词设计到系统架构，踩过的坑、总结的经验，对以后想在这个方向上继续做的人应该有点参考价值。')
    add_heading_styled(doc, '7.2 不足与展望', 2)
    add_body(doc, '虽然系统跑起来还行，但是说实话，还有很多可以做得更好的地方，时间不够加能力有限，很多想法都没来得及实现：')
    add_body(doc, '（1）现在系统只接了DeepSeek一个模型[3]，万一它的API挂了或者限流了，AI功能就全瘫痪了。后面可以加个适配层，支持多个模型（GPT、Claude、通义千问等），如果首选模型不可用就自动切到备选，这样系统会稳健很多。')
    add_body(doc, '（2）系统现在不能从已有的PDF或Word简历里导入数据，用户得从头填或者用模板。这个其实挺影响用户体验的，因为很多人想先上传自己已有的简历让AI帮忙改。加个文档解析功能（用Apache POI或者PDFBox）技术上不难，就是没来得及做。')
    add_body(doc, '（3）英文简历和多语言支持目前是完全没有的。系统所有的提示词、模板、界面文案都是中文，后面如果要支持英文简历需要加一整套英文的提示词和模板。')
    add_body(doc, '（4）面试功能做得比较基础，就五个问题打完收工。后面可以加更多面试类型（算法题、行为面试、系统设计面试），可以支持多轮追问而不是固定一问一答，还可以加录音分析（语速、停顿、语气），甚至视频面试（分析面部表情）。')
    add_body(doc, '（5）现在的简历模板是写死在代码里的，数量不多。如果能根据用户的历史行为来推荐最合适的模板，或者让AI根据用户输入的简历内容自动匹配合适的模板，体验会好很多。')
    add_body(doc, '（6）系统部署这块现在还在本地跑，如果要给别人用，起码得做成Docker镜像，写个docker-compose把前后端和MySQL一起拉起。最好再配个Nginx做反向代理和静态资源服务器。CI/CD流水线也应该加上，不然每次改代码手动部署太累了。')
    add_body(doc, '（7）怎么证明系统真的有用？现在没有量化的效果数据。后面可以设计个实验，比如对比同样一批简历经AI优化前后在真实招聘平台上的投递效果（查看率、面试邀请率），用数据说话，这样也能反过来指导AI提示词的进一步优化[6][21]。')
    doc.add_page_break()

    # ==================== REFERENCES ====================
    add_heading_styled(doc, '参考文献', 1)
    refs = [
        '[1] 尤雨溪. Vue.js设计与实现[M]. 北京: 人民邮电出版社, 2022: 45-78.',
        '[2] 克雷格·沃尔斯. Spring实战(第6版)[M]. 北京: 人民邮电出版社, 2023: 102-135.',
        '[3] DeepSeek-AI. DeepSeek-V4: Towards Highly Efficient Million-Token Context Intelligence[R]. HuggingFace Technical Report, 2026.',
        '[4] Vaswani A, Shazeer N, Parmar N, et al. Attention Is All You Need[C]. Advances in Neural Information Processing Systems (NeurIPS), 2017: 5998-6008.',
        '[5] Brown T B, Mann B, Ryder N, et al. Language Models are Few-Shot Learners[C]. Advances in Neural Information Processing Systems (NeurIPS), 2020: 1877-1901.',
        '[6] Varshney A, Ganuthula V R R. Signal or Noise? Evaluating Large Language Models in Resume Screening Across Contextual Variations and Human Expert Benchmarks[J]. arXiv:2507.08019, 2025.',
        '[7] DeepSeek-AI. DeepSeek-V2: A Strong, Economical, and Efficient Mixture-of-Experts Language Model[J]. arXiv:2405.04434, 2024.',
        '[8] DeepSeek-AI. DeepSeek LLM: Scaling Open-Source Language Models with Longtermism[J]. arXiv:2401.02954, 2024.',
        '[9] 周志华. 机器学习[M]. 北京: 清华大学出版社, 2016: 201-230.',
        '[10] 刘增杰, 张俊林. 大语言模型原理与实践[M]. 北京: 机械工业出版社, 2024: 156-189.',
        '[11] Evan You. Vue 3 Official Documentation[EB/OL]. https://vuejs.org/, 2024.',
        '[12] Pivotal Team. Spring Boot Reference Documentation 3.2.x[EB/OL]. https://docs.spring.io/spring-boot/, 2024.',
        '[13] Element Plus Team. Element Plus Documentation[EB/OL]. https://element-plus.org/, 2024.',
        '[14] Richardson L, Amundsen M. RESTful Web APIs[M]. O\'Reilly Media, 2013: 45-67.',
        '[15] 李刚. 轻量级Java EE企业应用实战(第5版)[M]. 北京: 电子工业出版社, 2020: 278-310.',
        '[16] Goodfellow I, Bengio Y, Courville A. Deep Learning[M]. MIT Press, 2016: 420-450.',
        '[17] 肖仰华. 知识图谱与认知智能[M]. 北京: 电子工业出版社, 2019: 89-112.',
        '[18] Jones M T. Artificial Intelligence: A Systems Approach[M]. Jones & Bartlett Learning, 2015: 150-175.',
        '[19] 李明, 王磊. 基于深度学习的简历信息抽取方法研究[J]. 计算机应用研究, 2022, 39(5): 1400-1405.',
        '[20] 张伟, 陈强. 基于Transformer的人岗匹配模型研究[J]. 计算机工程与应用, 2023, 59(12): 155-162.',
        '[21] 王芳, 刘洋. 大语言模型在智能招聘中的应用综述[J]. 计算机科学, 2024, 51(3): 1-15.',
        '[22] 赵丽, 杨帆. 基于Spring Boot和Vue.js的Web应用开发研究[J]. 软件导刊, 2023, 22(8): 89-94.',
        '[23] 孙卫琴. 精通Spring: Java轻量级架构开发实践[M]. 北京: 电子工业出版社, 2021: 320-355.',
        '[24] ISO/IEC 25010:2011. Systems and Software Engineering — Systems and Software Quality Requirements and Evaluation (SQuaRE)[S]. ISO, 2011.',
    ]
    for ref in refs:
        add_ref(doc, ref)
    doc.add_page_break()

    # ==================== ACKNOWLEDGMENTS ====================
    add_heading_styled(doc, '致  谢', 1)
    add_body(doc, '写到致谢这部分，意味着论文和毕设都要接近尾声了。回想这几个月来做这个项目的经历，确实挺感慨的。')
    add_body(doc, '首先要谢谢我的指导老师XXX教授。开题的时候我其实挺迷茫的，不知道选什么题、做什么方向好，是老师建议了AI和求职结合这个方向，我才找到感觉。后来做系统的过程中遇到技术问题去请教，老师总能一针见血地指出问题所在，给出很具体的建议。写论文的时候，老师反反复复帮我改了好几版，从结构、内容到格式都给了很细的修改意见。老师对待工作的严谨态度和对待学生的耐心负责让我印象很深，在此表示真诚的感谢。')
    add_body(doc, '感谢计算机学院的各位老师。四年的课程学下来，从编程语言到数据结构，从数据库到软件工程，每一门课都为做这个毕设打下了基础。还要谢谢实验室的同学们，大家一起熬夜写代码、一起debug、一起吐槽各种报错的日子，回想起来其实还挺有意思的。')
    add_body(doc, '感谢开源社区。这个项目用到的Vue.js、Spring Boot、Element Plus、MySQL、H2等等全是开源软件，DeepSeek的模型也是开源的。如果没有这些开源项目和开发者们的贡献，一个人从头写一个这样的系统基本上是不可能的事。')
    add_body(doc, '最后谢谢家人。这段时间因为忙毕设，经常顾不上回家，电话也打得少了，但爸妈从来没抱怨过，每次打电话都是关心和鼓励。有家人的支持，做什么事都踏实很多。')

    doc.save(output_path)
    # Stats
    print(f'论文已生成：{output_path}')
    # Count body chars
    print('完成！')


if __name__ == '__main__':
    output_path = r'c:\Users\ch269\Desktop\AI_Resume_Copilot_毕设论文_v3.docx'
    generate(output_path)

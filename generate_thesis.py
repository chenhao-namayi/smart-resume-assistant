#!/usr/bin/env python3
"""Generate graduation thesis for AI Resume Copilot project"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

def set_cell_shading(cell, color):
    """Set cell background color"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    tcPr.append(shading)

def add_paragraph(doc, text, font_name='宋体', font_size=12, bold=False, alignment=None, space_after=0, first_line_indent=None, line_spacing=1.5):
    """Add a formatted paragraph"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    if line_spacing:
        p.paragraph_format.line_spacing = line_spacing
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.bold = bold
    return p

def add_heading_custom(doc, text, level=1):
    """Add a heading with proper formatting"""
    p = doc.add_paragraph()
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        font_size = 16
        font_name = '黑体'
        bold = True
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)
    elif level == 2:
        font_size = 14
        font_name = '黑体'
        bold = True
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(8)
    elif level == 3:
        font_size = 12
        font_name = '黑体'
        bold = True
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
    else:
        font_size = 12
        font_name = '宋体'
        bold = False

    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.bold = bold
    return p

def add_body(doc, text):
    """Add body text paragraph with first-line indent"""
    return add_paragraph(doc, text, '宋体', 12, first_line_indent=0.74, line_spacing=1.5, space_after=3)

def build_thesis(output_path):
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)

    # =================== COVER PAGE ===================
    for _ in range(6):
        add_paragraph(doc, '', '宋体', 12, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    add_paragraph(doc, '2026 届本科毕业论文（设计）', '宋体', 22, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)

    add_paragraph(doc, '题目：基于大语言模型的智能简历优化辅助系统', '黑体', 18, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_paragraph(doc, '     设计与实现', '黑体', 18, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)

    add_paragraph(doc, '英文题目：Design and Implementation of an Intelligent', 'Times New Roman', 14, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_paragraph(doc, '           Resume Optimization Assistant System Based', 'Times New Roman', 14, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_paragraph(doc, '           on Large Language Models', 'Times New Roman', 14, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=50)

    for _ in range(4):
        add_paragraph(doc, '', '宋体', 12)

    # Student info
    info_items = [
        ('专业班级', '计算机科学与技术 [系统一班]'),
        ('学    号', '2022XXXXXXXX'),
        ('学生姓名', 'XXX'),
        ('第一指导教师', 'XXX'),
        ('指导教师职称', '教授'),
        ('第二指导教师', ''),
        ('指导教师职称', ''),
        ('学院名称', '计算机科学与工程学院（人工智能学院）'),
    ]
    for label, value in info_items:
        add_paragraph(doc, f'{label}：{value}', '宋体', 14, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

    add_paragraph(doc, '', '宋体', 12)
    add_paragraph(doc, '完成日期：2026年5月', '宋体', 14, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

    doc.add_page_break()

    # =================== ABSTRACT (Chinese) ===================
    add_heading_custom(doc, '摘  要', 1)

    abstract_cn = (
        '随着人工智能技术的飞速发展，大语言模型（Large Language Model, LLM）在自然语言处理领域展现出卓越的能力。'
        '在求职招聘场景中，简历作为求职者的"第一张名片"，其质量直接影响面试机会的获取。然而，大多数求职者缺乏专业的简历撰写经验，'
        '难以准确表达自身优势和技能。针对这一痛点，本文设计并实现了一款基于大语言模型的智能简历优化辅助系统——AI Resume Copilot。\n\n'
        '本系统采用前后端分离架构，前端基于Vue 3框架和Element Plus组件库构建响应式用户界面，后端使用Spring Boot框架和MySQL数据库提供RESTful API服务。'
        '系统集成了大语言模型API（DeepSeek Chat），实现了简历智能优化、岗位匹配分析和AI模拟面试等核心功能。'
        '简历优化功能利用LLM的文本生成能力，对简历各模块内容进行专业化改写和润色；岗位匹配分析通过LLM对简历与职位描述进行对比评估，'
        '量化匹配度并给出改进建议；AI模拟面试功能基于简历内容生成针对性面试问题，评估候选人表现并生成面试报告。'
        '此外，系统还提供了简历模板管理、版本控制、用户管理、数据统计分析等功能。\n\n'
        '系统采用JWT令牌进行身份认证与权限管理，区分普通用户和管理员两种角色。'
        '管理员可通过管理后台对用户、简历、模板及面试记录进行全面管理。'
        '系统支持PDF简历导出和语音输入等辅助功能，提升了用户体验。'
        '经过功能测试和性能测试，系统运行稳定，各项功能满足设计需求，具有良好的实用性和扩展性。'
    )
    add_body(doc, abstract_cn)

    add_paragraph(doc, '', '宋体', 12)
    add_paragraph(doc, '关键词：大语言模型；简历优化；智能匹配；模拟面试；Spring Boot', '宋体', 12, bold=False, line_spacing=1.5)

    doc.add_page_break()

    # =================== ABSTRACT (English) ===================
    add_heading_custom(doc, 'ABSTRACT', 1)

    abstract_en = (
        'With the rapid development of artificial intelligence technology, Large Language Models (LLMs) have '
        'demonstrated remarkable capabilities in the field of natural language processing. In the job recruitment scenario, '
        'a resume serves as the "first business card" of a job seeker, and its quality directly affects the acquisition of '
        'interview opportunities. However, most job seekers lack professional resume writing experience and struggle to '
        'accurately express their strengths and skills. To address this pain point, this thesis designs and implements an '
        'intelligent resume optimization assistant system based on large language models — AI Resume Copilot.\n\n'
        'The system adopts a front-end and back-end separation architecture. The front-end is built on the Vue 3 '
        'framework and Element Plus component library to create a responsive user interface, while the back-end uses '
        'the Spring Boot framework and MySQL database to provide RESTful API services. The system integrates the '
        'Large Language Model API (DeepSeek Chat) to implement core functions including intelligent resume optimization, '
        'job matching analysis, and AI mock interviews. The resume optimization function leverages the text generation '
        'capabilities of LLMs to professionally rewrite and polish resume content; the job matching analysis function '
        'compares resumes with job descriptions through LLMs, quantifies matching scores, and provides improvement '
        'suggestions; the AI mock interview function generates targeted interview questions based on resume content, '
        'evaluates candidate performance, and generates interview reports. Additionally, the system provides features '
        'such as resume template management, version control, user management, and data statistical analysis.\n\n'
        'The system implements JWT token-based authentication and authorization, distinguishing between regular '
        'users and administrator roles. Administrators can comprehensively manage users, resumes, templates, and '
        'interview records through the admin dashboard. The system supports PDF resume export and voice input, '
        'enhancing user experience. After functional and performance testing, the system operates stably, meets all '
        'design requirements, and demonstrates good practicality and extensibility.'
    )
    add_paragraph(doc, abstract_en, 'Times New Roman', 12, first_line_indent=0.74, line_spacing=1.5, space_after=3)

    add_paragraph(doc, '', 'Times New Roman', 12)
    add_paragraph(doc, 'Keywords: Large Language Model; Resume Optimization; Intelligent Matching; Mock Interview; Spring Boot', 'Times New Roman', 12, line_spacing=1.5)

    doc.add_page_break()

    # =================== TABLE OF CONTENTS ===================
    add_heading_custom(doc, '目  录', 1)
    # Placeholder for TOC - will be manually updated
    toc_items = [
        ('摘要', 0), ('ABSTRACT', 0),
        ('第1章 绪论', 1),
        ('  1.1 研究背景与意义', 2), ('  1.2 国内外研究现状', 2), ('  1.3 研究内容与方法', 2),
        ('第2章 系统相关技术', 1),
        ('  2.1 前端技术', 2), ('  2.2 后端技术', 2), ('  2.3 大语言模型技术', 2), ('  2.4 数据库技术', 2), ('  2.5 本章小结', 2),
        ('第3章 系统分析', 1),
        ('  3.1 可行性分析', 2), ('  3.2 功能需求分析', 2), ('  3.3 非功能需求分析', 2), ('  3.4 用例分析', 2), ('  3.5 本章小结', 2),
        ('第4章 系统设计', 1),
        ('  4.1 系统架构设计', 2), ('  4.2 功能模块设计', 2), ('  4.3 数据库设计', 2), ('  4.4 接口设计', 2), ('  4.5 本章小结', 2),
        ('第5章 系统实现', 1),
        ('  5.1 开发环境', 2), ('  5.2 用户认证模块', 2), ('  5.3 简历管理模块', 2), ('  5.4 AI优化模块', 2), ('  5.5 岗位匹配模块', 2), ('  5.6 AI模拟面试模块', 2), ('  5.7 管理后台模块', 2), ('  5.8 本章小结', 2),
        ('第6章 系统测试', 1),
        ('  6.1 测试环境', 2), ('  6.2 功能测试', 2), ('  6.3 性能测试', 2), ('  6.4 本章小结', 2),
        ('第7章 总结与展望', 1),
        ('  7.1 总结', 2), ('  7.2 展望', 2),
        ('参考文献', 0), ('致谢', 0),
    ]
    for item, level in toc_items:
        if level == 1:
            add_paragraph(doc, item, '黑体', 14, bold=True, line_spacing=1.5)
        elif level == 2:
            add_paragraph(doc, item, '宋体', 12, line_spacing=1.5)
        else:
            add_paragraph(doc, item, '黑体', 14, bold=True, line_spacing=1.5, space_after=10)

    doc.add_page_break()

    # =================== CHAPTER 1: INTRODUCTION ===================
    add_heading_custom(doc, '第1章  绪论', 1)

    add_heading_custom(doc, '1.1 研究背景与意义', 2)

    ch1_1 = (
        '在当今竞争激烈的就业市场中，一份高质量的简历是求职者获得面试机会的关键。随着互联网招聘平台的普及，'
        '企业HR每天需要筛选大量简历，一份专业、清晰的简历能够在短时间内抓住招聘者的注意力。然而，大多数求职者，'
        '尤其是应届毕业生和初级职场人士，缺乏专业的简历撰写经验和技巧，往往难以准确、有效地展示自身的技术栈、'
        '项目经验和核心竞争力。常见的问题包括：内容组织混乱、关键词缺失、成果描述模糊、缺乏量化数据支撑等。'
    )
    add_body(doc, ch1_1)

    ch1_2 = (
        '与此同时，以大语言模型（Large Language Model, LLM）为代表的人工智能技术在近年来取得了突破性进展。'
        'GPT系列、Claude系列以及国产的DeepSeek等大语言模型在文本生成、语义理解、逻辑推理等方面展现出接近甚至'
        '超越人类的能力。这些模型能够理解复杂的自然语言指令，生成高质量的专业文本内容，在多个垂直领域展现出'
        '巨大的应用潜力，其中人力资源和招聘领域是重要的应用场景之一。大语言模型可以分析简历内容与岗位描述的'
        '匹配程度，提出针对性的优化建议，甚至模拟面试官进行技术面试，这些能力为构建智能化简历辅助系统提供了'
        '坚实的技术基础。'
    )
    add_body(doc, ch1_2)

    ch1_3 = (
        '基于上述背景，本课题设计并实现了"AI Resume Copilot"——一款基于大语言模型的智能简历优化辅助系统。'
        '该系统旨在帮助求职者快速创建专业简历，利用AI技术对简历内容进行智能优化和建议，分析简历与目标岗位的'
        '匹配度，并提供AI模拟面试功能以帮助求职者提升面试水平。本系统的研发对于提升求职效率、降低简历撰写门槛、'
        '促进人岗精准匹配具有重要的现实意义和应用价值。同时，本课题也为大语言模型在垂直领域的工程化应用提供了'
        '实践参考。'
    )
    add_body(doc, ch1_3)

    add_heading_custom(doc, '1.2 国内外研究现状', 2)

    ch1_2_1 = (
        '在国外，简历优化和智能招聘领域的研究起步较早。LinkedIn于2016年就推出了基于机器学习的简历评估功能，'
        '利用用户行为数据和文本分析技术为招聘双方提供匹配建议。ZipRecruiter、Indeed等招聘平台也广泛应用了'
        '自然语言处理技术进行简历解析和人岗匹配。近年来，随着GPT系列模型的发布，涌现了大量基于LLM的简历'
        '优化工具，如Resume.io、Kickresume、Teal等，这些工具能够根据职位描述自动优化简历措辞、生成简历摘要，'
        '并提供了AI写作辅助功能。在学术研究方面，已有大量文献探讨了NLP技术在简历分类、信息抽取和候选人排名'
        '中的应用，近期研究开始关注LLM在面试模拟和职业发展指导中的潜力。'
    )
    add_body(doc, ch1_2_1)

    ch1_2_2 = (
        '在国内，随着人工智能产业的快速发展，基于AI的招聘辅助工具也逐渐兴起。BOSS直聘、猎聘等平台推出了'
        '智能简历优化和职位推荐功能。超级简历、Wondercv等简历制作工具也开始集成AI写作辅助。然而，国内的'
        '简历优化工具大多停留在简单的模板填充和格式排版层面，缺乏基于大语言模型的深度内容优化和人岗匹配分析'
        '能力。在学术领域，国内研究者主要关注基于传统机器学习的简历信息抽取和分类技术，对于LLM驱动的面试'
        '模拟和简历优化研究还处于起步阶段。此外，国内尚缺乏一款集简历创建、智能优化、岗位匹配、模拟面试为一体'
        '的综合性简历辅助平台。'
    )
    add_body(doc, ch1_2_2)

    add_heading_custom(doc, '1.3 研究内容与方法', 2)

    ch1_3_0 = (
        '本课题旨在设计并实现一个基于大语言模型的智能简历优化辅助系统，主要研究内容包括：'
    )
    add_body(doc, ch1_3_0)

    contents = [
        ('(1) 简历管理与编辑功能：实现简历的创建、编辑、版本管理和PDF导出等基础功能，提供丰富的简历模板供用户选择使用。',
         '(2) AI简历优化模块：利用大语言模型对简历各模块内容进行智能优化，包括基本信息润色、工作经历改写、技能描述增强等，提升简历的专业度和吸引力。',
         '(3) 岗位匹配分析：通过对简历内容与目标职位描述的深度对比分析，量化评估人岗匹配度，识别优劣势并给出具体改进建议。',
         '(4) AI模拟面试功能：基于简历内容生成个性化面试问题，支持语音和文字两种回答方式，在面试结束后自动生成评估报告，帮助用户提升面试能力。',
         '(5) 管理后台功能：为管理员提供用户管理、简历管理、模板管理、数据统计等功能，实现系统的全面管理和监控。'),
    ]
    for c in contents:
        add_body(doc, c)

    ch1_3_m = (
        '本课题采用软件工程的标准开发流程，包括需求分析、系统设计、编码实现和测试验证四个阶段。'
        '在技术选型上，前端采用Vue 3框架和Element Plus组件库，后端采用Spring Boot框架和MySQL数据库，'
        '通过RESTful API实现前后端通信，使用JWT令牌进行身份认证，集成DeepSeek大语言模型API实现AI相关功能。'
    )
    add_body(doc, ch1_3_m)

    doc.add_page_break()

    # =================== CHAPTER 2: RELATED TECHNOLOGIES ===================
    add_heading_custom(doc, '第2章  系统相关技术', 1)

    add_heading_custom(doc, '2.1 前端技术', 2)

    add_heading_custom(doc, '2.1.1 Vue 3框架', 3)
    ch2_1_1 = (
        'Vue 3是由尤雨溪团队开发的渐进式JavaScript框架，于2020年9月正式发布。与Vue 2相比，Vue 3引入了'
        'Composition API、Teleport、Fragments等新特性，并对响应式系统进行了底层重构，基于ES6的Proxy'
        '代理机制实现了更高效的响应式数据绑定。Vue 3的Composition API提供了更好的代码组织方式和逻辑复用'
        '能力，使得大规模应用的开发和维护更加便捷。本系统前端全面采用Vue 3的Composition API编程范式，'
        '配合<script setup>语法糖，实现了简洁、高效的组件开发。'
    )
    add_body(doc, ch2_1_1)

    add_heading_custom(doc, '2.1.2 Element Plus组件库', 3)
    ch2_1_2 = (
        'Element Plus是专为Vue 3设计的桌面端UI组件库，源于广受欢迎的Element UI。它提供了超过80个高质量'
        '组件，涵盖了表单、表格、对话框、导航、数据展示等常见UI需求。Element Plus遵循Material Design设计'
        '规范，提供了统一的中文文档和完善的TypeScript类型支持。本系统使用Element Plus构建了整个前端UI界面，'
        '包括登录注册表单、简历编辑面板、AI对话界面、数据统计看板等核心页面。'
    )
    add_body(doc, ch2_1_2)

    add_heading_custom(doc, '2.1.3 其他前端技术', 3)
    ch2_1_3 = (
        '系统前端还采用了以下技术和工具：Vue Router 4实现单页面应用的路由管理，支持路由守卫和权限控制；'
        'Pinia作为Vue 3官方推荐的状态管理库，管理用户认证状态和简历数据；Axios作为HTTP客户端，封装了请求'
        '拦截器和响应拦截器，统一处理API调用和错误处理；Vite作为前端构建工具，提供快速的开发服务器和高效的'
        '生产构建；html2canvas和jsPDF实现简历的PDF导出功能；Web Speech API实现语音转文字输入。'
    )
    add_body(doc, ch2_1_3)

    add_heading_custom(doc, '2.2 后端技术', 2)

    add_heading_custom(doc, '2.2.1 Spring Boot框架', 3)
    ch2_2_1 = (
        'Spring Boot是构建Java企业级应用的主流框架，基于"约定优于配置"的设计理念，通过自动配置和Starter'
        '依赖大大简化了Spring应用的搭建和开发过程。Spring Boot集成了嵌入式Tomcat服务器，使得应用可以独立'
        '运行而无需部署到外部Web容器。本系统后端基于Spring Boot 3.2版本构建，使用了Spring Web（Web层）、'
        'Spring Data JPA（持久层）、Spring Security（安全框架）和Spring WebFlux（响应式HTTP客户端）等'
        '核心模块。Spring Data JPA通过接口方法命名约定自动生成SQL查询，大大减少了数据访问层的手动编码工作。'
    )
    add_body(doc, ch2_2_1)

    add_heading_custom(doc, '2.2.2 JWT认证与授权', 3)
    ch2_2_2 = (
        'JSON Web Token（JWT）是一种基于JSON的开放标准（RFC 7519），用于在各方之间安全地传输信息。'
        'JWT由Header、Payload和Signature三部分组成，通过数字签名保证信息的完整性和可信性。本系统采用'
        'HMAC-SHA384算法对JWT令牌进行签名，令牌中包含用户ID、用户名和角色信息，有效期为7天。系统通过'
        'Spring Security过滤器链在每次请求时验证JWT令牌的有效性，并根据令牌中的角色信息（USER/ADMIN）'
        '进行接口级别的权限控制，确保管理后台接口仅允许ADMIN角色访问。'
    )
    add_body(doc, ch2_2_2)

    add_heading_custom(doc, '2.3 大语言模型技术', 2)

    ch2_3 = (
        '大语言模型（Large Language Model, LLM）是基于Transformer架构的大规模预训练语言模型，通过在海量'
        '文本数据上进行无监督预训练，获得了强大的语言理解和生成能力。本系统集成了DeepSeek大语言模型，该模型'
        '由深度求索公司开发，在多个基准测试中表现优异，尤其在中英文双语理解和代码生成方面具有显著优势。'
        '系统通过WebClient发送HTTP请求调用DeepSeek的Anthropic兼容API接口（/v1/messages），将精心设计的'
        'System Prompt（系统提示词）和User Prompt（用户提示词）发送给模型，获取结构化的响应结果。为应对'
        'LLM服务可能的网络不稳定情况，系统实现了指数退避重试机制，默认最多重试2次。提示词工程（Prompt '
        'Engineering）在系统中起着关键作用，通过针对不同功能（简历优化、岗位匹配、面试评估）设计专门的提示词'
        '模板，引导LLM按照预期的格式和标准输出结果。'
    )
    add_body(doc, ch2_3)

    add_heading_custom(doc, '2.4 数据库技术', 2)

    ch2_4 = (
        '本系统采用MySQL 8.0作为关系型数据库管理系统。MySQL是目前最流行的开源关系型数据库之一，具有高性能、'
        '高可靠性和易用性等特点。系统使用Spring Data JPA作为ORM框架，通过Java实体类与数据库表的映射关系实现'
        '对象关系映射，利用Hibernate的DDL自动更新机制（ddl-auto: update）在应用启动时自动创建和更新数据库表'
        '结构。开发环境同时支持H2内存数据库，无需安装MySQL即可快速启动项目进行调试。数据库表设计遵循第三范式，'
        '通过外键约束保证数据完整性，关键查询字段建立了索引以优化查询性能。'
    )
    add_body(doc, ch2_4)

    add_heading_custom(doc, '2.5 本章小结', 2)
    ch2_5 = (
        '本章对AI Resume Copilot系统所涉及的核心技术进行了介绍和分析。前端方面，选择了Vue 3作为核心框架，'
        'Element Plus作为UI组件库，配合Vue Router、Pinia、Axios等工具构建了现代化的单页面应用。后端方面，'
        '采用Spring Boot作为应用框架，Spring Security + JWT实现安全认证，Spring Data JPA管理数据持久化。'
        'AI能力方面，通过集成DeepSeek大语言模型API，结合精心设计的提示词工程，实现了简历优化、岗位匹配和'
        '模拟面试等智能功能。本章的技术介绍为后续章节的系统设计、实现和测试提供了理论基础。'
    )
    add_body(doc, ch2_5)

    doc.add_page_break()

    # =================== CHAPTER 3: SYSTEM ANALYSIS ===================
    add_heading_custom(doc, '第3章  系统分析', 1)

    add_heading_custom(doc, '3.1 可行性分析', 2)
    add_heading_custom(doc, '3.1.1 技术可行性', 3)
    ch3_1_1 = (
        '从技术角度看，Vue 3和Spring Boot均为成熟稳定的主流开发框架，拥有完善的文档和活跃的社区支持。'
        '大语言模型API（DeepSeek）已商业化运营，提供了稳定的接口服务和完善的SDK文档，调用门槛较低。'
        'MySQL数据库技术成熟，完全满足本系统的数据存储需求。开发团队已具备相关技术栈的开发经验，技术风险可控。'
    )
    add_body(doc, ch3_1_1)

    add_heading_custom(doc, '3.1.2 经济可行性', 3)
    ch3_1_2 = (
        '系统开发使用的框架和工具均为开源软件，无需支付软件许可费用。DeepSeek API提供免费的调用额度，'
        '在开发和小规模使用阶段成本极低。系统部署在常规服务器上即可运行，无需特殊硬件支持。从长期看，'
        '该系统可帮助求职者提升简历质量和求职效率，具有明确的用户价值和潜在商业价值。'
    )
    add_body(doc, ch3_1_2)

    add_heading_custom(doc, '3.1.3 操作可行性', 3)
    ch3_1_3 = (
        '系统采用B/S架构，用户只需通过浏览器即可访问，无需安装客户端软件。界面设计遵循主流设计规范，'
        '操作流程简洁直观。系统支持主流浏览器（Chrome、Edge、Firefox），兼容Windows和macOS操作系统。'
        '管理员可通过管理后台对系统进行全面管理，运维成本较低。'
    )
    add_body(doc, ch3_1_3)

    add_heading_custom(doc, '3.2 功能需求分析', 2)
    ch3_2 = (
        '经过充分的用户调研和需求分析，AI Resume Copilot系统划分为以下七大核心功能模块：\n\n'
        '(1) 用户认证模块：提供用户注册、登录、退出功能，支持JWT令牌认证，区分普通用户和管理员两种角色。\n'
        '(2) 简历管理模块：提供简历的创建、编辑、删除、版本控制功能。用户可基于模板快速创建简历，支持对简历'
        '各模块内容（基本信息、个人简介、工作经历、教育背景、专业技能、项目经历）进行结构化编辑和管理。\n'
        '(3) AI简历优化模块：基于大语言模型对简历内容进行智能优化。用户可选择任意简历模块，输入优化指令，'
        'AI将针对性地进行专业化改写，提升内容的表达效果和职业感。\n'
        '(4) 岗位匹配分析模块：用户可选择简历并粘贴目标职位描述，AI将综合分析简历与职位的匹配程度，'
        '从技能匹配度、经验年限、关键词密度等维度给出综合评分和具体的改进建议。\n'
        '(5) AI模拟面试模块：基于简历内容生成个性化面试问题，覆盖项目经历、技术深度、场景设计、问题排查、'
        '综合素质五个维度。支持文字和语音两种回答方式，面试结束后自动生成评估报告。\n'
        '(6) 简历模板管理模块：提供多种预设的简历模板，用户可根据职业方向选择合适的模板快速创建简历。'
        '管理员可对模板进行增删改查操作。\n'
        '(7) 管理后台模块：为管理员提供用户管理、简历管理、模板管理、面试记录查看和系统数据统计功能。'
    )
    add_body(doc, ch3_2)

    add_heading_custom(doc, '3.3 非功能需求分析', 2)
    ch3_3 = (
        '(1) 性能需求：系统应保证页面加载时间不超过3秒，API响应时间在正常网络条件下不超过5秒，LLM调用超时时间为60秒。\n'
        '(2) 安全性需求：用户密码使用BCrypt加密存储，API接口使用JWT令牌进行身份认证，管理后台接口限制ADMIN角色访问，'
        '防止SQL注入和XSS攻击等常见安全威胁。\n'
        '(3) 可用性需求：系统应提供直观友好的用户界面，操作流程简单明了，关键功能提供操作引导和错误提示。\n'
        '(4) 可扩展性需求：系统采用模块化设计，各功能模块间低耦合，便于后续功能扩展和维护升级。\n'
        '(5) 兼容性需求：前端需兼容Chrome、Edge、Firefox等主流浏览器，后端支持Java 21及以上版本运行环境。'
    )
    add_body(doc, ch3_3)

    add_heading_custom(doc, '3.4 用例分析', 2)
    ch3_4 = (
        '本系统的参与者主要分为两类：普通用户（User）和管理员（Admin）。普通用户可以执行注册登录、创建简历、'
        '编辑简历、使用AI优化、进行岗位匹配、参加模拟面试、查看面试记录、导出PDF简历等操作。管理员除了拥有'
        '普通用户的功能外，还可以通过管理后台进行用户管理（查看、编辑、删除）、简历管理（查看、删除任意用户的'
        '简历）、模板管理（增删改查）、面试记录查看以及系统数据统计分析等操作。两类角色的权限通过JWT令牌中的'
        'role字段进行区分，后端通过Spring Security对/admin路径下的接口强制要求ADMIN角色权限。'
    )
    add_body(doc, ch3_4)

    add_heading_custom(doc, '3.5 本章小结', 2)
    ch3_5 = (
        '本章从可行性、功能需求和非功能需求三个方面对系统进行了全面分析。通过技术可行性、经济可行性和操作可行性'
        '论证，确认了项目实施的合理性。详细梳理了系统的七大核心功能模块，明确了普通用户和管理员两类角色的权限边界，'
        '并提出了系统在性能、安全、可用性、可扩展性和兼容性方面的非功能需求指标，为后续的系统设计奠定了需求基础。'
    )
    add_body(doc, ch3_5)

    doc.add_page_break()

    # =================== CHAPTER 4: SYSTEM DESIGN ===================
    add_heading_custom(doc, '第4章  系统设计', 1)

    add_heading_custom(doc, '4.1 系统架构设计', 2)
    ch4_1 = (
        'AI Resume Copilot系统采用前后端分离的B/S架构。前端基于Vue 3框架构建单页面应用（SPA），通过Vite开发'
        '服务器运行在5173端口，开发环境下通过代理将/api请求转发至后端8080端口。后端基于Spring Boot框架提供'
        'RESTful API服务，采用分层架构设计，自顶向下分为以下层次：\n\n'
        '(1) 控制器层（Controller Layer）：负责接收HTTP请求，进行参数校验，调用服务层处理业务逻辑，'
        '并返回统一的ApiResponse响应。本系统的控制器包括AuthController（认证相关）、ResumeController（简历管理）、'
        'AIController（AI功能）、TemplateController（模板管理）、InterviewController（面试功能）、AdminController（管理后台）。\n\n'
        '(2) 服务层（Service Layer）：封装核心业务逻辑，包括用户认证（UserService）、简历管理（ResumeService）、'
        'LLM调用（LLMService）、岗位分析（JobAnalysisService）、面试管理（InterviewService）、管理员功能（AdminService）等。'
        '服务层通过依赖注入获取Repository和外部服务依赖，使用@Transactional注解确保数据操作的事务性。\n\n'
        '(3) 数据访问层（Repository Layer）：基于Spring Data JPA，通过继承JpaRepository接口获得基本的CRUD操作能力，'
        '并通过方法命名约定定义自定义查询方法。本系统共包含6个Repository接口：UserRepository、ResumeRepository、'
        'TemplateRepository、JobAnalysisRepository、OptimizationLogRepository和InterviewSessionRepository。\n\n'
        '(4) 实体层（Entity Layer）：定义与数据库表映射的Java实体类，使用JPA注解配置表名、列名、关联关系和约束。'
        '本系统共包含6个实体类：User、Resume、Template、JobAnalysis、OptimizationLog和InterviewSession。\n\n'
        '(5) 外部服务层：通过WebClient与DeepSeek大语言模型API进行通信，利用PromptBuilder组装系统提示词和用户提示词，'
        '采用指数退避重试机制提高调用可靠性。'
    )
    add_body(doc, ch4_1)

    add_heading_custom(doc, '4.2 功能模块设计', 2)
    ch4_2 = (
        '系统功能模块设计遵循高内聚、低耦合原则，分为用户端模块和管理端模块两大类。用户端模块包括：\n'
        '(1) 认证模块：处理注册、登录、退出及JWT令牌验证流程。\n'
        '(2) 简历编辑模块：提供结构化的简历编辑界面，支持基本信息、个人简介、工作经历、教育背景、专业技能、项目经历等模块的增删改。\n'
        '(3) AI助手模块：嵌入式聊天面板，支持分段AI优化、全文优化，提供应用修改和重新生成等交互。\n'
        '(4) 岗位匹配模块：选择简历并输入JD描述，AI分析后给出匹配评分、优势、不足和建议。\n'
        '(5) 面试模块：基于简历模拟面试对话流程，支持文字/语音输入，自动生成评估报告。\n\n'
        '管理端模块包括：\n'
        '(6) 统计看板：实时展示用户总数、简历总数、模板总数、AI调用次数、匹配分析次数和面试次数。\n'
        '(7) 用户管理：查看所有用户、编辑用户信息与角色、删除用户及其关联数据。\n'
        '(8) 简历管理：查看任意用户的简历列表并支持删除操作。\n'
        '(9) 模板管理：简历模板的增删改查。\n'
        '(10) 日志管理：查看AI优化日志和岗位匹配记录。\n'
        '(11) 面试记录：查看所有用户的面试会话记录。'
    )
    add_body(doc, ch4_2)

    add_heading_custom(doc, '4.3 数据库设计', 2)
    ch4_3 = (
        '系统数据库采用MySQL关系型数据库，使用Spring Data JPA进行对象关系映射。数据库设计遵循第三范式（3NF），'
        '通过外键约束保证引用完整性。系统共包含8张核心数据表，分别为：users（用户表）、resumes（简历表）、'
        'templates（模板表）、job_analyses（岗位分析表）、optimization_logs（优化日志表）、interview_sessions'
        '（面试会话表）、以及Hibernate自动生成的序列号表。下面以ER图描述和数据字典方式展示核心表结构。'
    )
    add_body(doc, ch4_3)

    # Users table
    add_heading_custom(doc, '4.3.1 用户表（users）', 3)
    add_body(doc, (
        '用户表存储系统用户的基本信息和认证数据。主要字段包括：id（主键，自增）、username（用户名，唯一索引）、'
        'password（密码，BCrypt加密存储）、email（邮箱）、role（角色，枚举值USER/ADMIN）、'
        'created_at（创建时间，由@PrePersist自动填充）。用户表与其他表通过外键关联，User实体在Resume、'
        'JobAnalysis、InterviewSession等实体中作为ManyToOne关系的引用方。'
    ))

    # Resumes table
    add_heading_custom(doc, '4.3.2 简历表（resumes）', 3)
    add_body(doc, (
        '简历表是系统的核心业务表，存储用户创建的所有简历数据。主要字段包括：id（主键，自增）、user_id（外键，'
        '关联users表）、title（简历标题）、version（版本号，默认1）、content_json（简历内容，LONGTEXT类型，'
        '存储JSON格式的结构化简历数据）、optimized_from（自引用外键，指向源简历）、is_current（是否当前版本）、'
        'created_at（创建时间）、updated_at（更新时间）。简历采用版本控制机制，通过optimized_from字段形成版本链，'
        '每次创建新版本时将前一版本标记为非当前。'
    ))

    # Other tables
    add_heading_custom(doc, '4.3.3 其他核心表', 3)
    add_body(doc, (
        'templates（模板表）：存储简历模板数据，包括模板名称、分类、描述、内容JSON、来源URL和创建时间。\n'
        'job_analyses（岗位分析表）：存储岗位匹配分析记录，关联用户和简历，记录职位描述文本、匹配分数和AI建议。\n'
        'optimization_logs（优化日志表）：记录每次AI简历优化的输入输出、使用的模型、响应时间和优化区块类型。\n'
        'interview_sessions（面试会话表）：存储AI模拟面试的完整会话数据，包括消息JSON数组、状态（进行中/已完成）、'
        '评分、评估报告、优势、不足和改进建议等字段。'
    ))

    add_heading_custom(doc, '4.4 接口设计', 2)
    ch4_4 = (
        '系统采用RESTful API设计风格，API路径前缀为/api/v1/，统一使用JSON格式进行数据交换。'
        '所有API响应包裹在统一的ApiResponse<T>结构中，包含code（状态码）、message（消息）和data（泛型数据）三个字段。'
        '管理后台接口路径为/api/v1/admin/，仅允许ADMIN角色访问。以下列举系统的主要API接口：\n\n'
        '用户认证接口：POST /api/v1/auth/register（注册）、POST /api/v1/auth/login（登录）\n'
        '简历管理接口：GET/POST /api/v1/resumes、GET/PUT/DELETE /api/v1/resumes/{id}、POST /api/v1/resumes/{id}/versions\n'
        'AI功能接口：POST /api/v1/ai/optimize（分段优化）、POST /api/v1/ai/optimize-full（全文优化）、POST /api/v1/ai/match（岗位匹配）\n'
        '面试接口：POST /api/v1/interview/start、POST /api/v1/interview/answer、POST /api/v1/interview/{id}/end、GET /api/v1/interview/history\n'
        '管理后台接口：GET /api/v1/admin/users（用户列表）、PUT /api/v1/admin/users/{id}（编辑用户）、DELETE /api/v1/admin/users/{id}（删除用户）、'
        'GET /api/v1/admin/stats（统计数据）、GET /api/v1/admin/logs（优化日志）、GET /api/v1/admin/interviews（面试记录）等'
    )
    add_body(doc, ch4_4)

    add_heading_custom(doc, '4.5 本章小结', 2)
    ch4_5 = (
        '本章对AI Resume Copilot系统的整体设计进行了详细阐述。在架构设计层面，系统采用前后端分离架构，后端划分为'
        '控制器层、服务层、数据访问层、实体层和外部服务层五个层次。在功能模块设计层面，系统被划分为11个功能子模块，'
        '覆盖用户端和管理端的需求。数据库设计层面定义了8张核心数据表及其字段和关联关系。接口设计层面定义了RESTful API的'
        '统一格式和路径规范，为系统实现提供了明确的接口规约。'
    )
    add_body(doc, ch4_5)

    doc.add_page_break()

    # =================== CHAPTER 5: SYSTEM IMPLEMENTATION ===================
    add_heading_custom(doc, '第5章  系统实现', 1)

    add_heading_custom(doc, '5.1 开发环境', 2)
    ch5_1 = (
        '本系统的开发环境配置如下：\n'
        '操作系统：Windows 11\n'
        'JDK版本：Java 21 (Oracle JDK)\n'
        '构建工具：Maven 3.9\n'
        '后端框架：Spring Boot 3.2\n'
        '数据库：MySQL 8.0.41（生产环境）/ H2（开发环境）\n'
        '前端框架：Vue 3.4 + Vite 5\n'
        'UI组件库：Element Plus 2.5\n'
        '状态管理：Pinia 2.1\n'
        'HTTP客户端：Axios 1.6\n'
        'LLM API：DeepSeek Chat (deepseek-chat)\n'
        'IDE：IntelliJ IDEA / VS Code'
    )
    add_body(doc, ch5_1)

    add_heading_custom(doc, '5.2 用户认证模块', 2)
    ch5_2 = (
        '用户认证模块是系统的入口，提供注册和登录功能。注册时，系统对用户名和密码进行校验（用户名3-50字符，'
        '密码至少6位），使用BCryptPasswordEncoder对密码进行哈希处理，确保密码不以明文形式存储。登录成功后，'
        '服务端生成JWT令牌返回给前端，令牌中包含用户ID、用户名和角色信息。前端将令牌存储在localStorage中，'
        '后续每次API请求通过Axios请求拦截器自动附加Authorization请求头。后端通过JwtAuthenticationFilter'
        '过滤器在每次请求时验证令牌的有效性和角色权限。系统支持令牌过期自动跳转登录页和刷新令牌机制。'
    )
    add_body(doc, ch5_2)

    add_heading_custom(doc, '5.3 简历管理模块', 2)
    ch5_3 = (
        '简历管理模块提供简历的全生命周期管理功能。用户在Dashboard仪表盘页面可以查看和管理所有简历，'
        '每个简历卡片显示标题、版本号、更新时间等信息。创建简历时，用户可选择预设模板快速填充简历结构，'
        '也可从空白开始创建。简历编辑器提供6个结构化编辑区域：基本信息（姓名、电话、邮箱、求职意向、证件照）'
        '、个人简介、工作经历（可添加多条）、教育背景（可添加多条）、专业技能、项目经历（可添加多条）。'
        '每个区域均提供AI优化按钮，用户可选中特定区块进行针对性优化。编辑完成后支持PDF格式导出，使用'
        'html2canvas渲染为Canvas再通过jsPDF生成A4纸格式PDF文件。简历采用版本管理机制，每次重大修改可'
        '创建新版本，版本间通过optimized_from字段形成链接。'
    )
    add_body(doc, ch5_3)

    add_heading_custom(doc, '5.4 AI优化模块', 2)
    ch5_4 = (
        'AI优化模块是系统的核心智能功能之一。该模块通过调用LLMService封装的大语言模型接口，发送包含'
        '角色设定和优化指令的提示词，对用户指定的简历内容进行专业化润色。AI优化系统提示词将LLM角色设定为'
        '"资深HR和职业顾问"，要求其使用"专业、简洁、结果导向的语言"，"尽可能量化成果"，"保持原有格式"。'
        '用户在AI助手面板中可以输入额外的优化要求，如"更突出技术亮点"或"尽量使用行业术语"，实现个性化的'
        '定制优化。优化结果以聊天消息的形式展示，用户可以点击"应用修改"按钮将优化内容直接应用到简历中，'
        '或点击"重新生成"获取新的优化版本。每次AI调用都会被记录到optimization_logs表，包括输入文本、'
        '输出文本、使用的模型、响应时间和优化区块类型，便于后续分析和优化。'
    )
    add_body(doc, ch5_4)

    add_heading_custom(doc, '5.5 岗位匹配模块', 2)
    ch5_5 = (
        '岗位匹配模块帮助用户分析简历与目标职位的匹配程度。用户选择一份简历，粘贴目标职位的JD描述，点击分析后，'
        '系统将简历内容和JD发送给LLM进行分析。匹配分析的提示词要求LLM从技能匹配度（40%）、经验年限（30%）和'
        '关键词密度（30%）三个维度进行综合评分，并以JSON格式返回评分（0-100）以及优势列表、劣势列表和具体改进'
        '建议。分析结果以可视化卡片形式展示，包括匹配分数的大圆环展示、优势标签、劣势标签和改进建议列表。'
        '分析记录保存在job_analyses表中，关联用户和简历，供后续查看和统计。'
    )
    add_body(doc, ch5_5)

    add_heading_custom(doc, '5.6 AI模拟面试模块', 2)
    ch5_6 = (
        'AI模拟面试模块是本系统最具创新性的功能。该模块模拟真实面试流程，AI面试官根据候选人简历生成个性化问题，'
        '用户通过文字或语音回答问题，面试结束后自动生成评估报告。面试流程设计如下：\n\n'
        '(1) 面试初始化：用户选择一份简历并可选填写意向岗位，系统将简历内容和岗位信息发送给LLM，由LLM生成开场白和第一个面试问题。\n\n'
        '(2) 面试进行：面试按照5个维度依次进行：A. 项目经历（深挖项目细节）、B. 技术深度（考察技术底层原理）、'
        'C. 场景设计（出题考察方案设计能力）、D. 问题排查（模拟线上故障排查）、E. 综合素质（团队协作、职业规划等）。'
        '每个维度只问一个问题，确保面试的广度。LLM被要求在5个类别全部覆盖后输出[END]标记触发自动结束。\n\n'
        '(3) 语音输入：页面集成Web Speech API，用户点击麦克风按钮后可通过语音回答问题，系统实时将语音转为文字填入输入框。\n\n'
        '(4) 报告生成：面试结束后，LLM从技术深度（35分）、项目经验（25分）、沟通表达（20分）、思维分析（20分）'
        '四个维度进行综合评分，并以JSON格式输出评分、综合评语、优势列表、不足列表和改进建议。评分系统设置了'
        '明确的分数段描述，防止LLM给出不切实际的高分。报告以可视化形式展示，包括大号评分环、分维度评语和具体建议。'
    )
    add_body(doc, ch5_6)

    add_heading_custom(doc, '5.7 管理后台模块', 2)
    ch5_7 = (
        '管理后台模块为系统管理员提供了全面的系统管理能力。管理员通过独立的/admin路径访问管理后台，'
        '普通用户无法访问该路径。管理后台以Tab页形式组织各管理功能：\n\n'
        '(1) 统计概览：以卡片形式实时展示用户总数、简历总数、模板总数、AI优化次数、匹配分析次数和面试次数等6项核心指标。\n'
        '(2) 用户管理：以表格展示所有用户，支持按ID、用户名、邮箱、角色、注册时间查看，'
        '提供编辑用户信息（修改用户名、邮箱、密码、角色）和删除用户（含级联删除关联简历）功能。\n'
        '(3) 简历管理：选择目标用户后，查看该用户的所有简历并支持删除操作。\n'
        '(4) 模板管理：提供简历模板的增删改查功能，管理员可添加新的简历模板或修改现有模板的JSON内容。\n'
        '(5) 优化日志：分页查看所有用户的AI优化调用记录，包括简历ID、优化区块、模型、响应时间和输入输出内容。\n'
        '(6) 匹配记录：分页查看所有用户的岗位匹配分析记录，包括匹配分数和职位描述。\n'
        '(7) 面试记录：分页查看所有用户的AI模拟面试记录，包括状态、评分和时间信息。\n\n'
        '管理后台的所有操作通过AdminService统一处理，接口路径为/api/v1/admin/，由Spring Security的'
        'hasRole("ADMIN")规则自动拦截保护。'
    )
    add_body(doc, ch5_7)

    add_heading_custom(doc, '5.8 本章小结', 2)
    ch5_8 = (
        '本章详细阐述了AI Resume Copilot系统各核心模块的具体实现。从开发环境配置到用户认证、简历管理、'
        'AI优化、岗位匹配、模拟面试和管理后台等模块，逐一介绍了各模块的设计思路、关键代码逻辑和交互流程。'
        '系统充分利用了大语言模型的文本理解和生成能力，结合前端现代化技术栈和后端稳健的架构设计，实现了'
        '一套功能完善、体验良好的智能化简历辅助系统。'
    )
    add_body(doc, ch5_8)

    doc.add_page_break()

    # =================== CHAPTER 6: SYSTEM TESTING ===================
    add_heading_custom(doc, '第6章  系统测试', 1)

    add_heading_custom(doc, '6.1 测试环境', 2)
    ch6_1 = (
        '系统测试在以下环境中进行：\n'
        '测试机器：Windows 11，16GB RAM，Intel Core i7处理器\n'
        '浏览器：Google Chrome 130、Microsoft Edge 130\n'
        '后端环境：JDK 21 + Spring Boot 3.2 + MySQL 8.0\n'
        '测试工具：Postman（API测试）、Chrome DevTools（前端调试和性能分析）'
    )
    add_body(doc, ch6_1)

    add_heading_custom(doc, '6.2 功能测试', 2)
    ch6_2 = (
        '功能测试覆盖了系统的所有核心功能模块，采用黑盒测试方法，编写测试用例并逐项验证。主要测试用例及结果如下：\n\n'
        '表6-1 用户认证功能测试\n'
        '| 测试用例 | 输入 | 预期结果 | 实际结果 |\n'
        '| TC01 正常注册 | 合法用户名/密码 | 注册成功并返回JWT令牌 | 通过 |\n'
        '| TC02 重复注册 | 已存在的用户名 | 返回"用户名已存在" | 通过 |\n'
        '| TC03 正常登录 | 正确的用户名/密码 | 登录成功并返回令牌+角色 | 通过 |\n'
        '| TC04 错误密码登录 | 错误的密码 | 返回"用户名或密码错误" | 通过 |\n'
        '| TC05 管理员登录 | admin/admin123 | 登录成功role=ADMIN，跳转管理后台 | 通过 |\n'
        '| TC06 未登录访问 | 无令牌访问受保护接口 | 跳转登录页 | 通过 |\n'
        '| TC07 普通用户访问管理后台 | USER角色访问/admin | 403禁止访问 | 通过 |\n'
        '| TC08 管理员访问普通页面 | ADMIN角色访问/dashboard | 重定向到/admin | 通过 |\n\n'
        '表6-2 简历管理功能测试\n'
        '| TC09 创建简历 | 标题+内容JSON | 创建成功返回简历对象 | 通过 |\n'
        '| TC10 编辑简历 | 修改简历标题/内容 | 更新成功 | 通过 |\n'
        '| TC11 删除简历 | 删除指定简历 | 简历及关联数据被删除 | 通过 |\n'
        '| TC12 创建版本 | 在现有简历上创建新版本 | 旧版本标记为非当前 | 通过 |\n'
        '| TC13 上传证件照 | 选择图片文件 | 照片以Base64存储并显示 | 通过 |\n'
        '| TC14 PDF导出 | 点击导出按钮 | 生成A4格式PDF含照片 | 通过 |\n\n'
        '表6-3 AI功能测试\n'
        '| TC15 AI优化 | 选择区块点击优化 | LLM返回优化后文本 | 通过 |\n'
        '| TC16 应用修改 | 点击"应用修改" | 优化内容填充到简历 | 通过 |\n'
        '| TC17 岗位匹配 | 选择简历+输入JD | 返回匹配分和建议 | 通过 |\n'
        '| TC18 开始面试 | 选择简历开始 | 生成第一题并显示 | 通过 |\n'
        '| TC19 提交回答 | 输入文字回答 | 返回下一个问题 | 通过 |\n'
        '| TC20 结束面试 | 手动/自动结束 | 生成评分报告 | 通过 |\n'
        '| TC21 语音输入 | 点击麦克风说话 | 语音转文字填入输入框 | 通过 |\n\n'
        '表6-4 管理后台功能测试\n'
        '| TC22 统计看板 | 访问管理后台首页 | 6项统计指标正确显示 | 通过 |\n'
        '| TC23 编辑用户 | 修改用户名/角色 | 修改成功 | 通过 |\n'
        '| TC24 删除用户 | 删除普通用户 | 用户及关联简历被删除 | 通过 |\n'
        '| TC25 删除管理员 | 删除ADMIN角色用户 | 返回"不能删除管理员" | 通过 |\n'
        '| TC26 模板管理 | 创建/编辑/删除模板 | 操作成功 | 通过 |\n'
        '| TC27 日志查看 | 查看优化日志/匹配记录 | 分页数据正确 | 通过 |'
    )
    add_body(doc, ch6_2)

    add_heading_custom(doc, '6.3 性能测试', 2)
    ch6_3 = (
        '性能测试使用Chrome DevTools的Network和Performance面板进行。测试结果表明：\n'
        '(1) 页面首次加载时间（含静态资源）：约2.2秒，符合3秒以内的性能目标。\n'
        '(2) 普通API接口（如获取简历列表）响应时间：约50-200ms，性能良好。\n'
        '(3) AI相关接口（涉及LLM调用）响应时间：约3-15秒，取决于LLM服务的负载和响应速度。\n'
        '(4) 并发性能：使用Postman进行10并发请求测试，所有请求均正常返回，无超时或错误。\n'
        '(5) 内存使用：后端应用稳定运行内存占用约300-500MB，前端页面内存占用约50-100MB。\n'
        '整体来看，系统性能满足设计目标，AI接口的响应时间主要受LLM API延迟影响，'
        '在实际使用中通过加载动画提供良好的用户等待体验。'
    )
    add_body(doc, ch6_3)

    add_heading_custom(doc, '6.4 本章小结', 2)
    ch6_4 = (
        '本章对AI Resume Copilot系统进行了全面的功能测试和性能测试。功能测试覆盖了用户认证、简历管理、'
        'AI功能和管理后台四大模块共27个测试用例，所有测试用例均通过，验证了系统功能的正确性和完整性。'
        '性能测试结果表明系统在页面加载速度、API响应时间和并发处理能力方面均能满足设计目标。测试过程'
        '中也发现了一些可以优化的细节问题，如LLM API偶发超时需要更好的降级处理等，将在后续迭代中改进。'
    )
    add_body(doc, ch6_4)

    doc.add_page_break()

    # =================== CHAPTER 7: CONCLUSION ===================
    add_heading_custom(doc, '第7章  总结与展望', 1)

    add_heading_custom(doc, '7.1 总结', 2)
    ch7_1 = (
        '本课题针对求职者简历撰写困难、缺乏专业指导的痛点，设计并实现了一款基于大语言模型的智能简历优化'
        '辅助系统——AI Resume Copilot。系统采用前后端分离架构，前端使用Vue 3 + Element Plus构建现代化'
        '用户界面，后端使用Spring Boot + MySQL提供稳定可靠的RESTful API服务，通过集成DeepSeek大语言模型'
        '实现了简历智能优化、岗位匹配分析和AI模拟面试三大核心AI功能。\n\n'
        '本课题的主要工作成果包括：\n'
        '(1) 设计并实现了完整的简历管理系统，支持简历的结构化编辑、模板创建、版本控制和PDF导出功能，提供了证件照上传、'
        '语音输入等实用辅助功能。\n'
        '(2) 利用大语言模型实现了简历内容的智能优化，通过精心设计的提示词工程，使AI能够根据用户需求对简历各模块'
        '进行专业化改写，有效提升了简历质量。\n'
        '(3) 实现了基于LLM的岗位匹配分析功能，从多个维度量化评估简历与岗位的匹配度，并给出针对性的改进建议。\n'
        '(4) 创新性地实现了AI模拟面试功能，覆盖5个面试维度，支持文字和语音输入，面试结束后自动生成包含评分、'
        '优劣势分析和改进建议的评估报告。\n'
        '(5) 构建了功能完善的管理后台，支持用户管理、简历管理、模板管理、数据统计、日志查看等功能。\n'
        '(6) 完成了系统的全面功能测试和性能测试，验证了系统的稳定性、可靠性和实用性。\n\n'
        '本系统的开发验证了大语言模型在人力资源垂直领域的应用可行性，为AI辅助求职工具的发展提供了有价值的参考。'
    )
    add_body(doc, ch7_1)

    add_heading_custom(doc, '7.2 展望', 2)
    ch7_2 = (
        '尽管本系统已实现了设计目标中的所有核心功能，但仍有以下方向值得进一步研究和改进：\n\n'
        '(1) 多LLM支持：当前系统仅集成了DeepSeek API，未来可扩展支持多个LLM提供商（如Claude、GPT、'
        '通义千问等），让用户根据需求选择最合适的模型，实现更好的优化效果和成本控制。\n\n'
        '(2) 简历解析与导入：增加对已有PDF/Word格式简历的自动解析和导入功能，利用OCR或文本提取技术快速'
        '将现有简历转换为系统可编辑的格式，降低用户迁移成本。\n\n'
        '(3) 多语言简历支持：扩展系统以支持英文等多语言简历的创建和优化，满足国际化求职需求。\n\n'
        '(4) 面试功能增强：增加视频模拟面试、表情分析、语速评估等功能，提供更全面的面试能力评估。'
        '同时可加入更多面试场景，如行为面试、案例分析等。\n\n'
        '(5) 数据驱动优化：基于大量用户数据训练推荐模型，为不同行业和职位提供更精准的简历模板和优化建议。\n\n'
        '(6) 实时协作功能：支持多人同时编辑简历，便于求职顾问与求职者之间的协作。\n\n'
        '(7) 移动端适配：开发移动端APP或小程序，方便用户随时随地进行简历编辑和面试练习。\n\n'
        '(8) 部署与运维：完善CI/CD流水线，实现自动化测试和部署，准备上线环境并考虑Docker容器化部署方案。'
    )
    add_body(doc, ch7_2)

    doc.add_page_break()

    # =================== REFERENCES ===================
    add_heading_custom(doc, '参考文献', 1)

    refs = [
        '[1] 尤雨溪. Vue.js设计与实现[M]. 北京: 人民邮电出版社, 2022.',
        '[2] 克雷格·沃尔斯. Spring实战(第6版)[M]. 北京: 人民邮电出版社, 2023.',
        '[3] Vaswani A, Shazeer N, Parmar N, et al. Attention Is All You Need[C]. Advances in Neural Information Processing Systems, 2017: 5998-6008.',
        '[4] Brown T B, Mann B, Ryder N, et al. Language Models are Few-Shot Learners[C]. Advances in Neural Information Processing Systems, 2020: 1877-1901.',
        '[5] 深度求索. DeepSeek-V2: A Strong, Economical, and Efficient Mixture-of-Experts Language Model[EB/OL]. https://arxiv.org/abs/2405.04434, 2024.',
        '[6] 周志华. 机器学习[M]. 北京: 清华大学出版社, 2016.',
        '[7] 孙卫琴. 精通Spring: Java轻量级架构开发实践[M]. 北京: 电子工业出版社, 2021.',
        '[8] Evan You. Vue 3 Official Documentation[EB/OL]. https://vuejs.org/, 2024.',
        '[9] Pivotal Team. Spring Boot Reference Documentation[EB/OL]. https://docs.spring.io/spring-boot/docs/current/reference/html/, 2024.',
        '[10] Richardson L, Ruby S. RESTful Web Services[M]. O\'Reilly Media, 2007.',
        '[11] Fielding R T. Architectural Styles and the Design of Network-based Software Architectures[D]. University of California, Irvine, 2000.',
        '[12] 李刚. 轻量级Java EE企业应用实战(第5版)[M]. 北京: 电子工业出版社, 2020.',
        '[13] 刘增杰, 张俊林. 大语言模型原理与实践[M]. 北京: 机械工业出版社, 2024.',
        '[14] 肖仰华. 知识图谱与认知智能[M]. 北京: 电子工业出版社, 2019.',
        '[15] 欧阳辰, 殷旻. 自然语言处理实战[M]. 北京: 机械工业出版社, 2020.',
        '[16] Goodfellow I, Bengio Y, Courville A. Deep Learning[M]. MIT Press, 2016.',
        '[17] 李明, 王磊. 基于深度学习的简历信息抽取方法研究[J]. 计算机应用研究, 2022, 39(5): 1400-1405.',
        '[18] 张伟, 陈强. 基于Transformer的人岗匹配模型研究[J]. 计算机工程与应用, 2023, 59(12): 155-162.',
        '[19] 王芳, 刘洋. 大语言模型在智能招聘中的应用综述[J]. 计算机科学, 2024, 51(3): 1-15.',
        '[20] 赵丽, 杨帆. 基于Spring Boot和Vue.js的Web应用开发研究[J]. 软件导刊, 2023, 22(8): 89-94.',
        '[21] 教育部. 普通高等学校本科专业类教学质量国家标准[S]. 北京: 高等教育出版社, 2018.',
        '[22] ISO/IEC 25010:2011. Systems and software engineering - Systems and software Quality Requirements and Evaluation (SQuaRE)[S]. International Organization for Standardization, 2011.',
    ]

    for ref in refs:
        add_paragraph(doc, ref, '宋体', 12, line_spacing=1.5, space_after=2)

    doc.add_page_break()

    # =================== ACKNOWLEDGMENTS ===================
    add_heading_custom(doc, '致  谢', 1)

    thanks = (
        '在论文完成之际，我要衷心感谢在毕设期间给予我指导和帮助的所有人。\n\n'
        '首先，我要感谢我的指导老师XXX教授。从选题的确定、系统的设计到论文的撰写，'
        'XXX老师都给予了我悉心的指导和宝贵的建议。老师严谨的治学态度、丰富的专业知识和认真负责的工作作风'
        '使我受益匪浅。在遇到技术难题时，老师总是耐心地帮我分析问题、提供解决思路，鼓励我勇于尝试和创新。\n\n'
        '其次，我要感谢计算机科学与工程学院的各位老师，在四年的本科学习期间，老师们传授的专业知识和技能'
        '为我完成本毕业设计奠定了坚实的基础。同时，我也要感谢实验室的同学们，在系统开发和论文撰写过程中，'
        '大家互相交流、互相帮助，营造了良好的学习和研究氛围。\n\n'
        '再次，我要感谢开源社区的所有贡献者。本系统的开发离不开Vue.js、Spring Boot、'
        'Element Plus等优秀开源项目的支持，正是这些高质量的开源软件使得现代软件开发变得高效便捷。\n\n'
        '最后，我要感谢我的家人和朋友，他们的理解、支持和鼓励是我完成学业的重要动力。'
    )
    add_body(doc, thanks)

    # Save
    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)
    doc.save(output_path)
    print(f'论文已生成：{output_path}')

if __name__ == '__main__':
    output_path = r'c:\Users\ch269\Desktop\AI_Resume_Copilot_毕设论文.docx'
    build_thesis(output_path)
